#!/usr/bin/env python3
import json
import base64
import email
import imaplib
import io
import mimetypes
import os
import re
import smtplib
import ssl
import sys
import threading
import time
import urllib.parse
import urllib.request
import zipfile
import shutil
import secrets
from copy import deepcopy
from datetime import datetime
from email.header import decode_header
from email.message import EmailMessage
from email.utils import parsedate_to_datetime
from html import escape
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from xml.etree import ElementTree as ET


BASE_DIR = Path(__file__).resolve().parent
WORKSPACE = BASE_DIR.parent
REPORT_DIR = WORKSPACE / "周报"
DRAFT_DIR = BASE_DIR / "drafts"
GENERATED_DIR = BASE_DIR / "generated"
USER_DATA_DIR = BASE_DIR / "user_data"
CONFIG_PATH = BASE_DIR / "config.json"
AGENT_CONFIG_PATH = BASE_DIR / "agent_config.json"
APP_RELATIVE_PATH = "/personal-office-assistant"
SESSIONS = {}
DEFAULT_ASSISTANT_PROMPT = (
    "请帮我优化下面的工作内容，要求：\n"
    "1、拆分成1、2、3、4这样的编号要点，每点单独换行。\n"
    "2、语言简洁明了，体现实际工作量和推进成果。\n"
    "3、修正错别字、病句和不通顺表达。\n"
    "4、不要编造不存在的事项，不要写空话套话。"
)

DEFAULT_EMAIL_SIGNATURE = """\n
周颖超

--------------------------------------------------------------------------------------
湖南承希科技有限公司  | Hunan Chency Technology Co.,Ltd.
--------------------------------------------------------------------------------------
Mobile: 185-2961-2716
E-mail: zhouyingchao@chencytech.com
Add: 湖南省长沙市岳麓区军民融合产业园1栋B座15A层
------------------------------------------------------------------------

本邮件及其附件含有保密信息，仅限于发送给上面地址中列出的个人或群组。禁止任何其他人以任何形式使用（包括但不限全部或部分地泄露、复制、或散发）本邮件中的信息。如果您错收了本邮件，请您立即电话或邮件通知发件人并删除本邮件！
This e-mail and its attachments contain confidential information and are intended only for the individual or group of individuals listed at the address above. Any other pere this email!"""


def safe_username(username):
    return re.sub(r"[^A-Za-z0-9_@.-]+", "_", str(username or "default")).strip("._") or "default"


def user_root(username):
    return USER_DATA_DIR / safe_username(username)


def user_report_dir(username):
    return user_root(username) / "reports"


def user_generated_dir(username):
    return user_root(username) / "generated"


def user_draft_dir(username):
    return user_root(username) / "drafts"


def user_profile_dir(username):
    return user_root(username) / "profile"


def user_diary_dir(username):
    return user_root(username) / "diary"


def user_forum_dir(username):
    return user_root(username) / "forum"


def read_config():
    config = {
        "leader_email": "",
        "cc": "",
        "sender_name": "周颖超",
        "weekly_greeting": "领导您好：",
        "trip_greeting": "领导您好：",
        "signature": "周颖超",
        "email_signature": DEFAULT_EMAIL_SIGNATURE,
        "assistant_prompt": DEFAULT_ASSISTANT_PROMPT,
        "assistant_api_url": "",
        "assistant_api_key": "",
        "assistant_model": "MiniMax-M2.7",
        "smtp_host": "smtp.263.net",
        "imap_host": "imap.263.net",
        "news_sources": [],
        "news_search_query": "轨道交通 OR 城市轨道 OR 地铁 OR 智慧轨交 OR 轨道交通安全",
        "news_auto_search": True,
        "news_auto_push": True,
        "news_push_time": "08:30",
        "users": [
            {"username": "admin", "password": "admin123", "role": "admin", "name": "管理员"},
            {"username": "member", "password": "member123", "role": "member", "name": "成员"},
        ],
        "user_mail_settings": {},
    }
    if CONFIG_PATH.exists():
        try:
            config.update(json.loads(CONFIG_PATH.read_text(encoding="utf-8")))
        except json.JSONDecodeError:
            pass
    return config


def write_config(config):
    CONFIG_PATH.write_text(json.dumps(config, ensure_ascii=False, indent=2), encoding="utf-8")


def read_agent_config():
    if not AGENT_CONFIG_PATH.exists():
        return {}
    try:
        return json.loads(AGENT_CONFIG_PATH.read_text(encoding="utf-8"))
    except Exception:
        return {}


def write_agent_config(cfg):
    AGENT_CONFIG_PATH.write_text(json.dumps(cfg, ensure_ascii=False, indent=2), encoding="utf-8")


def public_user(user):
    role = user.get("role", "member")
    username = user.get("username", "")
    avatar_url = user.get("avatar_url", "")
    if not avatar_url and username and (user_profile_dir(username) / "avatar.png").exists():
        avatar_url = f"/user-avatar/{urllib.parse.quote(username)}.png?v={int((user_profile_dir(username) / 'avatar.png').stat().st_mtime)}"
    return {
        "username": username,
        "role": role,
        "name": user.get("name") or username,
        "avatar_url": avatar_url,
        "bio": user.get("bio", ""),
        "hobbies": user.get("hobbies", ""),
        "is_admin": role in ("admin", "superadmin"),
        "is_superadmin": role == "superadmin",
    }


def ensure_user_space(username):
    for path in (user_report_dir(username), user_generated_dir(username), user_draft_dir(username), user_profile_dir(username)):
        path.mkdir(parents=True, exist_ok=True)


def find_user(username, password=None):
    for user in read_config().get("users", []):
        if user.get("username") != username:
            continue
        if password is None or str(user.get("password", "")) == str(password):
            return user
    return None


def default_email_signature_for_user(username):
    user = find_user(username)
    name = user.get("name") or username if user else username
    mail_cfg = read_config().get("user_mail_settings", {}).get(username or "", {})
    email = mail_cfg.get("user_email") or mail_cfg.get("smtp_user") or ""
    parts = [f"\n{name}"]
    if email:
        parts.append(f"E-mail: {email}")
    parts.append("")
    return "\n".join(parts)


def _reference_mail_config():
    """以 zhouyingchao 的邮件配置作为参考示例（去掉密码）"""
    config = read_config()
    ref = config.get("user_mail_settings", {}).get("zhouyingchao", {})
    return {
        "user_email": ref.get("user_email", "yourname@company.com"),
        "weekly_to": ref.get("weekly_to", ""),
        "weekly_cc": ref.get("weekly_cc", ""),
        "trip_to": ref.get("trip_to", ""),
        "trip_cc": ref.get("trip_cc", ""),
        "smtp_user": ref.get("smtp_user", "yourname@company.com"),
        "smtp_from": ref.get("smtp_from", "yourname@company.com"),
    }


def user_mail_config(username):
    """返回指定用户的邮件配置，含本人邮箱、周报/出差报告各自的收件人和抄送"""
    config = read_config()
    settings = config.get("user_mail_settings", {}).get(username or "", {})
    reference = _reference_mail_config()
    return {
        "user_email": settings.get("user_email", ""),
        "weekly_to": settings.get("weekly_to", ""),
        "weekly_cc": settings.get("weekly_cc", ""),
        "trip_to": settings.get("trip_to", ""),
        "trip_cc": settings.get("trip_cc", ""),
        "smtp_host": config.get("smtp_host", "smtp.263.net"),
        "smtp_port": int(config.get("smtp_port", 465) or 465),
        "smtp_user": settings.get("smtp_user", ""),
        "smtp_password": settings.get("smtp_password", ""),
        "smtp_from": settings.get("smtp_from", "") or settings.get("user_email", ""),
        "smtp_tls": bool(config.get("smtp_tls", False)),
        "smtp_ssl": bool(config.get("smtp_ssl", True)),
        "imap_host": config.get("imap_host", "imap.263.net"),
        "imap_port": int(config.get("imap_port", 993) or 993),
        "imap_user": settings.get("imap_user", "") or settings.get("smtp_user", ""),
        "imap_password": settings.get("imap_password", "") or settings.get("smtp_password", ""),
        "imap_ssl": bool(config.get("imap_ssl", True)),
        "email_signature": settings.get("email_signature", default_email_signature_for_user(username)),
        "reference": reference,
    }


def save_user_mail_config(username, payload):
    if not username:
        raise ValueError("请先登录")
    config = read_config()
    settings = config.setdefault("user_mail_settings", {}).setdefault(username, {})
    fields = ["user_email", "weekly_to", "weekly_cc", "trip_to", "trip_cc", "smtp_user", "smtp_from", "imap_user"]
    for field in fields:
        if field in payload:
            settings[field] = str(payload.get(field, "") or "").strip()
    if payload.get("smtp_password"):
        settings["smtp_password"] = str(payload.get("smtp_password") or "").strip()
    if payload.get("imap_password"):
        settings["imap_password"] = str(payload.get("imap_password") or "").strip()
    if payload.get("email_signature") is not None:
        settings["email_signature"] = str(payload.get("email_signature") or "").strip()
    if settings.get("user_email"):
        settings["smtp_from"] = settings.get("smtp_from") or settings["user_email"]
        settings["smtp_user"] = settings.get("smtp_user") or settings["user_email"]
    if settings.get("smtp_user"):
        settings["imap_user"] = settings.get("imap_user") or settings["smtp_user"]
    write_config(config)
    clear_mail_cache(username)
    result = user_mail_config(username)
    result["smtp_password_masked"] = "已配置" if result.get("smtp_password") else "未配置"
    result["imap_password_masked"] = "已配置" if result.get("imap_password") else "未配置"
    result.pop("smtp_password", None)
    result.pop("imap_password", None)
    return {"ok": True, "mail_config": result}


def admin_config_payload():
    config = read_config()
    return {
        "assistant_api_url": config.get("assistant_api_url", ""),
        "assistant_model": config.get("assistant_model", "MiniMax-M2.7"),
        "assistant_prompt": config.get("assistant_prompt", DEFAULT_ASSISTANT_PROMPT),
        "assistant_api_key_masked": "已配置" if config.get("assistant_api_key") else "未配置",
        "smtp_host": config.get("smtp_host", "smtp.263.net"),
        "smtp_port": int(config.get("smtp_port", 465) or 465),
        "smtp_tls": bool(config.get("smtp_tls", False)),
        "smtp_ssl": bool(config.get("smtp_ssl", True)),
        "imap_host": config.get("imap_host", "imap.263.net"),
        "imap_port": int(config.get("imap_port", 993) or 993),
        "imap_ssl": bool(config.get("imap_ssl", True)),
        "users": [public_user(user) for user in config.get("users", [])],
    }


def save_admin_config(payload):
    config = read_config()
    config["assistant_api_url"] = str(payload.get("assistant_api_url", "") or "").strip()
    config["assistant_model"] = str(payload.get("assistant_model", "") or "MiniMax-M2.7").strip()
    config["assistant_prompt"] = str(payload.get("assistant_prompt", "") or DEFAULT_ASSISTANT_PROMPT).strip()
    api_key = str(payload.get("assistant_api_key", "") or "").strip()
    if api_key:
        config["assistant_api_key"] = api_key
    write_config(config)
    return {"ok": True, "config": admin_config_payload()}


def save_server_config(payload):
    config = read_config()
    if payload.get("smtp_host") is not None:
        config["smtp_host"] = str(payload.get("smtp_host", "") or "smtp.263.net").strip()
    if payload.get("smtp_port") is not None:
        config["smtp_port"] = int(payload.get("smtp_port", 465) or 465)
    if payload.get("smtp_tls") is not None:
        config["smtp_tls"] = bool(payload.get("smtp_tls", False))
    if payload.get("smtp_ssl") is not None:
        config["smtp_ssl"] = bool(payload.get("smtp_ssl", True))
    if payload.get("imap_host") is not None:
        config["imap_host"] = str(payload.get("imap_host", "") or "imap.263.net").strip()
    if payload.get("imap_port") is not None:
        config["imap_port"] = int(payload.get("imap_port", 993) or 993)
    if payload.get("imap_ssl") is not None:
        config["imap_ssl"] = bool(payload.get("imap_ssl", True))
    write_config(config)
    return {"ok": True, "config": admin_config_payload()}


def _caller_role(caller_username):
    caller = find_user(caller_username)
    return caller.get("role", "member") if caller else "member"


def user_list(caller_username):
    config = read_config()
    caller_role = _caller_role(caller_username)
    users = config.get("users", [])
    if caller_role in ("superadmin", "admin"):
        return [public_user(u) for u in users]
    return []


def add_user(payload, caller_username):
    caller_role = _caller_role(caller_username)
    new_username = str(payload.get("username", "") or "").strip()
    password = str(payload.get("password", "") or "").strip()
    name = str(payload.get("name", "") or "").strip() or new_username
    role = str(payload.get("role", "") or "member").strip()
    if caller_role == "admin" and role != "member":
        raise ValueError("管理员只能添加普通成员")
    if not re.match(r"^[A-Za-z0-9_@.-]{2,40}$", new_username):
        raise ValueError("用户名只能包含字母、数字、下划线、点、@ 或横线，长度 2-40 位")
    if len(password) < 4:
        raise ValueError("密码至少 4 位")
    config = read_config()
    users = config.setdefault("users", [])
    if any(u.get("username") == new_username for u in users):
        raise ValueError("该用户名已存在")
    users.append({"username": new_username, "password": password, "role": role, "name": name})
    write_config(config)
    ensure_user_space(new_username)
    return {"ok": True, "users": user_list(caller_username)}


def delete_user(payload, caller_username):
    caller_role = _caller_role(caller_username)
    target_username = str(payload.get("username", "") or "").strip()
    config = read_config()
    users = config.get("users", [])
    target = next((u for u in users if u.get("username") == target_username), None)
    if not target:
        raise ValueError("用户不存在")
    target_role = target.get("role", "member")
    if caller_role == "admin" and target_role != "member":
        raise ValueError("管理员只能删除普通成员")
    if caller_role == "superadmin" and target_role == "superadmin":
        raise ValueError("不能删除超级管理员")
    config["users"] = [u for u in users if u.get("username") != target_username]
    write_config(config)
    return {"ok": True, "users": user_list(caller_username)}


def change_password(payload, username):
    old_password = str(payload.get("old_password", "") or "").strip()
    new_password = str(payload.get("new_password", "") or "").strip()
    if len(new_password) < 4:
        raise ValueError("新密码至少 4 位")
    user = find_user(username, old_password)
    if not user:
        raise ValueError("原密码错误")
    config = read_config()
    for u in config.get("users", []):
        if u.get("username") == username:
            u["password"] = new_password
            break
    write_config(config)
    return {"ok": True}


def update_user(payload, caller_username):
    caller_role = _caller_role(caller_username)
    target_username = str(payload.get("username", "") or "").strip()
    config = read_config()
    users = config.get("users", [])
    target = next((u for u in users if u.get("username") == target_username), None)
    if not target:
        raise ValueError("用户不存在")
    target_role = target.get("role", "member")
    new_role = payload.get("role")
    new_name = str(payload.get("name", "") or "").strip()
    new_password = str(payload.get("password", "") or "").strip()
    if caller_role == "admin":
        if target_role != "member":
            raise ValueError("管理员只能修改普通成员")
        if new_role is not None and new_role != target_role:
            raise ValueError("管理员不能修改用户权限")
    if caller_role == "superadmin" and target_role == "superadmin" and new_role is not None and new_role != "superadmin":
        # 防止超级管理员降级自己后系统没有超级管理员？暂不限制
        pass
    if new_name:
        target["name"] = new_name
    if new_password:
        target["password"] = new_password
    if new_role is not None and caller_role == "superadmin":
        target["role"] = new_role
    write_config(config)
    return {"ok": True, "users": user_list(caller_username)}


def save_user_profile(payload, username):
    if not username:
        raise ValueError("请先登录")
    config = read_config()
    target = None
    for user in config.get("users", []):
        if user.get("username") == username:
            target = user
            break
    if not target:
        raise ValueError("用户不存在")
    name = str(payload.get("name", "") or "").strip()
    bio = str(payload.get("bio", "") or "").strip()
    hobbies = str(payload.get("hobbies", "") or "").strip()
    if name:
        target["name"] = name[:40]
    target["bio"] = bio[:500]
    target["hobbies"] = hobbies[:300]
    avatar_preset = str(payload.get("avatar_preset", "") or "").strip()
    avatar_data = str(payload.get("avatar_data", "") or "").strip()
    profile_dir = user_profile_dir(username)
    profile_dir.mkdir(parents=True, exist_ok=True)
    avatar_path = profile_dir / "avatar.png"
    if avatar_preset == "assistant":
        src = BASE_DIR / "assets" / "ai-assistant-avatar.png"
        if src.exists():
            avatar_path.write_bytes(src.read_bytes())
    elif avatar_data:
        if "," in avatar_data:
            avatar_data = avatar_data.split(",", 1)[1]
        raw = base64.b64decode(avatar_data)
        if len(raw) > 3 * 1024 * 1024:
            raise ValueError("头像文件不能超过 3MB")
        avatar_path.write_bytes(raw)
    if avatar_path.exists():
        target["avatar_url"] = f"/user-avatar/{urllib.parse.quote(username)}.png?v={int(avatar_path.stat().st_mtime)}"
    write_config(config)
    return {"ok": True, "user": public_user(target)}


def merged_assistant_config(payload=None):
    config = read_config()
    payload = payload or {}
    api_key = str(payload.get("assistant_api_key", "") or "").strip() or config.get("assistant_api_key", "")
    return {
        "url": str(payload.get("assistant_api_url", "") or config.get("assistant_api_url", "")).strip().rstrip("/"),
        "key": api_key,
        "model": str(payload.get("assistant_model", "") or config.get("assistant_model", "MiniMax-M2.7")).strip(),
    }


def request_json(url, api_key, body=None, method="GET", timeout=20):
    data = None if body is None else json.dumps(body, ensure_ascii=False).encode("utf-8")
    request = urllib.request.Request(
        url,
        data=data,
        headers={
            "Content-Type": "application/json",
            "Authorization": "Bearer " + api_key,
        },
        method=method,
    )
    with urllib.request.urlopen(request, timeout=timeout) as response:
        return json.loads(response.read().decode("utf-8"))


def list_admin_models(payload):
    settings = merged_assistant_config(payload)
    if not settings["url"] or not settings["key"]:
        raise ValueError("请先填写 NewAPI 地址和 API Key")
    try:
        data = request_json(settings["url"] + "/v1/models", settings["key"], method="GET", timeout=15)
        models = []
        for item in data.get("data", []):
            model_id = item.get("id") if isinstance(item, dict) else str(item)
            if model_id:
                models.append(model_id)
        models = sorted(set(models))
        if models:
            return {"ok": True, "mode": "api", "models": models}
    except Exception as exc:
        fallback = ["MiniMax-M2.7", "MiniMax-M2.5", "MiniMax-M1", "minimax-m2.7", "minimax-m2.5"]
        return {"ok": True, "mode": "fallback", "warning": f"模型列表获取失败，已使用常用 MiniMax 选项：{exc}", "models": fallback}
    return {"ok": True, "mode": "fallback", "models": ["MiniMax-M2.7", "MiniMax-M2.5", "MiniMax-M1"]}


def test_admin_model(payload):
    settings = merged_assistant_config(payload)
    if not settings["url"] or not settings["key"] or not settings["model"]:
        raise ValueError("请填写 NewAPI 地址、API Key 和模型名称")
    data = request_json(
        settings["url"] + "/v1/chat/completions",
        settings["key"],
        body={
            "model": settings["model"],
            "messages": [
                {"role": "system", "content": "你只需要回复 OK。"},
                {"role": "user", "content": "请回复 OK，用于测试 API Key 和模型是否可用。"},
            ],
            "temperature": 0,
            "max_tokens": 16,
        },
        method="POST",
        timeout=30,
    )
    content = data.get("choices", [{}])[0].get("message", {}).get("content", "").strip()
    return {"ok": True, "model": settings["model"], "message": content or "测试请求已成功返回"}


def parse_weekly_date(name):
    match = re.search(r"(?:([0-9]{2})年)?([0-9]{1,2})月([0-9]{1,2})日-([0-9]{1,2})月([0-9]{1,2})日", name)
    if not match:
        return (0, 0, 0, 0)
    year = 2000 + int(match.group(1) or "25")
    month1 = int(match.group(2))
    day1 = int(match.group(3))
    month2 = int(match.group(4))
    day2 = int(match.group(5))
    return (year, month2, day2, month1 * 100 + day1)


def parse_trip_date(name):
    match = re.search(r"出差报告-([0-9]{8})-([0-9]{4})", name)
    if not match:
        return (0, 0)
    return (int(match.group(1)), int(match.group(2)))


def report_files(username=None):
    files = []
    base = user_report_dir(username) if username else REPORT_DIR
    if not base.exists():
        return files
    for path in base.iterdir():
        if not path.is_file() or path.name.startswith("."):
            continue
        lower = path.suffix.lower()
        if "工作周报" in path.name and lower in {".xlsx", ".xls"}:
            files.append(
                {
                    "kind": "weekly",
                    "name": path.name,
                    "path": str(path),
                    "mtime": path.stat().st_mtime,
                    "sort_key": parse_weekly_date(path.name),
                    "deletable": bool(username),
                }
            )
        elif path.name.startswith("出差报告") and lower in {".docx", ".md"}:
            files.append(
                {
                    "kind": "trip",
                    "name": path.name,
                    "path": str(path),
                    "mtime": path.stat().st_mtime,
                    "sort_key": parse_trip_date(path.name),
                    "deletable": bool(username),
                }
            )
    return files


def generated_files(username=None):
    files = []
    base = user_generated_dir(username) if username else GENERATED_DIR
    if not base.exists():
        return files
    for path in base.iterdir():
        if not path.is_file() or path.name.startswith("."):
            continue
        lower = path.suffix.lower()
        if "工作周报" in path.name and lower in {".xlsx", ".xls"}:
            kind = "weekly"
        elif "出差报告" in path.name and lower in {".docx", ".md"}:
            kind = "trip"
        else:
            continue
        files.append(
            {
                "kind": kind,
                "name": path.name,
                "path": str(path),
                "mtime": path.stat().st_mtime,
                "sort_key": (9999, int(path.stat().st_mtime)),
                "generated": True,
                "deletable": bool(username),
            }
        )
    return files


def all_files(username=None):
    user_files = generated_files(username) + report_files(username)
    if username and not any(not f.get("generated") for f in user_files):
        user_files = user_files + report_files(None)
    return user_files


def newest(kind, username=None, fallback_shared=False):
    items = [item for item in report_files(username) if item["kind"] == kind]
    if not items and fallback_shared and username:
        items = [item for item in report_files(None) if item["kind"] == kind]
    if not items:
        return None
    return sorted(items, key=lambda item: (item["sort_key"], item["mtime"]), reverse=True)[0]


def newest_any(kind, username=None, fallback_shared=False):
    items = [item for item in all_files(username) if item["kind"] == kind]
    if not items and fallback_shared and username:
        items = [item for item in all_files(None) if item["kind"] == kind]
    if not items and kind == "weekly":
        for fallback_user in ("zhouyingchao", "admin"):
            items = [item for item in all_files(fallback_user) if item["kind"] == kind]
            if items:
                break
    if not items:
        return None
    return sorted(items, key=lambda item: (item["sort_key"], item["mtime"]), reverse=True)[0]


def xml_text(node):
    if node is None:
        return ""
    return "".join(node.itertext())


def preview_docx(path, max_chars=900):
    try:
        with zipfile.ZipFile(path) as docx:
            xml = docx.read("word/document.xml")
        root = ET.fromstring(xml)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        paragraphs = [xml_text(p).strip() for p in root.findall(".//w:p", ns)]
        text = "\n".join(p for p in paragraphs if p)
        return text[:max_chars]
    except Exception:
        return ""


def preview_xlsx(path, max_chars=900):
    try:
        with zipfile.ZipFile(path) as book:
            shared = []
            if "xl/sharedStrings.xml" in book.namelist():
                root = ET.fromstring(book.read("xl/sharedStrings.xml"))
                ns = {"a": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
                shared = [xml_text(si).strip() for si in root.findall(".//a:si", ns)]
            sheet_names = [n for n in book.namelist() if n.startswith("xl/worksheets/sheet") and n.endswith(".xml")]
            texts = []
            for sheet in sheet_names[:2]:
                root = ET.fromstring(book.read(sheet))
                for cell in root.iter("{http://schemas.openxmlformats.org/spreadsheetml/2006/main}c"):
                    value = cell.find("{http://schemas.openxmlformats.org/spreadsheetml/2006/main}v")
                    if value is None or value.text is None:
                        continue
                    if cell.attrib.get("t") == "s":
                        idx = int(value.text)
                        if idx < len(shared):
                            texts.append(shared[idx])
                    else:
                        texts.append(value.text)
            return "\n".join(t for t in texts if t)[:max_chars]
    except Exception:
        return ""


def table_style_html():
    return (
        "width:100%;border-collapse:collapse;margin:8px 0 14px 0;"
        "table-layout:fixed;"
        "font-family:-apple-system,BlinkMacSystemFont,Segoe UI,PingFang SC,Microsoft YaHei,sans-serif;"
        "font-size:13px;background:#fff;"
    )


def xlsx_original_table_html(path):
    try:
        from openpyxl import load_workbook
        from openpyxl.utils import get_column_letter

        wb = load_workbook(path, data_only=True)
        ws = wb[wb.sheetnames[0]]
        min_row, max_row = 1, ws.max_row
        min_col, max_col = 1, ws.max_column

        while max_row > 1 and not any(ws.cell(max_row, col).value not in (None, "") for col in range(min_col, max_col + 1)):
            max_row -= 1
        while max_col > 1 and not any(ws.cell(row, max_col).value not in (None, "") for row in range(min_row, max_row + 1)):
            max_col -= 1

        merged_starts = {}
        merged_covered = set()
        for area in ws.merged_cells.ranges:
            if area.max_row < min_row or area.min_row > max_row or area.max_col < min_col or area.min_col > max_col:
                continue
            row_span = area.max_row - area.min_row + 1
            col_span = area.max_col - area.min_col + 1
            merged_starts[(area.min_row, area.min_col)] = (row_span, col_span)
            for r in range(area.min_row, area.max_row + 1):
                for c in range(area.min_col, area.max_col + 1):
                    if (r, c) != (area.min_row, area.min_col):
                        merged_covered.add((r, c))

        parts = [f'<table class="raw-table" style="{table_style_html()}">']
        colgroup = ["<colgroup>"]
        for col in range(min_col, max_col + 1):
            letter = get_column_letter(col)
            width = ws.column_dimensions[letter].width or 12
            colgroup.append(f'<col style="width:{min(max(width * 8, 72), 220)}px">')
        colgroup.append("</colgroup>")
        parts.extend(colgroup)

        for row in range(min_row, max_row + 1):
            parts.append("<tr>")
            for col in range(min_col, max_col + 1):
                if (row, col) in merged_covered:
                    continue
                cell = ws.cell(row, col)
                value = "" if cell.value is None else str(cell.value)
                row_span, col_span = merged_starts.get((row, col), (1, 1))
                attrs = []
                if row_span > 1:
                    attrs.append(f'rowspan="{row_span}"')
                if col_span > 1:
                    attrs.append(f'colspan="{col_span}"')
                align = cell.alignment.horizontal or "left"
                valign = cell.alignment.vertical or "middle"
                weight = "font-weight:700;" if cell.font and cell.font.bold else ""
                fill = ""
                fg = getattr(cell.fill, "fgColor", None)
                if fg and fg.type == "rgb" and fg.rgb and fg.rgb not in {"00000000", "FFFFFFFF"}:
                    fill = f"background:#{fg.rgb[-6:]};"
                style = (
                    "border:1px solid #cfd6df;padding:7px 9px;white-space:pre-wrap;word-break:break-word;"
                    f"text-align:{align};vertical-align:{valign};{weight}{fill}"
                )
                parts.append(f'<td {" ".join(attrs)} style="{style}">{html_escape(value)}</td>')
            parts.append("</tr>")
        parts.append("</table>")
        return "".join(parts)
    except Exception:
        return ""


def docx_original_table_html(path):
    try:
        ns = {"w": W_NS}
        parts = []
        with zipfile.ZipFile(path) as docx:
            root = ET.fromstring(docx.read("word/document.xml"))
        for table in root.findall(".//w:tbl", ns):
            parts.append(f'<table class="raw-table" style="{table_style_html()}">')
            grid_cols = []
            for grid_col in table.findall("w:tblGrid/w:gridCol", ns):
                try:
                    grid_cols.append(int(grid_col.attrib.get(w_tag("w"), "0") or 0))
                except ValueError:
                    grid_cols.append(0)
            if grid_cols and sum(grid_cols) > 0:
                total = sum(grid_cols)
                parts.append("<colgroup>")
                for width in grid_cols:
                    parts.append(f'<col style="width:{width / total * 100:.3f}%">')
                parts.append("</colgroup>")
            for row in table.findall("w:tr", ns):
                parts.append("<tr>")
                for cell in row.findall("w:tc", ns):
                    grid_span = cell.find("w:tcPr/w:gridSpan", ns)
                    colspan = grid_span.attrib.get(w_tag("val"), "1") if grid_span is not None else "1"
                    paragraphs = []
                    for paragraph in cell.findall("w:p", ns):
                        text = "".join(node.text or "" for node in paragraph.findall(".//w:t", ns)).strip()
                        if text:
                            paragraphs.append(text)
                    text = html_escape("\n".join(paragraphs)).replace("\n", "<br>")
                    span_attr = f' colspan="{html_escape(colspan)}"' if colspan and colspan != "1" else ""
                    cell_style = (
                        "border:1px solid #cfd6df;padding:7px 9px;vertical-align:top;"
                        "white-space:pre-wrap;word-break:break-word;"
                    )
                    parts.append(
                        f'<td{span_attr} style="{cell_style}">'
                        f"{text}</td>"
                    )
                parts.append("</tr>")
            parts.append("</table>")
        if parts:
            return "".join(parts)
        paragraphs = []
        for paragraph in root.findall(".//w:p", ns):
            text = "".join(node.text or "" for node in paragraph.findall(".//w:t", ns)).strip()
            if text:
                paragraphs.append(html_escape(text))
        return "".join(f"<p>{text}</p>" for text in paragraphs)
    except Exception:
        return ""


def preview_file_html(path):
    suffix = path.suffix.lower()
    if suffix in {".xlsx", ".xls"}:
        return xlsx_original_table_html(path)
    if suffix == ".docx":
        return docx_original_table_html(path)
    if suffix == ".md":
        text = path.read_text(encoding="utf-8", errors="ignore")
        return f'<div style="white-space:pre-wrap;">{html_escape(text)}</div>'
    return ""


def preview_file(path):
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return preview_docx(path)
    if suffix == ".xlsx":
        return preview_xlsx(path)
    if suffix == ".md":
        return path.read_text(encoding="utf-8", errors="ignore")[:900]
    return ""


def html_escape(text):
    if not text:
        return ""
    return (
        str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def format_weekly_body(path):
    """从 xlsx 周报提取结构化摘要，返回 (纯文本, HTML)"""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(path, data_only=True)
        ws = wb[wb.sheetnames[0]]

        def get(r, c):
            v = ws.cell(r, c).value
            return str(v).strip() if v else ""

        def find_row(label, fallback):
            for r in range(1, ws.max_row + 1):
                if label in get(r, 2):
                    return r
            return fallback

        def valid_data_row(cat, content):
            if not cat and not content:
                return False
            invalid = {"工作分类", "工作内容", "一、本周工作总结", "二、重点工作跟进", "三、下周工作计划"}
            return cat not in invalid and content not in invalid

        summary_title_row = find_row("一、本周工作总结", 3)
        follow_title_row = find_row("二、重点工作跟进", 9)
        next_title_row = find_row("三、下周工作计划", 16)
        summary_range = range(summary_title_row + 2, follow_title_row)
        follow_range = range(follow_title_row + 2, next_title_row)
        next_range = range(next_title_row + 2, ws.max_row + 1)

        # ===== 纯文本版本 =====
        lines = []
        title = get(2, 2)
        if title:
            lines.append(title)
            lines.append("")

        # 一、本周工作总结: B=工作分类 C=工作内容 E=完成情况 F=后续计划
        lines.append("一、本周工作总结")
        for row in summary_range:
            cat = get(row, 2)
            content = get(row, 3)
            status = get(row, 5)
            plan = get(row, 6)
            if not valid_data_row(cat, content):
                continue
            if cat:
                lines.append(f"【{cat}】")
            if content:
                for ln in content.split("\n"):
                    ln = ln.strip()
                    if ln:
                        lines.append(f"  {ln}")
            if status:
                lines.append(f"  完成情况：{status}")
            if plan:
                lines.append(f"  后续计划：{plan}")
            lines.append("")

        # 二、重点工作跟进: B=工作分类 C=工作内容 D=当前进展 F=困难与求助
        lines.append("二、重点工作跟进")
        for row in follow_range:
            cat = get(row, 2)
            content = get(row, 3)
            progress = get(row, 4)
            difficulty = get(row, 6)
            if not valid_data_row(cat, content):
                continue
            if cat:
                lines.append(f"【{cat}】")
            if content:
                for ln in content.split("\n"):
                    ln = ln.strip()
                    if ln:
                        lines.append(f"  {ln}")
            if progress:
                lines.append(f"  当前进展：{progress}")
            if difficulty:
                lines.append(f"  困难与求助：{difficulty}")
            lines.append("")

        # 三、下周工作计划: B=工作分类 C=工作内容 F=困难与求助
        lines.append("三、下周工作计划")
        for row in next_range:
            cat = get(row, 2)
            content = get(row, 3)
            difficulty = get(row, 6)
            if not valid_data_row(cat, content):
                continue
            if cat:
                lines.append(f"【{cat}】")
            if content:
                for ln in content.split("\n"):
                    ln = ln.strip()
                    if ln:
                        lines.append(f"  {ln}")
            if difficulty:
                lines.append(f"  困难与求助：{difficulty}")
            lines.append("")

        text_body = "\n".join(lines).strip()

        # ===== HTML 表格版本 =====
        h = []
        if title:
            h.append(f'<h3 style="margin:0 0 12px 0;font-size:16px;">{html_escape(title)}</h3>')

        th_style = "border:1px solid #ccc;padding:6px 8px;text-align:left;background:#f5f5f5;font-size:13px;"
        td_style = "border:1px solid #ccc;padding:6px 8px;text-align:left;font-size:13px;vertical-align:top;"
        table_style = "width:100%;border-collapse:collapse;margin-bottom:16px;font-family:-apple-system,BlinkMacSystemFont,Segoe UI,PingFang SC,Microsoft YaHei,sans-serif;"

        def make_table(headers, rows):
            parts = ['<table style="' + table_style + '">']
            parts.append('<tr>')
            for hd in headers:
                parts.append(f'<th style="{th_style}">{html_escape(hd)}</th>')
            parts.append('</tr>')
            for row in rows:
                if not any(row):
                    continue
                parts.append('<tr>')
                for cell in row:
                    cell_html = html_escape(cell).replace("\n", "<br>")
                    parts.append(f'<td style="{td_style}">{cell_html}</td>')
                parts.append('</tr>')
            parts.append('</table>')
            return "".join(parts)

        # 本周工作总结
        h.append('<p style="font-size:14px;font-weight:bold;margin:12px 0 6px 0;">一、本周工作总结</p>')
        summary_rows = []
        for row in summary_range:
            cat = get(row, 2)
            content = get(row, 3)
            status = get(row, 5)
            plan = get(row, 6)
            if valid_data_row(cat, content):
                summary_rows.append([cat, content, status, plan])
        h.append(make_table(["工作分类", "工作内容", "完成情况", "后续计划"], summary_rows))

        # 重点工作跟进
        h.append('<p style="font-size:14px;font-weight:bold;margin:12px 0 6px 0;">二、重点工作跟进</p>')
        follow_rows = []
        for row in follow_range:
            cat = get(row, 2)
            content = get(row, 3)
            progress = get(row, 4)
            difficulty = get(row, 6)
            if valid_data_row(cat, content):
                follow_rows.append([cat, content, progress, difficulty])
        h.append(make_table(["工作分类", "工作内容", "当前进展", "困难与求助"], follow_rows))

        # 下周工作计划
        h.append('<p style="font-size:14px;font-weight:bold;margin:12px 0 6px 0;">三、下周工作计划</p>')
        next_rows = []
        for row in next_range:
            cat = get(row, 2)
            content = get(row, 3)
            difficulty = get(row, 6)
            if valid_data_row(cat, content):
                next_rows.append([cat, content, difficulty])
        h.append(make_table(["工作分类", "工作内容", "困难与求助"], next_rows))

        html_body = "".join(h)
        return text_body, html_body
    except Exception:
        return "", ""


def format_trip_body(path):
    """从 docx 出差报告提取结构化摘要，返回 (纯文本, HTML)"""
    try:
        from docx import Document
        doc = Document(path)
        if not doc.tables:
            return "", ""
        table = doc.tables[0]

        def cell_text(r, c):
            try:
                return table.cell(r, c).text.strip()
            except Exception:
                return ""

        reporter = cell_text(0, 1)
        department = cell_text(0, 3)
        location = cell_text(0, 5)
        date_text = cell_text(1, 1)
        purpose = cell_text(2, 1)
        itinerary = cell_text(3, 1)
        details = cell_text(4, 1)
        issues = cell_text(5, 1)
        suggestions = cell_text(6, 1)

        # ===== 纯文本版本 =====
        lines = []
        if reporter:
            lines.append(f"报告人：{reporter}")
        if department:
            lines.append(f"部门：{department}")
        if location:
            lines.append(f"出差地点：{location}")
        if date_text:
            lines.append(f"出差时间：{date_text}")
        lines.append("")

        if purpose:
            lines.append("【出差目的】")
            lines.append(purpose)
            lines.append("")
        if itinerary:
            lines.append("【行程概览】")
            lines.append(itinerary)
            lines.append("")
        if details:
            lines.append("【工作详情】")
            lines.append(details)
            lines.append("")
        if issues:
            lines.append("【问题与困难】")
            lines.append(issues)
            lines.append("")
        if suggestions:
            lines.append("【改进建议】")
            lines.append(suggestions)
            lines.append("")

        text_body = "\n".join(lines).strip()

        # ===== HTML 表格版本 =====
        h = []
        h.append('<table style="width:100%;border-collapse:collapse;margin-bottom:16px;font-family:-apple-system,BlinkMacSystemFont,Segoe UI,PingFang SC,Microsoft YaHei,sans-serif;font-size:13px;">')

        info_rows = [
            ["报告人", reporter, "部门", department, "出差地点", location],
            ["出差时间", date_text, "", "", "", ""],
        ]
        td_label = "border:1px solid #ccc;padding:6px 8px;background:#f5f5f5;font-weight:bold;width:12%;"
        td_value = "border:1px solid #ccc;padding:6px 8px;"

        for row in info_rows:
            h.append('<tr>')
            h.append(f'<td style="{td_label}">{html_escape(row[0])}</td>')
            h.append(f'<td style="{td_value}" colspan="5">{html_escape(row[1])}</td>')
            h.append('</tr>')
        h.append('</table>')

        section_style = "font-size:14px;font-weight:bold;margin:12px 0 6px 0;"
        content_style = "font-size:13px;line-height:1.6;margin:0 0 12px 0;"

        if purpose:
            h.append(f'<p style="{section_style}">【出差目的】</p>')
            h.append(f'<p style="{content_style}">{html_escape(purpose).replace(chr(10), "<br>")}</p>')
        if itinerary:
            h.append(f'<p style="{section_style}">【行程概览】</p>')
            h.append(f'<p style="{content_style}">{html_escape(itinerary).replace(chr(10), "<br>")}</p>')
        if details:
            h.append(f'<p style="{section_style}">【工作详情】</p>')
            h.append(f'<p style="{content_style}">{html_escape(details).replace(chr(10), "<br>")}</p>')
        if issues:
            h.append(f'<p style="{section_style}">【问题与困难】</p>')
            h.append(f'<p style="{content_style}">{html_escape(issues).replace(chr(10), "<br>")}</p>')
        if suggestions:
            h.append(f'<p style="{section_style}">【改进建议】</p>')
            h.append(f'<p style="{content_style}">{html_escape(suggestions).replace(chr(10), "<br>")}</p>')

        html_body = "".join(h)
        return text_body, html_body
    except Exception:
        return "", ""


def compose_draft(kind, file_name=None, username=None):
    config = read_config()
    mail = user_mail_config(username)
    report = None
    if file_name:
        for item in all_files(username):
            if item["name"] == file_name:
                report = item
                break
    if report is None:
        report = newest(kind, username, fallback_shared=True)
    if report is None:
        return {"error": "没有找到对应报告"}

    path = Path(report["path"])
    preview = preview_file(path)
    preview_html = preview_file_html(path)
    sender = config.get("sender_name", "周颖超")
    if kind == "weekly":
        subject = f"【周报】{path.stem}"
        summary_text, summary_html = format_weekly_body(path)
        body = (
            f"{config.get('weekly_greeting', '领导您好：')}\n\n"
            f"附件为我的本周工作周报《{path.name}》，请查收。\n\n"
        )
        if summary_text:
            body += summary_text + "\n\n"
        body += (
            "本周主要工作内容已在附件中汇总，如有需要补充或调整的地方，我会及时完善。\n\n"
            f"{sender}\n"
            f"{datetime.now().strftime('%Y年%m月%d日')}"
        )
        body_html = f'<p>{html_escape(config.get("weekly_greeting", "领导您好："))}</p><p>附件为我的本周工作周报《{html_escape(path.name)}》，请查收。</p>'
        if summary_html:
            body_html += summary_html
        body_html += f'<p>本周主要工作内容已在附件中汇总，如有需要补充或调整的地方，我会及时完善。</p><p>{html_escape(sender)}<br>{datetime.now().strftime("%Y年%m月%d日")}</p>'
        to_addr = mail.get("weekly_to", "")
        cc_addr = mail.get("weekly_cc", "")
    else:
        subject = f"【出差报告】{path.stem}"
        summary_text, summary_html = format_trip_body(path)
        trip_table_html = preview_html or summary_html
        body = (
            f"{config.get('trip_greeting', '领导您好：')}\n\n"
            f"附件为我的出差报告《{path.name}》，请查收。\n\n"
        )
        if summary_text:
            body += summary_text + "\n\n"
        body += (
            "出差事项、过程记录和相关结论已在附件中说明，如需进一步补充材料，我会及时整理。\n\n"
            f"{sender}\n"
            f"{datetime.now().strftime('%Y年%m月%d日')}"
        )
        body_html = f'<p>{html_escape(config.get("trip_greeting", "领导您好："))}</p><p>附件为我的出差报告《{html_escape(path.name)}》，请查收。</p>'
        if trip_table_html:
            body_html += trip_table_html
        body_html += f'<p>出差事项、过程记录和相关结论已在附件中说明，如需进一步补充材料，我会及时整理。</p><p>{html_escape(sender)}<br>{datetime.now().strftime("%Y年%m月%d日")}</p>'
        to_addr = mail.get("trip_to", "")
        cc_addr = mail.get("trip_cc", "")

    return {
        "kind": kind,
        "subject": subject,
        "body": body,
        "body_html": body_html,
        "to": to_addr,
        "cc": cc_addr,
        "attachment": report["name"],
        "attachment_path": report["path"],
        "download_url": "/download?file=" + urllib.parse.quote(report["name"]),
        "preview": preview,
        "preview_html": preview_html,
    }

def smtp_settings(username=None):
    config = read_config()
    if username:
        mail = user_mail_config(username)
        return {
            "host": mail.get("smtp_host", ""),
            "port": int(mail.get("smtp_port") or 587),
            "user": mail.get("smtp_user", ""),
            "password": mail.get("smtp_password", ""),
            "from_addr": mail.get("smtp_from", "") or mail.get("user_email", ""),
            "use_ssl": bool(mail.get("smtp_ssl", False)),
            "use_tls": bool(mail.get("smtp_tls", True)),
        }
    return {
        "host": os.getenv("SMTP_HOST", config.get("smtp_host", "")),
        "port": int(os.getenv("SMTP_PORT", config.get("smtp_port", "587") or 587)),
        "user": os.getenv("SMTP_USER", config.get("smtp_user", "")),
        "password": os.getenv("SMTP_PASSWORD", config.get("smtp_password", "")),
        "from_addr": os.getenv("SMTP_FROM", config.get("smtp_from", "")),
        "use_ssl": str(os.getenv("SMTP_SSL", config.get("smtp_ssl", "false"))).lower() == "true",
        "use_tls": str(os.getenv("SMTP_TLS", config.get("smtp_tls", "true"))).lower() != "false",
    }


def imap_settings(username=None):
    mail = user_mail_config(username)
    return {
        "host": mail.get("imap_host", ""),
        "port": int(mail.get("imap_port") or 993),
        "user": mail.get("imap_user", "") or mail.get("smtp_user", ""),
        "password": mail.get("imap_password", "") or mail.get("smtp_password", ""),
        "use_ssl": bool(mail.get("imap_ssl", True)),
    }


def validate_smtp_ready(settings):
    missing = []
    if not settings.get("host"):
        missing.append("SMTP 服务器")
    if not settings.get("user"):
        missing.append("SMTP 用户名")
    if not settings.get("password"):
        missing.append("SMTP 授权码/密码")
    if missing:
        raise ValueError("当前账号邮件未配置完整，缺少：" + "、".join(missing) + "。请到左侧“邮件配置”补充后再发送。")


def validate_imap_ready(settings):
    missing = []
    if not settings.get("host"):
        missing.append("IMAP 服务器")
    if not settings.get("user"):
        missing.append("IMAP 用户名")
    if not settings.get("password"):
        missing.append("IMAP 授权码/密码")
    if missing:
        raise ValueError("当前账号收件箱未配置完整，缺少：" + "、".join(missing) + "。请到左侧“邮件配置”补充后再查看邮件。")


def test_user_mail_config(username):
    settings = smtp_settings(username)
    validate_smtp_ready(settings)
    if settings["use_ssl"]:
        with smtplib.SMTP_SSL(settings["host"], settings["port"], context=ssl.create_default_context(), timeout=20) as smtp:
            smtp.login(settings["user"], settings["password"])
    else:
        with smtplib.SMTP(settings["host"], settings["port"], timeout=20) as smtp:
            if settings["use_tls"]:
                smtp.starttls(context=ssl.create_default_context())
            smtp.login(settings["user"], settings["password"])
    return {"ok": True, "message": "邮箱配置测试成功，可以正常登录 SMTP。"}


def imap_connect(username):
    settings = imap_settings(username)
    validate_imap_ready(settings)
    if settings["use_ssl"]:
        box = imaplib.IMAP4_SSL(settings["host"], settings["port"], timeout=20)
    else:
        box = imaplib.IMAP4(settings["host"], settings["port"], timeout=20)
    box.login(settings["user"], settings["password"])
    return box


def user_mail_cache_dir(username):
    path = user_root(username) / "mail_cache"
    path.mkdir(parents=True, exist_ok=True)
    return path


def mail_cache_path(username, name):
    safe = re.sub(r"[^A-Za-z0-9_.-]+", "_", str(name or "cache"))
    return user_mail_cache_dir(username) / f"{safe}.json"


def read_mail_cache(username, name, max_age_seconds=600):
    path = mail_cache_path(username, name)
    if not path.exists():
        return None
    try:
        data = read_json_lenient(path)
        fetched_at = float(data.get("fetched_at", 0) or 0)
        if max_age_seconds and time.time() - fetched_at > max_age_seconds:
            return None
        return data
    except Exception:
        return None


def write_mail_cache(username, name, data):
    payload = dict(data or {})
    payload["fetched_at"] = time.time()
    mail_cache_path(username, name).write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return payload


def clear_mail_cache(username):
    path = user_mail_cache_dir(username)
    for item in path.glob("*.json"):
        try:
            item.unlink()
        except Exception:
            pass


def decode_mime_value(value):
    chunks = decode_header(value or "")
    output = []
    for chunk, charset in chunks:
        if isinstance(chunk, bytes):
            output.append(chunk.decode(charset or "utf-8", errors="replace"))
        else:
            output.append(chunk)
    return "".join(output).strip()


def normalize_mail_date(value):
    try:
        return parsedate_to_datetime(value).strftime("%Y-%m-%d %H:%M")
    except Exception:
        return value or ""


def extract_mail_text(msg, max_chars=800):
    parts = []
    html_parts = []
    if msg.is_multipart():
        for part in msg.walk():
            if part.get_content_disposition() == "attachment":
                continue
            if part.get_content_type() not in ("text/plain", "text/html"):
                continue
            payload = part.get_payload(decode=True)
            if not payload:
                continue
            charset = part.get_content_charset() or "utf-8"
            content = payload.decode(charset, errors="replace")
            if part.get_content_type() == "text/plain":
                parts.append(content)
            else:
                html_parts.append(content)
    else:
        payload = msg.get_payload(decode=True)
        if payload:
            content = payload.decode(msg.get_content_charset() or "utf-8", errors="replace")
            if msg.get_content_type() == "text/html":
                html_parts.append(content)
            else:
                parts.append(content)
    if not parts and html_parts:
        parts = [re.sub(r"<[^>]+>", " ", text) for text in html_parts]
    text = re.sub(r"\s+", " ", "\n".join(parts)).strip()
    return text[:max_chars]


def extract_mail_html(msg, max_chars=120000):
    html_parts = []
    if msg.is_multipart():
        for part in msg.walk():
            if part.get_content_disposition() == "attachment" or part.get_content_type() != "text/html":
                continue
            payload = part.get_payload(decode=True)
            if payload:
                html_parts.append(payload.decode(part.get_content_charset() or "utf-8", errors="replace"))
    elif msg.get_content_type() == "text/html":
        payload = msg.get_payload(decode=True)
        if payload:
            html_parts.append(payload.decode(msg.get_content_charset() or "utf-8", errors="replace"))
    html = "\n".join(html_parts).strip()
    html = re.sub(r"(?is)<script.*?>.*?</script>", "", html)
    html = re.sub(r"(?is)<style.*?>.*?</style>", "", html)
    html = re.sub(r"\son\w+\s*=\s*(['\"]).*?\1", "", html)
    html = re.sub(r"(?i)javascript:", "", html)
    return html[:max_chars]


def mail_summary_from_message(uid, msg):
    attachments = []
    for part in msg.walk() if msg.is_multipart() else []:
        name = part.get_filename()
        if name:
            payload = part.get_payload(decode=True) or b""
            attachments.append({
                "name": decode_mime_value(name),
                "size": len(payload),
                "type": part.get_content_type(),
            })
    body_html = extract_mail_html(msg)
    return {
        "uid": str(uid),
        "subject": decode_mime_value(msg.get("Subject", "")) or "无主题",
        "from": decode_mime_value(msg.get("From", "")),
        "to": decode_mime_value(msg.get("To", "")),
        "date": normalize_mail_date(msg.get("Date", "")),
        "preview": extract_mail_text(msg, 180),
        "body": extract_mail_text(msg, 6000),
        "body_html": body_html,
        "has_html": bool(body_html),
        "attachments": attachments,
    }


def mail_header_summary(uid, msg):
    return {
        "uid": str(uid),
        "subject": decode_mime_value(msg.get("Subject", "")) or "无主题",
        "from": decode_mime_value(msg.get("From", "")),
        "to": decode_mime_value(msg.get("To", "")),
        "date": normalize_mail_date(msg.get("Date", "")),
        "preview": "点击邮件查看正文",
        "attachments": [],
    }


def list_inbox_messages(username, limit=20, refresh=False):
    limit = max(1, min(int(limit or 20), 50))
    cache_name = f"inbox_{limit}"
    if not refresh:
        cached = read_mail_cache(username, cache_name, 1800)
        if cached:
            cached["ok"] = True
            cached["cached"] = True
            return cached
    box = imap_connect(username)
    try:
        box.select("INBOX", readonly=True)
        status, data = box.uid("search", None, "ALL")
        if status != "OK":
            raise ValueError("读取收件箱失败")
        uids = data[0].split()[-limit:][::-1]
        messages = []
        if uids:
            uid_set = b",".join(uids)
            status, fetched = box.uid("fetch", uid_set, "(BODY.PEEK[HEADER.FIELDS (SUBJECT FROM TO DATE)])")
            if status != "OK" or not fetched:
                raise ValueError("读取收件箱失败")
            for item in fetched:
                if not isinstance(item, tuple) or not item[1]:
                    continue
                meta = item[0].decode("utf-8", errors="ignore")
                match = re.search(r"UID\s+(\d+)", meta)
                uid = match.group(1) if match else ""
                if not uid:
                    continue
                messages.append(mail_header_summary(uid, email.message_from_bytes(item[1])))
        uid_order = {uid.decode("ascii", errors="ignore"): idx for idx, uid in enumerate(uids)}
        messages.sort(key=lambda item: uid_order.get(str(item.get("uid", "")), 999))
        result = write_mail_cache(username, cache_name, {"ok": True, "messages": messages, "cached": False})
        result["ok"] = True
        return result
    finally:
        try:
            box.logout()
        except Exception:
            pass


def get_inbox_message(username, uid, refresh=False):
    uid = str(uid or "").strip()
    if not uid:
        raise ValueError("缺少邮件 UID")
    cache_name = f"detail_{uid}"
    if not refresh:
        cached = read_mail_cache(username, cache_name, 7200)
        if cached:
            cached["ok"] = True
            cached["cached"] = True
            return cached
    box = imap_connect(username)
    try:
        box.select("INBOX", readonly=True)
        status, fetched = box.uid("fetch", uid, "(BODY.PEEK[])")
        if status != "OK" or not fetched:
            raise ValueError("邮件不存在或读取失败")
        raw = next((item[1] for item in fetched if isinstance(item, tuple)), None)
        if not raw:
            raise ValueError("邮件内容为空")
        result = write_mail_cache(username, cache_name, {"ok": True, "message": mail_summary_from_message(uid, email.message_from_bytes(raw)), "cached": False})
        result["ok"] = True
        return result
    finally:
        try:
            box.logout()
        except Exception:
            pass


def split_addresses(value):
    return [item.strip() for item in re.split(r"[;,]", value or "") if item.strip()]


def build_message(payload, username=None):
    settings = smtp_settings(username)
    mail_cfg = user_mail_config(username)
    user_email = mail_cfg.get("user_email", "")
    from_addr = payload.get("from") or user_email or settings["from_addr"] or settings["user"]
    if not from_addr:
        from_addr = "no-reply@local"

    config = read_config()
    user_mail = config.get("user_mail_settings", {}).get(username or "", {})
    signature = user_mail.get("email_signature", default_email_signature_for_user(username)) or ""
    body = payload.get("body", "") or ""
    if signature:
        if not body.endswith(signature):
            body = body + signature

    body_html = payload.get("body_html", "") or ""
    if signature and body_html:
        sig_html = html_escape(signature).replace("\n", "<br>")
        if sig_html not in body_html:
            body_html = body_html + "<br>" + sig_html

    msg = EmailMessage()
    msg["From"] = from_addr
    msg["To"] = ", ".join(split_addresses(payload.get("to", "")))
    if split_addresses(payload.get("cc", "")):
        msg["Cc"] = ", ".join(split_addresses(payload.get("cc", "")))
    msg["Subject"] = payload.get("subject", "")

    if body_html:
        # multipart/alternative: text/plain + text/html
        msg.set_content(body, subtype="plain", charset="utf-8")
        msg.add_alternative(body_html, subtype="html")
    else:
        msg.set_content(body, subtype="plain", charset="utf-8")

    attachment_name = payload.get("attachment", "")
    path = user_generated_dir(username) / attachment_name if username else GENERATED_DIR / attachment_name
    if not path.exists():
        path = user_report_dir(username) / attachment_name if username else REPORT_DIR / attachment_name
    if attachment_name and path.exists():
        ctype, encoding = mimetypes.guess_type(str(path))
        if ctype is None or encoding is not None:
            ctype = "application/octet-stream"
        maintype, subtype = ctype.split("/", 1)
        msg.add_attachment(path.read_bytes(), maintype=maintype, subtype=subtype, filename=path.name)
    for item in payload.get("attachments", []) or []:
        filename = Path(str(item.get("name", "attachment"))).name or "attachment"
        raw = base64.b64decode(item.get("content", "") or "")
        ctype = item.get("type") or mimetypes.guess_type(filename)[0] or "application/octet-stream"
        maintype, subtype = ctype.split("/", 1) if "/" in ctype else ("application", "octet-stream")
        msg.add_attachment(raw, maintype=maintype, subtype=subtype, filename=filename)
    return msg


def attachment_path_by_name(file_name, username=None):
    safe_name = Path(file_name or "").name
    bases = (user_generated_dir(username), user_report_dir(username)) if username else (GENERATED_DIR, REPORT_DIR)
    for base in bases:
        path = base / safe_name
        if path.exists() and path.is_file():
            return path
    return None


def history_path_by_name(file_name, username=None):
    safe_name = Path(file_name or "").name
    path = (user_report_dir(username) if username else REPORT_DIR) / safe_name
    if path.exists() and path.is_file():
        return path
    return None


def clean_lines(value):
    return [line.strip() for line in (value or "").splitlines() if line.strip()]


def parse_table_lines(value, columns):
    if isinstance(value, list):
        rows = []
        for row in value:
            if isinstance(row, dict):
                keys = ["category", "content", "status", "plan", "progress", "difficulty"]
                values = [str(row.get(key, "") or "").strip() for key in keys]
                if columns == 4:
                    rows.append([values[0], values[1], values[2] or values[4], values[3] or values[5]])
                else:
                    rows.append([values[0], values[1], values[5] or values[3]])
            elif isinstance(row, list):
                rows.append(([str(item or "").strip() for item in row] + [""] * columns)[:columns])
        return [row for row in rows if any(row)]

    rows = []
    for line in clean_lines(value):
        parts = [part.strip() for part in re.split(r"\s*[|｜]\s*", line)]
        parts = (parts + [""] * columns)[:columns]
        rows.append(parts)
    return rows


def style_xlsx_cell(cell, horizontal="left", vertical="top"):
    cell.alignment = cell.alignment.copy(
        wrap_text=True,
        horizontal=horizontal,
        vertical=vertical,
    )


def cell_value(ws, row, col):
    value = ws.cell(row, col).value
    return "" if value is None else str(value).strip()


def xlsx_row_text(values):
    return " | ".join(values).strip()


def normalize_numbered_text(value):
    text = str(value or "").strip()
    if not text:
        return ""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[；;]\s*(?=\d+[、.．])", "\n", text)
    text = re.sub(r"(?<!^)(?<!\n)\s*(?=\d+[、.．])", "\n", text)
    text = re.sub(r"[；;]\s*(?=[*＊·-]\s*)", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def local_optimize_text(text):
    text = normalize_numbered_text(text)
    if not text:
        return ""
    replacements = {
        "后段": "后端",
        "滑框": "滑动窗口",
        "扒取": "抓取",
        "api": "API",
        "Api": "API",
        "roi": "ROI",
        "mns": "MNS",
        "ui": "UI",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    parts = []
    for line in text.splitlines():
        line = re.sub(r"^\s*(?:\d+[、.．]|[-*＊·])\s*", "", line).strip(" 。；;")
        if line:
            split_parts = [p.strip(" 。；;") for p in re.split(r"[；;]\s*", line) if p.strip(" 。；;")]
            parts.extend(split_parts)
    if len(parts) <= 1:
        parts = [p.strip(" 。；;") for p in re.split(r"，(?=(?:实现|完成|推动|优化|设计|增加|完善|梳理|分析|开展|补充))", parts[0]) if p.strip(" 。；;")]
    return "\n".join(f"{idx}、{part}。" for idx, part in enumerate(parts, start=1))


def normalize_assistant_output(text):
    text = normalize_numbered_text(text)
    if not text:
        return ""
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    normalized = []
    for idx, line in enumerate(lines, start=1):
        line = re.sub(r"^\s*\d+\s*[、.．]\s*", f"{idx}、", line)
        if not re.match(r"^\d+、", line):
            line = f"{idx}、{line.strip(' 。；;')}。"
        normalized.append(line)
    return "\n".join(normalized)


def assistant_settings():
    config = read_config()
    return {
        "url": os.getenv("ASSISTANT_API_URL", config.get("assistant_api_url", "")).rstrip("/"),
        "key": os.getenv("ASSISTANT_API_KEY", config.get("assistant_api_key", "")),
        "model": os.getenv("ASSISTANT_MODEL", config.get("assistant_model", "MiniMax-M2.7")),
        "prompt": os.getenv("ASSISTANT_PROMPT", config.get("assistant_prompt", DEFAULT_ASSISTANT_PROMPT)),
    }


def optimize_text(payload):
    text = str(payload.get("text", "") or "").strip()
    prompt = str(payload.get("prompt", "") or "").strip() or assistant_settings()["prompt"]
    if not text:
        return {"ok": True, "mode": "empty", "text": ""}

    settings = assistant_settings()
    if settings["url"] and settings["key"]:
        try:
            body = json.dumps(
                {
                    "model": settings["model"],
                    "messages": [
                        {
                            "role": "system",
                            "content": (
                                "你是中文工作周报写作助手。只输出优化后的正文，不要解释，"
                                "不要使用Markdown标题，不要添加用户未提供的事项。"
                            ),
                        },
                        {
                            "role": "user",
                            "content": f"{prompt}\n\n待优化内容：\n{text}",
                        },
                    ],
                    "temperature": 0.2,
                },
                ensure_ascii=False,
            ).encode("utf-8")
            request = urllib.request.Request(
                settings["url"] + "/v1/chat/completions",
                data=body,
                headers={
                    "Content-Type": "application/json",
                    "Authorization": "Bearer " + settings["key"],
                },
                method="POST",
            )
            with urllib.request.urlopen(request, timeout=30) as response:
                data = json.loads(response.read().decode("utf-8"))
            content = data["choices"][0]["message"]["content"].strip()
            return {"ok": True, "mode": "api", "model": settings["model"], "text": normalize_assistant_output(content)}
        except Exception as exc:
            return {
                "ok": True,
                "mode": "local_fallback",
                "warning": f"智能接口调用失败，已使用本地规则优化：{exc}",
                "text": local_optimize_text(text),
            }

    return {"ok": True, "mode": "local", "text": local_optimize_text(text)}


WEEKLY_AGENT_SYSTEM = (
    "你是\"周报助手\"犇犇，通过对话帮用户填写和发送周报。\\n\\n"
    "工作流（必须按顺序执行）：\\n"
    "1. 获取日期：如需要当前日期或本周周期，调用 utils.get_date 获取。\\n"
    "2. 生成草稿：用户一旦提供了本周工作内容（哪怕只有几条），立即调用 weekly.compose 整理成周报结构。不要过度追问。\\n"
    "3. 生成预览：用户说\"预览/生成附件/确认\"后，调用 weekly.preview 生成 Excel、预览图和邮件草稿。\\n"
    "4. 发送邮件：用户明确说\"确认发送/可以发/发吧\"后，调用 weekly.send_confirmed 发送邮件。\\n\\n"
    "注意事项：\\n"
    "- 不要编造用户没有提供的信息；没有的信息留空即可。\\n"
    "- 用户只要说了工作内容，哪怕很简短，也立即调用 weekly.compose，不要反复追问。\\n"
    "- 每一步先向用户说明你在做什么，再调用 Skill。\\n"
    "- 调用 weekly.preview 时，period 字段格式为 YYYY.MM.DD-YYYY.MM.DD。\\n"
    "- 如果用户只说了\"帮我写周报\"但完全没提工作内容，才需要询问。"
)

TRIP_AGENT_SYSTEM = (
    "你是\"出差报告助手\"犇犇，通过对话帮用户填写和发送出差报告。\\n\\n"
    "工作流（必须按顺序执行）：\\n"
    "1. 收集信息：引导用户提供出差时间、地点、目的、行程等。如需要当前日期，先调用 utils.get_date。\\n"
    "2. 生成报告：信息完整后，调用 document.generate 生成正式出差报告 Word 文件。\\n"
    "3. 发送邮件：用户确认后，发送邮件（如系统支持）。\\n\\n"
    "注意事项：\\n"
    "- 不要编造用户没有提供的信息。\\n"
    "- 每一步都要先向用户说明你在做什么，再调用 Skill。"
)

DIARY_AGENT_SYSTEM = (
    "你是\"工作日记助手\"犇犇，通过对话帮用户记录每天的工作日记。\\n\\n"
    "工作流：\\n"
    "1. 收集信息：通过自然对话了解用户今天的工作内容、明日计划和想法。\\n"
    "2. 保存日记：信息完整后，调用 diary.save 保存到系统。\\n"
    "3. 查看历史：用户想查看时，调用 diary.get 或 diary.list。\\n\\n"
    "注意事项：\\n"
    "- 不要编造用户没有提供的信息。\\n"
    "- 日记内容保持自然语言，不要过度结构化。"
)

MAIL_AGENT_SYSTEM = (
    "你是\"智能邮件助手\"，帮助用户阅读邮件、提炼重点、起草普通邮件、优化邮件正文。\\n"
    "你会收到当前邮件助手页面上的收件箱、选中邮件或写信表单上下文。\\n"
    "你的任务不是写周报或出差报告；除非用户明确要求，否则不要输出报告表单 JSON。\\n"
    "请用中文直接回答，必要时给出可复制的邮件主题和正文。不要编造未提供的邮件内容。"
)

NEWS_AGENT_SYSTEM = (
    "你是\"每日资讯助手\"，帮助用户理解轨道交通每日资讯、提炼重点、总结影响和建议行动。\\n"
    "你会收到当前资讯页面内容。普通用户不能配置资讯源；只有超级管理员能配置和生成。\\n"
    "你的任务不是写周报或出差报告；不要输出报告表单 JSON。\\n"
    "请基于已提供资讯回答，不要编造最新新闻。"
)

FORUM_AGENT_SYSTEM = (
    "你是\"金点子论坛助手\"，帮助用户围绕当前话题提炼观点、生成评论、总结讨论热度和下一步行动。\\n"
    "你会收到当前论坛话题、评论输入框或话题创建草稿。\\n"
    "你的任务不是写周报或出差报告；不要输出报告表单 JSON。\\n"
    "请用中文给出适合论坛讨论的内容；如果用户要评论，优先给出一段可直接发布的评论。"
)

GENERAL_AGENT_SYSTEM = (
    "你是\"AI 办公总助手\"，根据用户当前所在页面提供工作建议。\\n"
    "可协助周报、出差报告、工作日记、邮件、论坛和每日资讯相关问题。\\n"
    "如果当前页面不是周报或出差报告，不要输出报告表单 JSON。\\n"
    "回答要简洁、可执行，不要编造不存在的信息。"
)

DEFAULT_AGENT_WORKFLOWS = {
    "weekly": "1) weekly.compose 规范化并填入周报结构 → 2) weekly.preview 生成周报文件、预览图片和邮件草稿 → 3) 用户确认后 weekly.send_confirmed 发送邮件",
    "trip": "1) trip.prefill 获取历史出差报告预填 → 2) 用户补充信息后 document.generate 生成正式报告 → 3) 发送邮件",
    "diary": "1) diary.save 保存工作日记 → 可随时 diary.get / diary.list 查看",
    "mailassistant": "1) mailbox.list 查看邮件 → 2) text.optimize 优化正文或起草回复",
    "forum": "1) forum.list 查看话题 → 2) forum.add_comment 发表评论或 forum.create 发起话题",
    "news": "1) news.list 查看每日资讯 → 2) text.optimize 提炼重点和影响",
}

AGENT_SYSTEM_PROMPTS = {
    "weekly": WEEKLY_AGENT_SYSTEM,
    "trip": TRIP_AGENT_SYSTEM,
    "diary": DIARY_AGENT_SYSTEM,
    "mailassistant": MAIL_AGENT_SYSTEM,
    "news": NEWS_AGENT_SYSTEM,
    "forum": FORUM_AGENT_SYSTEM,
    "dashboard": GENERAL_AGENT_SYSTEM,
}

AGENT_WORKFLOWS = dict(DEFAULT_AGENT_WORKFLOWS)


def load_agent_config():
    """启动时加载 agent_config.json 并覆盖默认提示词和工作流。"""
    global WEEKLY_AGENT_SYSTEM, TRIP_AGENT_SYSTEM, DIARY_AGENT_SYSTEM
    global MAIL_AGENT_SYSTEM, NEWS_AGENT_SYSTEM, FORUM_AGENT_SYSTEM
    global GENERAL_AGENT_SYSTEM, AGENT_SYSTEM_PROMPTS, AGENT_WORKFLOWS
    cfg = read_agent_config()
    prompts = cfg.get("prompts", {})
    if prompts.get("weekly"):
        WEEKLY_AGENT_SYSTEM = prompts["weekly"]
    if prompts.get("trip"):
        TRIP_AGENT_SYSTEM = prompts["trip"]
    if prompts.get("diary"):
        DIARY_AGENT_SYSTEM = prompts["diary"]
    if prompts.get("mailassistant"):
        MAIL_AGENT_SYSTEM = prompts["mailassistant"]
    if prompts.get("news"):
        NEWS_AGENT_SYSTEM = prompts["news"]
    if prompts.get("forum"):
        FORUM_AGENT_SYSTEM = prompts["forum"]
    if prompts.get("dashboard"):
        GENERAL_AGENT_SYSTEM = prompts["dashboard"]
    AGENT_SYSTEM_PROMPTS = {
        "weekly": WEEKLY_AGENT_SYSTEM,
        "trip": TRIP_AGENT_SYSTEM,
        "diary": DIARY_AGENT_SYSTEM,
        "mailassistant": MAIL_AGENT_SYSTEM,
        "news": NEWS_AGENT_SYSTEM,
        "forum": FORUM_AGENT_SYSTEM,
        "dashboard": GENERAL_AGENT_SYSTEM,
    }
    workflows = cfg.get("workflows", {})
    for key, val in workflows.items():
        if key in AGENT_WORKFLOWS and val:
            AGENT_WORKFLOWS[key] = val


def weekly_compose_skill_detail():
    return {
        "purpose": "把用户自然语言描述的工作内容整理成系统周报表单结构，支持从历史周报迁移、补全分类、优化表达、生成下周计划。",
        "when_to_use": [
            "用户说“帮我写周报、整理本周工作、根据这些内容生成周报”。",
            "用户粘贴零散工作内容，需要拆分成 1、2、3、4 点。",
            "用户要求把上周计划转成本周工作内容，或生成下一周计划。",
            "用户希望体现工作量很多、语言简洁、无错别字。"
        ],
        "input_schema": {
            "period": "周报周期，例如 2026.05.11-2026.05.15，可选",
            "raw_work": "本周原始工作内容，自然语言或多行文本",
            "last_week_plan": "上次周报的下周计划，可选，用于迁移到本周工作",
            "key_work": "重点工作/关键项目，可选",
            "next_plan": "下周计划原始描述，可选",
            "difficulties": "困难、风险、需要协调事项，可选",
            "style": "输出风格，可选，默认：简洁、具体、体现工作量"
        },
        "output_schema": {
            "weekly_summary": [{"category": "工作分类", "content": "工作内容", "status": "完成情况", "plan": "后续计划"}],
            "weekly_follow": [{"category": "重点工作", "content": "工作内容", "progress": "当前进展", "difficulty": "困难与求助"}],
            "weekly_next": [{"category": "工作分类", "content": "下周计划", "difficulty": "困难与求助"}]
        },
        "call_example": {
            "reply": "我来根据这些内容生成周报草稿。",
            "skill_call": {
                "name": "weekly.compose",
                "arguments": {
                    "period": "2026.05.11-2026.05.15",
                    "raw_work": "实现平台复杂页面层次拆解；实现多模型部署管理；训练安全帽反光衣手套检测模型；完善API授权；增加权限管理；优化UI。",
                    "key_work": "算法平台、模型服务、数据治理、多模型部署",
                    "next_plan": "推进业务与算法服务后端联调；实现ROI滑动窗口切片和结果去重；拆分微服务架构。",
                    "style": "体现工作量多，编号清晰，简洁明了"
                }
            }
        }
    }


def weekly_preview_skill_detail():
    return {
        "purpose": "根据已规范化的周报结构生成正式周报文件、邮件草稿和周报预览图片，供用户确认。",
        "when_to_use": [
            "用户已提供或确认周报结构化内容，需要生成预览。",
            "用户说“生成周报预览、让我先看看、预览没问题再发”。",
            "大模型已经通过 weekly.compose 得到 weekly_summary、weekly_follow、weekly_next。"
        ],
        "input_schema": {
            "period": "周报周期，例如 2026.05.11-2026.05.15",
            "weekly_summary": [{"category": "工作分类", "content": "工作内容", "status": "完成情况", "plan": "后续计划"}],
            "weekly_follow": [{"category": "重点工作", "content": "工作内容", "progress": "当前进展", "difficulty": "困难与求助"}],
            "weekly_next": [{"category": "工作分类", "content": "下周计划", "difficulty": "困难与求助"}]
        },
        "output_schema": {
            "file": "生成的周报 Excel 文件名",
            "download_url": "周报文件下载地址",
            "preview_image_url": "周报预览图片地址",
            "mail_draft": "待确认的邮件草稿"
        },
        "call_example": {
            "reply": "我先生成周报预览，请你确认内容和格式。",
            "skill_call": {
                "name": "weekly.preview",
                "arguments": {
                    "period": "2026.05.11-2026.05.15",
                    "weekly_summary": [{"category": "算法平台", "content": "完成平台复杂页面层次拆解。", "status": "已完成", "plan": "持续优化交互"}],
                    "weekly_follow": [],
                    "weekly_next": [{"category": "联调", "content": "推进业务与算法服务后端联调。", "difficulty": ""}]
                }
            }
        }
    }


def weekly_send_skill_detail():
    return {
        "purpose": "在用户明确确认周报预览无误后，发送周报邮件。",
        "when_to_use": [
            "用户已经查看周报预览图片并明确说“确认发送、没问题发送、可以发”。",
            "不得在未生成预览、未确认时调用。",
            "发送前需要已有 weekly.preview 返回的 file，或 arguments 中提供 attachment。"
        ],
        "input_schema": {
            "attachment": "weekly.preview 生成的周报文件名",
            "to": "收件人，可选；为空使用邮件配置中的周报收件人",
            "cc": "抄送，可选；为空使用邮件配置中的周报抄送",
            "subject": "主题，可选",
            "body": "正文，可选"
        },
        "output_schema": {
            "mode": "sent|draft",
            "message": "发送结果"
        },
        "call_example": {
            "reply": "已确认预览无误，我现在发送周报邮件。",
            "skill_call": {
                "name": "weekly.send_confirmed",
                "arguments": {"attachment": "周颖超工作周报2026.05.11-2026.05.15.xlsx"}
            }
        }
    }


def skill_defs():
    return [
        {
            "name": "reports.list",
            "module": "报告",
            "title": "查询报告列表",
            "description": "查询当前用户的周报、出差报告、生成文件和历史报告。",
            "parameters": {"kind": "all|weekly|trip，可选"},
            "safe": True,
        },
        {
            "name": "weekly.prefill",
            "module": "周报",
            "title": "获取最新周报预填",
            "description": "读取最新历史周报，将上次下周计划迁移到本次工作内容。",
            "parameters": {},
            "safe": True,
        },
        {
            "name": "weekly.compose",
            "module": "周报",
            "title": "编写/设计周报草稿",
            "description": "调用配置的大模型 API，将原始工作内容整理成周报三段式结构：本周工作总结、重点工作跟进、下周工作计划。",
            "parameters": {
                "period": "周报周期，可选",
                "raw_work": "本周原始工作内容",
                "last_week_plan": "上次周报下周计划，可选",
                "key_work": "重点工作，可选",
                "next_plan": "下周计划，可选",
                "difficulties": "困难与求助，可选",
                "style": "输出风格，可选"
            },
            "safe": True,
            "detail": weekly_compose_skill_detail(),
        },
        {
            "name": "weekly.preview",
            "module": "周报",
            "title": "生成周报预览",
            "description": "根据已整理的周报结构生成正式周报 Excel、邮件草稿和周报预览图片，供用户确认。",
            "parameters": {
                "period": "周报周期",
                "weekly_summary": "本周工作总结数组",
                "weekly_follow": "重点工作跟进数组",
                "weekly_next": "下周工作计划数组"
            },
            "safe": False,
            "detail": weekly_preview_skill_detail(),
        },
        {
            "name": "weekly.send_confirmed",
            "module": "周报",
            "title": "确认后发送周报邮件",
            "description": "仅在用户确认周报预览无误后，使用当前账号邮件配置发送周报邮件。",
            "parameters": {
                "attachment": "周报文件名，通常来自 weekly.preview",
                "to": "收件人，可选",
                "cc": "抄送，可选",
                "subject": "主题，可选",
                "body": "正文，可选"
            },
            "safe": False,
            "detail": weekly_send_skill_detail(),
        },
        {
            "name": "trip.prefill",
            "module": "出差报告",
            "title": "获取最新出差报告预填",
            "description": "读取最新历史出差报告并填充当前出差报告草稿。",
            "parameters": {},
            "safe": True,
        },
        {
            "name": "document.generate",
            "module": "报告",
            "title": "生成正式报告文件",
            "description": "按模板生成周报 Excel 或出差报告 Word。",
            "parameters": {"kind": "weekly|trip", "weekly_summary": "数组", "weekly_follow": "数组", "weekly_next": "数组", "trip fields": "出差报告字段"},
            "safe": False,
        },
        {
            "name": "text.optimize",
            "module": "通用",
            "title": "优化文本",
            "description": "用配置的大模型 API 优化工作内容、邮件正文、论坛内容等文本。",
            "parameters": {"text": "待优化文本", "prompt": "优化要求，可选"},
            "safe": True,
        },
        {
            "name": "diary.save",
            "module": "工作日记",
            "title": "保存工作日记",
            "description": "保存当前用户指定日期的工作日记。",
            "parameters": {"date": "YYYY-MM-DD", "today_work": "今日工作", "tomorrow_plan": "明日计划", "thoughts": "想法心得"},
            "safe": False,
        },
        {
            "name": "diary.get",
            "module": "工作日记",
            "title": "读取工作日记",
            "description": "读取当前用户指定日期的工作日记。",
            "parameters": {"date": "YYYY-MM-DD"},
            "safe": True,
        },
        {
            "name": "diary.list",
            "module": "工作日记",
            "title": "查询工作日记列表",
            "description": "查询当前用户的工作日记列表。",
            "parameters": {"start": "开始日期，可选", "end": "结束日期，可选"},
            "safe": True,
        },
        {
            "name": "mailbox.list",
            "module": "邮件",
            "title": "查看收件箱",
            "description": "通过当前用户 IMAP 配置读取最近邮件。",
            "parameters": {"limit": "10/20/50，可选"},
            "safe": True,
        },
        {
            "name": "mailbox.detail",
            "module": "邮件",
            "title": "查看邮件详情",
            "description": "通过邮件 UID 读取邮件正文、HTML 预览和附件信息。",
            "parameters": {"uid": "邮件 UID"},
            "safe": True,
        },
        {
            "name": "mail.send",
            "module": "邮件",
            "title": "发送普通邮件",
            "description": "使用当前用户 SMTP 配置发送普通邮件。",
            "parameters": {"to": "收件人", "cc": "抄送，可选", "subject": "主题", "body": "正文"},
            "safe": False,
        },
        {
            "name": "forum.list",
            "module": "金点子论坛",
            "title": "查看金点子话题",
            "description": "查看当前用户可见的金点子论坛话题。",
            "parameters": {},
            "safe": True,
        },
        {
            "name": "forum.create",
            "module": "金点子论坛",
            "title": "发布金点子话题",
            "description": "创建新的金点子论坛话题。",
            "parameters": {"title": "标题", "body": "内容", "tags": "标签数组，可选"},
            "safe": False,
        },
        {
            "name": "forum.comment",
            "module": "金点子论坛",
            "title": "发布论坛评论",
            "description": "给指定金点子话题添加评论。",
            "parameters": {"id": "话题 ID", "body": "评论内容"},
            "safe": False,
        },
        {
            "name": "news.latest",
            "module": "资讯",
            "title": "查看每日资讯",
            "description": "查看系统生成的最新每日资讯。",
            "parameters": {},
            "safe": True,
        },
        {
            "name": "utils.get_date",
            "module": "通用",
            "title": "获取日期信息",
            "description": "获取当前日期、今天星期几、本周起止日期、本月起止日期、本季度起止日期、当前年度等日期信息，用于填写周报周期、出差时间等。",
            "parameters": {"format": "日期格式，可选，如 YYYY.MM.DD 或 YYYY-MM-DD，默认 YYYY.MM.DD"},
            "safe": True,
        },
    ]


def skill_doc_markdown():
    lines = [
        "# 智能办公助手 Skill 文档",
        "",
        "本系统把周报、出差报告、工作日记、邮件、论坛、每日资讯等业务能力封装为可由大模型调用的 Skill。",
        "",
        "## 调用协议",
        "",
        "大模型需要调用功能时，请返回严格 JSON，不要包裹 Markdown 代码块：",
        "",
        '{"reply":"给用户看的说明","skill_call":{"name":"skill.name","arguments":{}}}',
        "",
        "如果只是普通问答，不需要调用 Skill，则直接自然语言回复即可。",
        "",
        "## 安全约束",
        "",
        "- safe=true 的 Skill 为查询、预览、优化类能力，可直接执行。",
        "- safe=false 的 Skill 会产生写入、生成文件、发邮件或发布内容等动作；前端应在关键场景增加确认。",
        "- 所有 Skill 默认在当前登录用户空间内执行，不能跨用户读取数据。",
        "",
        "## Skill 列表",
        "",
    ]
    for item in skill_defs():
        lines.extend([
            f"### {item['name']} - {item['title']}",
            "",
            item["description"],
            "",
            f"- 安全级别：{'查询/预览' if item['safe'] else '会产生写入或外部动作'}",
            f"- 参数：`{json.dumps(item['parameters'], ensure_ascii=False)}`",
            "",
        ])
        if item.get("detail"):
            detail = item["detail"]
            lines.extend([
                "#### 详细设计",
                "",
                f"- 用途：{detail.get('purpose', '')}",
                "",
                "适用场景：",
            ])
            lines.extend([f"- {line}" for line in detail.get("when_to_use", [])])
            lines.extend([
                "",
                "输入 Schema：",
                "",
                "```json",
                json.dumps(detail.get("input_schema", {}), ensure_ascii=False, indent=2),
                "```",
                "",
                "输出 Schema：",
                "",
                "```json",
                json.dumps(detail.get("output_schema", {}), ensure_ascii=False, indent=2),
                "```",
                "",
                "调用示例：",
                "",
                "```json",
                json.dumps(detail.get("call_example", {}), ensure_ascii=False, indent=2),
                "```",
                "",
            ])
    return "\n".join(lines)


def public_skill_docs():
    return {"ok": True, "skills": skill_defs(), "markdown": skill_doc_markdown()}


def skill_by_name(name):
    return next((item for item in skill_defs() if item.get("name") == name), None)


def parse_json_object(text):
    raw = (text or "").strip()
    raw = re.sub(r"^```(?:json)?\s*|\s*```$", "", raw, flags=re.I | re.S).strip()
    try:
        return json.loads(raw)
    except Exception:
        match = re.search(r"\{.*\}", raw, flags=re.S)
        if not match:
            raise
        return json.loads(match.group(0))


def normalize_weekly_rows(data):
    def clean_rows(rows, keys):
        result = []
        for row in rows or []:
            if not isinstance(row, dict):
                continue
            item = {key: str(row.get(key, "") or "").strip() for key in keys}
            if any(item.values()):
                result.append(item)
        return result

    return {
        "weekly_summary": clean_rows(data.get("weekly_summary", []), ["category", "content", "status", "plan"]),
        "weekly_follow": clean_rows(data.get("weekly_follow", []), ["category", "content", "progress", "difficulty"]),
        "weekly_next": clean_rows(data.get("weekly_next", []), ["category", "content", "difficulty"]),
    }


def weekly_payload_from_draft(args):
    draft = args.get("draft") if isinstance(args.get("draft"), dict) else args
    rows = normalize_weekly_rows(draft)
    return {
        "kind": "weekly",
        "period": str(args.get("period") or draft.get("period") or datetime.now().strftime("%Y.%m.%d-%Y.%m.%d")).strip(),
        "weekly_summary": rows["weekly_summary"],
        "weekly_follow": rows["weekly_follow"],
        "weekly_next": rows["weekly_next"],
    }


def text_wrap_units(text, max_units):
    lines = []
    current = ""
    units = 0
    for ch in str(text or ""):
        if ch == "\n":
            lines.append(current)
            current = ""
            units = 0
            continue
        weight = 2 if ord(ch) > 127 else 1
        if units + weight > max_units and current:
            lines.append(current)
            current = ch
            units = weight
        else:
            current += ch
            units += weight
    if current or not lines:
        lines.append(current)
    return lines


def load_preview_font(size=18, bold=False):
    try:
        from PIL import ImageFont
        candidates = [
            "/System/Library/Fonts/PingFang.ttc",
            "/System/Library/Fonts/STHeiti Light.ttc",
            "/Library/Fonts/Arial Unicode.ttf",
        ]
        for path in candidates:
            if Path(path).exists():
                return ImageFont.truetype(path, size=size)
        return ImageFont.load_default()
    except Exception:
        return None


def create_weekly_preview_image(payload, output_path):
    try:
        from PIL import Image, ImageDraw
    except Exception as exc:
        raise ValueError("缺少 Pillow，无法生成周报预览图片") from exc

    rows = normalize_weekly_rows(payload)
    title_font = load_preview_font(30, True)
    section_font = load_preview_font(22, True)
    text_font = load_preview_font(17)
    small_font = load_preview_font(14)
    width = 1400
    margin = 44
    y = 36

    def estimate_section_height(title, data, keys):
        height = 54
        for row in data or [{"category": "", "content": "暂无内容"}]:
            text = " | ".join(str(row.get(k, "") or "") for k in keys)
            height += max(72, len(text_wrap_units(text, 88)) * 23 + 34)
        return height

    total_height = 110
    total_height += estimate_section_height("本周工作总结", rows["weekly_summary"], ["category", "content", "status", "plan"])
    total_height += estimate_section_height("重点工作跟进", rows["weekly_follow"], ["category", "content", "progress", "difficulty"])
    total_height += estimate_section_height("下周工作计划", rows["weekly_next"], ["category", "content", "difficulty"])
    total_height = max(900, total_height + 80)

    image = Image.new("RGB", (width, total_height), "#f4f7fb")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((24, 24, width - 24, total_height - 24), radius=18, fill="#ffffff", outline="#dbe5f1", width=2)
    title = f"工作周报（{payload.get('period') or ''}）"
    draw.text((margin, y), title, fill="#172033", font=title_font)
    y += 58
    draw.text((margin, y), "预览图用于确认内容，正式格式以生成的 Excel 附件为准。", fill="#64748b", font=small_font)
    y += 38

    def section(title, data, columns):
        nonlocal y
        draw.rounded_rectangle((margin, y, width - margin, y + 38), radius=8, fill="#eff6ff", outline="#bfdbfe")
        draw.text((margin + 16, y + 8), title, fill="#1d4ed8", font=section_font)
        y += 50
        data = data or [{}]
        for index, row in enumerate(data, 1):
            line_parts = [f"{label}：{row.get(key, '')}" for key, label in columns if row.get(key, "")]
            if not line_parts:
                line_parts = ["暂无内容"]
            wrapped = []
            for part in line_parts:
                wrapped.extend(text_wrap_units(part, 92))
            card_h = max(70, len(wrapped) * 24 + 34)
            draw.rounded_rectangle((margin, y, width - margin, y + card_h), radius=8, fill="#ffffff", outline="#dbe5f1")
            draw.text((margin + 16, y + 14), f"{index}.", fill="#2563eb", font=text_font)
            text_y = y + 14
            for line in wrapped:
                draw.text((margin + 54, text_y), line, fill="#172033", font=text_font)
                text_y += 24
            y += card_h + 12
        y += 12

    section("一、本周工作总结", rows["weekly_summary"], [("category", "分类"), ("content", "内容"), ("status", "完成情况"), ("plan", "后续计划")])
    section("二、重点工作跟进", rows["weekly_follow"], [("category", "分类"), ("content", "内容"), ("progress", "进展"), ("difficulty", "困难")])
    section("三、下周工作计划", rows["weekly_next"], [("category", "分类"), ("content", "内容"), ("difficulty", "困难")])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    image.crop((0, 0, width, min(total_height, y + 36))).save(output_path)
    return output_path


def compose_weekly_skill(arguments):
    args = arguments or {}
    source_text = "\n".join(
        str(args.get(key, "") or "").strip()
        for key in ("last_week_plan", "raw_work", "key_work", "next_plan", "difficulties")
        if str(args.get(key, "") or "").strip()
    )
    if not source_text:
        raise ValueError("请提供 raw_work、last_week_plan、key_work 或 next_plan 中至少一项内容")

    settings = assistant_settings()
    prompt = (
        "你是企业周报 Skill，负责把原始工作描述整理成系统可写入的周报 JSON。"
        "要求：1、体现工作量和推进成果；2、语言简洁明了；3、修正错别字；4、不编造未提供事项；"
        "5、分类尽量贴近软件/算法/数据治理/联调/架构/UI/管理等实际工作；"
        "6、输出严格 JSON，不要 Markdown。"
        "JSON 格式："
        "{\"weekly_summary\":[{\"category\":\"\",\"content\":\"\",\"status\":\"\",\"plan\":\"\"}],"
        "\"weekly_follow\":[{\"category\":\"\",\"content\":\"\",\"progress\":\"\",\"difficulty\":\"\"}],"
        "\"weekly_next\":[{\"category\":\"\",\"content\":\"\",\"difficulty\":\"\"}]}"
    )
    user_content = json.dumps(
        {
            "period": args.get("period", ""),
            "raw_work": args.get("raw_work", ""),
            "last_week_plan": args.get("last_week_plan", ""),
            "key_work": args.get("key_work", ""),
            "next_plan": args.get("next_plan", ""),
            "difficulties": args.get("difficulties", ""),
            "style": args.get("style", "简洁、具体、体现工作量"),
        },
        ensure_ascii=False,
        indent=2,
    )

    if settings["url"] and settings["key"]:
        data = request_json(
            settings["url"] + "/v1/chat/completions",
            settings["key"],
            {
                "model": settings["model"],
                "messages": [
                    {"role": "system", "content": prompt},
                    {"role": "user", "content": user_content},
                ],
                "temperature": 0.25,
            },
            "POST",
            60,
        )
        content = data["choices"][0]["message"]["content"].strip()
        rows = normalize_weekly_rows(parse_json_object(content))
        return {"ok": True, "mode": "api", "model": settings["model"], "draft": rows}

    optimized = local_optimize_text(source_text)
    summary = [
        {"category": "工作推进", "content": line, "status": "已推进", "plan": ""}
        for line in optimized.splitlines()
        if line.strip()
    ]
    return {
        "ok": True,
        "mode": "local_fallback",
        "warning": "未配置大模型 API，已使用本地规则生成基础周报草稿。",
        "draft": {"weekly_summary": summary, "weekly_follow": [], "weekly_next": []},
    }


def weekly_preview_skill(arguments, username):
    payload = weekly_payload_from_draft(arguments or {})
    path = generate_weekly(payload, username)
    draft = compose_draft("weekly", path.name, username)
    preview_name = path.with_suffix(".png").name
    preview_path = user_generated_dir(username) / preview_name
    create_weekly_preview_image(payload, preview_path)
    return {
        "ok": True,
        "file": path.name,
        "download_url": "/download?file=" + urllib.parse.quote(path.name),
        "preview_image": preview_name,
        "preview_image_url": "/preview-image?file=" + urllib.parse.quote(preview_name),
        "mail_draft": draft,
        "message": "周报预览已生成，请确认预览图片和邮件草稿；确认无误后再调用 weekly.send_confirmed。",
    }


def weekly_send_confirmed_skill(arguments, username):
    args = arguments or {}
    attachment = str(args.get("attachment", "") or "").strip()
    if not attachment:
        raise ValueError("请提供 weekly.preview 生成的 attachment 文件名")
    draft = compose_draft("weekly", attachment, username)
    payload = {
        "to": args.get("to") or draft.get("to", ""),
        "cc": args.get("cc") or draft.get("cc", ""),
        "subject": args.get("subject") or draft.get("subject", ""),
        "body": args.get("body") or draft.get("body", ""),
        "body_html": args.get("body_html") or draft.get("body_html", ""),
        "attachment": attachment,
    }
    return send_mail(payload, username)


def get_date_skill(args):
    from datetime import datetime, timedelta

    fmt = str(args.get("format") or "YYYY.MM.DD").strip()
    if fmt == "YYYY-MM-DD":
        date_fmt = "%Y-%m-%d"
    elif fmt == "YYYY/MM/DD":
        date_fmt = "%Y/%m/%d"
    else:
        date_fmt = "%Y.%m.%d"

    today = datetime.now()
    weekday_names = ["星期一", "星期二", "星期三", "星期四", "星期五", "星期六", "星期日"]
    weekday = weekday_names[today.weekday()]

    # 本周起止（周一到周日）
    monday = today - timedelta(days=today.weekday())
    sunday = monday + timedelta(days=6)
    week_start = monday.strftime(date_fmt)
    week_end = sunday.strftime(date_fmt)

    # 本月起止
    month_start = today.replace(day=1)
    if today.month == 12:
        next_month = today.replace(year=today.year + 1, month=1, day=1)
    else:
        next_month = today.replace(month=today.month + 1, day=1)
    month_end = next_month - timedelta(days=1)

    # 本季度起止
    quarter = (today.month - 1) // 3 + 1
    quarter_start_month = (quarter - 1) * 3 + 1
    quarter_start = today.replace(month=quarter_start_month, day=1)
    if quarter == 4:
        next_q_start = today.replace(year=today.year + 1, month=1, day=1)
    else:
        next_q_start = today.replace(month=quarter_start_month + 3, day=1)
    quarter_end = next_q_start - timedelta(days=1)

    # 第几周（ISO week）
    iso_year, iso_week, _ = today.isocalendar()

    return {
        "ok": True,
        "today": today.strftime(date_fmt),
        "weekday": weekday,
        "week_range": f"{week_start}-{week_end}",
        "week_start": week_start,
        "week_end": week_end,
        "month_range": f"{month_start.strftime(date_fmt)}-{month_end.strftime(date_fmt)}",
        "quarter": f"Q{quarter}",
        "quarter_range": f"{quarter_start.strftime(date_fmt)}-{quarter_end.strftime(date_fmt)}",
        "year": str(today.year),
        "iso_week": f"{iso_year} 年第 {iso_week} 周",
    }


def execute_skill(name, arguments, username):
    args = arguments or {}
    if name == "utils.get_date":
        return get_date_skill(args)
    if name == "reports.list":
        kind = str(args.get("kind", "all") or "all")
        reports = all_files(username)
        if kind in ("weekly", "trip"):
            reports = [item for item in reports if item.get("kind") == kind]
        return {"ok": True, "reports": reports[:50]}
    if name == "weekly.prefill":
        return weekly_prefill(username)
    if name == "weekly.compose":
        return compose_weekly_skill(args)
    if name == "weekly.preview":
        return weekly_preview_skill(args, username)
    if name == "weekly.send_confirmed":
        return weekly_send_confirmed_skill(args, username)
    if name == "trip.prefill":
        return trip_prefill(username)
    if name == "document.generate":
        return generate_document(args, username)
    if name == "text.optimize":
        return optimize_text(args)
    if name == "diary.save":
        return save_diary(args, username)
    if name == "diary.get":
        return get_diary(str(args.get("date", "") or ""), username)
    if name == "diary.list":
        return list_diaries(args, username)
    if name == "mailbox.list":
        return list_inbox_messages(username, args.get("limit", 20), bool(args.get("refresh", False)))
    if name == "mailbox.detail":
        return get_inbox_message(username, str(args.get("uid", "") or ""), bool(args.get("refresh", False)))
    if name == "mail.send":
        args["body_html"] = ""
        return send_mail(args, username)
    if name == "forum.list":
        return forum_list_topics(username)
    if name == "forum.create":
        return forum_create_topic(args, username)
    if name == "forum.comment":
        return forum_add_comment(args, username)
    if name == "news.latest":
        return news_latest(False)
    raise ValueError("未知 Skill：" + str(name))


def skill_test(payload, username):
    name = str(payload.get("name", "") or "").strip()
    skill = skill_by_name(name)
    if not skill:
        raise ValueError("未知 Skill：" + name)
    if not skill.get("safe") and not payload.get("confirm_unsafe"):
        raise ValueError("该 Skill 会产生写入、生成文件、发邮件或发布内容，请勾选确认后再测试")

    arguments = payload.get("arguments") or {}
    if not isinstance(arguments, dict):
        raise ValueError("调用参数必须是 JSON 对象")

    instruction = str(payload.get("instruction", "") or "").strip()
    model_used = ""
    if instruction:
        settings = assistant_settings()
        if not settings["url"] or not settings["key"]:
            raise ValueError("未配置 AI 接口，请先在系统配置中设置 NewAPI 地址和 Key")
        system = (
            "你是智能办公助手的 Skill 测试器。"
            "请根据用户的自然语言测试要求，为指定 Skill 生成 arguments JSON 对象。"
            "只能返回 JSON 对象本身，不要返回 Markdown，不要调用其他 Skill。"
            "必须严格贴合 Skill 参数说明，不要编造用户没有提供的业务事实。"
        )
        user_prompt = (
            "待测试 Skill：\n"
            + json.dumps(skill, ensure_ascii=False, indent=2)
            + "\n\n已有参数草稿：\n"
            + json.dumps(arguments, ensure_ascii=False, indent=2)
            + "\n\n自然语言测试要求：\n"
            + instruction
        )
        data = request_json(
            settings["url"] + "/v1/chat/completions",
            settings["key"],
            {
                "model": settings["model"],
                "messages": [
                    {"role": "system", "content": system},
                    {"role": "user", "content": user_prompt},
                ],
                "temperature": 0.2,
            },
            "POST",
            60,
        )
        content = data["choices"][0]["message"]["content"].strip()
        generated = parse_json_object(content)
        if not isinstance(generated, dict):
            raise ValueError("大模型未返回有效的 arguments JSON 对象")
        arguments = generated
        model_used = settings["model"]

    result = execute_skill(name, arguments, username)
    return {
        "ok": True,
        "skill": name,
        "model": model_used,
        "arguments": arguments,
        "result": result,
    }


def _extract_balanced_json(text):
    start = text.find("{")
    if start == -1:
        return None
    count = 0
    for i in range(start, len(text)):
        if text[i] == "{":
            count += 1
        elif text[i] == "}":
            count -= 1
            if count == 0:
                return text[start : i + 1]
    return None


def _parse_xml_tool_call(text):
    """解析 <invoke name=\"...\">...<parameter>..</parameter></invoke> 格式的 tool call。"""
    invoke_match = re.search(r"<invoke\s+name=\"([^\"]+)\"[^>]*>(.*?)</invoke>", text, flags=re.S)
    if not invoke_match:
        # 也尝试 <minimax:tool_call>...<invoke>...</invoke>...</minimax:tool_call>
        invoke_match = re.search(r"<invoke\s+name=[\"']([^\"']+)[\"'][^>]*>(.*?)</invoke>", text, flags=re.S)
    if not invoke_match:
        return None
    name = invoke_match.group(1).strip()
    inner = invoke_match.group(2)
    arguments = {}
    for param_match in re.finditer(r"<parameter\s+name=\"([^\"]+)\">(.*?)</parameter>", inner, flags=re.S):
        key = param_match.group(1).strip()
        val = param_match.group(2).strip()
        try:
            arguments[key] = json.loads(val)
        except Exception:
            arguments[key] = val
    return {"name": name, "arguments": arguments}


def parse_skill_call(text):
    raw = (text or "").strip()
    raw = re.sub(r"^```(?:json)?\s*|\s*```$", "", raw, flags=re.I | re.S).strip()
    # Try direct parse first
    data = None
    try:
        data = parse_json_object(raw)
    except Exception:
        pass
    # Try extracting [TOOL_CALL] ... [/TOOL_CALL] format
    if not data:
        tool_match = re.search(r"\[TOOL_CALL\]\s*(.*?)\s*\[/TOOL_CALL\]", raw, flags=re.S)
        if tool_match:
            try:
                data = parse_json_object(tool_match.group(1))
            except Exception:
                pass
    # Try extracting JSON object from mixed text and fix unescaped control chars
    if not data:
        json_str = _extract_balanced_json(raw)
        if json_str:
            try:
                data = json.loads(json_str)
            except Exception:
                try:
                    fixed = json_str.replace("\n", "\\n").replace("\r", "\\r")
                    data = json.loads(fixed)
                except Exception:
                    pass
    if isinstance(data, dict):
        call = data.get("skill_call")
        if isinstance(call, dict) and call.get("name"):
            return {
                "reply": data.get("reply", ""),
                "name": str(call.get("name", "")).strip(),
                "arguments": call.get("arguments") or {},
            }
    # Try XML format (MiniMax etc.)
    xml_call = _parse_xml_tool_call(raw)
    if xml_call:
        return {
            "reply": "",
            "name": xml_call["name"],
            "arguments": xml_call["arguments"],
        }
    return None


def _clean_agent_reply(text):
    """从大模型回复中提取自然语言，去掉嵌入的 JSON/Skill 调用标记。"""
    if not text:
        return text
    # 去掉 [TOOL_CALL] ... [/TOOL_CALL]
    cleaned = re.sub(r"\[TOOL_CALL\].*?\[/TOOL_CALL\]", "", text, flags=re.S).strip()
    # 去掉 <minimax:tool_call> ... </minimax:tool_call>
    cleaned = re.sub(r"<minimax:tool_call>.*?</minimax:tool_call>", "", cleaned, flags=re.S).strip()
    # 去掉 <invoke>...</invoke>
    cleaned = re.sub(r"<invoke[^>]*>.*?</invoke>", "", cleaned, flags=re.S).strip()
    # 去掉独立的大括号 JSON 块（如果它占了多行或紧跟在文字后面）
    cleaned = re.sub(r"\n?\s*\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}\s*\n?", "\n", cleaned).strip()
    # 去掉 Markdown 代码块
    cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", cleaned, flags=re.I | re.S).strip()
    return cleaned


def agent_chat(payload, username=""):
    kind = payload.get("kind", "weekly")
    messages = payload.get("messages", [])
    settings = assistant_settings()
    if not settings["url"] or not settings["key"]:
        return {"ok": False, "error": "未配置 AI 接口，请在系统配置中设置 NewAPI 地址和 Key"}

    system = AGENT_SYSTEM_PROMPTS.get(kind) or AGENT_SYSTEM_PROMPTS.get("weekly")
    system = (
        system
        + "\\n\\n你现在运行在“智能办公助手 Skill 模式”。"
        + "\\n如果用户要求你操作软件功能，请从下面 Skill 中选择一个调用。"
        + "\\n周报工作流必须按顺序执行：1) weekly.compose 规范化并填入周报结构；2) weekly.preview 生成周报文件、预览图片和邮件草稿；3) 只有用户明确确认预览无误后，才能调用 weekly.send_confirmed 发送邮件。"
        + "\\n调用时必须且只能输出严格 JSON（禁止 Markdown、XML、\[TOOL_CALL\]）：{\"reply\":\"说明\",\"skill_call\":{\"name\":\"skill.name\",\"arguments\":{}}}"
        + "\\n如果不需要操作软件功能，直接自然语言回复。"
        + "\\n如需当前日期、本周起止日期、今天是星期几等时间信息，可调用 utils.get_date。"
        + "\\n可用 Skill：\\n"
        + json.dumps(skill_defs(), ensure_ascii=False)
    )
    api_messages = [{"role": "system", "content": system}]
    for m in messages:
        if not isinstance(m, dict):
            continue
        role = m.get("role")
        content = str(m.get("content", "") or "").strip()
        if role not in ("user", "assistant") or not content:
            continue
        api_messages.append({"role": role, "content": content})

    try:
        executed_calls = []
        max_rounds = 3
        for _round in range(max_rounds):
            data = request_json(
                settings["url"] + "/v1/chat/completions",
                settings["key"],
                {
                    "model": settings["model"],
                    "messages": api_messages,
                    "temperature": 0.6 if _round == 0 else 0.4,
                },
                "POST",
                60,
            )
            content = data["choices"][0]["message"]["content"].strip()
            skill_call = parse_skill_call(content)
            if not skill_call:
                # 没有新的 Skill 调用，直接返回最终回复
                final_reply = _clean_agent_reply(content)
                if executed_calls:
                    return {
                        "ok": True,
                        "reply": final_reply,
                        "skill_calls": executed_calls,
                    }
                return {"ok": True, "reply": final_reply}
            # 执行 Skill
            result = execute_skill(skill_call["name"], skill_call["arguments"], username)
            executed_calls.append({
                "name": skill_call["name"],
                "arguments": skill_call["arguments"],
                "result": result,
            })
            # 将结果反馈给大模型，继续下一轮
            api_messages.append({"role": "assistant", "content": content})
            api_messages.append({
                "role": "user",
                "content": (
                    "[系统通知：你刚才调用了 Skill '" + skill_call["name"] + "'，执行结果如下。"
                    "如果任务已完成，请用自然语言回复用户。"
                    "如果还需要继续操作（如先生成草稿再生成预览），可以继续调用下一个 Skill。]\n"
                    + json.dumps(result, ensure_ascii=False)
                ),
            })
        # 达到最大轮次，返回最后一次回复
        return {
            "ok": True,
            "reply": content,
            "skill_calls": executed_calls,
        }
    except Exception as exc:
        return {"ok": False, "error": str(exc)}


def save_agent_config(payload):
    cfg = read_agent_config()
    if "prompts" in payload:
        cfg["prompts"] = payload["prompts"]
    if "workflows" in payload:
        cfg["workflows"] = payload["workflows"]
    write_agent_config(cfg)
    load_agent_config()
    return {"ok": True, "message": "犇犇配置已保存并生效"}


def agent_orchestration():
    return {
        "ok": True,
        "agents": AGENT_SYSTEM_PROMPTS,
        "skill_mode_suffix": (
            "\\n\\n你现在运行在「智能办公助手 Skill 模式」。"
            "\\n如果用户要求你操作软件功能，请从下面 Skill 中选择一个调用。"
            '\\n调用时只输出严格 JSON：{\\"reply\\":\\"说明\\",\\"skill_call\\":{\\"name\\":\\"skill.name\\",\\"arguments\\":{}}}'
            "\\n如果不需要操作软件功能，直接自然语言回复。"
        ),
        "workflows": AGENT_WORKFLOWS,
        "skills": skill_defs(),
    }


# ===== 工作日记 =====

def read_json_lenient(path):
    text = path.read_text(encoding="utf-8")
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        data, _ = json.JSONDecoder().raw_decode(text)
        return data


def save_diary(payload, username):
    date_str = str(payload.get("date", "") or "").strip()
    if not date_str:
        raise ValueError("请选择日记日期")
    diary_dir = user_diary_dir(username)
    diary_dir.mkdir(parents=True, exist_ok=True)
    path = diary_dir / f"{date_str}.json"
    now = datetime.now().isoformat(timespec="seconds")
    data = {
        "date": date_str,
        "today_work": str(payload.get("today_work", "") or "").strip(),
        "tomorrow_plan": str(payload.get("tomorrow_plan", "") or "").strip(),
        "thoughts": str(payload.get("thoughts", "") or "").strip(),
        "updated_at": now,
    }
    if path.exists():
        old = read_json_lenient(path)
        data["created_at"] = old.get("created_at", now)
    else:
        data["created_at"] = now
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return {"ok": True, "diary": data}


def get_diary(date_str, username):
    if not date_str:
        return {"ok": True, "diary": None}
    path = user_diary_dir(username) / f"{date_str}.json"
    if not path.exists():
        return {"ok": True, "diary": None}
    return {"ok": True, "diary": read_json_lenient(path)}


def list_diaries(payload, username):
    diary_dir = user_diary_dir(username)
    if not diary_dir.exists():
        return {"ok": True, "diaries": []}
    items = []
    for p in sorted(diary_dir.glob("*.json"), reverse=True):
        try:
            d = read_json_lenient(p)
            items.append({
                "date": d.get("date", p.stem),
                "today_work_preview": d.get("today_work", "")[:80],
                "created_at": d.get("created_at", ""),
                "updated_at": d.get("updated_at", ""),
            })
        except Exception:
            continue
    limit = int(payload.get("limit", 0) or 0)
    if limit > 0:
        items = items[:limit]
    return {"ok": True, "diaries": items}


def delete_diary(date_str, username):
    if not date_str:
        raise ValueError("请选择要删除的日记日期")
    path = user_diary_dir(username) / f"{date_str}.json"
    if path.exists():
        path.unlink()
    return {"ok": True}


def summarize_diaries_for_weekly(payload, username):
    start_date = str(payload.get("start_date", "") or "").strip()
    end_date = str(payload.get("end_date", "") or "").strip()
    if not start_date or not end_date:
        raise ValueError("请选择日记日期范围")
    diary_dir = user_diary_dir(username)
    contents = []
    cur = datetime.strptime(start_date, "%Y-%m-%d")
    end = datetime.strptime(end_date, "%Y-%m-%d")
    while cur <= end:
        date_str = cur.strftime("%Y-%m-%d")
        path = diary_dir / f"{date_str}.json"
        if path.exists():
            try:
                d = json.loads(path.read_text(encoding="utf-8"))
                parts = [f"【{date_str}】"]
                if d.get("today_work"):
                    parts.append(f"今日工作：{d['today_work']}")
                if d.get("tomorrow_plan"):
                    parts.append(f"明日计划：{d['tomorrow_plan']}")
                if d.get("thoughts"):
                    parts.append(f"思路想法：{d['thoughts']}")
                contents.append("\n".join(parts))
            except Exception:
                pass
        cur += __import__("datetime").timedelta(days=1)
    if not contents:
        return {"ok": True, "mode": "empty", "summary": "", "warning": "该日期范围内没有工作日记"}
    full_text = "\n\n".join(contents)
    settings = assistant_settings()
    prompt = (
        "请根据下面的工作日记内容，总结生成周报的三个部分。"
        "要求：1. 本周工作总结：概括已完成的主要工作，按要点列出；"
        "2. 重点工作跟进：提炼需要持续跟进的重大事项和当前进展；"
        "3. 下周工作计划：基于日记中的明日计划和思路，整理出下周工作安排。"
        "输出格式：\n本周工作总结：\n1. ...\n2. ...\n\n重点工作跟进：\n1. ...\n2. ...\n\n下周工作计划：\n1. ...\n2. ..."
    )
    if settings["url"] and settings["key"]:
        try:
            body = json.dumps(
                {
                    "model": settings["model"],
                    "messages": [
                        {"role": "system", "content": "你是中文工作周报写作助手。只输出总结内容，不要解释。"},
                        {"role": "user", "content": f"{prompt}\n\n工作日记内容：\n{full_text}"},
                    ],
                    "temperature": 0.3,
                },
                ensure_ascii=False,
            ).encode("utf-8")
            req = urllib.request.Request(
                settings["url"] + "/v1/chat/completions",
                data=body,
                headers={"Content-Type": "application/json", "Authorization": "Bearer " + settings["key"]},
                method="POST",
            )
            with urllib.request.urlopen(req, timeout=60) as resp:
                data = json.loads(resp.read().decode("utf-8"))
            content = data["choices"][0]["message"]["content"].strip()
            return {"ok": True, "mode": "api", "summary": content}
        except Exception as exc:
            return {"ok": False, "error": f"AI 总结失败：{exc}"}
    return {"ok": False, "error": "未配置 AI 接口，请在系统配置中设置 API 地址和 Key"}


# ===== 金点子论坛 =====

def forum_store_path(username):
    forum_dir = USER_DATA_DIR / "_forum"
    forum_dir.mkdir(parents=True, exist_ok=True)
    return forum_dir / "topics.json"


def load_forum_topics(username):
    path = forum_store_path(username)
    if not path.exists():
        return []
    try:
        data = read_json_lenient(path)
        return data if isinstance(data, list) else []
    except Exception:
        return []


def save_forum_topics(username, topics):
    forum_store_path(username).write_text(json.dumps(topics, ensure_ascii=False, indent=2), encoding="utf-8")


def forum_public_topic(topic, with_comments=False):
    item = dict(topic)
    comments = item.get("comments", [])
    likes = item.get("likes", [])
    item["comment_count"] = len(comments)
    item["like_count"] = len(likes)
    item["view_count"] = int(item.get("view_count", 0) or 0)
    item["heat"] = item["view_count"] + item["like_count"] * 5 + item["comment_count"] * 3
    if not with_comments:
        item.pop("comments", None)
    item.pop("likes", None)
    return item


def forum_list_topics(username):
    topics = sorted(load_forum_topics(username), key=lambda item: item.get("updated_at", ""), reverse=True)
    return {"ok": True, "topics": [forum_public_topic(topic) for topic in topics]}


def forum_get_topic(payload, username):
    topic_id = str(payload.get("id", "") or "").strip()
    topics = load_forum_topics(username)
    for topic in topics:
        if topic.get("id") == topic_id:
            topic["view_count"] = int(topic.get("view_count", 0) or 0) + 1
            save_forum_topics(username, topics)
            return {"ok": True, "topic": forum_public_topic(topic, with_comments=True)}
    raise ValueError("话题不存在")


def forum_create_topic(payload, username):
    title = str(payload.get("title", "") or "").strip()
    body = str(payload.get("body", "") or "").strip()
    source = str(payload.get("source", "") or "user").strip() or "user"
    if not title:
        raise ValueError("请填写话题标题")
    if not body:
        raise ValueError("请填写话题内容")
    now = datetime.now().isoformat(timespec="seconds")
    user = find_user(username) or {}
    topic = {
        "id": secrets.token_urlsafe(10),
        "title": title,
        "body": body,
        "source": source,
        "author": user.get("name") or username,
        "created_at": now,
        "updated_at": now,
        "view_count": 0,
        "likes": [],
        "comments": [],
    }
    topics = load_forum_topics(username)
    topics.append(topic)
    save_forum_topics(username, topics)
    return {"ok": True, "topic": forum_public_topic(topic, with_comments=True)}


def forum_add_comment(payload, username):
    topic_id = str(payload.get("topic_id", "") or "").strip()
    content = str(payload.get("content", "") or "").strip()
    parent_id = str(payload.get("parent_id", "") or "").strip()
    if not topic_id:
        raise ValueError("请选择话题")
    if not content:
        raise ValueError("请填写讨论内容")
    topics = load_forum_topics(username)
    user = find_user(username) or {}
    now = datetime.now().isoformat(timespec="seconds")
    for topic in topics:
        if topic.get("id") == topic_id:
            topic.setdefault("comments", []).append({
                "id": secrets.token_urlsafe(8),
                "parent_id": parent_id,
                "author": user.get("name") or username,
                "content": content,
                "created_at": now,
            })
            topic["updated_at"] = now
            save_forum_topics(username, topics)
            return {"ok": True, "topic": forum_public_topic(topic, with_comments=True)}
    raise ValueError("话题不存在")


def forum_toggle_like(payload, username):
    topic_id = str(payload.get("topic_id", "") or "").strip()
    if not topic_id:
        raise ValueError("请选择话题")
    topics = load_forum_topics(username)
    now = datetime.now().isoformat(timespec="seconds")
    for topic in topics:
        if topic.get("id") == topic_id:
            likes = topic.setdefault("likes", [])
            if username in likes:
                likes.remove(username)
                liked = False
            else:
                likes.append(username)
                liked = True
            topic["updated_at"] = now
            save_forum_topics(username, topics)
            item = forum_public_topic(topic, with_comments=True)
            item["liked"] = liked
            return {"ok": True, "topic": item}
    raise ValueError("话题不存在")


def extract_uploaded_text(file_item):
    name = Path(file_item.get("name", "") or "upload.txt").name
    raw = base64.b64decode(file_item.get("data", "") or "")
    suffix = Path(name).suffix.lower()
    if suffix in {".txt", ".md", ".csv"}:
        return raw.decode("utf-8", errors="ignore")
    if suffix == ".docx":
        try:
            with zipfile.ZipFile(io.BytesIO(raw)) as docx:
                root = ET.fromstring(docx.read("word/document.xml"))
            ns = {"w": W_NS}
            parts = []
            for paragraph in root.findall(".//w:p", ns):
                text = "".join(node.text or "" for node in paragraph.findall(".//w:t", ns)).strip()
                if text:
                    parts.append(text)
            return "\n".join(parts)
        except Exception:
            return ""
    if suffix == ".xlsx":
        try:
            from openpyxl import load_workbook
            wb = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
            parts = []
            for ws in wb.worksheets[:3]:
                for row in ws.iter_rows(values_only=True):
                    values = [str(value).strip() for value in row if value not in (None, "")]
                    if values:
                        parts.append(" | ".join(values))
            return "\n".join(parts)
        except Exception:
            return ""
    return raw.decode("utf-8", errors="ignore")


def forum_ai_topic(payload, username):
    seed = str(payload.get("seed", "") or "").strip()
    chat = str(payload.get("chat", "") or "").strip()
    files = payload.get("files", []) or []
    document_texts = []
    for item in files[:5]:
        text = extract_uploaded_text(item)
        if text:
            document_texts.append(f"【{Path(item.get('name', '')).name}】\n{text[:6000]}")
    if not seed and not chat and not document_texts:
        raise ValueError("请输入灵感信息、聊天内容，或上传文档")

    settings = assistant_settings()
    if not settings["url"] or not settings["key"]:
        raise ValueError("未配置 AI 接口，请在系统配置中设置 NewAPI 地址和 Key 后再使用智能生成话题")

    content = "\n\n".join([
        f"输入信息：\n{seed}" if seed else "",
        f"聊天内容：\n{chat}" if chat else "",
        f"文档内容：\n{chr(10).join(document_texts)}" if document_texts else "",
    ]).strip()
    try:
        data = request_json(
            settings["url"] + "/v1/chat/completions",
            settings["key"],
            {
                "model": settings["model"],
                "messages": [
                    {
                        "role": "system",
                        "content": (
                            "请根据输入信息，为团队金点子论坛生成一个适合当天讨论的话题。"
                            "话题要具体、有讨论价值，能启发大家围绕工作改进、产品创新、效率提升或项目机会发表观点。"
                            "严格返回 JSON：{\"title\":\"话题标题\",\"body\":\"话题介绍和讨论引导\"}，不要返回 Markdown。"
                        ),
                    },
                    {"role": "user", "content": content},
                ],
                "temperature": 0.7,
            },
            "POST",
            60,
        )
        result_text = data["choices"][0]["message"]["content"].strip()
        result_text = re.sub(r"^```(?:json)?|```$", "", result_text, flags=re.I).strip()
        result = json.loads(result_text)
        title = str(result.get("title", "") or "").strip()
        body = str(result.get("body", "") or "").strip()
    except Exception as exc:
        raise ValueError(f"智能生成话题调用失败：{exc}")

    if not title or not body:
        raise ValueError("智能生成话题没有返回完整标题和内容，请检查模型配置")
    return forum_create_topic({"title": title, "body": body, "source": "ai"}, username)


def forum_ai_comment(payload, username):
    topic_id = str(payload.get("topic_id", "") or "").strip()
    if not topic_id:
        raise ValueError("请选择话题")

    topics = load_forum_topics(username)
    topic = next((item for item in topics if item.get("id") == topic_id), None)
    if not topic:
        raise ValueError("话题不存在")

    comments = topic.get("comments", [])
    discussion = "\n".join(
        f"- {comment.get('author', '成员')}：{comment.get('content', '')}"
        for comment in comments[-20:]
    ) or "暂无评论"
    settings = assistant_settings()
    if not settings["url"] or not settings["key"]:
        raise ValueError("未配置 AI 接口，请在系统配置中设置 NewAPI 地址和 Key 后再使用 AI 潜水评论")

    user = find_user(username) or {}
    display_name = user.get("name") or username or "成员"
    try:
        data = request_json(
            settings["url"] + "/v1/chat/completions",
            settings["key"],
            {
                "model": settings["model"],
                "messages": [
                    {
                        "role": "system",
                        "content": (
                            f"你是{display_name}的 AI 潜水员，正在金点子论坛里协助他阅读话题和已有讨论，"
                            "然后以他的助手身份发表一条简短、有启发的评论。"
                            "评论要像团队讨论里的自然发言：可以总结共识、补充一个角度、提出一个可落地追问或下一步建议。"
                            "不要居高临下，不要写长文，不要使用 Markdown 标题，控制在 120-220 字。"
                        ),
                    },
                    {
                        "role": "user",
                        "content": (
                            f"话题：{topic.get('title', '')}\n\n"
                            f"话题说明：\n{topic.get('body', '')}\n\n"
                            f"已有讨论：\n{discussion}"
                        ),
                    },
                ],
                "temperature": 0.65,
            },
            "POST",
            60,
        )
        content = data["choices"][0]["message"]["content"].strip()
    except Exception as exc:
        raise ValueError(f"AI 潜水评论调用失败：{exc}")

    if not content:
        raise ValueError("AI 潜水评论没有返回内容，请检查模型配置")

    return forum_add_comment({"topic_id": topic_id, "content": content}, f"{display_name}的 AI 潜水员")


# ===== 每日资讯 =====

def news_dir():
    path = USER_DATA_DIR / "_news"
    path.mkdir(parents=True, exist_ok=True)
    return path


def news_config_payload():
    config = read_config()
    return {
        "sources": config.get("news_sources", []),
        "search_query": config.get("news_search_query", "轨道交通 OR 城市轨道 OR 地铁"),
        "auto_search": bool(config.get("news_auto_search", True)),
        "auto_push": bool(config.get("news_auto_push", True)),
        "push_time": config.get("news_push_time", "08:30"),
    }


def save_news_config(payload):
    config = read_config()
    sources = payload.get("sources", [])
    if isinstance(sources, str):
        sources = [{"name": "", "url": line.strip()} for line in sources.splitlines() if line.strip()]
    clean_sources = []
    for item in sources:
        if not isinstance(item, dict):
            continue
        url = str(item.get("url", "") or "").strip()
        if not url:
            continue
        if not re.match(r"^https?://", url, re.I):
            raise ValueError("资讯网页路径必须以 http:// 或 https:// 开头")
        clean_sources.append({
            "name": str(item.get("name", "") or "").strip(),
            "url": url,
        })
    config["news_sources"] = clean_sources
    config["news_search_query"] = str(payload.get("search_query", "") or "").strip()
    config["news_auto_search"] = bool(payload.get("auto_search", True))
    config["news_auto_push"] = bool(payload.get("auto_push", True))
    config["news_push_time"] = str(payload.get("push_time", "") or "08:30").strip()
    write_config(config)
    return {"ok": True, "config": news_config_payload()}


def strip_html_text(text):
    text = re.sub(r"(?is)<script.*?>.*?</script>", " ", text)
    text = re.sub(r"(?is)<style.*?>.*?</style>", " ", text)
    text = re.sub(r"(?is)<[^>]+>", " ", text)
    text = html_escape_unescape(text)
    return re.sub(r"\s+", " ", text).strip()


def html_escape_unescape(text):
    import html
    return html.unescape(str(text or ""))


def fetch_news_source(source):
    url = source.get("url", "")
    req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0 PersonalWorkSite/1.0"})
    with urllib.request.urlopen(req, timeout=20) as resp:
        raw = resp.read(800000)
        ctype = resp.headers.get_content_charset() or "utf-8"
    text = raw.decode(ctype, errors="replace")
    return {
        "name": source.get("name") or urllib.parse.urlparse(url).netloc,
        "url": url,
        "text": strip_html_text(text)[:12000],
    }


def latest_news_issue():
    path = news_dir() / "latest.json"
    if not path.exists():
        return None
    try:
        return read_json_lenient(path)
    except Exception:
        return None


def save_news_issue(issue):
    issue_path = news_dir() / f"{issue['date']}.json"
    latest_path = news_dir() / "latest.json"
    raw = json.dumps(issue, ensure_ascii=False, indent=2)
    issue_path.write_text(raw, encoding="utf-8")
    latest_path.write_text(raw, encoding="utf-8")


def public_news_issue(issue, include_config_details=False):
    if not issue:
        return None
    item = dict(issue)
    if not include_config_details:
        for key in ("sources", "errors", "search_query", "auto_search"):
            item.pop(key, None)
    return item


def news_latest(include_config=False):
    result = {"ok": True, "issue": public_news_issue(latest_news_issue(), include_config)}
    if include_config:
        result["config"] = news_config_payload()
    return result


def generate_news_issue(payload=None, username="system"):
    payload = payload or {}
    config_payload = news_config_payload()
    settings = assistant_settings()
    if not settings["url"] or not settings["key"]:
        raise ValueError("未配置 AI 接口，请在系统配置中设置 NewAPI 地址和 Key 后再生成每日资讯")

    sources = config_payload["sources"]
    source_results = []
    errors = []
    for source in sources[:12]:
        try:
            source_results.append(fetch_news_source(source))
        except Exception as exc:
            errors.append(f"{source.get('name') or source.get('url')}：{exc}")

    search_query = str(payload.get("search_query") or config_payload["search_query"] or "").strip()
    auto_search = bool(payload.get("auto_search", config_payload["auto_search"]))
    if not source_results and not (auto_search and search_query):
        raise ValueError("请先配置资讯网页路径，或开启大模型自动网络搜索并填写搜索关键词")

    source_text = "\n\n".join(
        f"【{item['name']}】{item['url']}\n{item['text']}"
        for item in source_results
    )
    search_instruction = (
        f"请调用平台可用的联网搜索/工具能力，围绕以下关键词检索今天或近期轨道交通关键资讯：{search_query}"
        if auto_search and search_query
        else "未启用自动网络搜索，只基于已抓取网页内容生成。"
    )
    today = datetime.now().strftime("%Y-%m-%d")
    prompt = (
        "你是轨道交通行业资讯分析智能体。请生成一份每日关键资讯简报，聚焦轨道交通、城市轨道、地铁、智慧运维、施工安全、客流、低空/AI/数字化等与轨交相关的信息。"
        "大模型能力必须来自当前平台配置的 API。若可用工具支持联网搜索，请使用工具补充最新信息；若不可用，请只基于提供的网页内容，不要编造事实。"
        "严格返回 JSON，不要 Markdown："
        "{\"title\":\"今日标题\",\"summary\":\"总览摘要\",\"items\":[{\"title\":\"资讯标题\",\"source\":\"来源\",\"url\":\"链接\",\"impact\":\"影响/价值\",\"action\":\"建议动作\"}],\"keywords\":[\"关键词\"]}"
    )
    user_content = (
        f"日期：{today}\n\n"
        f"自动搜索要求：\n{search_instruction}\n\n"
        f"已配置网页抓取内容：\n{source_text or '无'}"
    )
    try:
        data = request_json(
            settings["url"] + "/v1/chat/completions",
            settings["key"],
            {
                "model": settings["model"],
                "messages": [
                    {"role": "system", "content": prompt},
                    {"role": "user", "content": user_content},
                ],
                "temperature": 0.35,
            },
            "POST",
            90,
        )
        result_text = data["choices"][0]["message"]["content"].strip()
        result_text = re.sub(r"^```(?:json)?|```$", "", result_text, flags=re.I).strip()
        parsed = json.loads(result_text)
    except Exception as exc:
        raise ValueError(f"每日资讯生成失败：{exc}")

    issue = {
        "date": today,
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "generated_by": username,
        "title": str(parsed.get("title", "") or f"{today} 轨道交通每日资讯").strip(),
        "summary": str(parsed.get("summary", "") or "").strip(),
        "items": parsed.get("items", []) if isinstance(parsed.get("items", []), list) else [],
        "keywords": parsed.get("keywords", []) if isinstance(parsed.get("keywords", []), list) else [],
        "sources": [{"name": item["name"], "url": item["url"]} for item in source_results],
        "errors": errors,
        "search_query": search_query,
        "auto_search": auto_search,
    }
    save_news_issue(issue)
    return {"ok": True, "issue": issue}


def news_auto_worker():
    while True:
        try:
            config = read_config()
            if config.get("news_auto_push", True):
                latest = latest_news_issue()
                now = datetime.now()
                today = now.strftime("%Y-%m-%d")
                push_time = str(config.get("news_push_time", "08:30") or "08:30")
                if re.match(r"^\d{2}:\d{2}$", push_time) and now.strftime("%H:%M") < push_time:
                    time.sleep(1800)
                    continue
                if not latest or latest.get("date") != today:
                    generate_news_issue({"auto_search": config.get("news_auto_search", True)}, "system")
        except Exception:
            pass
        time.sleep(1800)


def safe_uploaded_name(name, kind):
    name = Path(name or "").name.strip()
    if not name:
        raise ValueError("上传文件缺少文件名")
    suffix = Path(name).suffix.lower()
    if kind == "weekly":
        if suffix not in {".xlsx", ".xls"}:
            raise ValueError("周报只支持 .xlsx 或 .xls 文件")
        if "工作周报" not in name:
            name = "工作周报-" + name
    elif kind == "trip":
        if suffix not in {".docx", ".md"}:
            raise ValueError("出差报告只支持 .docx 或 .md 文件")
        if not name.startswith("出差报告"):
            name = "出差报告-" + name
    else:
        raise ValueError("请选择上传类型")
    return re.sub(r"[/:\\\\]+", "_", name)


def unique_report_path(name, username=None):
    base = user_report_dir(username) if username else REPORT_DIR
    base.mkdir(parents=True, exist_ok=True)
    path = base / name
    if not path.exists():
        return path
    stem = path.stem
    suffix = path.suffix
    for idx in range(1, 1000):
        candidate = base / f"{stem}-上传{idx}{suffix}"
        if not candidate.exists():
            return candidate
    raise ValueError("同名文件过多，请修改文件名后再上传")


def upload_history_reports(payload, username=None):
    kind = payload.get("kind")
    files = payload.get("files") or []
    if not files:
        raise ValueError("请选择需要上传的文件")
    uploaded = []
    for item in files:
        name = safe_uploaded_name(item.get("name", ""), kind)
        data = item.get("data", "")
        if "," in data:
            data = data.split(",", 1)[1]
        raw = base64.b64decode(data)
        if not raw:
            raise ValueError(f"{name} 文件内容为空")
        path = unique_report_path(name, username)
        path.write_bytes(raw)
        uploaded.append({"name": path.name, "path": str(path), "size": len(raw)})
    return {"ok": True, "uploaded": uploaded}


def delete_history_report(payload, username=None):
    path = history_path_by_name(payload.get("name", ""), username)
    if path is None:
        raise ValueError("历史报告文件不存在")
    name = path.name
    path.unlink()
    return {"ok": True, "deleted": name}


def delete_report_file(payload, username=None):
    path = attachment_path_by_name(payload.get("name", ""), username)
    if path is None:
        raise ValueError("报告文件不存在，或不是当前账号的可删除文件")
    name = path.name
    path.unlink()
    return {"ok": True, "deleted": name}


def estimated_text_lines(text, width_chars):
    lines = 0
    for part in str(text or "").splitlines() or [""]:
        visual_len = sum(2 if ord(ch) > 127 else 1 for ch in part)
        lines += max(1, (visual_len + width_chars - 1) // width_chars)
    return lines


def adjust_row_height(ws, row_idx, columns):
    width_chars = {
        2: 40,
        3: 30,
        4: 72,
        5: 24,
        6: 72,
        (3, 4): 68,
        (4, 5): 92,
        (3, 4, 5): 112,
    }
    max_lines = 1
    for col in columns:
        width = width_chars.get(col, 32)
        for merged in ws.merged_cells.ranges:
            if merged.min_row == row_idx and merged.max_row == row_idx and merged.min_col == col:
                key = tuple(range(merged.min_col, merged.max_col + 1))
                width = width_chars.get(key, width)
                break
        max_lines = max(max_lines, estimated_text_lines(ws.cell(row_idx, col).value, width))
    ws.row_dimensions[row_idx].height = min(320, max(34, max_lines * 16 + 8))


def weekly_prefill(username=None):
    from openpyxl import load_workbook

    template = Path((newest_any("weekly", username, fallback_shared=True) or {}).get("path", ""))
    if not template.exists():
        return {"weekly_summary": "", "weekly_follow": "", "weekly_next": ""}

    wb = load_workbook(template, data_only=True)
    ws = wb[wb.sheetnames[0]]

    summary = []
    summary_rows = []
    for row in range(18, min(ws.max_row, 26) + 1):
        category = cell_value(ws, row, 2)
        content = normalize_numbered_text(cell_value(ws, row, 3))
        if category or content:
            # 本周总结继承上周计划：工作分类 | 工作内容 | 完成情况 | 后续计划
            summary.append(xlsx_row_text([category, content, "", ""]))
            summary_rows.append({"category": category, "content": content, "status": "", "plan": ""})

    follow = []
    follow_rows = []
    for row in range(11, min(ws.max_row, 15) + 1):
        category = cell_value(ws, row, 2)
        content = normalize_numbered_text(cell_value(ws, row, 3))
        progress = normalize_numbered_text(cell_value(ws, row, 4))
        difficulty = normalize_numbered_text(cell_value(ws, row, 6))
        if category or content or progress or difficulty:
            follow.append(xlsx_row_text([category, content, progress, difficulty]))
            follow_rows.append(
                {
                    "category": category,
                    "content": content,
                    "progress": progress,
                    "difficulty": difficulty,
                }
            )

    return {
        "weekly_summary": "\n".join(summary),
        "weekly_follow": "\n".join(follow),
        "weekly_next": "",
        "summary_rows": summary_rows,
        "follow_rows": follow_rows,
        "next_rows": [],
        "source": template.name,
    }


def table_cell_text(table, row, col):
    try:
        return normalize_numbered_text(table.cell(row, col).text.strip())
    except Exception:
        return ""


def split_trip_date(date_text):
    text = date_text or ""
    match = re.search(r"(\d{4})年(\d{1,2})月(\d{1,2})日\s*至\s*(\d{4})年(\d{1,2})月(\d{1,2})日", text)
    if match:
        start = f"{int(match.group(1)):04d}-{int(match.group(2)):02d}-{int(match.group(3)):02d}"
        end = f"{int(match.group(4)):04d}-{int(match.group(5)):02d}-{int(match.group(6)):02d}"
        return start, end
    match = re.search(r"(\d{4})[-/.](\d{1,2})[-/.](\d{1,2}).*?(\d{4})[-/.](\d{1,2})[-/.](\d{1,2})", text)
    if match:
        start = f"{int(match.group(1)):04d}-{int(match.group(2)):02d}-{int(match.group(3)):02d}"
        end = f"{int(match.group(4)):04d}-{int(match.group(5)):02d}-{int(match.group(6)):02d}"
        return start, end
    return "", ""


def format_trip_date_text(start, end):
    def fmt(value):
        value = str(value or "").strip()
        match = re.match(r"(\d{4})[-/.]?(\d{1,2})[-/.]?(\d{1,2})$", value)
        if not match:
            return value
        return f"{int(match.group(1)):04d}年{int(match.group(2)):02d}月{int(match.group(3)):02d}日"
    start_text = fmt(start)
    end_text = fmt(end)
    if start_text and end_text:
        return f"{start_text}至{end_text}"
    return start_text or end_text


def trip_prefill(username=None):
    from docx import Document

    template = Path((newest("trip", username, fallback_shared=True) or {}).get("path", ""))
    if not template.exists() or template.suffix.lower() != ".docx":
        return {"source": ""}

    doc = Document(template)
    if not doc.tables:
        return {"source": template.name}
    table = doc.tables[0]
    date_text = table_cell_text(table, 1, 1)
    start, end = split_trip_date(date_text)
    return {
        "source": template.name,
        "reporter": table_cell_text(table, 0, 1) or "周颖超",
        "department": table_cell_text(table, 0, 3) or "场景研究院",
        "location": table_cell_text(table, 0, 5),
        "trip_start": start,
        "trip_end": end,
        "trip_date_text": date_text,
        "purpose": table_cell_text(table, 2, 1),
        "itinerary": table_cell_text(table, 3, 1),
        "details": table_cell_text(table, 4, 1),
        "issues": table_cell_text(table, 5, 1),
        "suggestions": table_cell_text(table, 6, 1),
    }


def generate_weekly(payload, username=None):
    from openpyxl import load_workbook
    from openpyxl.styles import Border, Side

    template = Path((newest_any("weekly", username, fallback_shared=True) or {}).get("path", ""))
    if not template.exists() or not template.is_file():
        raise ValueError("没有找到周报模板")

    output_dir = user_generated_dir(username) if username else GENERATED_DIR
    output_dir.mkdir(parents=True, exist_ok=True)
    period = (payload.get("period") or datetime.now().strftime("%Y.%m.%d-%Y.%m.%d")).strip()
    safe_period = re.sub(r"[^0-9A-Za-z.\-\u4e00-\u9fff]+", "", period)
    output = output_dir / f"周颖超工作周报{safe_period}.xlsx"
    if template.resolve() == output.resolve():
        fallback = newest("weekly", username, fallback_shared=True)
        if fallback and Path(fallback.get("path", "")).resolve() != output.resolve():
            template = Path(fallback["path"])
    shutil.copy2(template, output)

    wb = load_workbook(output)
    ws = wb[wb.sheetnames[0]]

    # 取消数据区域的合并单元格，方便写入
    for mc in list(ws.merged_cells.ranges):
        if not (mc.max_row < 5 or mc.min_row > 20):
            ws.unmerge_cells(str(mc))

    ws["B2"] = f"工作周报（{period}）"

    summary_rows = parse_table_lines(payload.get("weekly_summary", ""), 4) or [
        ["", "", "", ""]
    ]
    follow_rows = parse_table_lines(payload.get("weekly_follow", ""), 4) or [
        ["", "", "", ""]
    ]
    next_rows = parse_table_lines(payload.get("weekly_next", ""), 3) or [
        ["", "", ""]
    ]

    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    def merge_safe(min_r, min_c, max_r, max_c):
        try:
            ws.merge_cells(start_row=min_r, start_column=min_c, end_row=max_r, end_column=max_c)
        except Exception:
            pass

    def write_and_style(row_idx, col_idx, value, halign="left"):
        cell = ws.cell(row_idx, col_idx)
        cell.value = normalize_numbered_text(value)
        cell.border = thin_border
        style_xlsx_cell(cell, horizontal=halign, vertical="center")

    def write_section_title(row_idx, title):
        write_and_style(row_idx, 2, title, halign="center")
        merge_safe(row_idx, 2, row_idx, 6)

    def write_header(row_idx, labels, merges=None):
        for col_idx in range(2, 7):
            write_and_style(row_idx, col_idx, "")
        for col_idx, label in labels:
            write_and_style(row_idx, col_idx, label, halign="center")
        for merge in merges or []:
            merge_safe(row_idx, merge[0], row_idx, merge[1])

    max_row = max(ws.max_row, 24)
    if max_row >= 3:
        ws.delete_rows(3, max_row - 2)

    summary_title = 3
    summary_header = 4
    summary_start = 5
    write_section_title(summary_title, "一、本周工作总结")
    write_header(summary_header, [(2, "工作分类"), (3, "工作内容"), (5, "上周内容完成情况"), (6, "后续计划")], [(3, 4)])

    # 本周工作总结: B=工作分类 C=工作内容 E=完成情况 F=后续计划
    for idx, row in enumerate(summary_rows, start=summary_start):
        write_and_style(idx, 2, row[0], halign="center")
        write_and_style(idx, 3, row[1])
        write_and_style(idx, 5, row[2])
        write_and_style(idx, 6, row[3])
        adjust_row_height(ws, idx, (2, 3, 5, 6))
        merge_safe(idx, 3, idx, 4)

    follow_title = summary_start + len(summary_rows)
    follow_header = follow_title + 1
    follow_start = follow_header + 1
    write_section_title(follow_title, "二、重点工作跟进")
    write_header(follow_header, [(2, "工作分类"), (3, "工作内容"), (4, "当前进展"), (6, "困难与求助")], [(4, 5)])

    # 重点工作跟进: B=工作分类 C=工作内容 D=当前进展 F=困难与求助
    for idx, row in enumerate(follow_rows, start=follow_start):
        write_and_style(idx, 2, row[0], halign="center")
        write_and_style(idx, 3, row[1])
        write_and_style(idx, 4, row[2])
        write_and_style(idx, 6, row[3])
        adjust_row_height(ws, idx, (2, 3, 4, 6))
        merge_safe(idx, 4, idx, 5)

    next_title = follow_start + len(follow_rows)
    next_header = next_title + 1
    next_start = next_header + 1
    write_section_title(next_title, "三、下周工作计划")
    write_header(next_header, [(2, "工作分类"), (3, "工作内容"), (6, "困难与求助")], [(3, 5)])

    # 下周工作计划: B=工作分类 C=工作内容 F=困难与求助
    for idx, row in enumerate(next_rows, start=next_start):
        write_and_style(idx, 2, row[0], halign="center")
        write_and_style(idx, 3, row[1])
        write_and_style(idx, 6, row[2])
        adjust_row_height(ws, idx, (2, 3, 6))
        merge_safe(idx, 3, idx, 5)

    merge_safe(2, 2, 2, 6)

    wb.save(output)
    return output


def clear_docx_row_height(row):
    from docx.oxml.ns import qn

    tr_pr = row._tr.get_or_add_trPr()
    for height in list(tr_pr.findall(qn("w:trHeight"))):
        tr_pr.remove(height)


def set_cell_text(cell, text, bold=False, compact=True):
    from docx.shared import Pt
    from docx.oxml.ns import qn

    text = normalize_numbered_text(text)
    parts = text.splitlines() or [""]
    while len(cell.paragraphs) < len(parts):
        cell.add_paragraph()
    for index, part in enumerate(parts):
        paragraph = cell.paragraphs[index]
        for run in list(paragraph.runs):
            paragraph._p.remove(run._r)
        run = paragraph.add_run(part)
        run.bold = bold
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(11)
        paragraph.paragraph_format.line_spacing = 1.08 if compact else 1.2
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2 if compact else 4)
    for paragraph in cell.paragraphs[len(parts):]:
        paragraph._element.getparent().remove(paragraph._element)


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_NS = "http://www.w3.org/XML/1998/namespace"
ET.register_namespace("w", W_NS)


def w_tag(name):
    return f"{{{W_NS}}}{name}"


def set_docx_xml_cell_text(cell, text):
    text = normalize_numbered_text(text)
    lines = text.splitlines() or [""]
    paragraphs = cell.findall(w_tag("p"))
    paragraph_template = deepcopy(paragraphs[0]) if paragraphs else ET.Element(w_tag("p"))
    paragraph_props = paragraph_template.find(w_tag("pPr"))
    run_template = paragraph_template.find(".//" + w_tag("r"))
    run_props = deepcopy(run_template.find(w_tag("rPr"))) if run_template is not None and run_template.find(w_tag("rPr")) is not None else None

    for paragraph in paragraphs:
        cell.remove(paragraph)

    for line in lines:
        paragraph = ET.Element(w_tag("p"))
        if paragraph_props is not None:
            paragraph.append(deepcopy(paragraph_props))
        run = ET.SubElement(paragraph, w_tag("r"))
        if run_props is not None:
            run.append(deepcopy(run_props))
        text_node = ET.SubElement(run, w_tag("t"))
        text_node.set(f"{{{XML_NS}}}space", "preserve")
        text_node.text = line
        cell.append(paragraph)


def generate_trip_from_docx_template(template, output, values):
    with zipfile.ZipFile(template, "r") as source:
        document_xml = source.read("word/document.xml")
        root = ET.fromstring(document_xml)
        table = root.find(".//" + w_tag("tbl"))
        if table is None:
            raise ValueError("出差报告模板中没有找到表格")
        rows = table.findall(w_tag("tr"))

        def cell(row_idx, actual_col_idx):
            return rows[row_idx].findall(w_tag("tc"))[actual_col_idx]

        cell_map = {
            (0, 1): values.get("reporter", "周颖超"),
            (0, 3): values.get("department", "场景研究院"),
            (0, 5): values.get("location", ""),
            (1, 1): values.get("date_text", ""),
            (2, 1): values.get("purpose", ""),
            (3, 1): values.get("itinerary", ""),
            (4, 1): values.get("details", ""),
            (5, 1): values.get("issues", ""),
            (6, 1): values.get("suggestions", ""),
        }
        for (row_idx, col_idx), value in cell_map.items():
            set_docx_xml_cell_text(cell(row_idx, col_idx), value)

        output.parent.mkdir(parents=True, exist_ok=True)
        temp_output = output.with_suffix(output.suffix + ".tmp")
        with zipfile.ZipFile(temp_output, "w", zipfile.ZIP_DEFLATED) as target:
            for item in source.infolist():
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True) if item.filename == "word/document.xml" else source.read(item.filename)
                target.writestr(item, data)
        temp_output.replace(output)


def generate_trip(payload, username=None):
    template = Path((newest("trip", username, fallback_shared=True) or {}).get("path", ""))
    if not template.exists() or template.suffix.lower() != ".docx":
        raise ValueError("没有找到 Word 出差报告模板")

    output_dir = user_generated_dir(username) if username else GENERATED_DIR
    output_dir.mkdir(parents=True, exist_ok=True)
    start = (payload.get("trip_start") or datetime.now().strftime("%Y%m%d")).replace("-", "")
    end = (payload.get("trip_end") or start).replace("-", "")[-4:]
    output = output_dir / f"出差报告-{start}-{end}-周颖超.docx"
    date_text = payload.get("trip_date_text") or format_trip_date_text(payload.get("trip_start", ""), payload.get("trip_end", ""))
    values = {
        "reporter": payload.get("reporter", "周颖超"),
        "department": payload.get("department", "场景研究院"),
        "location": payload.get("location", ""),
        "date_text": date_text,
        "purpose": payload.get("purpose", ""),
        "itinerary": payload.get("itinerary", ""),
        "details": payload.get("details", ""),
        "issues": payload.get("issues", ""),
        "suggestions": payload.get("suggestions", ""),
    }
    generate_trip_from_docx_template(template, output, values)
    return output


def generate_document(payload, username=None):
    kind = payload.get("kind")
    if kind == "weekly":
        path = generate_weekly(payload, username)
    elif kind == "trip":
        path = generate_trip(payload, username)
    else:
        raise ValueError("未知生成类型")
    draft = compose_draft(kind, path.name, username)
    return {
        "ok": True,
        "file": path.name,
        "path": str(path),
        "draft": draft,
    }


def save_draft(payload, username=None):
    draft_dir = user_draft_dir(username) if username else DRAFT_DIR
    draft_dir.mkdir(parents=True, exist_ok=True)
    msg = build_message(payload, username)
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    safe_subject = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff_-]+", "_", payload.get("subject", "draft"))[:40]
    path = draft_dir / f"{stamp}-{safe_subject}.eml"
    path.write_bytes(bytes(msg))
    return path


def send_mail(payload, username=None):
    settings = smtp_settings(username)
    if not settings["host"]:
        draft = save_draft(payload, username)
        return {"ok": True, "mode": "draft", "message": f"邮件未发出：当前账号未配置 SMTP 服务器，已生成邮件草稿：{draft}"}
    validate_smtp_ready(settings)

    msg = build_message(payload, username)
    recipients = split_addresses(payload.get("to", "")) + split_addresses(payload.get("cc", ""))
    if not recipients:
        raise ValueError("请填写收件人邮箱")

    if settings["use_ssl"]:
        with smtplib.SMTP_SSL(settings["host"], settings["port"], context=ssl.create_default_context(), timeout=20) as smtp:
            if settings["user"]:
                smtp.login(settings["user"], settings["password"])
            smtp.send_message(msg, to_addrs=recipients)
    else:
        with smtplib.SMTP(settings["host"], settings["port"], timeout=20) as smtp:
            if settings["use_tls"]:
                smtp.starttls(context=ssl.create_default_context())
            if settings["user"]:
                smtp.login(settings["user"], settings["password"])
            smtp.send_message(msg, to_addrs=recipients)
    return {"ok": True, "mode": "sent", "message": "邮件已发送"}


def app_html():
    return """<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>犇犇 | 智能办公平台</title>
  <style>
    :root {
      color-scheme: dark;
      --bg: #0b1120;
      --surface: #111827;
      --panel: #1e293b;
      --elevated: #243447;
      --ink: #f1f5f9;
      --muted: #94a3b8;
      --line: #334155;
      --accent: #3b82f6;
      --accent-2: #8b5cf6;
      --accent-gradient: linear-gradient(135deg, #3b82f6, #8b5cf6);
      --danger: #ef4444;
      --success: #10b981;
      --warning: #f59e0b;
      --shadow: 0 8px 32px rgba(2, 8, 20, .45);
      --glow: 0 0 20px rgba(59, 130, 246, .15);
      --radius: 12px;
      --radius-sm: 8px;
    }
    * { box-sizing: border-box; }
    body {
      margin: 0;
      font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", "PingFang SC", "Microsoft YaHei", sans-serif;
      background: var(--bg);
      color: var(--ink);
    }
    header {
      background: #123c3b;
      color: #fff;
      padding: 22px 28px;
      border-bottom: 4px solid #d69839;
    }
    .header-row {
      display: flex;
      justify-content: space-between;
      gap: 16px;
      align-items: center;
    }
    .userbar {
      display: none;
      align-items: center;
      gap: 10px;
      color: #d7e5e3;
      font-size: 13px;
    }
    .user-avatar {
      width: 30px; height: 30px; border-radius: 50%;
      object-fit: cover; background: #eaf1fb; border: 1px solid var(--line);
    }
    .userbar button { background: rgba(255,255,255,.16); }
    h1 { margin: 0; font-size: 24px; font-weight: 720; letter-spacing: 0; }
    .sub { margin-top: 6px; color: #d7e5e3; font-size: 14px; }
    main {
      display: grid;
      grid-template-columns: 240px minmax(0, 1fr);
      gap: 16px;
      padding: 20px;
      max-width: 1320px;
      margin: 0 auto;
      align-items: start;
    }
    section {
      background: var(--panel);
      border: 1px solid var(--line);
      border-radius: 8px;
      box-shadow: var(--shadow);
    }
    .task-board {
      padding: 14px;
      position: sticky;
      top: 14px;
    }
    .task-grid {
      display: grid;
      grid-template-columns: 1fr;
      gap: 8px;
    }
    .task-card {
      width: 100%;
      min-height: 64px;
      text-align: left;
      border: 1px solid var(--line);
      border-radius: 8px;
      padding: 11px 12px;
      background: #fff;
      color: var(--ink);
      box-shadow: none;
      cursor: pointer;
    }
    .task-card.active {
      border-color: var(--accent);
      background: #eef8f6;
      box-shadow: inset 0 0 0 1px #b9ded8;
    }
    .task-card .task-name {
      display: block;
      font-size: 15px;
      font-weight: 760;
      margin-bottom: 4px;
    }
    .task-card .task-desc {
      display: block;
      color: var(--muted);
      font-size: 12px;
      line-height: 1.45;
      font-weight: 520;
    }
    .task-panel.hidden { display: none; }
    .auth-panel {
      max-width: 460px;
      margin: 34px auto 0;
      padding: 18px;
    }
    .auth-panel.hidden { display: none; }
    .mail-layout {
      display: grid;
      grid-template-columns: minmax(250px, 340px) minmax(0, 1fr);
      gap: 16px;
    }
    .side {
      padding: 14px;
      border: 1px solid var(--line);
      border-radius: 8px;
      background: #fbfcfd;
    }
    .composer { padding: 18px; }
    .sub-tabs {
      display: flex;
      gap: 8px;
      margin: 0 0 18px;
      padding: 4px;
      background: #eef1f5;
      border-radius: 10px;
      width: fit-content;
    }
    .sub-tab {
      padding: 8px 18px;
      border-radius: 8px;
      font-size: 14px;
      font-weight: 680;
      background: transparent;
      color: var(--muted);
      border: 0;
      cursor: pointer;
      transition: all .15s ease;
    }
    .sub-tab.active {
      background: #fff;
      color: var(--accent);
      box-shadow: 0 2px 8px rgba(26,34,47,.08);
    }
    .sub-tab:hover:not(.active) {
      color: var(--ink);
      background: rgba(255,255,255,.5);
    }
    h2 { margin: 0 0 14px; font-size: 17px; }
    label { display: block; margin: 12px 0 6px; font-size: 13px; color: var(--muted); }
    input, select, textarea {
      width: 100%;
      border: 1px solid var(--line);
      border-radius: 6px;
      padding: 10px 11px;
      font: inherit;
      color: var(--ink);
      background: #fff;
    }
    textarea { min-height: 260px; resize: vertical; line-height: 1.55; }
    .row { display: grid; grid-template-columns: 1fr 1fr; gap: 12px; }
    .toolbar { display: flex; gap: 10px; align-items: center; flex-wrap: wrap; margin-top: 14px; }
    button {
      border: 0;
      border-radius: 6px;
      padding: 10px 14px;
      font-weight: 680;
      cursor: pointer;
      background: var(--accent);
      color: #fff;
    }
    button.secondary { background: #2f3a4a; }
    button.warn { background: var(--accent-2); }
    button.action-button {
      padding: 10px 14px;
      font-size: 14px;
      background: var(--accent);
      color: #fff;
      border: 0;
      white-space: nowrap;
    }
    .download-link {
      display: inline-flex;
      align-items: center;
      min-height: 36px;
      padding: 8px 12px;
      border-radius: 6px;
      background: #eef8f6;
      color: var(--accent);
      border: 1px solid #b9ded8;
      text-decoration: none;
      font-weight: 680;
      font-size: 14px;
    }
    .report {
      border: 1px solid var(--line);
      border-radius: 8px;
      padding: 12px;
      margin-bottom: 10px;
      cursor: pointer;
      background: #fff;
    }
    .report.active { border-color: var(--accent); background: #eef8f6; }
    .report .name { font-weight: 680; word-break: break-all; }
    .report .meta { color: var(--muted); font-size: 12px; margin-top: 5px; }
    .preview {
      white-space: pre-wrap;
      background: #f8fafb;
      border: 1px solid var(--line);
      border-radius: 6px;
      padding: 12px;
      min-height: 130px;
      max-height: 260px;
      overflow: auto;
      color: #354052;
      font-size: 13px;
      line-height: 1.55;
    }
    .rich-preview {
      white-space: normal;
      background: #f8fafb;
      border: 1px solid var(--line);
      border-radius: 6px;
      padding: 12px;
      min-height: 130px;
      max-height: 360px;
      overflow: auto;
      color: #354052;
      font-size: 13px;
      line-height: 1.55;
    }
    .rich-preview table {
      min-width: 760px;
    }
    .rich-preview td,
    .rich-preview th {
      box-sizing: border-box;
    }
    .rich-preview p {
      margin: 8px 0;
    }
    .status { margin-top: 12px; font-size: 14px; color: var(--muted); }
    .status.ok { color: var(--accent); }
    .status.err { color: var(--danger); }
    .upload-list {
      margin-top: 12px;
      display: grid;
      gap: 8px;
    }
    .config-grid {
      display: grid;
      grid-template-columns: 1fr 1fr;
      gap: 12px;
    }
    .upload-item {
      border: 1px solid var(--line);
      border-radius: 8px;
      padding: 10px 12px;
      background: #fff;
      color: #354052;
      font-size: 13px;
      word-break: break-all;
    }
    .history-tools {
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 10px;
      margin-top: 16px;
      flex-wrap: wrap;
    }
    .history-list {
      display: grid;
      gap: 8px;
      margin-top: 10px;
    }
    .history-item {
      display: grid;
      grid-template-columns: minmax(0, 1fr) auto;
      gap: 10px;
      align-items: center;
      border: 1px solid var(--line);
      border-radius: 8px;
      padding: 10px 12px;
      background: #fff;
    }
    .history-name { font-weight: 680; word-break: break-all; }
    .history-meta { margin-top: 4px; color: var(--muted); font-size: 12px; }
    .history-actions { display: flex; gap: 8px; flex-wrap: wrap; justify-content: flex-end; }
    .user-list {
      display: grid;
      gap: 8px;
      margin-top: 10px;
    }
    .user-item {
      display: flex;
      justify-content: space-between;
      gap: 10px;
      border: 1px solid var(--line);
      border-radius: 8px;
      padding: 10px 12px;
      background: #fff;
    }
    .guide {
      border: 1px solid var(--line);
      border-radius: 8px;
      padding: 14px;
      margin-bottom: 16px;
      background: #fbfcfd;
    }
    .guide textarea { min-height: 94px; }
    .guide .hint { color: var(--muted); font-size: 12px; margin-top: 5px; line-height: 1.45; }
    .guide-grid { display: grid; grid-template-columns: repeat(3, 1fr); gap: 12px; }
    .hidden { display: none; }
    .weekly-workspace {
      display: grid;
      gap: 12px;
    }
    .weekly-period-card {
      border: 1px solid var(--line);
      border-radius: 8px;
      background: #fff;
      padding: 12px;
    }
    .weekly-period-head {
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 12px;
      margin-bottom: 10px;
    }
    .weekly-period-title {
      font-weight: 760;
      color: #263142;
    }
    .weekly-period-grid {
      display: grid;
      grid-template-columns: 1fr 1fr;
      gap: 10px;
    }
    .weekly-section-grid {
      display: grid;
      grid-template-columns: repeat(3, minmax(0, 1fr));
      gap: 12px;
      align-items: start;
    }
    .weekly-section-card {
      border: 1px solid var(--line);
      border-radius: 8px;
      background: #fff;
      min-height: 430px;
      display: flex;
      flex-direction: column;
      overflow: hidden;
    }
    .weekly-section-head {
      display: flex;
      align-items: flex-start;
      justify-content: space-between;
      gap: 10px;
      padding: 12px;
      border-bottom: 1px solid var(--line);
      background: #fbfcfd;
    }
    .weekly-section-title {
      font-weight: 760;
      color: #263142;
      line-height: 1.3;
    }
    .weekly-section-meta {
      color: var(--muted);
      font-size: 12px;
      margin-top: 4px;
    }
    .weekly-rows {
      padding: 10px;
      overflow: auto;
      max-height: 370px;
      flex: 1;
    }
    .edit-card-grid {
      display: grid;
      grid-template-columns: repeat(2, minmax(0, 1fr));
      gap: 10px;
      margin: 8px 0 12px;
    }
    .work-block {
      border: 1px solid var(--line);
      border-radius: 8px;
      background: #fff;
      margin: 0 0 8px 0;
      padding: 10px;
      cursor: pointer;
      transition: border-color .15s ease, background .15s ease;
    }
    .work-block:hover, .edit-card:hover {
      border-color: var(--accent);
      background: #f5fbfa;
    }
    .work-head {
      display: flex;
      justify-content: space-between;
      gap: 10px;
      align-items: center;
      margin-bottom: 6px;
    }
    .work-title { font-weight: 700; color: #263142; }
    .work-summary {
      color: var(--muted);
      font-size: 13px;
      line-height: 1.42;
      white-space: pre-wrap;
      display: -webkit-box;
      -webkit-line-clamp: 4;
      -webkit-box-orient: vertical;
      overflow: hidden;
    }
    .edit-card {
      min-height: 92px;
      border: 1px solid var(--line);
      border-radius: 8px;
      background: #fff;
      padding: 12px;
      cursor: pointer;
    }
    .edit-card-title {
      font-weight: 720;
      color: #263142;
      margin-bottom: 7px;
    }
    .edit-card-preview {
      color: var(--muted);
      font-size: 13px;
      line-height: 1.45;
      white-space: pre-wrap;
      display: -webkit-box;
      -webkit-line-clamp: 2;
      -webkit-box-orient: vertical;
      overflow: hidden;
    }
    .mini {
      padding: 6px 9px;
      font-size: 12px;
      background: #e8edf3;
      color: #263142;
    }
    .danger {
      background: #fff0ed;
      color: var(--danger);
      border: 1px solid #f3bbb2;
    }
    .field-grid {
      display: grid;
      grid-template-columns: minmax(130px, .75fr) minmax(260px, 1.7fr) minmax(150px, 1fr) minmax(150px, 1fr);
      gap: 10px;
    }
    .field-grid.three {
      grid-template-columns: minmax(150px, .85fr) minmax(320px, 2fr) minmax(170px, 1fr);
    }
    .field-grid textarea {
      min-height: 78px;
      line-height: 1.55;
      white-space: pre-wrap;
    }
    .field-store, .trip-store { display: none; }
    .field-grid label { margin-top: 0; }
    .section-actions { margin-top: 8px; }
    .modal {
      position: fixed;
      inset: 0;
      display: flex;
      align-items: center;
      justify-content: center;
      padding: 22px;
      background: rgba(15, 23, 42, .42);
      z-index: 20;
    }
    .modal.hidden { display: none; }
    .modal-box {
      width: min(960px, 96vw);
      max-height: 90vh;
      overflow: auto;
      background: #fff;
      border-radius: 8px;
      border: 1px solid var(--line);
      box-shadow: 0 22px 60px rgba(15, 23, 42, .24);
      padding: 18px;
    }
    .modal-head {
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 12px;
      margin-bottom: 12px;
    }
    .modal-title { font-size: 18px; font-weight: 760; }
    .modal-fields {
      display: grid;
      grid-template-columns: 1fr 1fr;
      gap: 12px;
    }
    .modal-fields .wide { grid-column: 1 / -1; }
    .modal-fields textarea { min-height: 190px; }
    .field-assist {
      margin: 7px 0 4px;
      padding: 9px;
      border: 1px solid #c8d8d5;
      border-radius: 8px;
      background: #f4fbfa;
    }
    .field-assist-head {
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 10px;
      flex-wrap: wrap;
    }
    .field-assist details {
      flex: 1;
      min-width: 220px;
    }
    .field-assist summary {
      cursor: pointer;
      color: #174541;
      font-size: 13px;
      font-weight: 700;
    }
    .field-assist textarea {
      min-height: 104px;
      margin-top: 8px;
      font-size: 13px;
    }
    .assist-status {
      min-height: 18px;
      margin-top: 7px;
      color: var(--muted);
      font-size: 13px;
    }
    .assist-status.ok { color: var(--accent); }
    .assist-status.err { color: var(--danger); }
    .modal-actions {
      display: flex;
      justify-content: flex-end;
      gap: 10px;
      margin-top: 14px;
    }
    @media (max-width: 860px) {
      main { grid-template-columns: 1fr; padding: 12px; }
      .task-board { position: static; }
      .row { grid-template-columns: 1fr; }
      .task-grid, .mail-layout, .edit-card-grid, .modal-fields, .config-grid, .weekly-section-grid, .weekly-period-grid { grid-template-columns: 1fr; }
      .guide-grid, .field-grid, .field-grid.three { grid-template-columns: 1fr; }
      .weekly-section-card { min-height: auto; }
      .weekly-rows { max-height: none; }
      .modal-fields .wide { grid-column: auto; }
      header { padding: 18px; }
    }
    .agent-float { position: fixed; bottom: 24px; right: 24px; z-index: 1000; }
    .agent-toggle {
      width: 112px; height: 112px; border-radius: 0;
      background: transparent;
      color: #fff; font-size: 24px; border: 0;
      box-shadow: none;
      cursor: pointer; transition: transform .2s;
      display: grid; place-items: center; padding: 0;
      touch-action: none;
    }
    .agent-avatar {
      width: 96px; height: 96px; object-fit: contain;
      filter: drop-shadow(0 4px 10px rgba(15, 23, 42, .22));
      pointer-events: none;
    }
    .agent-header .agent-avatar { width: 30px; height: 30px; margin-right: 8px; }
    .agent-toggle:hover { transform: scale(1.08); }
    .agent-window {
      position: absolute; bottom: 70px; right: 0;
      width: 420px; max-height: 640px;
      background: #fff; border-radius: 14px;
      box-shadow: 0 20px 50px rgba(26,34,47,.18);
      display: flex; flex-direction: column;
      overflow: hidden; border: 1px solid var(--line);
      resize: both;
    }
    .agent-window.hidden { display: none; }
    .agent-header {
      display: flex; justify-content: space-between; align-items: center;
      padding: 14px 16px; background: #123c3b; color: #fff;
      cursor: move; user-select: none;
    }
    .agent-header span { font-weight: 720; font-size: 15px; }
    .agent-close { background: rgba(255,255,255,.18); color: #fff; padding: 4px 10px; border-radius: 6px; font-size: 13px; }
    .agent-body { display: flex; flex-direction: column; flex: 1; min-height: 0; }
    .agent-messages { flex: 1; overflow-y: auto; padding: 14px; display: flex; flex-direction: column; gap: 10px; }
    .agent-msg { max-width: 88%; padding: 10px 14px; border-radius: 14px; font-size: 14px; line-height: 1.7; white-space: pre-wrap; word-break: break-word; }
    .agent-msg.user { align-self: flex-end; background: #17736a; color: #fff; border-bottom-right-radius: 4px; }
    .agent-msg.assistant { align-self: flex-start; background: #f1f5f9; color: var(--ink); border-bottom-left-radius: 4px; }
    .agent-actions { display: flex; gap: 8px; padding: 0 14px 10px; }
    .agent-action { flex: 1; padding: 8px; font-size: 13px; background: #eef8f6; color: var(--accent); border: 1px solid #b9ded8; border-radius: 8px; cursor: pointer; font-weight: 680; }
    .agent-action:hover { background: #dcefea; }
    .agent-input-wrap { display: flex; gap: 8px; padding: 10px 14px 14px; border-top: 1px solid var(--line); align-items: flex-end; }
    .agent-input-wrap textarea { flex: 1; border: 1px solid var(--line); border-radius: 8px; padding: 9px 12px; font-size: 14px; min-height: 56px; max-height: 140px; resize: vertical; line-height: 1.5; }
    .agent-input-wrap button { padding: 9px 16px; background: var(--accent); color: #fff; border-radius: 8px; font-size: 14px; border: 0; font-weight: 680; height: fit-content; }

    /* ===== 工作日记 ===== */
    .diary-tabs {
      display: flex; gap: 6px; margin: 0 0 20px;
      padding: 4px; background: var(--surface);
      border-radius: var(--radius); width: fit-content;
      border: 1px solid var(--line);
    }
    .diary-tab {
      padding: 8px 18px; border-radius: var(--radius-sm);
      font-size: 13px; font-weight: 700;
      background: transparent; color: var(--muted); border: 0;
      cursor: pointer; transition: all .15s ease;
      display: inline-flex; align-items: center; gap: 7px;
    }
    .diary-tab.active {
      background: var(--accent-gradient); color: #fff;
      box-shadow: 0 4px 12px rgba(59, 130, 246, .25);
    }
    .diary-tab:hover:not(.active) { color: var(--ink); background: var(--elevated); }
    .diary-view.hidden { display: none; }
    .diary-card {
      background: var(--panel); border: 1px solid var(--line);
      border-radius: var(--radius); padding: 24px;
    }
    .diary-card textarea { min-height: 120px; }
    .diary-field { cursor: pointer; background: var(--surface); transition: border-color .2s; }
    .diary-field:hover { border-color: var(--accent); }
    .diary-two-col {
      display: grid;
      grid-template-columns: 1fr 1fr;
      gap: 16px;
      margin-bottom: 4px;
    }
    .diary-list-header {
      display: flex; justify-content: space-between; align-items: center;
      margin-bottom: 14px;
    }
    .diary-list { display: grid; gap: 10px; }
    .diary-item {
      display: grid;
      grid-template-columns: auto minmax(0, 1fr) auto;
      gap: 14px; align-items: center;
      padding: 14px 16px;
      background: var(--surface);
      border: 1px solid var(--line);
      border-radius: var(--radius-sm);
      cursor: pointer;
      transition: all .15s ease;
    }
    .diary-item:hover {
      border-color: var(--accent);
      background: rgba(59, 130, 246, .06);
    }
    .diary-item-date {
      font-weight: 800; font-size: 13px;
      color: var(--accent); white-space: nowrap;
      background: rgba(59, 130, 246, .1);
      padding: 4px 10px; border-radius: 6px;
    }
    .diary-item-preview {
      color: var(--muted); font-size: 13px;
      white-space: nowrap; overflow: hidden; text-overflow: ellipsis;
    }
    .diary-empty {
      text-align: center; padding: 40px 20px;
      color: var(--muted); font-size: 14px;
    }
    .diary-detail-body { display: grid; gap: 16px; }
    .diary-detail-section {
      background: var(--surface); border: 1px solid var(--line);
      border-radius: var(--radius-sm); padding: 14px;
    }
    .diary-detail-label {
      font-size: 12px; font-weight: 700; color: var(--muted);
      text-transform: uppercase; letter-spacing: .06em;
      margin-bottom: 8px;
    }
    .diary-detail-content {
      font-size: 14px; line-height: 1.65; color: var(--ink);
      white-space: pre-wrap;
    }
    .diary-summarize-box {
      background: rgba(59, 130, 246, .08);
      border: 1px dashed rgba(59, 130, 246, .35);
      border-radius: var(--radius-sm);
      padding: 14px 16px;
      margin-bottom: 16px;
      display: flex; align-items: center; gap: 14px;
      flex-wrap: wrap;
    }
    .diary-summarize-box label { margin: 0; font-size: 12px; }
    .diary-summarize-box input {
      width: auto; min-width: 140px; padding: 8px 12px; font-size: 13px;
    }
    .diary-summarize-box button { padding: 8px 16px; font-size: 13px; }

    @media (max-width: 520px) {
      .agent-window { width: calc(100vw - 32px); right: -8px; }
    }

    /* ===== Modern AI Office Design System Overrides ===== */
    body {
      background: var(--bg);
      color: var(--ink);
      font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "PingFang SC", "Microsoft YaHei", sans-serif;
      line-height: 1.6;
      -webkit-font-smoothing: antialiased;
    }
    header {
      background: rgba(17, 24, 39, .85);
      backdrop-filter: blur(12px);
      border-bottom: 1px solid var(--line);
      position: sticky;
      top: 0;
      z-index: 100;
      padding: 14px 24px;
    }
    .header-row { display: flex; justify-content: space-between; align-items: center; gap: 20px; }
    .brand { display: flex; align-items: center; gap: 12px; font-size: 18px; font-weight: 700; letter-spacing: .3px; }
    .brand-icon {
      width: 36px; height: 36px; border-radius: 10px;
      background: var(--accent-gradient); display: grid; place-items: center; color: #fff; font-size: 18px;
      box-shadow: var(--glow);
    }
    .icon, .nav-icon, .card-icon {
      display: inline-flex;
      align-items: center;
      justify-content: center;
      flex-shrink: 0;
    }
    .icon svg, .nav-icon svg, .card-icon svg {
      width: 1.15em;
      height: 1.15em;
      stroke-width: 2;
    }
    .global-search {
      flex: 1; max-width: 420px; position: relative;
    }
    .global-search input {
      width: 100%; background: var(--surface); border: 1px solid var(--line); border-radius: 999px;
      padding: 10px 16px 10px 38px; color: var(--ink); font-size: 14px;
      transition: all .2s;
    }
    .global-search input:focus { outline: none; border-color: var(--accent); box-shadow: 0 0 0 3px rgba(59,130,246,.2); }
    .global-search .search-icon { position: absolute; left: 14px; top: 50%; transform: translateY(-50%); color: var(--muted); }
    .global-search .search-icon svg { width: 16px; height: 16px; }
    .userbar { display: none; align-items: center; gap: 12px; }
    .user-avatar {
      width: 32px; height: 32px; border-radius: 50%;
      object-fit: cover; background: #eaf1fb; border: 1px solid var(--line);
    }
    .userbar button {
      background: var(--elevated); color: var(--muted); border: 1px solid var(--line);
      padding: 8px 16px; border-radius: var(--radius-sm); font-size: 13px; font-weight: 600;
      transition: all .2s;
    }
    .userbar button:hover { background: var(--panel); color: var(--ink); border-color: var(--accent); }

    main {
      display: grid;
      grid-template-columns: 260px minmax(0, 1fr);
      gap: 0;
      padding: 0;
      max-width: none;
      margin: 0;
      align-items: start;
      min-height: calc(100vh - 70px);
    }
    .sidebar {
      background: var(--surface);
      border-right: 1px solid var(--line);
      padding: 20px 14px;
      position: sticky;
      top: 70px;
      height: calc(100vh - 70px);
      overflow-y: auto;
    }
    .sidebar-title {
      font-size: 11px; font-weight: 700; text-transform: uppercase;
      letter-spacing: .08em; color: var(--muted); margin: 0 10px 14px;
    }
    .nav-group { margin-bottom: 20px; }
    .nav-group + .nav-group { border-top: 1px solid var(--line); padding-top: 20px; }
    .task-grid { display: flex; flex-direction: column; gap: 6px; }
    .task-card {
      width: 100%; text-align: left; border: 1px solid transparent;
      border-radius: var(--radius-sm); padding: 10px 12px;
      background: transparent; color: var(--muted); box-shadow: none;
      cursor: pointer; transition: all .15s; font-size: 13px;
      display: flex; align-items: center; gap: 10px;
    }
    .task-card:hover {
      background: var(--panel); color: var(--ink); border-color: var(--line);
    }
    .task-card.active {
      background: rgba(59, 130, 246, .12); color: var(--accent); border-color: rgba(59, 130, 246, .35);
      box-shadow: inset 0 0 0 1px rgba(59, 130, 246, .15);
    }
    .task-card .task-name { font-weight: 700; font-size: 13.5px; margin: 0; }
    .task-card .task-desc { display: none; }
    .task-card .nav-icon { width: 20px; height: 20px; display: grid; place-items: center; font-size: 15px; flex-shrink: 0; }

    .composer { padding: 28px; background: var(--bg); }
    .page-header { margin-bottom: 24px; }
    .page-header h2 { font-size: 22px; font-weight: 800; margin: 0 0 6px; letter-spacing: -.3px; }
    .page-header p { color: var(--muted); font-size: 14px; margin: 0; }

    .sub-tabs {
      display: flex; gap: 6px; margin: 0 0 24px;
      padding: 4px; background: var(--surface);
      border-radius: var(--radius); width: fit-content;
      border: 1px solid var(--line);
    }
    .sub-tabs.hidden { display: none; }
    .sub-tab {
      padding: 8px 18px; border-radius: var(--radius-sm);
      font-size: 13px; font-weight: 700;
      background: transparent; color: var(--muted); border: 0;
      cursor: pointer; transition: all .15s ease;
      display: inline-flex; align-items: center; gap: 7px;
    }
    .sub-tab.active {
      background: var(--accent-gradient); color: #fff;
      box-shadow: 0 4px 12px rgba(59, 130, 246, .25);
    }
    .sub-tab:hover:not(.active) { color: var(--ink); background: var(--elevated); }

    section, .task-panel {
      background: var(--panel);
      border: 1px solid var(--line);
      border-radius: var(--radius);
      box-shadow: var(--shadow);
    }
    .task-panel { overflow: hidden; }

    h2 { margin: 0 0 14px; font-size: 17px; color: var(--ink); }
    label { display: block; margin: 12px 0 6px; font-size: 12px; color: var(--muted); font-weight: 600; text-transform: uppercase; letter-spacing: .04em; }
    input, select, textarea {
      width: 100%; background: var(--surface); border: 1px solid var(--line);
      border-radius: var(--radius-sm); padding: 10px 12px;
      font: inherit; color: var(--ink); transition: all .2s;
    }
    input:focus, select:focus, textarea:focus { outline: none; border-color: var(--accent); box-shadow: 0 0 0 3px rgba(59,130,246,.12); }
    textarea { min-height: 120px; resize: vertical; line-height: 1.6; background: var(--surface); }
    .row { display: grid; grid-template-columns: 1fr 1fr; gap: 16px; }
    .toolbar { display: flex; gap: 10px; align-items: center; flex-wrap: wrap; margin-top: 18px; }
    #generateToolbar.hidden { display: none !important; }
    button {
      border: 0; border-radius: var(--radius-sm); padding: 10px 18px;
      font-weight: 700; cursor: pointer; background: var(--accent-gradient);
      color: #fff; font-size: 13px; transition: all .2s; letter-spacing: .2px;
      box-shadow: 0 4px 14px rgba(59, 130, 246, .25);
    }
    button:hover { transform: translateY(-1px); box-shadow: 0 6px 20px rgba(59, 130, 246, .35); }
    button:active { transform: translateY(0); }
    button.secondary { background: var(--elevated); color: var(--ink); box-shadow: none; border: 1px solid var(--line); }
    button.secondary:hover { background: var(--panel); border-color: var(--accent); }
    button.warn { background: linear-gradient(135deg, #f59e0b, #ef4444); }
    button.action-button { padding: 10px 16px; font-size: 13px; background: var(--accent-gradient); color: #fff; border: 0; white-space: nowrap; }
    .download-link {
      display: inline-flex; align-items: center; min-height: 36px;
      padding: 8px 14px; border-radius: var(--radius-sm);
      background: rgba(59, 130, 246, .1); color: var(--accent);
      border: 1px solid rgba(59, 130, 246, .25); text-decoration: none;
      font-weight: 700; font-size: 13px; transition: all .2s;
    }
    .download-link:hover { background: rgba(59, 130, 246, .18); }

    .report {
      border: 1px solid var(--line); border-radius: var(--radius-sm);
      padding: 14px; margin-bottom: 10px; cursor: pointer;
      background: var(--surface); transition: all .15s;
    }
    .report:hover { border-color: var(--accent); background: rgba(59, 130, 246, .06); }
    .report.active { border-color: var(--accent); background: rgba(59, 130, 246, .1); box-shadow: 0 0 0 1px rgba(59, 130, 246, .15); }
    .report-head { display: flex; justify-content: space-between; gap: 10px; align-items: flex-start; }
    .report-actions { display: flex; gap: 6px; flex-shrink: 0; }
    .report-actions .mini { padding: 5px 8px; font-size: 12px; }

    .weekly-period-card, .weekly-section-card {
      background: var(--surface); border: 1px solid var(--line); border-radius: var(--radius);
      padding: 18px; margin-bottom: 16px;
    }
    .weekly-period-head { display: flex; justify-content: space-between; align-items: flex-start; gap: 12px; margin-bottom: 14px; }
    .weekly-period-title { font-size: 15px; font-weight: 800; color: var(--ink); }
    .weekly-section-title { font-size: 14px; font-weight: 700; color: var(--ink); margin-bottom: 4px; }
    .weekly-section-meta { font-size: 12px; color: var(--muted); }
    .work-block {
      background: var(--elevated); border: 1px solid var(--line);
      border-radius: var(--radius-sm); padding: 14px; margin-bottom: 10px;
      cursor: pointer; transition: all .15s;
    }
    .work-block:hover { border-color: var(--accent); }

    .trip-store input, .trip-store textarea {
      background: var(--surface); margin-bottom: 10px;
    }

    .auth-panel {
      max-width: 420px; margin: 80px auto;
      background: var(--panel); border: 1px solid var(--line);
      border-radius: var(--radius); padding: 36px;
      box-shadow: var(--shadow);
    }
    .auth-panel h2 { font-size: 20px; margin-bottom: 20px; text-align: center; }

    /* Dashboard */
    .dashboard-grid { display: grid; gap: 24px; }
    .welcome-card {
      background: var(--panel); border: 1px solid var(--line);
      border-radius: var(--radius); padding: 28px;
      position: relative; overflow: hidden;
    }
    .welcome-card::before {
      content: ""; position: absolute; top: -40%; right: -10%;
      width: 300px; height: 300px; border-radius: 50%;
      background: radial-gradient(circle, rgba(59,130,246,.15), transparent 70%);
      pointer-events: none;
    }
    .welcome-card h1 { font-size: 24px; font-weight: 800; margin: 0 0 8px; position: relative; z-index: 1; }
    .welcome-card p { color: var(--muted); margin: 0 0 20px; position: relative; z-index: 1; }
    .quick-actions { display: flex; gap: 10px; flex-wrap: wrap; position: relative; z-index: 1; }
    .quick-actions button { padding: 10px 20px; font-size: 13px; }

    .shortcut-grid { display: grid; grid-template-columns: repeat(3, minmax(0, 1fr)); gap: 16px; }
    .shortcut-card {
      background: var(--panel); border: 1px solid var(--line);
      border-radius: var(--radius); padding: 20px; cursor: pointer;
      transition: all .2s; position: relative; overflow: hidden;
      display: flex; align-items: center; gap: 16px;
    }
    .shortcut-card:hover {
      transform: translateY(-2px);
      border-color: var(--accent);
      box-shadow: 0 8px 24px rgba(2, 8, 20, .5);
    }
    .shortcut-card .icon {
      width: 52px; height: 52px; border-radius: 14px;
      background: var(--accent-gradient); display: grid; place-items: center;
      color: #fff; font-size: 22px; flex-shrink: 0;
      box-shadow: var(--glow);
    }
    .shortcut-card .title { font-weight: 800; font-size: 16px; margin-bottom: 6px; }
    .shortcut-card .desc { font-size: 13px; color: var(--muted); line-height: 1.55; }
    @media (max-width: 1180px) {
      .shortcut-grid { grid-template-columns: repeat(2, minmax(0, 1fr)); }
    }

    .recent-docs { background: var(--panel); border: 1px solid var(--line); border-radius: var(--radius); padding: 20px; }
    .recent-docs h3 { font-size: 14px; font-weight: 800; margin: 0 0 14px; text-transform: uppercase; letter-spacing: .06em; color: var(--muted); }
    .doc-item { display: flex; align-items: center; justify-content: space-between; padding: 10px 0; border-bottom: 1px solid var(--line); }
    .doc-item:last-child { border-bottom: 0; }
    .doc-name { font-size: 13px; font-weight: 600; }
    .doc-meta { font-size: 11px; color: var(--muted); }
    .ai-command {
      display: grid;
      grid-template-columns: minmax(0, 1fr) auto;
      gap: 10px;
      position: relative;
      z-index: 1;
      max-width: 760px;
    }
    .ai-command input {
      min-height: 46px;
      border-radius: 999px;
      padding: 12px 18px;
      background: #fff;
      color: var(--ink);
      border-color: #cbd8e8;
      box-shadow: inset 0 1px 0 rgba(255, 255, 255, .8);
    }
    .ai-command input:focus {
      background: #fff;
      border-color: var(--accent);
      box-shadow: 0 0 0 4px rgba(37, 99, 235, .12);
    }
    .ai-command button { border-radius: 999px; }
    .module-grid {
      display: grid;
      grid-template-columns: 1fr 1fr;
      gap: 20px;
    }
    .recommend-list { display: grid; gap: 10px; }
    .recommend-item {
      display: grid;
      grid-template-columns: auto minmax(0, 1fr);
      gap: 10px;
      align-items: start;
      padding: 12px;
      background: var(--surface);
      border: 1px solid var(--line);
      border-radius: var(--radius-sm);
    }
    .recommend-item .card-icon { color: var(--accent); margin-top: 2px; }
    .recommend-item strong { display: block; font-size: 13px; margin-bottom: 2px; }
    .recommend-item span { color: var(--muted); font-size: 12px; line-height: 1.45; }
    .forum-layout {
      display: grid;
      grid-template-columns: minmax(320px, .9fr) minmax(520px, 1.5fr);
      gap: 16px;
    }
    .forum-card {
      background: #fff;
      border: 1px solid var(--line);
      border-radius: var(--radius);
      padding: 16px;
      box-shadow: var(--shadow);
    }
    .forum-card h3 { margin: 0 0 12px; font-size: 15px; }
    .forum-create-panel.hidden { display: none; }
    .forum-topic-list { display: grid; gap: 10px; margin-top: 12px; max-height: 660px; overflow: auto; }
    .forum-topic-item {
      border: 1px solid var(--line);
      border-radius: var(--radius-sm);
      padding: 12px;
      background: #f8fbff;
      cursor: pointer;
      transition: all .16s ease;
    }
    .forum-topic-item:hover,
    .forum-topic-item.active {
      border-color: #93c5fd;
      background: #eff6ff;
    }
    .forum-topic-title { font-weight: 760; font-size: 14px; line-height: 1.45; }
    .forum-topic-meta { margin-top: 6px; color: var(--muted); font-size: 12px; }
    .forum-topic-stats {
      display: flex;
      gap: 8px;
      flex-wrap: wrap;
      margin-top: 8px;
    }
    .forum-stat {
      display: inline-flex;
      align-items: center;
      gap: 4px;
      padding: 3px 8px;
      border-radius: 999px;
      background: #eef6ff;
      color: #1d4ed8;
      font-size: 12px;
      font-weight: 680;
    }
    .forum-topic-body,
    .forum-comment-body {
      white-space: pre-wrap;
      color: var(--ink);
      line-height: 1.65;
      font-size: 14px;
    }
    .forum-comments { display: grid; gap: 10px; margin-top: 12px; }
    .forum-comment {
      border: 1px solid var(--line);
      border-radius: var(--radius-sm);
      padding: 12px;
      background: #f8fafc;
    }
    .forum-comment.reply {
      margin-left: 28px;
      border-left: 3px solid #93c5fd;
    }
    .forum-comment-meta { color: var(--muted); font-size: 12px; margin-bottom: 6px; }
    .forum-comment-actions { margin-top: 8px; display: flex; gap: 8px; }
    .forum-pagination {
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 12px;
      margin-top: 12px;
      padding-top: 10px;
      border-top: 1px solid var(--line);
      color: var(--muted);
      font-size: 12px;
    }
    .forum-empty {
      color: var(--muted);
      border: 1px dashed var(--line);
      border-radius: var(--radius-sm);
      padding: 18px;
      text-align: center;
      background: #fbfdff;
    }
    .news-layout {
      display: grid;
      grid-template-columns: minmax(320px, .85fr) minmax(480px, 1.4fr);
      gap: 16px;
    }
    .news-layout.reader { grid-template-columns: 1fr; }
    .news-card {
      background: #fff;
      border: 1px solid var(--line);
      border-radius: var(--radius);
      padding: 16px;
      box-shadow: var(--shadow);
    }
    .news-card h3 { margin: 0 0 12px; font-size: 15px; }
    .news-source-row {
      display: grid;
      grid-template-columns: minmax(100px, .45fr) minmax(160px, 1fr) auto;
      gap: 8px;
      align-items: center;
      margin-bottom: 8px;
    }
    .news-issue-head {
      display: flex;
      justify-content: space-between;
      gap: 12px;
      align-items: flex-start;
      margin-bottom: 12px;
    }
    .news-issue-title { font-size: 20px; font-weight: 820; color: var(--ink); }
    .news-issue-meta { color: var(--muted); font-size: 12px; margin-top: 4px; }
    .news-summary {
      white-space: pre-wrap;
      line-height: 1.65;
      padding: 12px;
      border: 1px solid var(--line);
      border-radius: var(--radius-sm);
      background: #f8fbff;
      margin-bottom: 12px;
    }
    .news-list { display: grid; gap: 10px; }
    .news-item {
      border: 1px solid var(--line);
      border-radius: var(--radius-sm);
      padding: 12px;
      background: #fff;
    }
    .news-item-title { font-weight: 760; margin-bottom: 6px; }
    .news-item-meta { color: var(--muted); font-size: 12px; margin-bottom: 6px; }
    .news-item p { margin: 6px 0; line-height: 1.6; }
    .news-keywords { display: flex; flex-wrap: wrap; gap: 6px; margin-top: 12px; }
    .news-keyword {
      padding: 4px 8px;
      border-radius: 999px;
      background: #eff6ff;
      color: #1d4ed8;
      font-size: 12px;
      font-weight: 680;
    }
    @media (max-width: 1100px) {
      .forum-layout { grid-template-columns: 1fr; }
      .news-layout { grid-template-columns: 1fr; }
    }

    /* Agent */
    .agent-float { position: fixed; bottom: 28px; right: 28px; z-index: 3000; touch-action: none; }
    .agent-toggle {
      width: 112px; height: 112px; border-radius: 0;
      background: transparent; color: #fff; font-size: 24px;
      border: 0; box-shadow: none;
      cursor: grab; transition: all .2s;
      display: grid; place-items: center; padding: 0;
      touch-action: none;
    }
    .agent-toggle:hover { transform: scale(1.04); box-shadow: none; }
    .agent-toggle:active { cursor: grabbing; }
    .agent-window {
      position: absolute; bottom: 72px; right: 0;
      width: 420px; max-height: 640px; min-height: 400px;
      background: var(--panel); border-radius: var(--radius);
      box-shadow: 0 24px 60px rgba(2, 8, 20, .55);
      display: flex; flex-direction: column; overflow: hidden;
      border: 1px solid var(--line); resize: both;
    }
    .agent-header {
      display: flex; justify-content: space-between; align-items: center;
      padding: 14px 18px; background: var(--accent-gradient); color: #fff;
      cursor: move; user-select: none;
    }
    .agent-close { background: rgba(255,255,255,.2); color: #fff; padding: 4px 10px; border-radius: 6px; font-size: 13px; border: 0; cursor: pointer; }
    .agent-close:hover { background: rgba(255,255,255,.3); }
    .agent-msg { max-width: 88%; padding: 10px 14px; border-radius: 14px; font-size: 14px; line-height: 1.7; white-space: pre-wrap; word-break: break-word; }
    .agent-msg.user { align-self: flex-end; background: var(--accent); color: #fff; border-bottom-right-radius: 4px; }
    .agent-msg.assistant { align-self: flex-start; background: var(--surface); color: var(--ink); border-bottom-left-radius: 4px; border: 1px solid var(--line); }
    .agent-action { flex: 1; padding: 8px; font-size: 12px; background: var(--surface); color: var(--accent); border: 1px solid var(--line); border-radius: var(--radius-sm); cursor: pointer; font-weight: 700; }
    .agent-action:hover { background: rgba(59, 130, 246, .1); border-color: var(--accent); }
    .agent-input-wrap textarea { background: var(--surface); border: 1px solid var(--line); border-radius: var(--radius-sm); padding: 10px 12px; font-size: 14px; min-height: 56px; max-height: 140px; resize: vertical; line-height: 1.5; flex: 1; color: var(--ink); }
    .agent-input-wrap textarea:focus { outline: none; border-color: var(--accent); box-shadow: 0 0 0 3px rgba(59,130,246,.12); }
    .agent-input-wrap button { padding: 10px 18px; font-size: 13px; }
    .agent-lightbox {
      position: fixed; inset: 0; z-index: 4000;
      background: rgba(0,0,0,.85);
      display: flex; align-items: center; justify-content: center;
      cursor: zoom-out; padding: 20px;
    }
    .agent-lightbox.hidden { display: none; }
    .agent-lightbox img {
      max-width: 90vw; max-height: 90vh;
      border-radius: 12px; box-shadow: 0 24px 64px rgba(0,0,0,.5);
      background: #fff; object-fit: contain;
    }

    @media (max-width: 900px) {
      main { grid-template-columns: 1fr; }
      .sidebar { display: none; }
      .module-grid { grid-template-columns: 1fr; }
      .ai-command { grid-template-columns: 1fr; }
    }

    /* ===== Comfortable Light AI Office Theme ===== */
    :root {
      color-scheme: light;
      --bg: #f4f7fb;
      --surface: #f8fafc;
      --panel: #ffffff;
      --elevated: #eef4ff;
      --ink: #172033;
      --muted: #64748b;
      --line: #dbe5f1;
      --accent: #2563eb;
      --accent-2: #0ea5e9;
      --accent-gradient: linear-gradient(135deg, #2563eb, #0ea5e9);
      --danger: #dc2626;
      --success: #059669;
      --warning: #d97706;
      --shadow: 0 12px 34px rgba(15, 23, 42, .08);
      --glow: 0 12px 28px rgba(37, 99, 235, .18);
    }
    body {
      background:
        radial-gradient(circle at top left, rgba(37, 99, 235, .08), transparent 30%),
        linear-gradient(180deg, #f7fbff 0%, #eef4fb 100%);
      color: var(--ink);
    }
    header {
      background: rgba(255, 255, 255, .92);
      border-bottom: 1px solid var(--line);
      color: var(--ink);
      box-shadow: 0 8px 24px rgba(15, 23, 42, .05);
    }
    .global-search input,
    input,
    select,
    textarea,
    .agent-input-wrap textarea {
      background: #ffffff;
      color: var(--ink);
      border-color: #cbd8e8;
    }
    input::placeholder,
    textarea::placeholder { color: #94a3b8; }
    input:focus,
    select:focus,
    textarea:focus,
    .agent-input-wrap textarea:focus {
      background: #fff;
      border-color: var(--accent);
      box-shadow: 0 0 0 4px rgba(37, 99, 235, .12);
    }
    .sidebar {
      background: rgba(255, 255, 255, .82);
      backdrop-filter: blur(14px);
      border-right: 1px solid var(--line);
    }
    .task-card {
      color: #526174;
    }
    .task-card:hover {
      background: #eef5ff;
      color: var(--ink);
      border-color: #bfdbfe;
    }
    .task-card.active {
      background: linear-gradient(135deg, rgba(37, 99, 235, .12), rgba(14, 165, 233, .1));
      color: #1d4ed8;
      border-color: #93c5fd;
      box-shadow: inset 3px 0 0 #2563eb;
    }
    .composer { background: transparent; }
    section,
    .task-panel,
    .welcome-card,
    .shortcut-card,
    .insight-card,
    .recent-docs,
    .weekly-period-card,
    .weekly-section-card,
    .guide,
    .side,
    .auth-panel,
    .modal-box,
    .agent-window {
      background: rgba(255, 255, 255, .96);
      border-color: var(--line);
      color: var(--ink);
      box-shadow: var(--shadow);
    }
    .sub-tabs {
      background: #eaf1fb;
      border-color: #d4e2f3;
    }
    .sub-tab { color: #5f7189; }
    .sub-tab:hover:not(.active) {
      background: #f7fbff;
      color: var(--ink);
    }
    .work-block,
    .edit-card,
    .report,
    .history-item,
    .upload-item,
    .user-item,
    .recommend-item,
    .rich-preview,
    .preview {
      background: #ffffff;
      color: var(--ink);
      border-color: var(--line);
    }
    .weekly-section-head {
      background: #f7fbff;
      border-bottom-color: var(--line);
    }
    .work-title,
    .edit-card-title,
    .weekly-period-title,
    .weekly-section-title,
    .doc-name,
    .history-name { color: var(--ink); }
    .work-summary,
    .edit-card-preview,
    .history-meta,
    .doc-meta,
    .insight-desc,
    .shortcut-card .desc,
    .recommend-item span,
    .page-header p,
    label,
    .sidebar-title { color: var(--muted); }
    button.secondary {
      background: #ffffff;
      color: #1e3a8a;
      border: 1px solid #bfdbfe;
    }
    button.secondary:hover {
      background: #eff6ff;
      border-color: #60a5fa;
    }
    .agent-header {
      background: var(--accent-gradient);
      color: #fff;
    }
    .agent-msg.assistant {
      background: #f8fafc;
      color: var(--ink);
      border: 1px solid var(--line);
    }
    .agent-progress {
      display: grid;
      grid-template-columns: repeat(4, minmax(0, 1fr));
      gap: 6px;
      padding: 10px 12px;
      border-bottom: 1px solid var(--line);
      background: #f8fafc;
    }
    .agent-step {
      border: 1px solid #dbe5f1;
      border-radius: 8px;
      padding: 6px 4px;
      text-align: center;
      font-size: 11px;
      line-height: 1.2;
      color: #64748b;
      background: #fff;
      font-weight: 700;
    }
    .agent-step.active {
      color: #1d4ed8;
      border-color: #93c5fd;
      background: #eff6ff;
    }
    .agent-step.done {
      color: #047857;
      border-color: #86efac;
      background: #ecfdf5;
    }
    .agent-card {
      display: grid;
      gap: 8px;
      min-width: min(320px, 78vw);
    }
    .agent-card-title {
      font-size: 13px;
      font-weight: 850;
      color: var(--ink);
    }
    .agent-card-grid {
      display: grid;
      grid-template-columns: repeat(3, minmax(0, 1fr));
      gap: 6px;
    }
    .agent-card-metric {
      border: 1px solid var(--line);
      border-radius: 8px;
      padding: 7px;
      background: #fff;
      text-align: center;
    }
    .agent-card-metric strong {
      display: block;
      font-size: 16px;
      color: #1d4ed8;
    }
    .agent-card-note {
      color: var(--muted);
      font-size: 12px;
      line-height: 1.5;
    }
    .agent-intro {
      display: grid;
      gap: 10px;
      min-width: min(340px, 78vw);
    }
    .agent-intro-title {
      font-size: 14px;
      font-weight: 850;
      color: var(--ink);
    }
    .agent-intro-copy {
      color: #475569;
      font-size: 13px;
      line-height: 1.55;
    }
    .agent-quick-grid {
      display: grid;
      grid-template-columns: 1fr;
      gap: 7px;
    }
    .agent-quick {
      border: 1px solid #bfdbfe;
      background: #eff6ff;
      color: #1d4ed8;
      border-radius: 8px;
      padding: 8px 10px;
      font-size: 12px;
      font-weight: 750;
      text-align: left;
      cursor: pointer;
    }
    .agent-quick:hover {
      background: #dbeafe;
      border-color: #60a5fa;
    }
    .agent-help-list {
      margin: 0;
      padding-left: 16px;
      color: #64748b;
      font-size: 12px;
      line-height: 1.55;
    }
    .agent-card-list {
      margin: 0;
      padding-left: 16px;
      color: #475569;
      font-size: 12px;
      line-height: 1.55;
    }
    .send-review {
      margin-top: 12px;
      border: 1px solid #bfdbfe;
      border-radius: 12px;
      background: #eff6ff;
      padding: 12px;
      display: grid;
      gap: 10px;
    }
    .send-review.hidden { display: none; }
    .send-review-title {
      font-size: 14px;
      font-weight: 850;
      color: #1e3a8a;
    }
    .send-review-grid {
      display: grid;
      grid-template-columns: 96px minmax(0, 1fr);
      gap: 6px 10px;
      font-size: 13px;
      color: #334155;
    }
    .send-review-grid .label {
      color: #64748b;
      font-weight: 700;
    }
    .send-review-actions {
      display: flex;
      gap: 8px;
      justify-content: flex-end;
      flex-wrap: wrap;
    }
    .send-review-warning {
      color: #b45309;
      font-size: 13px;
      font-weight: 700;
    }
    .agent-action {
      background: #eff6ff;
      color: #1d4ed8;
      border-color: #bfdbfe;
    }
    .skill-hero,
    .skill-protocol {
      background: #fff;
      border: 1px solid var(--line);
      border-radius: var(--radius);
      box-shadow: var(--shadow);
      padding: 16px;
      margin-bottom: 16px;
    }
    .skill-hero {
      display: flex;
      justify-content: space-between;
      align-items: center;
      gap: 16px;
    }
    .skill-orchestration {
      background: #fff;
      border: 1px solid var(--line);
      border-radius: var(--radius);
      box-shadow: var(--shadow);
      padding: 16px;
      margin-bottom: 16px;
    }
    .skill-orchestration.hidden { display: none; }
    .skill-orchestration .orchestration-section {
      margin-top: 12px;
    }
    .skill-orchestration .orchestration-section-title {
      font-size: 13px;
      font-weight: 800;
      color: var(--ink);
      margin-bottom: 6px;
    }
    .skill-orchestration pre {
      white-space: pre-wrap;
      background: #f8fbff;
      border: 1px solid #dbe5f1;
      border-radius: var(--radius-sm);
      padding: 10px 12px;
      font-size: 12px;
      color: #334155;
      overflow: auto;
      max-height: 320px;
    }
    .skill-protocol pre,
    .skill-call-example {
      white-space: pre-wrap;
      background: #f8fbff;
      border: 1px solid #dbe5f1;
      border-radius: var(--radius-sm);
      padding: 12px;
      margin: 10px 0 0;
      color: #1e293b;
      font-size: 12px;
      line-height: 1.6;
      overflow: auto;
    }
    .skill-filter-row {
      display: grid;
      grid-template-columns: minmax(280px, 1fr) 200px;
      gap: 12px;
      margin-bottom: 16px;
    }
    .skill-installed-head {
      margin: 6px 0 10px;
      display: flex;
      justify-content: space-between;
      align-items: center;
      flex-wrap: wrap;
      gap: 6px;
    }
    .skill-total-count {
      font-size: 13px;
      color: var(--muted);
      font-weight: 700;
    }
    .skill-module-summary {
      display: grid;
      grid-template-columns: repeat(auto-fit, minmax(180px, 1fr));
      gap: 12px;
      margin-bottom: 16px;
    }
    .skill-module-card {
      background: #fff;
      border: 1px solid var(--line);
      border-radius: var(--radius);
      box-shadow: var(--shadow);
      padding: 14px;
      cursor: pointer;
      transition: all .16s ease;
    }
    .skill-module-card:hover,
    .skill-module-card.active {
      transform: translateY(-1px);
      border-color: #93c5fd;
      background: #eff6ff;
    }
    .skill-module-name {
      color: var(--ink);
      font-weight: 850;
      font-size: 14px;
    }
    .skill-module-count {
      color: var(--muted);
      font-size: 12px;
      margin-top: 4px;
    }
    .skill-list {
      display: grid;
      grid-template-columns: repeat(auto-fit, minmax(310px, 1fr));
      gap: 14px;
      align-items: start;
    }
    .skill-card {
      background: #fff;
      border: 1px solid var(--line);
      border-radius: var(--radius);
      box-shadow: var(--shadow);
      padding: 16px;
      display: grid;
      gap: 10px;
      cursor: pointer;
      transition: transform .16s ease, box-shadow .16s ease, border-color .16s ease;
      align-self: start;
    }
    .skill-card:hover {
      transform: translateY(-2px);
      border-color: #93c5fd;
      box-shadow: 0 18px 36px rgba(15, 23, 42, .11);
    }
    .skill-card-head {
      display: flex;
      justify-content: space-between;
      align-items: flex-start;
      gap: 10px;
    }
    .skill-name {
      font-size: 15px;
      font-weight: 850;
      color: var(--ink);
    }
    .skill-title {
      color: var(--muted);
      font-size: 12px;
      margin-top: 2px;
    }
    .skill-badge {
      border-radius: 999px;
      padding: 5px 9px;
      background: #eff6ff;
      border: 1px solid #bfdbfe;
      color: #1d4ed8;
      font-size: 12px;
      font-weight: 760;
      white-space: nowrap;
    }
    .skill-badge.warn {
      background: #fff7ed;
      border-color: #fed7aa;
      color: #c2410c;
    }
    .skill-desc {
      color: #475569;
      font-size: 13px;
      line-height: 1.6;
      display: -webkit-box;
      -webkit-line-clamp: 2;
      -webkit-box-orient: vertical;
      overflow: hidden;
    }
    .skill-card.compact { padding: 12px; gap: 6px; }
    .skill-card-actions {
      display: flex;
      gap: 6px;
      align-items: center;
      flex-shrink: 0;
    }
    .skill-help-btn {
      background: #f1f5f9;
      border: 1px solid #e2e8f0;
      border-radius: 6px;
      padding: 4px 8px;
      font-size: 12px;
      color: #475569;
      cursor: pointer;
      white-space: nowrap;
      line-height: 1;
    }
    .skill-help-btn:hover { background: #e2e8f0; color: #0f172a; }
    .skill-detail {
      display: none;
      margin-top: 4px;
      padding: 10px;
      border: 1px solid #e2e8f0;
      border-radius: 8px;
      background: #f8fafc;
      color: #475569;
      font-size: 12px;
      line-height: 1.5;
      max-height: 300px;
      overflow: auto;
      cursor: default;
    }
    .skill-detail.open { display: block; }
    .skill-detail .meta-label {
      font-size: 11px;
      font-weight: 700;
      color: #94a3b8;
      text-transform: uppercase;
      letter-spacing: .3px;
      margin: 8px 0 4px;
    }
    .skill-detail pre {
      background: #fff;
      border: 1px solid #e2e8f0;
      border-radius: 6px;
      padding: 8px;
      overflow: auto;
      font-size: 11px;
      margin: 0;
      max-height: 180px;
    }
    .skill-test-box {
      width: min(1080px, 96vw);
      max-height: 92vh;
      overflow: auto;
    }
    .skill-test-grid {
      display: grid;
      grid-template-columns: minmax(360px, 1fr) minmax(360px, 1fr);
      gap: 16px;
    }
    .skill-test-panel {
      border: 1px solid var(--line);
      border-radius: var(--radius);
      background: #f8fbff;
      padding: 14px;
      display: grid;
      gap: 10px;
    }
    .skill-test-panel textarea {
      min-height: 150px;
      font-family: ui-monospace, SFMono-Regular, Menlo, Monaco, Consolas, monospace;
      font-size: 13px;
      line-height: 1.6;
      resize: vertical;
    }
    .skill-test-panel textarea#skillTestInstruction {
      min-height: 110px;
      font-family: inherit;
    }
    .skill-test-result {
      min-height: 330px;
      max-height: 520px;
      overflow: auto;
      white-space: pre-wrap;
      word-break: break-word;
      background: #0f172a;
      color: #e2e8f0;
      border: 1px solid #1e293b;
      border-radius: 14px;
      padding: 14px;
      font-size: 12px;
      line-height: 1.65;
    }
    .skill-confirm-line {
      display: flex;
      align-items: center;
      gap: 8px;
      color: #b45309;
      font-size: 13px;
      font-weight: 700;
    }
    .skill-confirm-line input {
      width: 16px;
      height: 16px;
    }
    .skill-test-links {
      display: grid;
      gap: 10px;
    }
    .skill-preview-image {
      width: 100%;
      border-radius: 12px;
      border: 1px solid var(--line);
      background: #fff;
    }
    @media (max-width: 900px) {
      .skill-test-grid {
        grid-template-columns: 1fr;
      }
    }
    .mail-assistant-grid {
      display: grid;
      grid-template-columns: minmax(340px, .9fr) minmax(480px, 1.4fr);
      gap: 16px;
      margin-bottom: 16px;
    }
    .mailbox-card,
    .mail-detail-card,
    .mail-compose-card {
      background: #fff;
      border: 1px solid var(--line);
      border-radius: var(--radius);
      overflow: hidden;
      box-shadow: var(--shadow);
    }
    .mailbox-list {
      max-height: 520px;
      overflow: auto;
      padding: 12px;
      display: grid;
      gap: 10px;
    }
    .mailbox-item {
      border: 1px solid var(--line);
      border-radius: var(--radius-sm);
      padding: 10px 12px;
      background: #f8fbff;
      cursor: pointer;
      transition: all .16s ease;
    }
    .mailbox-item:hover,
    .mailbox-item.active {
      background: #eff6ff;
      border-color: #93c5fd;
      transform: translateY(-1px);
    }
    .mailbox-subject {
      color: var(--ink);
      font-weight: 760;
      font-size: 13.5px;
      margin-bottom: 4px;
    }
    .mailbox-meta,
    .mailbox-preview {
      color: var(--muted);
      font-size: 12px;
      line-height: 1.5;
    }
    .mail-section-head {
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 12px;
      padding: 14px 16px;
      background: #f7fbff;
      border-bottom: 1px solid var(--line);
    }
    .mail-section-title {
      color: var(--ink);
      font-size: 15px;
      font-weight: 800;
      margin-bottom: 3px;
    }
    .mail-section-meta {
      color: var(--muted);
      font-size: 12px;
      line-height: 1.45;
    }
    .mail-detail {
      min-height: 520px;
      padding: 0;
      background: #f8fbff;
      color: var(--ink);
      line-height: 1.65;
    }
    .mail-detail-head {
      border-bottom: 1px solid var(--line);
      padding: 16px 18px;
      background: #fff;
    }
    .mail-detail-body {
      margin: 16px;
      padding: 18px;
      border: 1px solid var(--line);
      border-radius: var(--radius-sm);
      background: #fff;
      min-height: 260px;
      overflow: auto;
      white-space: normal;
    }
    .mail-detail-body.plain {
      white-space: pre-wrap;
    }
    .mail-attachment-list {
      display: flex;
      flex-wrap: wrap;
      gap: 8px;
      margin-top: 10px;
    }
    .mail-attachment {
      display: inline-flex;
      align-items: center;
      gap: 6px;
      padding: 7px 10px;
      border: 1px solid #bfdbfe;
      border-radius: 999px;
      background: #eff6ff;
      color: #1d4ed8;
      font-size: 12px;
      font-weight: 700;
    }
    .mail-compose-card {
      padding-bottom: 16px;
    }
    .mail-compose-layout {
      display: grid;
      grid-template-columns: minmax(420px, 1fr) minmax(320px, .85fr);
      gap: 16px;
      padding: 16px;
      align-items: start;
    }
    .mail-compose-form,
    .mail-compose-preview {
      display: grid;
      gap: 12px;
    }
    .mail-compose-form textarea {
      min-height: 220px;
    }
    .mail-compose-preview .rich-preview {
      min-height: 300px;
      max-height: 460px;
      overflow: auto;
      background: #fff;
    }
    .mail-file-list {
      display: grid;
      gap: 8px;
    }
    .mail-file-item {
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 10px;
      padding: 8px 10px;
      border: 1px solid var(--line);
      border-radius: var(--radius-sm);
      background: #f8fbff;
      font-size: 12px;
      color: var(--muted);
    }
    .profile-grid { display: grid; grid-template-columns: 96px minmax(0, 1fr); gap: 16px; align-items: start; }
    .profile-avatar-large {
      width: 86px; height: 86px; border-radius: 50%; object-fit: cover;
      background: #eaf1fb; border: 1px solid var(--line);
      box-shadow: var(--shadow);
    }
    .profile-actions { display: flex; gap: 8px; flex-wrap: wrap; margin-top: 8px; }
    @media (max-width: 1100px) {
      .mail-assistant-grid { grid-template-columns: 1fr; }
      .mail-detail { min-height: 280px; }
      .mail-compose-layout { grid-template-columns: 1fr; }
      .skill-hero { align-items: flex-start; flex-direction: column; }
      .skill-filter-row { grid-template-columns: 1fr; }
    }

  </style>
</head>
<body>
  <header>
    <div class="header-row">
      <div class="brand">
        <div class="brand-icon"><span class="icon" data-icon="sparkles"></span></div>
        <div>智能办公助手 / AI 办公平台</div>
      </div>
      <div class="global-search">
        <span class="search-icon icon" data-icon="search"></span>
        <input type="text" placeholder="搜索文档、邮件、知识..." />
      </div>
      <div class="userbar" id="userbar">
        <img class="user-avatar" id="userAvatar" src="/assets/ai-assistant-avatar.png" alt="" />
        <span id="userInfo"></span>
        <button type="button" id="profileButton" class="mini">个人资料</button>
        <button type="button" id="changePassButton" class="mini">修改密码</button>
        <button type="button" id="logoutButton">退出</button>
      </div>
    </div>
  </header>
  <section class="auth-panel" id="authPanel">
    <h2>登录</h2>
    <label>用户名</label>
    <input id="loginUser" />
    <label>密码</label>
    <input id="loginPass" type="password" />
    <div class="toolbar">
      <button id="loginButton" type="button">登录</button>
    </div>
    <div id="loginStatus" class="status"></div>
  </section>
  <main id="appMain" class="hidden">
    <aside class="sidebar">
      <div class="nav-group">
        <div class="sidebar-title">工作台</div>
        <div class="task-grid">
          <button class="task-card active" type="button" data-task="dashboard">
            <span class="nav-icon" data-icon="layout-dashboard"></span>
            <span class="task-name">首页</span>
          </button>
        </div>
      </div>
      <div class="nav-group">
        <div class="sidebar-title">犇犇</div>
        <div class="task-grid">
          <button class="task-card" type="button" data-task="weekly">
            <span class="nav-icon" data-icon="file-spreadsheet"></span>
            <span class="task-name">周报助手</span>
          </button>
          <button class="task-card" type="button" data-task="trip">
            <span class="nav-icon" data-icon="briefcase-business"></span>
            <span class="task-name">出差报告助手</span>
          </button>
          <button class="task-card" type="button" data-task="diary">
            <span class="nav-icon" data-icon="book-open"></span>
            <span class="task-name">工作日记</span>
          </button>
          <button class="task-card" type="button" data-task="forum">
            <span class="nav-icon" data-icon="messages-square"></span>
            <span class="task-name">金点子论坛</span>
          </button>
          <button class="task-card" type="button" data-task="news">
            <span class="nav-icon" data-icon="newspaper"></span>
            <span class="task-name">每日资讯</span>
          </button>
          <button class="task-card" type="button" data-task="mailassistant">
            <span class="nav-icon" data-icon="inbox"></span>
            <span class="task-name">邮件助手</span>
          </button>
        </div>
      </div>
      <div class="nav-group">
        <div class="sidebar-title">系统</div>
        <div class="task-grid">
          <button class="task-card" type="button" data-task="mailconfig">
            <span class="nav-icon" data-icon="mail-check"></span>
            <span class="task-name">邮件配置</span>
          </button>
          <button class="task-card admin-only hidden" type="button" data-task="config">
            <span class="nav-icon" data-icon="settings"></span>
            <span class="task-name">系统配置</span>
          </button>
          <button class="task-card superadmin-only hidden" type="button" data-task="skills">
            <span class="nav-icon" data-icon="puzzle"></span>
            <span class="task-name">系统 Skill</span>
          </button>
          <button class="task-card superadmin-only hidden" type="button" data-task="usermanage">
            <span class="nav-icon" data-icon="users"></span>
            <span class="task-name">用户管理</span>
          </button>
        </div>
      </div>
    </aside>
    <section class="composer">
      <div class="task-panel" id="dashboardPanel">
        <div class="dashboard-grid">
          <div class="welcome-card">
            <h1>欢迎回来，<span id="dashUserName"></span></h1>
            <p>围绕日常办公闭环工作：记录工作、生成报告、处理邮件、共创金点子、查看每日轨交资讯。</p>
            <div class="ai-command">
              <input id="dashboardAsk" type="text" placeholder="输入问题，例如：根据本周日记生成周报、整理出差报告、提炼邮件重点..." />
              <button type="button" id="dashboardAskButton"><span class="icon" data-icon="send"></span> 提问助手</button>
            </div>
          </div>
          <div class="shortcut-grid">
            <div class="shortcut-card" onclick="navigateTo('weekly','edit')">
              <div class="icon" data-icon="file-spreadsheet"></div>
              <div>
                <div class="title">填写周报</div>
                <div class="desc">按模板快速生成本周工作总结、重点跟进和下周计划</div>
              </div>
            </div>
            <div class="shortcut-card" onclick="navigateTo('diary')">
              <div class="icon" data-icon="book-open"></div>
              <div>
                <div class="title">记录工作日记</div>
                <div class="desc">沉淀每日工作、明日计划和想法，周报可直接智能总结</div>
              </div>
            </div>
            <div class="shortcut-card" onclick="navigateTo('trip','edit')">
              <div class="icon" data-icon="briefcase-business"></div>
              <div>
                <div class="title">填写出差报告</div>
                <div class="desc">录入出差地点、时间、目的、行程和总结内容</div>
              </div>
            </div>
            <div class="shortcut-card" onclick="navigateTo('forum')">
              <div class="icon" data-icon="messages-square"></div>
              <div>
                <div class="title">金点子论坛</div>
                <div class="desc">浏览历史话题、查看热度点赞评论，围绕话题分级讨论</div>
              </div>
            </div>
            <div class="shortcut-card" onclick="navigateTo('news')">
              <div class="icon" data-icon="newspaper"></div>
              <div>
                <div class="title">每日资讯</div>
                <div class="desc">查看轨道交通关键资讯，配置和生成由超级管理员维护</div>
              </div>
            </div>
            <div class="shortcut-card" onclick="navigateTo('mailassistant')">
              <div class="icon" data-icon="inbox"></div>
              <div>
                <div class="title">邮件助手</div>
                <div class="desc">读取收件箱、查看邮件详情、普通发信；收件箱已支持缓存加速</div>
              </div>
            </div>
          </div>
          <div class="module-grid">
            <div class="recent-docs">
              <h3>最近文档</h3>
              <div id="recentDocsList"><div class="doc-item"><span class="doc-name">暂无最近文档</span></div></div>
            </div>
            <div class="recent-docs">
              <h3>当前能力</h3>
              <div class="recommend-list">
                <div class="recommend-item">
                  <span class="card-icon" data-icon="calendar-check"></span>
                  <div><strong>报告只在专属模块生成</strong><span>“按标准模板生成文件”仅在周报助手和出差报告助手中出现。</span></div>
                </div>
                <div class="recommend-item">
                  <span class="card-icon" data-icon="messages-square"></span>
                  <div><strong>论坛以浏览评论为主</strong><span>发起话题默认收起，历史话题展示热度、点赞、评论和浏览。</span></div>
                </div>
                <div class="recommend-item">
                  <span class="card-icon" data-icon="newspaper"></span>
                  <div><strong>资讯配置受权限保护</strong><span>普通用户只看每日资讯，只有超级管理员能配置来源和立即生成。</span></div>
                </div>
                <div class="recommend-item">
                  <span class="card-icon" data-icon="inbox"></span>
                  <div><strong>邮件读取已缓存</strong><span>收件箱列表优先读缓存，点击刷新才强制重新拉取最新邮件。</span></div>
                </div>
              </div>
            </div>
          </div>
        </div>
      </div>
      <div class="page-header hidden" id="pageHeader">
        <h2 id="taskTitle"></h2>
        <p id="taskDesc"></p>
      </div>
      <div class="sub-tabs hidden" id="subTabs">
        <button class="sub-tab active" type="button" data-sub="edit"><span class="icon" data-icon="clipboard-pen"></span> 填写报告</button>
        <button class="sub-tab" type="button" data-sub="mail"><span class="icon" data-icon="mail"></span> 发送邮件</button>
        <button class="sub-tab" type="button" data-sub="history"><span class="icon" data-icon="folder-clock"></span> 历史管理</button>
      </div>
      <input id="kind" type="hidden" value="weekly" />
      <div class="task-panel hidden" id="diaryPanel">
        <div class="diary-tabs">
          <button class="diary-tab active" type="button" data-diarytab="write"><span class="icon" data-icon="pen-line"></span> 记录日记</button>
          <button class="diary-tab" type="button" data-diarytab="browse"><span class="icon" data-icon="library"></span> 浏览日记</button>
        </div>
        <div class="diary-view" id="diaryWriteView">
          <div class="diary-card">
            <div class="row">
              <div>
                <label>日期</label>
                <input id="diaryDate" type="date" />
              </div>
              <div style="display:flex;align-items:flex-end;">
                <button type="button" class="secondary" id="diaryLoadToday">载入今天</button>
              </div>
            </div>
            <div class="diary-two-col">
              <div>
                <label>今日工作内容</label>
                <textarea id="diaryTodayWork" class="diary-field" readonly placeholder="记录今天完成的主要工作、遇到的问题、解决方案..."></textarea>
              </div>
              <div>
                <label>明日工作计划</label>
                <textarea id="diaryTomorrowPlan" class="diary-field" readonly placeholder="计划明天要做的事项..."></textarea>
              </div>
            </div>
            <label>思路与想法</label>
            <textarea id="diaryThoughts" class="diary-field" readonly placeholder="灵光一闪的想法、改进建议、学习心得..."></textarea>
            <div class="toolbar">
              <button type="button" id="diarySaveButton">保存日记</button>
              <button type="button" class="secondary" id="diaryClearButton">清空</button>
              <button type="button" id="diaryAgentButton"><span class="icon" data-icon="sparkles"></span> AI 智能记录</button>
              <span id="diaryStatus" class="status"></span>
            </div>
          </div>
        </div>
        <div class="diary-view hidden" id="diaryBrowseView">
          <div class="diary-card">
            <div class="diary-list-header">
              <h3 style="margin:0;font-size:15px;">日记列表</h3>
              <button type="button" class="mini" id="diaryRefreshList">刷新</button>
            </div>
            <div id="diaryList" class="diary-list"><div class="diary-empty">暂无日记，去记录第一篇吧～</div></div>
          </div>
        </div>
        <div class="modal hidden" id="diaryDetailModal">
          <div class="modal-box" style="width:min(720px,96vw)">
            <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:16px;">
              <h3 style="margin:0" id="diaryDetailDate">日记详情</h3>
              <button type="button" class="secondary" onclick="el('diaryDetailModal').classList.add('hidden')">关闭</button>
            </div>
            <div class="diary-detail-body">
              <div class="diary-detail-section">
                <div class="diary-detail-label">今日工作内容</div>
                <div class="diary-detail-content" id="diaryDetailToday"></div>
              </div>
              <div class="diary-detail-section">
                <div class="diary-detail-label">明日工作计划</div>
                <div class="diary-detail-content" id="diaryDetailTomorrow"></div>
              </div>
              <div class="diary-detail-section">
                <div class="diary-detail-label">思路与想法</div>
                <div class="diary-detail-content" id="diaryDetailThoughts"></div>
              </div>
            </div>
            <div class="toolbar" style="justify-content:flex-end;margin-top:16px;">
              <button type="button" class="secondary" id="diaryDetailEdit">编辑</button>
              <button type="button" class="danger" id="diaryDetailDelete">删除</button>
            </div>
          </div>
        </div>
      </div>
      <div class="task-panel hidden" id="forumPanel">
        <div class="forum-layout">
          <div class="forum-card">
            <div style="display:flex;justify-content:space-between;gap:12px;align-items:center;">
              <h3 style="margin:0;">全部历史话题</h3>
              <span class="history-meta" id="forumTopicCount">0 个话题</span>
            </div>
            <div class="toolbar" style="margin-top:12px;">
              <button type="button" class="secondary" id="forumToggleCreate">发起话题</button>
              <button type="button" class="secondary" id="forumRefreshButton">刷新</button>
            </div>
            <div class="forum-create-panel hidden" id="forumCreatePanel">
              <h3 style="margin-top:18px;">发起金点子话题</h3>
              <label>话题标题</label>
              <input id="forumTitle" placeholder="例如：如何让工作日记自动沉淀成项目经验？" />
              <label>话题内容</label>
              <textarea id="forumBody" placeholder="写下背景、想讨论的问题、希望大家从哪些角度给建议..."></textarea>
              <div class="toolbar">
                <button type="button" id="forumCreateButton">发布话题</button>
              </div>
              <div id="forumCreateStatus" class="status"></div>
              <h3 style="margin-top:24px;">智能体每日起题</h3>
              <label>输入信息</label>
              <textarea id="forumAiSeed" placeholder="输入今天的工作背景、灵感、项目机会或想让大家讨论的方向..."></textarea>
              <label>聊天内容</label>
              <textarea id="forumAiChat" placeholder="可粘贴群聊、会议纪要、用户反馈等内容..."></textarea>
              <label>传入文档</label>
              <input id="forumAiFiles" type="file" multiple accept=".txt,.md,.csv,.docx,.xlsx" />
              <div class="toolbar">
                <button type="button" class="warn" id="forumAiButton"><span class="icon" data-icon="sparkles"></span> 智能生成话题</button>
              </div>
              <div id="forumAiStatus" class="status"></div>
            </div>
            <div class="forum-topic-list" id="forumTopicList">
              <div class="forum-empty">暂无话题，先发起一个金点子吧。</div>
            </div>
          </div>
          <div class="forum-card">
            <div id="forumTopicDetail" style="margin-top:16px;">
              <div class="forum-empty">选择一个话题查看详情和讨论。</div>
            </div>
          </div>
        </div>
      </div>
      <div class="task-panel hidden" id="newsPanel">
        <div class="news-layout" id="newsLayout">
          <div class="news-card superadmin-only hidden" id="newsConfigCard">
            <h3>资讯收集配置</h3>
            <label>网页路径</label>
            <div id="newsSources"></div>
            <div class="toolbar">
              <button type="button" class="secondary" id="newsAddSource">新增网页</button>
            </div>
            <label>大模型自动搜索关键词</label>
            <textarea id="newsSearchQuery" placeholder="轨道交通 OR 城市轨道 OR 地铁 OR 智慧轨交"></textarea>
            <div class="row">
              <label><input id="newsAutoSearch" type="checkbox" style="width:auto" /> 允许大模型自动网络搜索</label>
              <label><input id="newsAutoPush" type="checkbox" style="width:auto" /> 每日自动更新推送</label>
            </div>
            <label>推送时间</label>
            <input id="newsPushTime" type="time" value="08:30" />
            <div class="toolbar">
              <button type="button" id="newsSaveConfig">保存配置</button>
              <button type="button" class="warn" id="newsGenerateNow"><span class="icon" data-icon="sparkles"></span> 立即生成今日资讯</button>
            </div>
            <div id="newsConfigStatus" class="status"></div>
            <div class="hint">自动搜索会通过平台配置的大模型 API 发起；若模型平台不支持联网工具，将主要基于配置网页内容生成。</div>
          </div>
          <div class="news-card">
            <div class="news-issue-head">
              <div>
                <div class="news-issue-title" id="newsTitle">暂无每日资讯</div>
                <div class="news-issue-meta" id="newsMeta">配置资讯源后生成今日简报。</div>
              </div>
              <button type="button" class="secondary" id="newsRefresh">刷新</button>
            </div>
            <div class="news-summary" id="newsSummary">暂无摘要。</div>
            <div class="news-list" id="newsItems"></div>
            <div class="news-keywords" id="newsKeywords"></div>
          </div>
        </div>
      </div>
      <div class="task-panel" id="weeklyPanel">
      <div class="guide" id="weeklyGuide">
        <div class="weekly-workspace">
          <div class="weekly-period-card">
            <div class="weekly-period-head">
              <div>
                <div class="weekly-period-title">周报时段</div>
                <div class="weekly-section-meta">默认上周一至上周五，可用日历快速调整。</div>
              </div>
              <div class="toolbar" style="margin-top:0">
                <span class="history-meta" id="weeklyPeriodText"></span>
                <button class="action-button" id="loadLatestWeekly" type="button">获取最新历史报告</button>
              </div>
            </div>
            <div class="weekly-period-grid">
              <div>
                <label>开始日期</label>
                <input id="weeklyStart" type="date" />
              </div>
              <div>
                <label>结束日期</label>
                <input id="weeklyEnd" type="date" />
              </div>
            </div>
          </div>
          <input id="weeklyPeriod" type="hidden" />
          <div class="diary-summarize-box" id="diarySummarizeBox">
            <span class="icon" data-icon="sparkles" style="color:var(--accent);font-size:18px;"></span>
            <div style="flex:1;min-width:200px;">
              <div style="font-weight:700;font-size:13px;margin-bottom:4px;">从工作日记智能总结</div>
              <div style="font-size:12px;color:var(--muted);">选择日期范围，AI 自动总结日记内容填充周报</div>
            </div>
            <div style="display:flex;gap:8px;align-items:center;flex-wrap:wrap;">
              <input id="diarySumStart" type="date" style="width:auto;min-width:130px;" />
              <span style="color:var(--muted);font-size:13px;">至</span>
              <input id="diarySumEnd" type="date" style="width:auto;min-width:130px;" />
              <button type="button" id="diarySummarizeBtn">智能总结</button>
            </div>
          </div>
          <div class="weekly-section-grid">
            <section class="weekly-section-card">
              <div class="weekly-section-head">
                <div>
                  <div class="weekly-section-title">一、本周工作总结</div>
                  <div class="weekly-section-meta"><span id="summaryCount">0</span> 项，点击卡片填写详情</div>
                </div>
                <button class="mini" id="addSummary" type="button">新增</button>
              </div>
              <div class="weekly-rows" id="summaryRows"></div>
            </section>
            <section class="weekly-section-card">
              <div class="weekly-section-head">
                <div>
                  <div class="weekly-section-title">二、重点工作跟进</div>
                  <div class="weekly-section-meta"><span id="followCount">0</span> 项，记录进展和困难</div>
                </div>
                <button class="mini" id="addFollow" type="button">新增</button>
              </div>
              <div class="weekly-rows" id="followRows"></div>
            </section>
            <section class="weekly-section-card">
              <div class="weekly-section-head">
                <div>
                  <div class="weekly-section-title">三、下周工作计划</div>
                  <div class="weekly-section-meta"><span id="nextCount">0</span> 项，安排后续计划</div>
                </div>
                <button class="mini" id="addNext" type="button">新增</button>
              </div>
              <div class="weekly-rows" id="nextRows"></div>
            </section>
          </div>
        </div>
      </div>
      </div>
      <div class="task-panel hidden" id="tripPanel">
      <div class="guide" id="tripGuide">
        <div class="toolbar" style="margin-top:0;margin-bottom:12px">
          <button class="action-button" id="loadLatestTrip" type="button">获取最新历史报告</button>
          <span class="history-meta">使用最近一份出差报告自动填充当前模板。</span>
        </div>
        <div class="trip-store">
          <input id="tripReporter" value="周颖超" />
          <input id="tripDepartment" value="场景研究院" />
          <input id="tripLocation" placeholder="青岛" />
          <input id="tripStart" placeholder="2026-05-06" />
          <input id="tripEnd" placeholder="2026-05-08" />
          <textarea id="tripPurpose"></textarea>
          <textarea id="tripItinerary"></textarea>
          <textarea id="tripDetails"></textarea>
          <textarea id="tripIssues"></textarea>
          <textarea id="tripSuggestions"></textarea>
        </div>
        <div class="edit-card-grid" id="tripCards"></div>
      </div>
      </div>
      <div class="task-panel hidden" id="configPanel">
      <div class="guide">
        <div class="superadmin-only hidden">
        <h2 style="margin:0 0 14px">大模型 API 配置</h2>
        <div class="config-grid">
          <div>
            <label>NewAPI 地址</label>
            <input id="configApiUrl" placeholder="http://host:port" />
          </div>
          <div>
            <label>模型名称</label>
            <select id="configModelSelect"></select>
          </div>
        </div>
        <label>手动模型名称</label>
        <input id="configModel" placeholder="MiniMax-M2.7" />
        <label>API Key</label>
        <input id="configApiKey" type="password" placeholder="留空则保持原 Key 不变" />
        <div class="hint" id="configKeyHint"></div>
        <label>默认优化提示词</label>
        <textarea id="configPrompt"></textarea>
        <div class="toolbar">
          <button class="secondary" id="loadModels" type="button">获取模型列表</button>
          <button class="warn" id="testModel" type="button">测试 API Key</button>
          <button id="saveConfig" type="button">保存系统配置</button>
        </div>
        <div id="configTestStatus" class="status"></div>
        </div>
        <div class="superadmin-only hidden">
        <h2 style="margin-top:18px">邮件服务器（全局）</h2>
        <div class="config-grid">
          <div>
            <label>SMTP 服务器</label>
            <input id="configSmtpHost" placeholder="smtp.263.net" />
          </div>
          <div>
            <label>SMTP 端口</label>
            <input id="configSmtpPort" type="number" placeholder="465" />
          </div>
        </div>
        <div class="row">
          <label><input id="configSmtpTls" type="checkbox" style="width:auto" /> 使用 TLS</label>
          <label><input id="configSmtpSsl" type="checkbox" style="width:auto" /> 使用 SSL</label>
        </div>
        <div class="config-grid">
          <div>
            <label>IMAP 服务器</label>
            <input id="configImapHost" placeholder="imap.263.net" />
          </div>
          <div>
            <label>IMAP 端口</label>
            <input id="configImapPort" type="number" placeholder="993" />
          </div>
        </div>
        <div class="row">
          <label><input id="configImapSsl" type="checkbox" style="width:auto" /> 使用 SSL 连接 IMAP</label>
        </div>
        <div class="toolbar">
          <button id="saveServerConfig" type="button">保存邮件服务器配置</button>
        </div>
        <div id="configServerStatus" class="status"></div>
        </div>
      </div>
      </div>
      <div class="task-panel hidden" id="userManagePanel">
      <div class="guide">
        <h2 style="margin:0 0 14px">用户管理</h2>
        <div class="config-grid">
          <div>
            <label>用户名</label>
            <input id="newUserName" placeholder="zhangsan" />
          </div>
          <div>
            <label>显示名称</label>
            <input id="newDisplayName" placeholder="张三" />
          </div>
        </div>
        <div class="config-grid" id="newUserRoleBox" style="display:none;">
          <div>
            <label>角色权限</label>
            <select id="newUserRole">
              <option value="member">普通成员</option>
              <option value="admin">管理员</option>
              <option value="superadmin">超级管理员</option>
            </select>
          </div>
        </div>
        <label>初始密码</label>
        <input id="newUserPassword" type="password" placeholder="至少 4 位" />
        <div class="toolbar">
          <button id="addUser" type="button">新增用户</button>
        </div>
        <div id="userManageStatus" class="status"></div>
        <div class="user-list" id="userList"><div class="upload-item">正在加载用户列表...</div></div>
      </div>
      </div>
      <div class="task-panel hidden" id="mailConfigPanel">
      <div class="guide">
        <h2>我的邮箱账户</h2>
        <div class="upload-item" style="margin-bottom:14px;">
          发送邮件需要 SMTP；读取收件箱需要 IMAP。263 邮箱常用配置是 SMTP 465/SSL、IMAP 993/SSL。用户名通常填写完整邮箱地址；授权码不是网页登录密码，需要在邮箱后台单独生成。
        </div>
        <div class="config-grid">
          <div>
            <label>本人邮箱</label>
            <input id="mailUserEmail" placeholder="your-email@example.com" />
          </div>
          <div>
            <label>SMTP 发件地址</label>
            <input id="mailSmtpFrom" placeholder="默认使用本人邮箱" />
          </div>
        </div>
        <div class="config-grid">
          <div>
            <label>SMTP 用户名</label>
            <input id="mailSmtpUser" placeholder="通常为邮箱账号" />
            <div class="hint">留空保存时会自动使用“本人邮箱”。</div>
          </div>
          <div>
            <label>SMTP 授权码/密码</label>
            <input id="mailSmtpPassword" type="password" placeholder="留空则保持原密码不变" />
            <div class="hint" id="mailPasswordHint"></div>
          </div>
        </div>
        <div class="config-grid" style="margin-top:8px;">
          <div style="background:var(--surface);border:1px solid var(--line);border-radius:var(--radius-sm);padding:10px 12px;">
            <div style="font-size:12px;color:var(--muted);margin-bottom:2px;">SMTP 服务器（全局）</div>
            <div id="mailSmtpHostDisplay" style="font-weight:700;font-size:14px;">smtp.263.net</div>
          </div>
          <div style="background:var(--surface);border:1px solid var(--line);border-radius:var(--radius-sm);padding:10px 12px;">
            <div style="font-size:12px;color:var(--muted);margin-bottom:2px;">SMTP 端口 / SSL</div>
            <div id="mailSmtpPortDisplay" style="font-weight:700;font-size:14px;">465 / SSL</div>
          </div>
        </div>
        <h2 style="margin-top:18px">收件箱 IMAP</h2>
        <div class="config-grid">
          <div>
            <label>IMAP 用户名</label>
            <input id="mailImapUser" placeholder="默认使用 SMTP 用户名" />
            <div class="hint">只收邮件时需要；留空保存时会自动使用 SMTP 用户名。</div>
          </div>
          <div>
            <label>IMAP 授权码/密码</label>
            <input id="mailImapPassword" type="password" placeholder="留空则保持原密码不变" />
            <div class="hint" id="mailImapPasswordHint"></div>
          </div>
        </div>
        <div class="config-grid" style="margin-top:8px;">
          <div style="background:var(--surface);border:1px solid var(--line);border-radius:var(--radius-sm);padding:10px 12px;">
            <div style="font-size:12px;color:var(--muted);margin-bottom:2px;">IMAP 服务器（全局）</div>
            <div id="mailImapHostDisplay" style="font-weight:700;font-size:14px;">imap.263.net</div>
          </div>
          <div style="background:var(--surface);border:1px solid var(--line);border-radius:var(--radius-sm);padding:10px 12px;">
            <div style="font-size:12px;color:var(--muted);margin-bottom:2px;">IMAP 端口 / SSL</div>
            <div id="mailImapPortDisplay" style="font-weight:700;font-size:14px;">993 / SSL</div>
          </div>
        </div>
        <h2 style="margin-top:18px">周报邮件</h2>
        <div class="config-grid">
          <div>
            <label>周报收件人</label>
            <input id="mailWeeklyTo" placeholder="leader@example.com" />
          </div>
          <div>
            <label>周报抄送</label>
            <input id="mailWeeklyCc" placeholder="可选，多个用分号分隔" />
          </div>
        </div>
        <h2 style="margin-top:18px">出差报告邮件</h2>
        <div class="config-grid">
          <div>
            <label>出差报告收件人</label>
            <input id="mailTripTo" placeholder="leader@example.com" />
          </div>
          <div>
            <label>出差报告抄送</label>
            <input id="mailTripCc" placeholder="可选，多个用分号分隔" />
          </div>
        </div>
        <label>邮件签名模板</label>
        <textarea id="mailEmailSignature" placeholder="留空则使用默认签名"></textarea>
        <div class="toolbar" style="margin-top:16px">
          <button id="saveMailConfig" type="button">保存我的邮件配置</button>
          <button class="warn" id="testMailConfig" type="button">测试我的邮箱配置</button>
        </div>
        <div id="mailConfigStatus" class="status"></div>
      </div>
      </div>
      <div class="task-panel hidden" id="skillsPanel">
      <div class="guide">
        <div class="skill-hero">
          <div>
            <div class="mail-section-title">系统 Skill 管理</div>
            <div class="mail-section-meta">查看当前已安装的 Skill、能力说明、调用参数和调用示例。此页面仅超级管理员可见。</div>
          </div>
          <div class="toolbar" style="margin-top:0">
            <button class="secondary" id="openAgentOrchestration" type="button"><span class="icon" data-icon="bot"></span> 犇犇编排逻辑</button>
            <button class="secondary" id="openAgentConfig" type="button"><span class="icon" data-icon="settings"></span> 编辑犇犇配置</button>
            <button class="secondary" id="openSkillDocs" type="button"><span class="icon" data-icon="file-text"></span> 查看完整文档</button>
            <button class="warn" id="downloadSkillDocs" type="button"><span class="icon" data-icon="archive"></span> 下载文档</button>
          </div>
        </div>
        <div class="skill-protocol">
          <div class="mail-section-title">大模型调用协议</div>
          <div class="mail-section-meta">当大模型需要操作软件时，返回下面这种 JSON；普通问答直接自然语言回复即可。</div>
          <pre>{"reply":"给用户看的说明","skill_call":{"name":"skill.name","arguments":{}}}</pre>
        </div>
        <div class="skill-orchestration hidden" id="agentOrchestrationPanel">
          <div class="mail-section-title">犇犇编排逻辑</div>
          <div class="mail-section-meta">当前犇犇各模块角色定义、系统提示词与工作流编排。</div>
          <div id="agentOrchestrationContent"></div>
        </div>
        <div class="skill-installed-head">
          <div>
            <div class="mail-section-title">已安装 Skill 模块</div>
            <div class="mail-section-meta">点击下面的模块卡片查看对应 Skill 能力。</div>
          </div>
          <div class="skill-total-count" id="skillTotalCount"></div>
        </div>
        <div class="skill-module-summary" id="skillModuleSummary">
          <div class="skill-module-card active" data-module="周报">
            <div class="skill-module-name">周报 Skill</div>
            <div class="skill-module-count">已安装 4 个能力，点击查看</div>
          </div>
          <div class="skill-module-card" data-module="出差报告">
            <div class="skill-module-name">出差报告 Skill</div>
            <div class="skill-module-count">已安装能力，点击查看</div>
          </div>
          <div class="skill-module-card" data-module="工作日记">
            <div class="skill-module-name">日记 Skill</div>
            <div class="skill-module-count">已安装能力，点击查看</div>
          </div>
          <div class="skill-module-card" data-module="金点子论坛">
            <div class="skill-module-name">金点子论坛 Skill</div>
            <div class="skill-module-count">已安装能力，点击查看</div>
          </div>
          <div class="skill-module-card" data-module="邮件">
            <div class="skill-module-name">邮件 Skill</div>
            <div class="skill-module-count">已安装能力，点击查看</div>
          </div>
          <div class="skill-module-card" data-module="资讯">
            <div class="skill-module-name">资讯 Skill</div>
            <div class="skill-module-count">已安装能力，点击查看</div>
          </div>
        </div>
        <div class="skill-filter-row">
          <input id="skillSearch" placeholder="搜索 Skill 名称、模块或说明..." />
          <select id="skillModuleFilter">
            <option value="all">全部模块</option>
          </select>
        </div>
        <div class="skill-list" id="skillList">
          <div class="skill-card" data-skill-name="weekly.compose">
            <div class="skill-card-head">
              <div>
                <div class="skill-name">weekly.compose</div>
                <div class="skill-title">周报 Skill · 编写/设计周报草稿</div>
              </div>
              <span class="skill-badge">查询/预览</span>
            </div>
            <div class="skill-desc">调用配置的大模型 API，将原始工作内容整理成周报三段式结构：本周工作总结、重点工作跟进、下周工作计划。</div>
            <div class="mailbox-meta">调用示例</div>
            <pre class="skill-call-example">{"reply":"我来根据这些内容生成周报草稿。","skill_call":{"name":"weekly.compose","arguments":{"raw_work":"本周原始工作内容","next_plan":"下周计划","style":"体现工作量多，编号清晰，简洁明了"}}}</pre>
          </div>
          <div class="skill-card" data-skill-name="weekly.preview">
            <div class="skill-card-head">
              <div>
                <div class="skill-name">weekly.preview</div>
                <div class="skill-title">周报 Skill · 生成周报预览</div>
              </div>
              <span class="skill-badge warn">生成文件/预览图</span>
            </div>
            <div class="skill-desc">把结构化周报内容写入 Excel 模板，生成周报预览图片和邮件草稿，等待用户确认。</div>
            <div class="mailbox-meta">调用示例</div>
            <pre class="skill-call-example">{"reply":"我先生成周报预览，请你确认。","skill_call":{"name":"weekly.preview","arguments":{"period":"2026.05.11-2026.05.15","weekly_summary":[],"weekly_follow":[],"weekly_next":[]}}}</pre>
          </div>
          <div class="skill-card" data-skill-name="weekly.send_confirmed">
            <div class="skill-card-head">
              <div>
                <div class="skill-name">weekly.send_confirmed</div>
                <div class="skill-title">周报 Skill · 确认后发送周报邮件</div>
              </div>
              <span class="skill-badge warn">发送邮件</span>
            </div>
            <div class="skill-desc">只有用户确认周报预览无误后，才使用 weekly.preview 生成的附件发送邮件。</div>
            <div class="mailbox-meta">调用示例</div>
            <pre class="skill-call-example">{"reply":"确认预览无误后发送。","skill_call":{"name":"weekly.send_confirmed","arguments":{"attachment":"周颖超工作周报2026.05.11-2026.05.15.xlsx"}}}</pre>
          </div>
          <div class="skill-card" data-skill-name="weekly.prefill">
            <div class="skill-card-head">
              <div>
                <div class="skill-name">weekly.prefill</div>
                <div class="skill-title">周报 Skill · 获取最新周报预填</div>
              </div>
              <span class="skill-badge">查询/预览</span>
            </div>
            <div class="skill-desc">读取最新历史周报，将上次“下周计划”迁移到本次“本周工作内容”。</div>
            <div class="mailbox-meta">调用示例</div>
            <pre class="skill-call-example">{"reply":"我来获取最新历史周报。","skill_call":{"name":"weekly.prefill","arguments":{}}}</pre>
          </div>
        </div>
      </div>
      </div>
      <div class="task-panel hidden" id="mailAssistantPanel">
      <div class="guide">
        <div class="mail-assistant-grid">
          <div class="mailbox-card">
            <div class="mail-section-head">
              <div>
                <div class="mail-section-title">收件箱</div>
                <div class="mail-section-meta">查看当前账号最近邮件，支持正文和附件信息预览。</div>
              </div>
              <div class="toolbar" style="margin-top:0">
                <select id="mailboxLimit" style="width:110px">
                  <option value="10">最近 10 封</option>
                  <option value="20" selected>最近 20 封</option>
                  <option value="50">最近 50 封</option>
                </select>
                <button class="secondary" id="refreshMailbox" type="button"><span class="icon" data-icon="refresh-cw"></span> 刷新</button>
              </div>
            </div>
            <div class="mailbox-list" id="mailboxList"></div>
          </div>
          <div class="mail-detail-card">
            <div class="mail-section-head">
              <div>
                <div class="mail-section-title">邮件详情</div>
                <div class="mail-section-meta">按邮件原始版式预览正文，附件在头部集中展示。</div>
              </div>
            </div>
            <div class="mail-detail" id="mailDetail">暂无选中邮件。</div>
          </div>
        </div>
        <div class="mail-compose-card">
          <div class="mail-section-head">
            <div>
              <div class="mail-section-title">发送邮件</div>
              <div class="mail-section-meta">使用当前账号 SMTP 发送普通邮件，可添加附件并实时预览正文。</div>
            </div>
          </div>
          <div class="mail-compose-layout">
            <div class="mail-compose-form">
              <div class="row">
                <div>
                  <label>收件人</label>
                  <input id="assistantMailTo" placeholder="recipient@example.com" />
                </div>
                <div>
                  <label>抄送</label>
                  <input id="assistantMailCc" placeholder="可选，多个邮箱用分号分隔" />
                </div>
              </div>
              <label>主题</label>
              <input id="assistantMailSubject" placeholder="请输入邮件主题" />
              <label>正文</label>
              <textarea id="assistantMailBody" placeholder="请输入邮件正文"></textarea>
              <label>附件</label>
              <input id="assistantMailFiles" type="file" multiple />
              <div class="mail-file-list" id="assistantMailFileList"></div>
              <div class="toolbar">
                <button class="secondary" id="clearAssistantMail" type="button">清空</button>
                <button class="warn" id="sendAssistantMail" type="button"><span class="icon" data-icon="send"></span> 发送邮件</button>
              </div>
              <div id="assistantMailStatus" class="status"></div>
            </div>
            <div class="mail-compose-preview">
              <label>正文预览</label>
              <div class="rich-preview" id="assistantMailPreview">暂无正文内容</div>
            </div>
          </div>
        </div>
      </div>
      </div>
      <div class="toolbar hidden" id="generateToolbar">
        <button id="generate">按标准模板生成文件</button>
      </div>
      <div class="task-panel hidden" id="mailPanel">
      <div class="mail-layout">
      <div class="side">
        <h2>报告文件</h2>
        <label>类型</label>
        <select id="mailKind">
          <option value="weekly">周报</option>
          <option value="trip">出差报告</option>
        </select>
        <div id="reports"></div>
      </div>
      <div>
      <h2>邮件内容</h2>
      <div class="row">
        <div>
          <label>收件人</label>
          <input id="to" placeholder="leader@example.com" />
        </div>
        <div>
          <label>抄送</label>
          <input id="cc" placeholder="可选，多个邮箱用分号分隔" />
        </div>
      </div>
      <label>主题</label>
      <input id="subject" />
      <label>正文</label>
      <textarea id="body"></textarea>
      <label>发送邮件正文预览</label>
      <div class="rich-preview" id="bodyPreview"></div>
      <label>附件</label>
      <input id="attachment" readonly />
      <div class="toolbar" style="margin-top:8px">
        <a id="downloadLink" class="download-link hidden" href="#" target="_blank">下载查看当前附件</a>
      </div>
      <label>模板内容预览</label>
      <div class="rich-preview" id="preview"></div>
      <div class="toolbar">
        <button id="refresh">重新生成正文</button>
        <button class="secondary" id="copy">复制正文</button>
        <button class="warn" id="send">发送/生成草稿</button>
      </div>
      <div id="sendReview" class="send-review hidden"></div>
      </div>
      </div>
      </div>
      <div class="task-panel hidden" id="uploadPanel">
      <div class="guide">
        <div class="row">
          <div>
            <label>上传类型</label>
            <select id="uploadKind">
              <option value="weekly">历史周报（.xlsx / .xls）</option>
              <option value="trip">历史出差报告（.docx / .md）</option>
            </select>
          </div>
          <div>
            <label>选择文件</label>
            <input id="uploadFiles" type="file" multiple accept=".xlsx,.xls,.docx,.md" />
          </div>
        </div>
        <div class="toolbar">
          <button id="uploadButton" type="button">上传到历史报告库</button>
        </div>
        <div class="hint">上传后的文件会保存到当前用户的个人历史报告空间，并自动出现在“发送报告邮件”的报告列表中。</div>
        <div class="upload-list" id="uploadList"></div>
        <div class="history-tools">
          <h2 style="margin:0">历史报告管理</h2>
          <select id="historyKind">
            <option value="all">全部</option>
            <option value="weekly">周报</option>
            <option value="trip">出差报告</option>
          </select>
        </div>
        <div class="history-list" id="historyList"></div>
      </div>
      </div>
      <div id="status" class="status"></div>
    </section>
  </main>
  <div class="modal hidden" id="passModal" role="dialog" aria-modal="true">
    <div class="modal-box" style="width:min(420px,96vw)">
      <div class="modal-head">
        <div class="modal-title">修改密码</div>
        <button class="mini secondary" type="button" onclick="el('passModal').classList.add('hidden')">取消</button>
      </div>
      <div class="modal-fields">
        <label>原密码</label>
        <input id="passOld" type="password" placeholder="输入当前密码" />
        <label>新密码</label>
        <input id="passNew" type="password" placeholder="至少 4 位" />
        <label>确认新密码</label>
        <input id="passConfirm" type="password" placeholder="再次输入新密码" />
        <div id="passStatus" class="status"></div>
      </div>
      <div class="modal-actions">
        <button class="secondary" type="button" onclick="el('passModal').classList.add('hidden')">取消</button>
        <button type="button" id="passSave">确认修改</button>
      </div>
    </div>
  </div>
  <div class="modal hidden" id="editModal" role="dialog" aria-modal="true">
    <div class="modal-box">
      <div class="modal-head">
        <div class="modal-title" id="modalTitle">编辑内容</div>
        <button class="mini secondary" type="button" id="modalClose">退出并保存</button>
      </div>
      <div class="modal-fields" id="modalFields"></div>
      <div class="modal-actions">
        <button class="secondary" type="button" id="modalCancel">退出并保存</button>
        <button type="button" id="modalSave">保存</button>
      </div>
    </div>
  </div>
  <div class="modal hidden" id="skillTestModal" role="dialog" aria-modal="true">
    <div class="modal-box skill-test-box">
      <div class="modal-head">
        <div>
          <div class="modal-title" id="skillTestTitle">Skill 测试</div>
          <div class="history-meta" id="skillTestMeta">使用平台配置的大模型 API 和当前登录用户空间执行测试。</div>
        </div>
        <button class="mini secondary" type="button" id="skillTestClose">关闭</button>
      </div>
      <div class="skill-test-grid">
        <div class="skill-test-panel">
          <label>调用参数 JSON</label>
          <textarea id="skillTestArgs" spellcheck="false"></textarea>
          <label>自然语言测试要求（可选）</label>
          <textarea id="skillTestInstruction" placeholder="例如：把下面的原始工作内容梳理成周报草稿，要求体现工作量多、编号清晰、简洁明了。"></textarea>
          <label class="skill-confirm-line hidden" id="skillConfirmWrap">
            <input id="skillConfirmUnsafe" type="checkbox" />
            我确认执行该 Skill 测试，可能会生成文件、写入数据或发送内容
          </label>
          <div class="toolbar" style="margin-top:0">
            <button type="button" id="skillRunTest"><span class="icon" data-icon="play"></span> 运行测试</button>
          </div>
          <div id="skillTestStatus" class="status"></div>
        </div>
        <div class="skill-test-panel">
          <label>测试结果</label>
          <pre class="skill-test-result" id="skillTestResult">点击“运行测试”后，这里会显示 Skill 返回结果。</pre>
          <div class="skill-test-links" id="skillTestLinks"></div>
        </div>
      </div>
    </div>
  </div>
  <div class="modal hidden" id="agentConfigModal" role="dialog" aria-modal="true">
    <div class="modal-box" style="width:min(960px,96vw);max-height:92vh;overflow:auto;">
      <div class="modal-head">
        <div>
          <div class="modal-title">犇犇配置</div>
          <div class="history-meta">编辑系统提示词和工作流编排，保存后立即生效。</div>
        </div>
        <button class="mini secondary" type="button" id="agentConfigClose">关闭</button>
      </div>
      <div class="config-tabs" style="display:flex;gap:8px;margin-bottom:12px;">
        <button type="button" class="secondary agent-config-tab active" data-tab="prompts">系统提示词</button>
        <button type="button" class="secondary agent-config-tab" data-tab="workflows">工作流</button>
      </div>
      <div id="agentConfigPrompts"></div>
      <div id="agentConfigWorkflows" class="hidden"></div>
      <div class="toolbar" style="margin-top:12px;">
        <button type="button" id="agentConfigSave" class="primary"><span class="icon" data-icon="save"></span> 保存配置</button>
        <span id="agentConfigStatus" class="status" style="margin-left:8px;"></span>
      </div>
    </div>
  </div>
  <script>
    let agentKind = null;
    let agentMessages = [];
    let defaultAssistantPrompt = `请帮我优化下面的工作内容，要求：
1、拆分成1、2、3、4这样的编号要点，每点单独换行。
2、语言简洁明了，体现实际工作量和推进成果。
3、修正错别字、病句和不通顺表达。
4、不要编造不存在的事项，不要写空话套话。`;
    let state = { reports: [], selected: null, task: 'weekly', subTab: 'edit', user: null, weeklyPrefilled: false, tripPrefilled: false, modalSave: null, restoringDraft: false, assistantMailFiles: [], forumSelected: null, forumCommentPage: 1, currentSkill: null, agentStage: 0 };
    const FORM_DRAFT_PREFIX = 'personalWorkSite.formDraft.v2';
    const el = id => document.getElementById(id);
    const lucideIcons = {
      sparkles: '<path d="M9.94 14.5 8.5 18.06 7.06 14.5 3.5 13.06 7.06 11.62 8.5 8.06l1.44 3.56 3.56 1.44-3.56 1.44Z"/><path d="M18 8.5 17.2 10.7 15 11.5l2.2.8L18 14.5l.8-2.2 2.2-.8-2.2-.8L18 8.5Z"/><path d="M15 2l-.9 2.1L12 5l2.1.9L15 8l.9-2.1L18 5l-2.1-.9L15 2Z"/>',
      search: '<circle cx="11" cy="11" r="7"/><path d="m20 20-3.5-3.5"/>',
      'layout-dashboard': '<rect width="7" height="9" x="3" y="3" rx="1"/><rect width="7" height="5" x="14" y="3" rx="1"/><rect width="7" height="9" x="14" y="12" rx="1"/><rect width="7" height="5" x="3" y="16" rx="1"/>',
      'file-spreadsheet': '<path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8Z"/><path d="M14 2v6h6"/><path d="M8 13h8"/><path d="M8 17h8"/><path d="M10 9v8"/><path d="M14 9v8"/>',
      'briefcase-business': '<path d="M12 12h.01"/><path d="M16 6V4a2 2 0 0 0-2-2h-4a2 2 0 0 0-2 2v2"/><path d="M22 13a18.15 18.15 0 0 1-20 0"/><rect width="20" height="14" x="2" y="6" rx="2"/>',
      'mail-check': '<path d="m22 7-8.97 5.7a1.94 1.94 0 0 1-2.06 0L2 7"/><path d="M22 12.5V6a2 2 0 0 0-2-2H4a2 2 0 0 0-2 2v12c0 1.1.9 2 2 2h8"/><path d="m16 19 2 2 4-4"/>',
      inbox: '<polyline points="22 12 16 12 14 15 10 15 8 12 2 12"/><path d="M5.45 5.11 2 12v6a2 2 0 0 0 2 2h16a2 2 0 0 0 2-2v-6l-3.45-6.89A2 2 0 0 0 16.76 4H7.24a2 2 0 0 0-1.79 1.11Z"/>',
      paperclip: '<path d="m21.44 11.05-9.19 9.19a6 6 0 0 1-8.49-8.49l9.19-9.19a4 4 0 0 1 5.66 5.66l-9.2 9.19a2 2 0 0 1-2.83-2.83l8.49-8.48"/>',
      'refresh-cw': '<path d="M3 12a9 9 0 0 1 15.1-6.6L21 8"/><path d="M21 3v5h-5"/><path d="M21 12a9 9 0 0 1-15.1 6.6L3 16"/><path d="M3 21v-5h5"/>',
      settings: '<path d="M12.22 2h-.44a2 2 0 0 0-2 2v.18a2 2 0 0 1-1 1.73l-.43.25a2 2 0 0 1-2 0l-.15-.08a2 2 0 0 0-2.73.73l-.22.38a2 2 0 0 0 .73 2.73l.15.1a2 2 0 0 1 1 1.72v.51a2 2 0 0 1-1 1.74l-.15.09a2 2 0 0 0-.73 2.73l.22.38a2 2 0 0 0 2.73.73l.15-.08a2 2 0 0 1 2 0l.43.25a2 2 0 0 1 1 1.73V20a2 2 0 0 0 2 2h.44a2 2 0 0 0 2-2v-.18a2 2 0 0 1 1-1.73l.43-.25a2 2 0 0 1 2 0l.15.08a2 2 0 0 0 2.73-.73l.22-.38a2 2 0 0 0-.73-2.73l-.15-.09a2 2 0 0 1-1-1.74v-.51a2 2 0 0 1 1-1.72l.15-.1a2 2 0 0 0 .73-2.73l-.22-.38a2 2 0 0 0-2.73-.73l-.15.08a2 2 0 0 1-2 0l-.43-.25a2 2 0 0 1-1-1.73V4a2 2 0 0 0-2-2Z"/><circle cx="12" cy="12" r="3"/>',
      send: '<path d="m22 2-7 20-4-9-9-4Z"/><path d="M22 2 11 13"/>',
      'file-plus-2': '<path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8Z"/><path d="M14 2v6h6"/><path d="M12 18v-6"/><path d="M9 15h6"/>',
      bot: '<path d="M12 8V4H8"/><rect width="16" height="12" x="4" y="8" rx="2"/><path d="M2 14h2"/><path d="M20 14h2"/><path d="M15 13v2"/><path d="M9 13v2"/>',
      'bell-ring': '<path d="M10.27 21a2 2 0 0 0 3.46 0"/><path d="M4 8a8 8 0 0 1 16 0c0 7 3 7 3 9H1c0-2 3-2 3-9"/><path d="M18.75 3.2A10 10 0 0 1 22 8"/><path d="M1.99 8a10 10 0 0 1 3.26-4.8"/>',
      'file-text': '<path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8Z"/><path d="M14 2v6h6"/><path d="M16 13H8"/><path d="M16 17H8"/><path d="M10 9H8"/>',
      'messages-square': '<path d="M14 9a2 2 0 0 1-2 2H6l-4 4V5a2 2 0 0 1 2-2h8a2 2 0 0 1 2 2Z"/><path d="M18 9h2a2 2 0 0 1 2 2v10l-4-4h-6a2 2 0 0 1-2-2v-1"/>',
      puzzle: '<path d="M19.4 13.5a1.8 1.8 0 0 0 0-3 1.8 1.8 0 0 0-2.4 1.7V9a2 2 0 0 0-2-2h-3.2a1.8 1.8 0 0 0-3.4-1.1A1.8 1.8 0 0 0 10.1 8H7a2 2 0 0 0-2 2v3.1a1.8 1.8 0 0 1 0 3.8V20a2 2 0 0 0 2 2h3.1a1.8 1.8 0 0 1 3.8 0H17a2 2 0 0 0 2-2v-3.1a1.8 1.8 0 0 0 .4-3.4Z"/>',
      'thumbs-up': '<path d="M7 10v12"/><path d="M15 5.88 14 10h5.83a2 2 0 0 1 1.92 2.56l-2.33 8A2 2 0 0 1 17.5 22H4a2 2 0 0 1-2-2v-8a2 2 0 0 1 2-2h3.28a2 2 0 0 0 1.7-.94L13 2a2.3 2.3 0 0 1 2 3.88Z"/>',
      newspaper: '<path d="M4 22h14a2 2 0 0 0 2-2V4H6a2 2 0 0 0-2 2v16Z"/><path d="M18 14h-8"/><path d="M15 18h-5"/><path d="M10 6h6v4h-6z"/><path d="M4 8H2v12a2 2 0 0 0 2 2"/>',
      mail: '<rect width="20" height="16" x="2" y="4" rx="2"/><path d="m22 7-8.97 5.7a1.94 1.94 0 0 1-2.06 0L2 7"/>',
      'folder-clock': '<path d="M10 20H4a2 2 0 0 1-2-2V6a2 2 0 0 1 2-2h5l2 3h9a2 2 0 0 1 2 2v1.5"/><circle cx="16" cy="16" r="5"/><path d="M16 13v3l2 1"/>',
      archive: '<rect width="20" height="5" x="2" y="3" rx="1"/><path d="M4 8v11a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8"/><path d="M10 12h4"/>',
      'calendar-check': '<path d="M8 2v4"/><path d="M16 2v4"/><rect width="18" height="18" x="3" y="4" rx="2"/><path d="M3 10h18"/><path d="m9 16 2 2 4-4"/>',
      'wand-sparkles': '<path d="m21.64 3.64-1.28-1.28a1.21 1.21 0 0 0-1.72 0L2.36 18.64a1.21 1.21 0 0 0 0 1.72l1.28 1.28a1.21 1.21 0 0 0 1.72 0L21.64 5.36a1.21 1.21 0 0 0 0-1.72Z"/><path d="m14 7 3 3"/><path d="M5 6v4"/><path d="M19 14v4"/><path d="M10 2v2"/><path d="M7 8H3"/><path d="M21 16h-4"/><path d="M11 3H9"/>',
      database: '<ellipse cx="12" cy="5" rx="9" ry="3"/><path d="M3 5v14c0 1.66 4.03 3 9 3s9-1.34 9-3V5"/><path d="M3 12c0 1.66 4.03 3 9 3s9-1.34 9-3"/>',
      'clipboard-pen': '<rect width="8" height="4" x="8" y="2" rx="1"/><path d="M16 4h2a2 2 0 0 1 2 2v9"/><path d="M8 4H6a2 2 0 0 0-2 2v14a2 2 0 0 0 2 2h7"/><path d="M17.5 22 22 17.5 20.5 16 16 20.5V22Z"/>',
      'book-open': '<path d="M2 3h6a4 4 0 0 1 4 4v14a3 3 0 0 0-3-3H2z"/><path d="M22 3h-6a4 4 0 0 0-4 4v14a3 3 0 0 1 3-3h7z"/>',
      'pen-line': '<path d="M12 20h9"/><path d="M16.5 3.5a2.12 2.12 0 0 1 3 3L7 19l-4 1 1-4Z"/>',
      users: '<path d="M16 21v-2a4 4 0 0 0-4-4H6a4 4 0 0 0-4 4v2"/><circle cx="9" cy="7" r="4"/><path d="M22 21v-2a4 4 0 0 0-3-3.87"/><path d="M16 3.13a4 4 0 0 1 0 7.75"/>',
      'library': '<path d="m16 6 4 14"/><path d="M12 6v14"/><path d="M8 8v12"/><path d="M4 4v16"/>',
      'chevron-right': '<path d="m9 18 6-6-6-6"/>',
      play: '<polygon points="6 3 20 12 6 21 6 3"/>'
    };

    function renderIcons(root = document) {
      root.querySelectorAll('[data-icon]').forEach(node => {
        const name = node.dataset.icon;
        if (!lucideIcons[name] || node.dataset.iconRendered === 'true') return;
        node.innerHTML = `<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-linecap="round" stroke-linejoin="round" aria-hidden="true">${lucideIcons[name]}</svg>`;
        node.dataset.iconRendered = 'true';
      });
    }

    async function api(path, options) {
      const res = await fetch(path, options);
      const data = await res.json();
      if (!res.ok) throw new Error(data.error || '请求失败');
      return data;
    }
    async function apiPost(path, body) {
      const res = await fetch(path, {
        method: 'POST',
        headers: { 'Content-Type': 'application/json' },
        body: JSON.stringify(body || {})
      });
      const data = await res.json();
      if (!res.ok) throw new Error(data.error || '请求失败');
      return data;
    }

    function applyUser(user) {
      state.user = user;
      const authed = !!user;
      el('authPanel').classList.toggle('hidden', authed);
      el('appMain').classList.toggle('hidden', !authed);
      el('userbar').style.display = authed ? 'flex' : 'none';
      if (el('agentFloat')) el('agentFloat').classList.toggle('hidden', !authed);
      if (authed) {
        const roleText = user.role === 'superadmin' ? '超级管理员' : user.role === 'admin' ? '管理员' : '成员';
        el('userInfo').textContent = `${user.name || user.username} · ${roleText}`;
        el('userAvatar').src = user.avatar_url || '/assets/ai-assistant-avatar.png';
      }
      document.querySelectorAll('.admin-only').forEach(node => {
        node.classList.toggle('hidden', !user?.is_admin);
      });
      document.querySelectorAll('.superadmin-only').forEach(node => {
        node.classList.toggle('hidden', !user?.is_superadmin);
      });
      el('newsLayout')?.classList.toggle('reader', !user?.is_superadmin);
    }

    function readFileAsBase64(file) {
      return new Promise((resolve, reject) => {
        const reader = new FileReader();
        reader.onload = () => resolve(String(reader.result || '').split(',')[1] || '');
        reader.onerror = () => reject(reader.error || new Error('读取文件失败'));
        reader.readAsDataURL(file);
      });
    }

    function openProfileModal() {
      const user = state.user || {};
      el('profileName').value = user.name || user.username || '';
      el('profileBio').value = user.bio || '';
      el('profileHobbies').value = user.hobbies || '';
      el('profileAvatarPreview').src = user.avatar_url || '/assets/ai-assistant-avatar.png';
      el('profileAvatarFile').value = '';
      el('profileStatus').textContent = '';
      el('profileStatus').className = 'status';
      el('profileModal').classList.remove('hidden');
    }

    function closeProfileModal() {
      el('profileModal').classList.add('hidden');
    }

    async function saveProfile(avatarPreset = '') {
      try {
        el('profileStatus').textContent = '保存中...';
        el('profileStatus').className = 'status';
        const file = el('profileAvatarFile').files[0];
        const avatarData = file ? await readFileAsBase64(file) : '';
        const result = await apiPost('/api/profile', {
          name: el('profileName').value,
          bio: el('profileBio').value,
          hobbies: el('profileHobbies').value,
          avatar_data: avatarData,
          avatar_preset: avatarPreset
        });
        applyUser(result.user);
        el('profileAvatarPreview').src = result.user.avatar_url || '/assets/ai-assistant-avatar.png';
        el('profileAvatarFile').value = '';
        el('profileStatus').textContent = '个人资料已保存';
        el('profileStatus').className = 'status ok';
      } catch (err) {
        el('profileStatus').textContent = err.message;
        el('profileStatus').className = 'status err';
      }
    }

    function renderReports() {
      const kind = el('mailKind').value;
      const list = state.reports.filter(r => r.kind === kind);
      el('reports').innerHTML = list.map(r => `
        <div class="report ${state.selected === r.name ? 'active' : ''}" data-name="${encodeURIComponent(r.name)}">
          <div class="report-head">
            <div>
              <div class="name">${escapeHtml(r.name)}</div>
              <div class="meta">${r.generated ? '新生成' : (r.kind === 'weekly' ? '周报模板' : '出差报告模板')} · ${new Date(r.mtime * 1000).toLocaleString()}</div>
            </div>
            <div class="report-actions">
              ${r.deletable ? `<button class="mini danger delete-report-file" type="button">删除</button>` : ''}
            </div>
          </div>
        </div>
      `).join('');
      document.querySelectorAll('.report').forEach(node => {
        node.addEventListener('click', () => loadDraft(kind, decodeURIComponent(node.dataset.name)));
      });
      el('reports').querySelectorAll('.delete-report-file').forEach(button => {
        button.addEventListener('click', event => {
          event.stopPropagation();
          deleteReportFile(kind, decodeURIComponent(button.closest('.report').dataset.name));
        });
      });
    }

    function clearMailDraft() {
      ['to', 'cc', 'subject', 'body', 'attachment'].forEach(id => {
        if (el(id)) el(id).value = '';
      });
      state.bodyHtml = '';
      renderBodyPreview();
      el('downloadLink')?.classList.add('hidden');
      if (el('preview')) el('preview').innerHTML = '暂无可预览内容';
      clearSendReview();
    }

    async function deleteReportFile(kind, name) {
      if (!confirm('确定删除这个报告文件吗？\\n' + name)) return;
      const wasSelected = state.selected === name;
      try {
        const result = await api('/api/delete-report', {
          method: 'POST',
          headers: { 'Content-Type': 'application/json' },
          body: JSON.stringify({ name })
        });
        el('status').textContent = '已删除报告文件：' + result.deleted;
        el('status').className = 'status ok';
        if (wasSelected) state.selected = '';
        await loadReports({ preserveSelection: true });
        if (wasSelected) {
          const next = state.reports.find(r => r.kind === kind)?.name || '';
          state.selected = next;
          renderReports();
          if (next) await loadDraft(kind, next);
          else clearMailDraft();
        }
      } catch (err) {
        el('status').textContent = err.message;
        el('status').className = 'status err';
      }
    }

    function reportKindName(kind) {
      return kind === 'weekly' ? '周报' : '出差报告';
    }

    function renderRecentDocs() {
      const list = state.reports.slice(0, 6);
      const box = el('recentDocsList');
      if (!box) return;
      box.innerHTML = list.length ? list.map(r => `
        <div class="doc-item">
          <div>
            <div class="doc-name">${escapeHtml(r.name)}</div>
            <div class="doc-meta">${r.kind === 'weekly' ? '周报' : '出差报告'} · ${new Date(r.mtime * 1000).toLocaleDateString()}</div>
          </div>
          <a class="download-link" style="padding:6px 12px;font-size:12px;" href="/download?file=${encodeURIComponent(r.name)}" target="_blank">查看</a>
        </div>
      `).join('') : '<div class="doc-item"><span class="doc-name">暂无最近文档</span></div>';
    }

    function renderHistoryReports() {
      const kind = el('historyKind')?.value || 'all';
      const list = state.reports.filter(r => !r.generated && (kind === 'all' || r.kind === kind));
      el('historyList').innerHTML = list.length ? list.map(r => `
        <div class="history-item" data-name="${encodeURIComponent(r.name)}">
          <div>
            <div class="history-name">${escapeHtml(r.name)}</div>
            <div class="history-meta">${reportKindName(r.kind)} · ${new Date(r.mtime * 1000).toLocaleString()}</div>
          </div>
          <div class="history-actions">
            <a class="download-link" href="/download?file=${encodeURIComponent(r.name)}" target="_blank">下载</a>
            <button class="mini danger delete-history" type="button">删除</button>
          </div>
        </div>
      `).join('') : '<div class="upload-item">暂无历史报告。</div>';
      el('historyList').querySelectorAll('.delete-history').forEach(button => {
        button.addEventListener('click', async () => {
          const name = decodeURIComponent(button.closest('.history-item').dataset.name);
          if (!confirm('确定删除这个历史报告吗？\\n' + name)) return;
          try {
            const result = await api('/api/delete-history', {
              method: 'POST',
              headers: { 'Content-Type': 'application/json' },
              body: JSON.stringify({ name })
            });
            el('status').textContent = '已删除历史报告：' + result.deleted;
            el('status').className = 'status ok';
            await loadReports({ preserveSelection: true });
          } catch (err) {
            el('status').textContent = err.message;
            el('status').className = 'status err';
          }
        });
      });
    }

    function setSubTab(sub) {
      state.subTab = sub;
      document.querySelectorAll('.sub-tab').forEach(tab => {
        tab.classList.toggle('active', tab.dataset.sub === sub);
      });
      const task = state.task;
      const isAssistant = task === 'weekly' || task === 'trip';
      const isWeekly = task === 'weekly';
      const isEdit = sub === 'edit';
      const isMail = sub === 'mail';
      const isHistory = sub === 'history';
      el('weeklyPanel').classList.toggle('hidden', !isAssistant || !isWeekly || !isEdit);
      el('tripPanel').classList.toggle('hidden', !isAssistant || isWeekly || !isEdit);
      el('generateToolbar').classList.toggle('hidden', !isAssistant || !isEdit);
      el('mailPanel').classList.toggle('hidden', !isAssistant || !isMail);
      el('uploadPanel').classList.toggle('hidden', !isAssistant || !isHistory);
      if (isMail) renderReports();
      if (isHistory) renderHistoryReports();
      el('status').textContent = '';
      el('status').className = 'status';
    }

    function navigateTo(task, sub) {
      state.subTab = sub || 'edit';
      setTask(task);
    }

    function setTask(task) {
      state.task = task;
      document.querySelectorAll('.task-card').forEach(card => {
        card.classList.toggle('active', card.dataset.task === task);
      });
      const isDashboard = task === 'dashboard';
      const isAssistant = task === 'weekly' || task === 'trip';
      el('dashboardPanel').classList.toggle('hidden', !isDashboard);
      el('pageHeader').classList.toggle('hidden', isDashboard);
      el('subTabs').classList.toggle('hidden', !isAssistant);
      el('weeklyPanel').classList.add('hidden');
      el('tripPanel').classList.add('hidden');
      el('mailPanel').classList.add('hidden');
      el('uploadPanel').classList.add('hidden');
      el('mailAssistantPanel').classList.toggle('hidden', task !== 'mailassistant');
      el('diaryPanel').classList.toggle('hidden', task !== 'diary');
      el('forumPanel').classList.toggle('hidden', task !== 'forum');
      el('newsPanel').classList.toggle('hidden', task !== 'news');
      el('generateToolbar').classList.add('hidden');
      el('configPanel').classList.toggle('hidden', task !== 'config');
      el('mailConfigPanel').classList.toggle('hidden', task !== 'mailconfig');
      el('skillsPanel').classList.toggle('hidden', task !== 'skills');
      el('userManagePanel').classList.toggle('hidden', task !== 'usermanage');
      if (!isAssistant) {
        state.subTab = 'edit';
        document.querySelectorAll('.sub-tab').forEach(tab => tab.classList.toggle('active', tab.dataset.sub === 'edit'));
      }
      const titles = { dashboard: '工作台', weekly: '周报助手', trip: '出差报告助手', diary: '工作日记', forum: '金点子论坛', news: '每日资讯', mailassistant: '邮件助手', config: '系统配置', mailconfig: '邮件配置', skills: '系统 Skill', usermanage: '用户管理' };
      const descs = {
        dashboard: '智能办公一站式工作台',
        weekly: '填写周报、发送邮件、管理历史周报',
        trip: '填写出差报告、发送邮件、管理历史出差报告',
        diary: '记录每日工作、查看历史日记，写周报时可智能总结',
        forum: '智能体或成员发起每日话题，大家围绕创意、改进和机会展开讨论',
        news: '收集轨道交通关键资讯，调用平台大模型生成每日简报',
        mailassistant: '查看收件箱、阅读邮件、发送普通邮件',
        config: '管理员配置 AI 接口和系统参数',
        mailconfig: '配置发件邮箱、收件人和抄送地址',
        skills: '查看已安装 Skill、能力说明、调用参数和示例',
        usermanage: '管理系统用户、角色权限和密码'
      };
      el('taskTitle').textContent = titles[task] || '';
      el('taskDesc').textContent = descs[task] || '';
      syncAgentToTask(task);
      if (isDashboard) {
        renderRecentDocs();
        const u = state.user;
        el('dashUserName').textContent = u ? (u.name || u.username) : '';
      }
      if (isAssistant) {
        el('kind').value = task;
        el('mailKind').value = task;
        el('uploadKind').value = task;
        el('historyKind').value = task;
        const hideEl = (sel) => { if (sel) sel.style.display = 'none'; };
        const hidePrevLabel = (sel) => { const lab = sel?.previousElementSibling; if (lab && lab.tagName === 'LABEL') lab.style.display = 'none'; };
        hideEl(el('mailKind')); hidePrevLabel(el('mailKind'));
        hideEl(el('uploadKind')); hidePrevLabel(el('uploadKind'));
        const historyTools = el('historyKind')?.closest('.history-tools');
        if (historyTools) historyTools.style.display = 'none';
        setSubTab(state.subTab || 'edit');
      }
      if (task === 'mailconfig') {
        loadMailConfig();
      }
      if (task === 'skills') {
        loadSkills();
      }
      if (task === 'mailassistant') {
        loadMailbox();
      }
      if (task === 'diary') {
        initDiaryPanel();
      }
      if (task === 'forum') {
        loadForumTopics();
      }
      if (task === 'news') {
        loadNews();
      }
    }

    function workFields(section) {
      const isNext = section === 'next';
      return isNext
        ? [
            ['category', '工作分类'],
            ['content', '工作内容'],
            ['difficulty', '困难与求助']
          ]
        : [
            ['category', '工作分类'],
            ['content', '工作内容'],
            [section === 'summary' ? 'status' : 'progress', section === 'summary' ? '完成情况' : '当前进展'],
            [section === 'summary' ? 'plan' : 'difficulty', section === 'summary' ? '后续计划' : '困难与求助']
          ];
    }

    function sectionName(section) {
      return { summary: '本周工作总结', follow: '重点工作跟进', next: '下周工作计划' }[section] || '工作内容';
    }

    function sectionTarget(section) {
      return section === 'summary' ? el('summaryRows') : section === 'follow' ? el('followRows') : el('nextRows');
    }

    function updateWeeklyCounts() {
      [['summary', 'summaryCount'], ['follow', 'followCount'], ['next', 'nextCount']].forEach(([section, id]) => {
        const count = collectWorkRows(section).length;
        if (el(id)) el(id).textContent = count;
      });
    }

    function rowPreview(row) {
      const parts = [row.category, row.content, row.status, row.progress, row.plan, row.difficulty]
        .map(value => String(value || '').trim())
        .filter(Boolean);
      return parts.join('\\n') || '点击填写任务内容';
    }

    function rowTemplate(section, row, index) {
      const fields = workFields(section);
      const title = row.category || `${sectionName(section)} ${index + 1}`;
      return `
        <div class="work-block" data-section="${section}" data-index="${index}">
          <div class="work-head">
            <div class="work-title">${escapeHtml(title)}</div>
            <button class="mini danger remove-row" type="button">删除</button>
          </div>
          <div class="work-summary">${escapeHtml(rowPreview(row))}</div>
          <div class="field-store">
            ${fields.map(([key, label]) => `
              <textarea data-key="${key}" data-label="${label}">${escapeHtml(row[key] || '')}</textarea>
            `).join('')}
          </div>
        </div>`;
    }

    function escapeHtml(value) {
      return String(value ?? '').replace(/[&<>"']/g, ch => ({
        '&': '&amp;', '<': '&lt;', '>': '&gt;', '"': '&quot;', "'": '&#39;'
      }[ch]));
    }

    function renderWorkRows(section, rows) {
      const target = sectionTarget(section);
      target.innerHTML = (rows && rows.length ? rows : [{}]).map((row, index) => rowTemplate(section, row, index)).join('');
      attachWorkEvents(target);
      updateWeeklyCounts();
    }

    function attachWorkEvents(target) {
      target.querySelectorAll('.remove-row').forEach(button => {
        button.addEventListener('click', event => {
          event.stopPropagation();
          button.closest('.work-block').remove();
          renumberRows(target);
          updateWeeklyCounts();
        });
      });
      target.querySelectorAll('.work-block').forEach(block => {
        block.addEventListener('click', () => openWorkModal(block));
      });
    }

    function renumberRows(container) {
      container.querySelectorAll('.work-block').forEach((block, index) => {
        block.dataset.index = index;
        refreshWorkCard(block);
      });
    }

    function addWorkRow(section, row = {}) {
      const target = sectionTarget(section);
      if (target.children.length === 1 && !target.querySelector('textarea')?.value.trim()) {
        target.innerHTML = '';
      }
      const wrapper = document.createElement('div');
      wrapper.innerHTML = rowTemplate(section, row, target.children.length).trim();
      const block = wrapper.firstElementChild;
      target.appendChild(block);
      attachWorkEvents(target);
      updateWeeklyCounts();
      openWorkModal(block);
    }

    function collectWorkRows(section) {
      const target = sectionTarget(section);
      return [...target.querySelectorAll('.work-block')].map(block => {
        const row = {};
        block.querySelectorAll('textarea[data-key]').forEach(input => {
          row[input.dataset.key] = input.value.trim();
        });
        return row;
      }).filter(row => Object.values(row).some(Boolean));
    }

    function refreshWorkCard(block) {
      const row = {};
      block.querySelectorAll('textarea[data-key]').forEach(input => {
        row[input.dataset.key] = input.value.trim();
      });
      const section = block.dataset.section;
      const index = Number(block.dataset.index || 0);
      block.querySelector('.work-title').textContent = row.category || `${sectionName(section)} ${index + 1}`;
      block.querySelector('.work-summary').textContent = rowPreview(row);
      updateWeeklyCounts();
      saveFormDraft();
    }

    function openModal(title, fields, onSave) {
      el('modalTitle').textContent = title;
      el('modalFields').innerHTML = fields.map(field => {
        const wide = field.multiline ? ' wide' : '';
        const control = field.multiline
          ? `<textarea data-modal-key="${field.key}" placeholder="${field.label}">${escapeHtml(field.value || '')}</textarea>`
          : `<input data-modal-key="${field.key}" type="${field.type || 'text'}" placeholder="${field.label}" value="${escapeHtml(field.value || '')}" />`;
        const assist = field.multiline ? `
          <div class="field-assist">
            <div class="field-assist-head">
              <details>
                <summary>展开/修改提示词</summary>
                <textarea class="field-assist-prompt">${escapeHtml(defaultAssistantPrompt)}</textarea>
              </details>
              <button class="warn assist-optimize" type="button">辅助优化</button>
            </div>
            <div class="assist-status">待优化</div>
          </div>` : '';
        return `
          <div class="${wide}">
            <label>${field.label}</label>
            ${assist}
            ${control}
          </div>`;
      }).join('');
      state.modalSave = onSave;
      el('editModal').classList.remove('hidden');
      const first = el('modalFields').querySelector('input, textarea');
      if (first) first.focus();
    }

    function closeModal() {
      el('editModal').classList.add('hidden');
      state.modalSave = null;
    }

    function saveAndCloseModal() {
      saveCurrentModal();
      closeModal();
    }

    function modalValues() {
      const values = {};
      el('modalFields').querySelectorAll('[data-modal-key]').forEach(input => {
        values[input.dataset.modalKey] = input.value.trim();
      });
      return values;
    }

    function draftStorageKey() {
      return state.user?.username ? `${FORM_DRAFT_PREFIX}:${state.user.username}` : '';
    }

    function tripFormPayload() {
      return {
        reporter: el('tripReporter').value,
        department: el('tripDepartment').value,
        location: el('tripLocation').value,
        trip_start: el('tripStart').value,
        trip_end: el('tripEnd').value,
        purpose: el('tripPurpose').value,
        itinerary: el('tripItinerary').value,
        details: el('tripDetails').value,
        issues: el('tripIssues').value,
        suggestions: el('tripSuggestions').value
      };
    }

    function saveFormDraft() {
      const key = draftStorageKey();
      if (!key || state.restoringDraft) return;
      try {
        const draft = {
          updatedAt: Date.now(),
          weekly: {
            start: el('weeklyStart').value,
            end: el('weeklyEnd').value,
            period: el('weeklyPeriod').value,
            summary: collectWorkRows('summary'),
            follow: collectWorkRows('follow'),
            next: collectWorkRows('next')
          },
          trip: tripFormPayload()
        };
        localStorage.setItem(key, JSON.stringify(draft));
      } catch (err) {
        console.warn('保存本地草稿失败', err);
      }
    }

    function restoreSavedFormDraft(kind) {
      const key = draftStorageKey();
      if (!key) return false;
      let draft = null;
      try {
        draft = JSON.parse(localStorage.getItem(key) || 'null');
      } catch (err) {
        return false;
      }
      if (!draft) return false;
      state.restoringDraft = true;
      try {
        if (kind === 'weekly' && draft.weekly) {
          if (draft.weekly.start) el('weeklyStart').value = draft.weekly.start;
          if (draft.weekly.end) el('weeklyEnd').value = draft.weekly.end;
          syncWeeklyPeriod();
          if ((draft.weekly.summary || []).length || (draft.weekly.follow || []).length || (draft.weekly.next || []).length) {
            renderWorkRows('summary', draft.weekly.summary || []);
            renderWorkRows('follow', draft.weekly.follow || []);
            renderWorkRows('next', draft.weekly.next || []);
            state.weeklyPrefilled = true;
            return true;
          }
        }
        if (kind === 'trip' && draft.trip) {
          Object.entries({
            tripReporter: draft.trip.reporter,
            tripDepartment: draft.trip.department,
            tripLocation: draft.trip.location,
            tripStart: draft.trip.trip_start,
            tripEnd: draft.trip.trip_end,
            tripPurpose: draft.trip.purpose,
            tripItinerary: draft.trip.itinerary,
            tripDetails: draft.trip.details,
            tripIssues: draft.trip.issues,
            tripSuggestions: draft.trip.suggestions
          }).forEach(([id, value]) => {
            if (value !== undefined && value !== null) el(id).value = value;
          });
          renderTripCards();
          state.tripPrefilled = true;
          return true;
        }
      } finally {
        state.restoringDraft = false;
      }
      return false;
    }

    function clearSavedFormDraft(kind) {
      const key = draftStorageKey();
      if (!key) return;
      try {
        const draft = JSON.parse(localStorage.getItem(key) || '{}');
        if (kind === 'weekly') delete draft.weekly;
        if (kind === 'trip') delete draft.trip;
        localStorage.setItem(key, JSON.stringify({ ...draft, updatedAt: Date.now() }));
      } catch (err) {
        localStorage.removeItem(key);
      }
    }

    function saveCurrentModal() {
      if (state.modalSave) {
        state.modalSave(modalValues());
        saveFormDraft();
      }
    }

    function openWorkModal(block) {
      const section = block.dataset.section;
      const index = Number(block.dataset.index || 0) + 1;
      const fields = workFields(section).map(([key, label]) => {
        const input = block.querySelector(`[data-key="${key}"]`);
        return { key, label, value: input?.value || '', multiline: key !== 'category' };
      });
      openModal(`${sectionName(section)} ${index}`, fields, values => {
        fields.forEach(field => {
          const input = block.querySelector(`[data-key="${field.key}"]`);
          if (input) input.value = values[field.key] || '';
        });
        refreshWorkCard(block);
      });
    }

    const tripGroups = [
      {
        key: 'base',
        title: '基础信息',
        fields: [
          ['tripReporter', '报告人', false],
          ['tripDepartment', '部门', false],
          ['tripLocation', '出差地点', false],
          ['tripStart', '开始日期', false, 'date'],
          ['tripEnd', '结束日期', false, 'date']
        ]
      },
      { key: 'purpose', title: '出差目的', fields: [['tripPurpose', '出差目的', true]] },
      { key: 'itinerary', title: '行程概览', fields: [['tripItinerary', '行程概览', true]] },
      { key: 'details', title: '工作详情', fields: [['tripDetails', '工作详情', true]] },
      { key: 'issues', title: '问题与反馈', fields: [['tripIssues', '问题与反馈', true]] },
      { key: 'suggestions', title: '总结与建议', fields: [['tripSuggestions', '总结与建议', true]] }
    ];

    function tripPreview(group) {
      const values = group.fields.map(([id, label]) => {
        const value = el(id).value.trim();
        return value ? (group.key === 'base' ? `${label}：${value}` : value) : '';
      }).filter(Boolean);
      return values.join('\\n') || '点击填写内容';
    }

    function formatPeriodDate(value) {
      if (!value) return '';
      const [year, month, day] = value.split('-').map(Number);
      if (!year || !month || !day) return value;
      return `${year}.${month}.${day}`;
    }

    function toDateInputValue(date) {
      const year = date.getFullYear();
      const month = String(date.getMonth() + 1).padStart(2, '0');
      const day = String(date.getDate()).padStart(2, '0');
      return `${year}-${month}-${day}`;
    }

    function setDefaultWeeklyDates(force = false) {
      if (!force && (el('weeklyStart').value || el('weeklyEnd').value)) {
        syncWeeklyPeriod();
        return;
      }
      const today = new Date();
      const day = today.getDay() || 7;
      const thisMonday = new Date(today);
      thisMonday.setHours(0, 0, 0, 0);
      thisMonday.setDate(today.getDate() - day + 1);
      const thisFriday = new Date(thisMonday);
      thisFriday.setDate(thisMonday.getDate() + 4);
      el('weeklyStart').value = toDateInputValue(thisMonday);
      el('weeklyEnd').value = toDateInputValue(thisFriday);
      el('diarySumStart').value = toDateInputValue(thisMonday);
      el('diarySumEnd').value = toDateInputValue(thisFriday);
      syncWeeklyPeriod();
    }

    function syncWeeklyPeriod() {
      const start = formatPeriodDate(el('weeklyStart').value);
      const end = formatPeriodDate(el('weeklyEnd').value);
      el('weeklyPeriod').value = start && end ? `${start}-${end}` : (start || end);
      if (el('weeklyPeriodText')) {
        el('weeklyPeriodText').textContent = el('weeklyPeriod').value || '未选择时段';
      }
      saveFormDraft();
    }

    function renderTripCards() {
      el('tripCards').innerHTML = tripGroups.map(group => `
        <div class="edit-card" data-trip-group="${group.key}">
          <div class="edit-card-title">${group.title}</div>
          <div class="edit-card-preview">${escapeHtml(tripPreview(group))}</div>
        </div>
      `).join('');
      el('tripCards').querySelectorAll('.edit-card').forEach(card => {
        card.addEventListener('click', () => openTripModal(card.dataset.tripGroup));
      });
    }

    function openTripModal(groupKey) {
      const group = tripGroups.find(item => item.key === groupKey);
      if (!group) return;
      const fields = group.fields.map(([id, label, multiline, type]) => ({
        key: id,
        label,
        multiline,
        type,
        value: el(id).value
      }));
      openModal(group.title, fields, values => {
        fields.forEach(field => {
          el(field.key).value = values[field.key] || '';
        });
        renderTripCards();
        saveFormDraft();
      });
    }

    function sampleSkillCall(skill) {
      const args = {};
      Object.keys(skill.parameters || {}).forEach(key => {
        const desc = String(skill.parameters[key] || '');
        if (desc.includes('YYYY-MM-DD')) args[key] = '2026-05-14';
        else if (desc.includes('数组')) args[key] = [];
        else if (key === 'kind') args[key] = 'weekly';
        else if (key === 'limit') args[key] = 20;
        else args[key] = desc.replace(/，可选/g, '') || '';
      });
      return JSON.stringify({
        reply: `准备调用 ${skill.title}`,
        skill_call: { name: skill.name, arguments: args }
      }, null, 2);
    }

    function skillFallback(name) {
      const module = name.split('.')[0] || '其他';
      return {
        name,
        module: module === 'weekly' ? '周报' : module,
        title: name,
        description: '系统 Skill 测试',
        parameters: {},
        safe: !/send|preview|generate|save|create|comment/.test(name)
      };
    }

    function skillDefaultArguments(skill) {
      const example = skill?.detail?.call_example?.skill_call?.arguments;
      if (example && typeof example === 'object') return example;
      const args = {};
      Object.keys(skill?.parameters || {}).forEach(key => {
        const desc = String(skill.parameters[key] || '');
        if (desc.includes('YYYY-MM-DD')) args[key] = '2026-05-14';
        else if (desc.includes('数组')) args[key] = [];
        else if (key === 'kind') args[key] = 'weekly';
        else if (key === 'limit') args[key] = 20;
        else if (key === 'uid') args[key] = '';
        else args[key] = desc.includes('可选') ? '' : desc;
      });
      return args;
    }

    function openSkillTest(name) {
      const skill = (state.skills || []).find(item => item.name === name) || skillFallback(name);
      state.currentSkill = skill;
      el('skillTestTitle').textContent = `${skill.name} 测试`;
      el('skillTestMeta').textContent = `${skill.module || '其他'} Skill · ${skill.title || ''} · ${skill.safe ? '查询/预览类' : '写入/外部动作类'}`;
      el('skillTestArgs').value = JSON.stringify(skillDefaultArguments(skill), null, 2);
      el('skillTestInstruction').value = '';
      el('skillConfirmUnsafe').checked = false;
      el('skillConfirmWrap').classList.toggle('hidden', !!skill.safe);
      el('skillTestStatus').textContent = '';
      el('skillTestResult').textContent = '点击“运行测试”后，这里会显示 Skill 返回结果。';
      el('skillTestLinks').innerHTML = '';
      el('skillTestModal').classList.remove('hidden');
      renderIcons(el('skillTestModal'));
    }

    function closeSkillTest() {
      el('skillTestModal').classList.add('hidden');
      state.currentSkill = null;
    }

    function attachSkillCardTests(root = document) {
      root.querySelectorAll('.skill-card[data-skill-name]').forEach(card => {
        if (card.dataset.testBound === 'true') return;
        card.dataset.testBound = 'true';
        card.addEventListener('click', () => openSkillTest(card.dataset.skillName));
      });
    }

    function renderSkillTestArtifacts(result) {
      const data = result?.result || {};
      const links = [];
      if (data.download_url) {
        links.push(`<a class="secondary" href="${escapeHtml(data.download_url)}" target="_blank">打开生成文件</a>`);
      }
      if (data.preview_image_url) {
        links.push(`<img class="skill-preview-image" src="${escapeHtml(data.preview_image_url)}" alt="Skill 预览图片" />`);
      }
      el('skillTestLinks').innerHTML = links.join('');
    }

    async function runSkillTest() {
      const skill = state.currentSkill;
      if (!skill) return;
      let args = {};
      try {
        args = JSON.parse(el('skillTestArgs').value || '{}');
      } catch (err) {
        el('skillTestStatus').textContent = '调用参数不是有效 JSON。';
        return;
      }
      el('skillRunTest').disabled = true;
      el('skillTestStatus').textContent = el('skillTestInstruction').value.trim() ? '正在调用平台配置的大模型 API 生成参数并执行 Skill...' : '正在执行 Skill 测试...';
      el('skillTestResult').textContent = '测试运行中...';
      el('skillTestLinks').innerHTML = '';
      try {
        const result = await api('/api/skill-test', {
          method: 'POST',
          headers: { 'Content-Type': 'application/json' },
          body: JSON.stringify({
            name: skill.name,
            arguments: args,
            instruction: el('skillTestInstruction').value.trim(),
            confirm_unsafe: el('skillConfirmUnsafe').checked
          })
        });
        el('skillTestStatus').textContent = result.model ? `测试完成，已使用模型：${result.model}` : '测试完成。';
        el('skillTestResult').textContent = JSON.stringify(result, null, 2);
        renderSkillTestArtifacts(result);
      } catch (err) {
        el('skillTestStatus').textContent = err.message;
        el('skillTestResult').textContent = JSON.stringify({ ok: false, error: err.message }, null, 2);
      } finally {
        el('skillRunTest').disabled = false;
      }
    }

    function renderSkills(skills) {
      const moduleOrder = ['周报', '出差报告', '工作日记', '金点子论坛', '邮件', '资讯', '报告', '通用'];
      const moduleLabels = {
        周报: '周报 Skill',
        出差报告: '出差报告 Skill',
        工作日记: '日记 Skill',
        金点子论坛: '金点子论坛 Skill',
        邮件: '邮件 Skill',
        资讯: '资讯 Skill',
        报告: '报告 Skill',
        通用: '通用 Skill'
      };
      const foundModules = [...new Set((skills || []).map(item => item.module || '其他'))];
      const orderedModules = [
        ...moduleOrder.filter(module => foundModules.includes(module)),
        ...foundModules.filter(module => !moduleOrder.includes(module))
      ];
      const modules = ['all', ...orderedModules];
      const filter = el('skillModuleFilter');
      const previous = filter.value || 'all';
      filter.innerHTML = modules.map(module => `<option value="${escapeHtml(module)}">${module === 'all' ? '全部模块' : escapeHtml(module)}</option>`).join('');
      filter.value = modules.includes(previous) ? previous : 'all';
      const keyword = (el('skillSearch').value || '').trim().toLowerCase();
      const moduleName = filter.value || 'all';
      const counts = {};
      (skills || []).forEach(skill => {
        const module = skill.module || '其他';
        counts[module] = (counts[module] || 0) + 1;
      });
      el('skillModuleSummary').innerHTML = orderedModules.map(module => `
        <div class="skill-module-card ${moduleName === module ? 'active' : ''}" data-module="${escapeHtml(module)}">
          <div class="skill-module-name">${escapeHtml(moduleLabels[module] || (module + ' Skill'))}</div>
          <div class="skill-module-count">已安装 ${counts[module] || 0} 个能力，点击查看</div>
        </div>
      `).join('');
      el('skillModuleSummary').querySelectorAll('.skill-module-card').forEach(card => {
        card.addEventListener('click', () => {
          el('skillModuleFilter').value = card.dataset.module;
          renderSkills(state.skills || []);
        });
      });
      el('skillTotalCount').textContent = `共 ${(skills || []).length} 个 Skill`;
      const list = (skills || []).filter(skill => {
        const text = `${skill.name} ${skill.title} ${skill.module} ${skill.description}`.toLowerCase();
        return (moduleName === 'all' || skill.module === moduleName) && (!keyword || text.includes(keyword));
      });
      const summaryText = (text, max = 72) => {
        const compact = String(text || '').replace(/\s+/g, ' ').trim();
        return compact.length > max ? compact.slice(0, max) + '...' : compact;
      };
      el('skillList').innerHTML = list.length ? list.map(skill => `
        <div class="skill-card compact" data-skill-name="${escapeHtml(skill.name)}">
          <div class="skill-card-head">
            <div>
              <div class="skill-name">${escapeHtml(skill.name)}</div>
              <div class="skill-title">${escapeHtml(skill.module || '其他')} Skill · ${escapeHtml(skill.title || '')}</div>
            </div>
            <div class="skill-card-actions">
              <span class="skill-badge ${skill.safe ? '' : 'warn'}">${skill.safe ? '查询/预览' : '写入/外部动作'}</span>
              <button type="button" class="skill-help-btn" data-help="${escapeHtml(skill.name)}" aria-expanded="false">详情</button>
            </div>
          </div>
          <div class="skill-desc">${escapeHtml(summaryText(skill.description || '点击详情查看适用场景、参数和调用示例。'))}</div>
          <div class="skill-detail" id="skill-detail-${escapeHtml(skill.name)}">
            <div class="meta-label">说明</div>
            <div>${escapeHtml(skill.description || '')}</div>
            ${skill.detail && skill.detail.when_to_use && skill.detail.when_to_use.length ? `
              <div class="meta-label">适用场景</div>
              <div>${(skill.detail.when_to_use || []).map(item => `· ${escapeHtml(item)}`).join('<br>')}</div>
            ` : ''}
            <div class="meta-label">参数</div>
            <pre>${escapeHtml(JSON.stringify((skill.detail && skill.detail.input_schema) || skill.parameters || {}, null, 2))}</pre>
            ${skill.detail && skill.detail.output_schema ? `
              <div class="meta-label">输出结构</div>
              <pre>${escapeHtml(JSON.stringify(skill.detail.output_schema || {}, null, 2))}</pre>
            ` : ''}
            <div class="meta-label">调用示例</div>
            <pre>${escapeHtml(JSON.stringify((skill.detail && skill.detail.call_example) || {
              reply: `准备调用 ${skill.title}`,
              skill_call: { name: skill.name, arguments: {} }
            }, null, 2))}</pre>
          </div>
        </div>
      `).join('') : '<div class="upload-item">没有匹配的 Skill。</div>';
      attachSkillCardTests(el('skillList'));
      el('skillList').querySelectorAll('.skill-help-btn').forEach(btn => {
        btn.addEventListener('click', (e) => {
          e.stopPropagation();
          const detail = el(`skill-detail-${btn.dataset.help}`);
          if (!detail) return;
          const open = detail.classList.toggle('open');
          btn.setAttribute('aria-expanded', open ? 'true' : 'false');
          btn.textContent = open ? '收起' : '详情';
        });
      });
    }

    async function loadSkills() {
      if (!state.user?.is_superadmin) return;
      el('skillList').innerHTML = '<div class="upload-item">正在读取系统 Skill...</div>';
      try {
        const data = await api('/api/skills');
        state.skills = data.skills || [];
        renderSkills(state.skills);
      } catch (err) {
        el('skillList').innerHTML = `<div class="upload-item">${escapeHtml(err.message)}</div>`;
      }
    }

    async function loadReports(options = {}) {
      const data = await api('/api/reports');
      state.reports = data.reports;
      window.latestWeekly = data.latest_weekly;
      window.latestTrip = data.latest_trip;
      const task = state.task || 'weekly';
      if (!options.preserveSelection) {
        state.selected = task === 'weekly' ? data.latest_weekly : data.latest_trip;
        el('mailKind').value = task;
      }
      setTask(task);
      renderReports();
      renderHistoryReports();
      if (!options.preserveSelection && state.selected) {
        await loadDraft(task, state.selected);
      }
    }

    async function loadAdminConfig() {
      if (!state.user?.is_admin) return;
      const data = await api('/api/admin-config');
      el('configApiUrl').value = data.assistant_api_url || '';
      el('configModel').value = data.assistant_model || 'MiniMax-M2.7';
      renderModelOptions([data.assistant_model || 'MiniMax-M2.7'], data.assistant_model || 'MiniMax-M2.7');
      el('configPrompt').value = data.assistant_prompt || defaultAssistantPrompt;
      el('configApiKey').value = '';
      el('configKeyHint').textContent = 'API Key 状态：' + (data.assistant_api_key_masked || '未配置');
      el('configSmtpHost').value = data.smtp_host || 'smtp.263.net';
      el('configSmtpPort').value = data.smtp_port || 465;
      el('configSmtpTls').checked = !!data.smtp_tls;
      el('configSmtpSsl').checked = !!data.smtp_ssl;
      el('configImapHost').value = data.imap_host || 'imap.263.net';
      el('configImapPort').value = data.imap_port || 993;
      el('configImapSsl').checked = !!data.imap_ssl;
    }
    async function loadUserManage() {
      console.log('loadUserManage called, is_superadmin:', state.user?.is_superadmin);
      if (!state.user?.is_superadmin) { console.log('loadUserManage skipped: not superadmin'); return; }
      el('newUserRoleBox').style.display = 'grid';
      await loadUserList();
    }
    async function loadUserList() {
      if (!state.user?.is_admin) { console.log('loadUserList skipped: not admin'); return; }
      try {
        const data = await api('/api/admin-users-list');
        console.log('loadUserList success, users count:', (data.users || []).length);
        renderUsers(data.users || []);
      } catch (err) {
        console.error('loadUserList error:', err);
        el('userList').innerHTML = '<div class="upload-item">加载用户列表失败: ' + escapeHtml(err.message) + '</div>';
      }
    }

    async function loadMailConfig() {
      const data = await api('/api/mail-config');
      const ref = data.reference || {};
      const setVal = (id, val, refVal) => {
        const node = el(id);
        if (!node) return;
        node.value = val || '';
        node.placeholder = refVal || '';
      };
      setVal('mailUserEmail', data.user_email, ref.user_email);
      setVal('mailSmtpFrom', data.smtp_from, ref.smtp_from);
      setVal('mailSmtpUser', data.smtp_user, ref.smtp_user);
      el('mailSmtpPassword').value = '';
      el('mailPasswordHint').textContent = 'SMTP 密码/授权码状态：' + (data.smtp_password_masked || '未配置');
      setVal('mailImapUser', data.imap_user, ref.imap_user);
      el('mailImapPassword').value = '';
      el('mailImapPasswordHint').textContent = 'IMAP 密码/授权码状态：' + (data.imap_password_masked || '未配置');
      setVal('mailWeeklyTo', data.weekly_to, ref.weekly_to);
      setVal('mailWeeklyCc', data.weekly_cc, ref.weekly_cc);
      setVal('mailTripTo', data.trip_to, ref.trip_to);
      setVal('mailTripCc', data.trip_cc, ref.trip_cc);
      el('mailEmailSignature').value = data.email_signature || '';
      // 显示全局服务器配置
      el('mailSmtpHostDisplay').textContent = data.smtp_host || 'smtp.263.net';
      el('mailSmtpPortDisplay').textContent = (data.smtp_port || 465) + ' / ' + (data.smtp_ssl ? 'SSL' : 'TLS');
      el('mailImapHostDisplay').textContent = data.imap_host || 'imap.263.net';
      el('mailImapPortDisplay').textContent = (data.imap_port || 993) + ' / ' + (data.imap_ssl ? 'SSL' : 'TLS');
    }

    function mailConfigPayload() {
      const userEmail = el('mailUserEmail').value.trim();
      const smtpUser = el('mailSmtpUser').value.trim() || userEmail;
      return {
        user_email: userEmail,
        smtp_from: el('mailSmtpFrom').value.trim() || userEmail,
        smtp_user: smtpUser,
        smtp_password: el('mailSmtpPassword').value,
        imap_user: el('mailImapUser').value.trim() || smtpUser,
        imap_password: el('mailImapPassword').value,
        weekly_to: el('mailWeeklyTo').value,
        weekly_cc: el('mailWeeklyCc').value,
        trip_to: el('mailTripTo').value,
        trip_cc: el('mailTripCc').value,
        email_signature: el('mailEmailSignature').value
      };
    }

    function explainMailLoginError(message) {
      const text = String(message || '邮箱测试失败');
      if (text.includes('SMTP 用户名')) return text + ' 请填写完整邮箱地址；留空保存时系统会自动使用“本人邮箱”。';
      if (text.includes('SMTP 授权码') || text.includes('SMTP 密码')) return text + ' 请填写邮箱后台生成的 SMTP 授权码，不是网页登录密码。';
      if (text.includes('IMAP 用户名')) return text + ' 如需读取收件箱，IMAP 用户名通常与 SMTP 用户名相同。';
      if (text.includes('IMAP 授权码') || text.includes('IMAP 密码')) return text + ' 如需读取收件箱，可复用邮箱授权码或单独生成 IMAP 授权码。';
      if (text.includes('timed out') || text.includes('超时')) return text + ' 请检查服务器、端口和 SSL 设置是否匹配。263 邮箱通常是 SMTP 465/SSL。';
      return text;
    }

    function textToHtml(text) {
      return escapeHtml(text || '').split('\\n').join('<br>');
    }

    function renderBodyPreview() {
      if (state.bodyHtml) {
        el('bodyPreview').innerHTML = state.bodyHtml;
      } else {
        el('bodyPreview').innerHTML = textToHtml(el('body').value || '暂无正文内容');
      }
    }

    function renderModelOptions(models, selected) {
      const unique = [...new Set((models || []).filter(Boolean))];
      if (!unique.includes(selected) && selected) unique.unshift(selected);
      el('configModelSelect').innerHTML = unique.map(model => `
        <option value="${escapeHtml(model)}" ${model === selected ? 'selected' : ''}>${escapeHtml(model)}</option>
      `).join('');
    }

    function renderUsers(users) {
      const roleLabel = (role) => {
        if (role === 'superadmin') return '超级管理员';
        if (role === 'admin') return '管理员';
        return '普通成员';
      };
      const isSuper = state.user?.is_superadmin;
      const isAdmin = state.user?.is_admin;
      console.log('renderUsers called, isSuper:', isSuper, 'users count:', (users || []).length);
      el('userList').innerHTML = (users || []).map((user, idx) => {
        const canEdit = isSuper;
        const canDelete = isSuper && user.username !== state.user?.username;
        const roleSelect = canEdit
          ? `<select class="user-role-select mini" data-user="${escapeHtml(user.username)}" style="margin-right:8px;">
              <option value="member" ${user.role === 'member' ? 'selected' : ''}>普通成员</option>
              <option value="admin" ${user.role === 'admin' ? 'selected' : ''}>管理员</option>
              <option value="superadmin" ${user.role === 'superadmin' ? 'selected' : ''}>超级管理员</option>
            </select>`
          : `<span class="mini" style="margin-right:8px;color:var(--muted);">${roleLabel(user.role)}</span>`;
        const editBtn = canEdit
          ? `<button class="mini secondary edit-user-toggle" data-user="${escapeHtml(user.username)}" data-idx="${idx}" type="button" style="margin-right:8px;">编辑</button>`
          : '';
        const deleteBtn = canDelete
          ? `<button class="mini danger delete-user" data-user="${escapeHtml(user.username)}" type="button">删除</button>`
          : '';
        const editForm = canEdit ? `
          <div class="user-edit-form hidden" id="userEdit_${idx}" style="grid-column:1/-1;margin-top:8px;padding-top:10px;border-top:1px solid var(--line);">
            <div class="config-grid">
              <div>
                <label style="margin-top:0;">显示名称</label>
                <input class="edit-name" data-user="${escapeHtml(user.username)}" value="${escapeHtml(user.name || '')}" placeholder="${escapeHtml(user.username)}" />
              </div>
              <div>
                <label style="margin-top:0;">重置密码（留空则不修改）</label>
                <input class="edit-password" data-user="${escapeHtml(user.username)}" type="password" placeholder="不修改则留空" />
              </div>
            </div>
            <div class="toolbar" style="margin-top:10px;">
              <button class="mini save-user-edit" data-user="${escapeHtml(user.username)}" data-idx="${idx}" type="button">保存修改</button>
              <button class="mini secondary cancel-user-edit" data-idx="${idx}" type="button">取消</button>
            </div>
          </div>
        ` : '';
        return `
        <div class="user-item" style="display:grid;grid-template-columns:minmax(0,1fr) auto auto;gap:10px;align-items:center;" data-idx="${idx}">
          <div>
            <strong>${escapeHtml(user.name || user.username)}</strong>
            <div class="history-meta">${escapeHtml(user.username)} · ${roleLabel(user.role)}</div>
          </div>
          <div style="display:flex;align-items:center;">${roleSelect}${editBtn}${deleteBtn}</div>
          ${editForm}
        </div>
      `}).join('') || '<div class="upload-item">暂无用户。</div>';
      // 绑定删除按钮
      el('userList').querySelectorAll('.delete-user').forEach(btn => {
        btn.addEventListener('click', async () => {
          const username = btn.dataset.user;
          if (!confirm(`确定删除用户 ${username} 吗？`)) return;
          try {
            const result = await apiPost('/api/admin-users-delete', { username });
            renderUsers(result.users || []);
            el('userManageStatus').textContent = '用户已删除';
            el('userManageStatus').className = 'status ok';
          } catch (err) {
            el('userManageStatus').textContent = err.message;
            el('userManageStatus').className = 'status err';
          }
        });
      });
      // 绑定角色修改
      if (isSuper) {
        el('userList').querySelectorAll('.user-role-select').forEach(sel => {
          sel.addEventListener('change', async () => {
            const username = sel.dataset.user;
            const newRole = sel.value;
            try {
              const result = await apiPost('/api/admin-users-update', { username, role: newRole });
              renderUsers(result.users || []);
              el('userManageStatus').textContent = '权限已更新';
              el('userManageStatus').className = 'status ok';
            } catch (err) {
              el('userManageStatus').textContent = err.message;
              el('userManageStatus').className = 'status err';
              await loadUserList();
            }
          });
        });
      }
      // 绑定编辑展开/收起
      el('userList').querySelectorAll('.edit-user-toggle').forEach(btn => {
        btn.addEventListener('click', () => {
          const idx = btn.dataset.idx;
          const form = el('userEdit_' + idx);
          if (form) {
            const wasHidden = form.classList.contains('hidden');
            // 先关闭所有编辑表单
            el('userList').querySelectorAll('.user-edit-form').forEach(f => f.classList.add('hidden'));
            el('userList').querySelectorAll('.edit-user-toggle').forEach(b => b.textContent = '编辑');
            if (wasHidden) {
              form.classList.remove('hidden');
              btn.textContent = '收起';
            }
          }
        });
      });
      // 绑定取消编辑
      el('userList').querySelectorAll('.cancel-user-edit').forEach(btn => {
        btn.addEventListener('click', () => {
          const idx = btn.dataset.idx;
          const form = el('userEdit_' + idx);
          if (form) {
            form.classList.add('hidden');
            const toggle = el('userList').querySelector(`.edit-user-toggle[data-idx="${idx}"]`);
            if (toggle) toggle.textContent = '编辑';
          }
        });
      });
      // 绑定保存编辑
      el('userList').querySelectorAll('.save-user-edit').forEach(btn => {
        btn.addEventListener('click', async () => {
          const username = btn.dataset.user;
          const idx = btn.dataset.idx;
          const row = el('userList').querySelector(`.user-item[data-idx="${idx}"]`);
          const name = row?.querySelector('.edit-name')?.value?.trim() || '';
          const password = row?.querySelector('.edit-password')?.value?.trim() || '';
          const payload = { username };
          if (name) payload.name = name;
          if (password) payload.password = password;
          try {
            const result = await apiPost('/api/admin-users-update', payload);
            renderUsers(result.users || []);
            el('userManageStatus').textContent = '用户信息已更新';
            el('userManageStatus').className = 'status ok';
          } catch (err) {
            el('userManageStatus').textContent = err.message;
            el('userManageStatus').className = 'status err';
          }
        });
      });
    }

    function renderMailbox(messages) {
      const box = el('mailboxList');
      box.innerHTML = (messages || []).length ? messages.map(item => `
        <div class="mailbox-item" data-uid="${escapeHtml(item.uid)}">
          <div class="mailbox-subject">${escapeHtml(item.subject || '无主题')}</div>
          <div class="mailbox-meta">${escapeHtml(item.from || '未知发件人')}</div>
          <div class="mailbox-meta">${escapeHtml(item.date || '')}</div>
          ${(item.attachments || []).length ? `<div class="mailbox-meta">${(item.attachments || []).length} 个附件</div>` : ''}
          <div class="mailbox-preview">${escapeHtml(item.preview || '暂无正文预览')}</div>
        </div>
      `).join('') : '<div class="upload-item">暂无邮件，或当前邮箱没有可读取的收件箱邮件。</div>';
      box.querySelectorAll('.mailbox-item').forEach(item => {
        item.addEventListener('click', () => loadMailDetail(item.dataset.uid));
      });
    }

    async function loadMailbox(forceRefresh = false) {
      if (!el('mailboxList')) return;
      el('mailboxList').innerHTML = `<div class="upload-item">${forceRefresh ? '正在刷新收件箱...' : '正在读取收件箱缓存...'}</div>`;
      el('mailDetail').textContent = '请选择左侧邮件查看详情。';
      try {
        const query = new URLSearchParams({ limit: el('mailboxLimit').value || '20' });
        if (forceRefresh) query.set('refresh', '1');
        const data = await api('/api/mailbox?' + query.toString());
        renderMailbox(data.messages || []);
        if (data.cached) {
          el('mailDetail').textContent = '已从本地缓存加载。需要最新邮件时点击“刷新”。';
        }
      } catch (err) {
        el('mailboxList').innerHTML = `<div class="upload-item">${escapeHtml(err.message)}</div>`;
      }
    }

    async function loadMailDetail(uid) {
      if (!uid) return;
      document.querySelectorAll('.mailbox-item').forEach(item => item.classList.toggle('active', item.dataset.uid === uid));
      el('mailDetail').textContent = '正在读取邮件详情...';
      try {
        const query = new URLSearchParams({ uid });
        const data = await api('/api/mailbox-detail?' + query.toString());
        const msg = data.message || {};
        const attachments = msg.attachments || [];
        const bodyHtml = msg.body_html
          ? `<div class="mail-detail-body">${msg.body_html}</div>`
          : `<div class="mail-detail-body plain">${escapeHtml(msg.body || msg.preview || '暂无可读取的文本正文')}</div>`;
        el('mailDetail').innerHTML = `
          <div class="mail-detail-head">
            <div class="mailbox-subject">${escapeHtml(msg.subject || '无主题')}</div>
            <div class="mailbox-meta">发件人：${escapeHtml(msg.from || '')}</div>
            <div class="mailbox-meta">收件人：${escapeHtml(msg.to || '')}</div>
            <div class="mailbox-meta">时间：${escapeHtml(msg.date || '')}</div>
            <div class="mail-attachment-list">
              ${attachments.length ? attachments.map(file => `
                <span class="mail-attachment">
                  <span class="icon" data-icon="paperclip"></span>
                  ${escapeHtml(file.name || '附件')}
                  ${file.size ? ` · ${formatFileSize(file.size)}` : ''}
                </span>
              `).join('') : '<span class="mailbox-meta">无附件</span>'}
            </div>
          </div>
          ${bodyHtml}
        `;
        renderIcons(el('mailDetail'));
      } catch (err) {
        el('mailDetail').textContent = err.message;
      }
    }

    function formatFileSize(bytes) {
      const size = Number(bytes || 0);
      if (size < 1024) return size + ' B';
      if (size < 1024 * 1024) return (size / 1024).toFixed(1) + ' KB';
      return (size / 1024 / 1024).toFixed(1) + ' MB';
    }

    function renderAssistantMailFiles() {
      const box = el('assistantMailFileList');
      box.innerHTML = state.assistantMailFiles.length ? state.assistantMailFiles.map((file, index) => `
        <div class="mail-file-item">
          <span>${escapeHtml(file.name)} · ${formatFileSize(file.size)}</span>
          <button class="mini secondary remove-assistant-file" type="button" data-index="${index}">移除</button>
        </div>
      `).join('') : '<div class="mailbox-meta">未添加附件</div>';
      box.querySelectorAll('.remove-assistant-file').forEach(button => {
        button.addEventListener('click', () => {
          state.assistantMailFiles.splice(Number(button.dataset.index), 1);
          renderAssistantMailFiles();
        });
      });
    }

    function renderAssistantMailPreview() {
      el('assistantMailPreview').innerHTML = textToHtml(el('assistantMailBody').value || '暂无正文内容');
    }

    function clearAssistantMail() {
      el('assistantMailTo').value = '';
      el('assistantMailCc').value = '';
      el('assistantMailSubject').value = '';
      el('assistantMailBody').value = '';
      el('assistantMailFiles').value = '';
      state.assistantMailFiles = [];
      renderAssistantMailFiles();
      renderAssistantMailPreview();
      el('assistantMailStatus').textContent = '';
      el('assistantMailStatus').className = 'status';
    }

    function adminConfigPayload() {
      return {
        assistant_api_url: el('configApiUrl').value,
        assistant_api_key: el('configApiKey').value,
        assistant_model: el('configModel').value || el('configModelSelect').value,
        assistant_prompt: el('configPrompt').value,
      };
    }

    async function boot() {
      renderIcons();
      try {
        const session = await api('/api/session');
        defaultAssistantPrompt = session.assistant_prompt || defaultAssistantPrompt;
        if (session.authenticated) {
          applyUser(session.user);
          await loadReports();
        } else {
          applyUser(null);
        }
      } catch (err) {
        applyUser(null);
        el('loginStatus').textContent = err.message;
        el('loginStatus').className = 'status err';
      }
    }

    async function loadDraft(kind, name) {
      state.selected = name;
      renderReports();
      const query = new URLSearchParams({ kind, file: name || '' });
      const draft = await api('/api/draft?' + query.toString());
      el('to').value = draft.to || '';
      el('cc').value = draft.cc || '';
      el('subject').value = draft.subject || '';
      el('body').value = draft.body || '';
      state.bodyHtml = draft.body_html || '';
      renderBodyPreview();
      el('attachment').value = draft.attachment || '';
      if (draft.download_url) {
        el('downloadLink').href = draft.download_url;
        el('downloadLink').classList.remove('hidden');
      } else {
        el('downloadLink').classList.add('hidden');
      }
      el('preview').innerHTML = draft.preview_html || textToHtml(draft.preview || '暂无可预览内容');
      el('status').textContent = '';
      el('status').className = 'status';
      if (kind === 'weekly' && !state.weeklyPrefilled) {
        await loadWeeklyPrefill();
      }
      if (kind === 'trip' && !state.tripPrefilled) {
        await loadTripPrefill();
      }
    }

    async function loadWeeklyPrefill(force = false) {
      if (!force && state.weeklyPrefilled) return;
      if (force) clearSavedFormDraft('weekly');
      setDefaultWeeklyDates();
      const prefill = await api('/api/weekly-prefill');
      renderWorkRows('summary', prefill.summary_rows || []);
      renderWorkRows('follow', prefill.follow_rows || []);
      renderWorkRows('next', prefill.next_rows || []);
      state.weeklyPrefilled = true;
      const restored = force ? false : restoreSavedFormDraft('weekly');
      if (prefill.source) {
        el('status').textContent = restored ? '已恢复上次未生成的周报草稿。' : `已获取最新历史周报：${prefill.source}。上次“下周计划”已写入本次“本周工作内容”，重点工作已复制，下周计划保持为空。`;
        el('status').className = 'status ok';
      } else {
        el('status').textContent = '没有找到可用于预填的历史周报。';
        el('status').className = 'status err';
      }
      saveFormDraft();
    }

    async function loadTripPrefill(force = false) {
      if (!force && state.tripPrefilled) return;
      if (force) clearSavedFormDraft('trip');
      const prefill = await api('/api/trip-prefill');
      const today = new Date();
      const fmt = d => d.getFullYear() + '-' + String(d.getMonth()+1).padStart(2,'0') + '-' + String(d.getDate()).padStart(2,'0');
      const defaultStart = fmt(today);
      const defaultEnd = fmt(new Date(today.getTime() + 2*24*60*60*1000));
      el('tripReporter').value = prefill.reporter || '周颖超';
      el('tripDepartment').value = prefill.department || '场景研究院';
      el('tripLocation').value = prefill.location || '';
      el('tripStart').value = prefill.trip_start || defaultStart;
      el('tripEnd').value = prefill.trip_end || defaultEnd;
      el('tripPurpose').value = prefill.purpose || '';
      el('tripItinerary').value = prefill.itinerary || '';
      el('tripDetails').value = prefill.details || '';
      el('tripIssues').value = prefill.issues || '';
      el('tripSuggestions').value = prefill.suggestions || '';
      renderTripCards();
      state.tripPrefilled = true;
      const restored = force ? false : restoreSavedFormDraft('trip');
      if (prefill.source) {
        el('status').textContent = restored ? '已恢复上次未生成的出差报告草稿。' : `已获取最新历史出差报告：${prefill.source}，并自动填入当前模板。`;
        el('status').className = 'status ok';
      } else {
        el('status').textContent = '没有找到可用于预填的历史出差报告。';
        el('status').className = 'status err';
      }
      saveFormDraft();
    }

    el('mailKind').addEventListener('change', async () => {
      const kind = el('mailKind').value;
      const generated = state.reports.find(r => r.kind === kind && r.generated);
      const latest = generated?.name || (kind === 'weekly' ? window.latestWeekly : window.latestTrip);
      await loadDraft(kind, latest || '');
    });

    async function loadAgentOrchestration() {
      const panel = el('agentOrchestrationPanel');
      const content = el('agentOrchestrationContent');
      if (!panel.classList.contains('hidden')) {
        panel.classList.add('hidden');
        return;
      }
      content.innerHTML = '<div class="upload-item">正在加载编排逻辑...</div>';
      panel.classList.remove('hidden');
      try {
        const data = await api('/api/agent-orchestration');
        if (!data.ok) throw new Error(data.error || '加载失败');
        const agents = data.agents || {};
        const workflows = data.workflows || {};
        const skills = data.skills || [];
        let html = '';
        html += '<div class="orchestration-section"><div class="orchestration-section-title">🧠 犇犇角色定义与系统提示词</div>';
        Object.entries(agents).forEach(([key, val]) => {
          const label = { weekly: '周报助手', trip: '出差报告助手', diary: '日记助手', mailassistant: '邮件助手', news: '资讯助手', forum: '论坛助手', dashboard: '总助手' }[key] || key;
          html += `<div style="margin-bottom:8px;font-size:12px;font-weight:700;color:#475569;">${label}</div><pre>${escapeHtml(val)}</pre>`;
        });
        html += '</div>';
        html += `<div class="orchestration-section"><div class="orchestration-section-title">🔄 Skill 模式追加提示词</div><pre>${escapeHtml(data.skill_mode_suffix || '')}</pre></div>`;
        html += '<div class="orchestration-section"><div class="orchestration-section-title">📋 工作流编排</div>';
        Object.entries(workflows).forEach(([key, val]) => {
          const label = { weekly: '周报', trip: '出差报告', diary: '工作日记', mailassistant: '邮件', news: '资讯', forum: '金点子论坛' }[key] || key;
          html += `<div style="margin-bottom:6px;font-size:12px;"><strong>${label}：</strong>${escapeHtml(val)}</div>`;
        });
        html += '</div>';
        html += `<div class="orchestration-section"><div class="orchestration-section-title">🛠 可用 Skill 列表</div><pre>${escapeHtml(JSON.stringify(skills.map(s => ({ name: s.name, module: s.module, title: s.title, safe: s.safe })), null, 2))}</pre></div>`;
        content.innerHTML = html;
      } catch (err) {
        content.innerHTML = `<div class="upload-item">${escapeHtml(err.message)}</div>`;
      }
    }
    let agentConfigData = {};
    function openAgentConfigModal() {
      el('agentConfigModal').classList.remove('hidden');
      loadAgentConfigEditor();
    }
    function closeAgentConfigModal() {
      el('agentConfigModal').classList.add('hidden');
    }
    async function loadAgentConfigEditor() {
      el('agentConfigStatus').textContent = '正在加载...';
      try {
        const data = await api('/api/agent-config');
        if (!data.ok) throw new Error(data.error || '加载失败');
        agentConfigData = data.config || {};
        const prompts = agentConfigData.prompts || {};
        const workflows = agentConfigData.workflows || {};
        const labels = { weekly: '周报助手', trip: '出差报告助手', diary: '日记助手', mailassistant: '邮件助手', news: '资讯助手', forum: '论坛助手', dashboard: '总助手' };
        let pHtml = '';
        Object.entries(labels).forEach(([key, label]) => {
          pHtml += `<div style="margin-bottom:10px;"><label style="font-size:12px;font-weight:700;color:#475569;">${label}</label><textarea id="agentPrompt-${key}" rows="6" spellcheck="false" style="width:100%;margin-top:4px;font-size:12px;">${escapeHtml(prompts[key] || '')}</textarea></div>`;
        });
        el('agentConfigPrompts').innerHTML = pHtml;
        let wHtml = '';
        Object.entries(labels).forEach(([key, label]) => {
          if (key === 'dashboard') return;
          wHtml += `<div style="margin-bottom:10px;"><label style="font-size:12px;font-weight:700;color:#475569;">${label}</label><textarea id="agentWorkflow-${key}" rows="3" spellcheck="false" style="width:100%;margin-top:4px;font-size:12px;">${escapeHtml(workflows[key] || '')}</textarea></div>`;
        });
        el('agentConfigWorkflows').innerHTML = wHtml;
        el('agentConfigStatus').textContent = '';
      } catch (err) {
        el('agentConfigStatus').textContent = err.message;
      }
    }
    async function saveAgentConfig() {
      el('agentConfigStatus').textContent = '保存中...';
      const prompts = {};
      const workflows = {};
      const keys = ['weekly','trip','diary','mailassistant','news','forum','dashboard'];
      keys.forEach(key => {
        const ta = el(`agentPrompt-${key}`);
        if (ta && ta.value.trim()) prompts[key] = ta.value.trim();
      });
      ['weekly','trip','diary','mailassistant','news','forum'].forEach(key => {
        const ta = el(`agentWorkflow-${key}`);
        if (ta && ta.value.trim()) workflows[key] = ta.value.trim();
      });
      try {
        const result = await api('/api/agent-config', {
          method: 'POST',
          headers: { 'Content-Type': 'application/json' },
          body: JSON.stringify({ prompts, workflows })
        });
        if (!result.ok) throw new Error(result.error || '保存失败');
        el('agentConfigStatus').textContent = result.message || '保存成功';
      } catch (err) {
        el('agentConfigStatus').textContent = err.message;
      }
    }
    if (el('openAgentOrchestration')) el('openAgentOrchestration').addEventListener('click', loadAgentOrchestration);
    if (el('openAgentConfig')) el('openAgentConfig').addEventListener('click', openAgentConfigModal);
    if (el('agentConfigClose')) el('agentConfigClose').addEventListener('click', closeAgentConfigModal);
    if (el('agentConfigModal')) el('agentConfigModal').addEventListener('click', event => { if (event.target === el('agentConfigModal')) closeAgentConfigModal(); });
    if (el('agentConfigSave')) el('agentConfigSave').addEventListener('click', saveAgentConfig);
    document.querySelectorAll('.agent-config-tab').forEach(tab => {
      tab.addEventListener('click', () => {
        document.querySelectorAll('.agent-config-tab').forEach(t => t.classList.remove('active'));
        tab.classList.add('active');
        const name = tab.dataset.tab;
        el('agentConfigPrompts').classList.toggle('hidden', name !== 'prompts');
        el('agentConfigWorkflows').classList.toggle('hidden', name !== 'workflows');
      });
    });
    el('openSkillDocs').addEventListener('click', () => window.open('/skill-docs', '_blank'));
    el('downloadSkillDocs').addEventListener('click', () => window.open('/download-skill-doc', '_blank'));
    el('skillSearch').addEventListener('input', () => renderSkills(state.skills || []));
    el('skillModuleFilter').addEventListener('change', () => renderSkills(state.skills || []));
    el('skillTestClose').addEventListener('click', closeSkillTest);
    el('skillRunTest').addEventListener('click', runSkillTest);
    el('skillTestModal').addEventListener('click', event => {
      if (event.target === el('skillTestModal')) closeSkillTest();
    });
    attachSkillCardTests(el('skillsPanel'));

    el('refreshMailbox').addEventListener('click', () => loadMailbox(true));
    el('mailboxLimit').addEventListener('change', loadMailbox);
    el('assistantMailBody').addEventListener('input', renderAssistantMailPreview);
    el('assistantMailFiles').addEventListener('change', () => {
      state.assistantMailFiles = [...el('assistantMailFiles').files].map(file => ({
        file,
        name: file.name,
        size: file.size,
        type: file.type || 'application/octet-stream'
      }));
      renderAssistantMailFiles();
    });
    el('clearAssistantMail').addEventListener('click', clearAssistantMail);
    el('sendAssistantMail').addEventListener('click', async () => {
      el('assistantMailStatus').textContent = '正在发送邮件...';
      el('assistantMailStatus').className = 'status';
      try {
        const attachments = [];
        for (const item of state.assistantMailFiles) {
          attachments.push({
            name: item.name,
            type: item.type,
            content: await readFileAsBase64(item.file)
          });
        }
        const result = await api('/api/mail-send', {
          method: 'POST',
          headers: { 'Content-Type': 'application/json' },
          body: JSON.stringify({
            to: el('assistantMailTo').value,
            cc: el('assistantMailCc').value,
            subject: el('assistantMailSubject').value,
            body: el('assistantMailBody').value,
            attachments
          })
        });
        el('assistantMailStatus').textContent = result.message || '邮件已发送。';
        el('assistantMailStatus').className = 'status ok';
      } catch (err) {
        el('assistantMailStatus').textContent = err.message;
        el('assistantMailStatus').className = 'status err';
      }
    });

    el('generate').addEventListener('click', async () => {
      const kind = el('kind').value;
      if (!['weekly', 'trip'].includes(kind) || !['weekly', 'trip'].includes(state.task)) {
        el('generateToolbar').classList.add('hidden');
        return;
      }
      syncWeeklyPeriod();
      el('status').textContent = '正在按模板生成文件...';
      el('status').className = 'status';
      const payload = kind === 'weekly' ? {
        kind,
        period: el('weeklyPeriod').value,
        weekly_summary: collectWorkRows('summary'),
        weekly_follow: collectWorkRows('follow'),
        weekly_next: collectWorkRows('next')
      } : {
        kind,
        reporter: el('tripReporter').value,
        department: el('tripDepartment').value,
        location: el('tripLocation').value,
        trip_start: el('tripStart').value,
        trip_end: el('tripEnd').value,
        purpose: el('tripPurpose').value,
        itinerary: el('tripItinerary').value,
        details: el('tripDetails').value,
        issues: el('tripIssues').value,
        suggestions: el('tripSuggestions').value
      };
      try {
        const result = await api('/api/generate', {
          method: 'POST',
          headers: { 'Content-Type': 'application/json' },
          body: JSON.stringify(payload)
        });
        clearSavedFormDraft(kind);
        await loadReports();
        el('kind').value = kind;
        await loadDraft(kind, result.file);
        el('status').textContent = '已生成标准文件，并设为当前邮件附件：' + result.file;
        el('status').className = 'status ok';
      } catch (err) {
        el('status').textContent = err.message;
        el('status').className = 'status err';
      }
    });

    document.querySelectorAll('.task-card').forEach(card => {
      card.addEventListener('click', async () => {
        const task = card.dataset.task;
        state.subTab = 'edit';
        setTask(task);
        if (task === 'weekly') {
          await loadWeeklyPrefill();
          const generated = state.reports.find(r => r.kind === 'weekly' && r.generated);
          const latest = generated?.name || window.latestWeekly;
          await loadDraft('weekly', latest || '');
        } else if (task === 'trip') {
          await loadTripPrefill();
          const generated = state.reports.find(r => r.kind === 'trip' && r.generated);
          const latest = generated?.name || window.latestTrip;
          await loadDraft('trip', latest || '');
        } else if (task === 'mailconfig') {
          loadMailConfig();
        } else if (task === 'mailassistant') {
          loadMailbox();
        } else if (task === 'config') {
          await loadAdminConfig();
        } else if (task === 'skills') {
          await loadSkills();
        } else if (task === 'usermanage') {
          await loadUserManage();
        }
      });
    });

    document.querySelectorAll('.sub-tab').forEach(tab => {
      tab.addEventListener('click', () => {
        setSubTab(tab.dataset.sub);
      });
    });

    el('loginButton').addEventListener('click', async () => {
      el('loginStatus').textContent = '正在登录...';
      el('loginStatus').className = 'status';
      try {
        const result = await api('/api/login', {
          method: 'POST',
          headers: { 'Content-Type': 'application/json' },
          body: JSON.stringify({ username: el('loginUser').value, password: el('loginPass').value })
        });
        defaultAssistantPrompt = result.assistant_prompt || defaultAssistantPrompt;
        applyUser(result.user);
        await loadReports();
      } catch (err) {
        el('loginStatus').textContent = err.message;
        el('loginStatus').className = 'status err';
      }
    });

    el('changePassButton').addEventListener('click', () => {
      el('passOld').value = '';
      el('passNew').value = '';
      el('passConfirm').value = '';
      el('passStatus').textContent = '';
      el('passModal').classList.remove('hidden');
    });
    el('passSave').addEventListener('click', async () => {
      const oldPass = el('passOld').value;
      const newPass = el('passNew').value;
      const confirm = el('passConfirm').value;
      if (!oldPass || !newPass) {
        el('passStatus').textContent = '请填写原密码和新密码';
        el('passStatus').className = 'status err';
        return;
      }
      if (newPass.length < 4) {
        el('passStatus').textContent = '新密码至少 4 位';
        el('passStatus').className = 'status err';
        return;
      }
      if (newPass !== confirm) {
        el('passStatus').textContent = '两次输入的新密码不一致';
        el('passStatus').className = 'status err';
        return;
      }
      try {
        await apiPost('/api/change-password', { old_password: oldPass, new_password: newPass });
        el('passStatus').textContent = '密码修改成功，请重新登录';
        el('passStatus').className = 'status ok';
        setTimeout(() => { el('passModal').classList.add('hidden'); }, 1500);
      } catch (err) {
        el('passStatus').textContent = err.message;
        el('passStatus').className = 'status err';
      }
    });
    el('logoutButton').addEventListener('click', async () => {
      await api('/api/logout', { method: 'POST', headers: { 'Content-Type': 'application/json' }, body: '{}' });
      applyUser(null);
    });

    el('saveConfig').addEventListener('click', async () => {
      el('status').textContent = '正在保存系统配置...';
      el('status').className = 'status';
      try {
        const result = await api('/api/admin-config', {
          method: 'POST',
          headers: { 'Content-Type': 'application/json' },
          body: JSON.stringify(adminConfigPayload())
        });
        defaultAssistantPrompt = result.config.assistant_prompt || defaultAssistantPrompt;
        el('configApiKey').value = '';
        el('configKeyHint').textContent = 'API Key 状态：' + (result.config.assistant_api_key_masked || '未配置');
        el('status').textContent = '系统配置已保存。';
        el('status').className = 'status ok';
      } catch (err) {
        el('status').textContent = err.message;
        el('status').className = 'status err';
      }
    });

    el('saveServerConfig').addEventListener('click', async () => {
      el('configServerStatus').textContent = '正在保存邮件服务器配置...';
      el('configServerStatus').className = 'status';
      try {
        await apiPost('/api/server-config', {
          smtp_host: el('configSmtpHost').value,
          smtp_port: el('configSmtpPort').value,
          smtp_tls: el('configSmtpTls').checked,
          smtp_ssl: el('configSmtpSsl').checked,
          imap_host: el('configImapHost').value,
          imap_port: el('configImapPort').value,
          imap_ssl: el('configImapSsl').checked,
        });
        el('configServerStatus').textContent = '邮件服务器配置已保存（对所有用户生效）。';
        el('configServerStatus').className = 'status ok';
      } catch (err) {
        el('configServerStatus').textContent = err.message;
        el('configServerStatus').className = 'status err';
      }
    });

    el('saveMailConfig').addEventListener('click', async () => {
      el('mailConfigStatus').textContent = '正在保存邮件配置...';
      el('mailConfigStatus').className = 'status';
      try {
        const result = await api('/api/mail-config', {
          method: 'POST',
          headers: { 'Content-Type': 'application/json' },
          body: JSON.stringify(mailConfigPayload())
        });
        el('mailSmtpPassword').value = '';
        el('mailImapPassword').value = '';
        el('mailPasswordHint').textContent = 'SMTP 密码/授权码状态：' + (result.mail_config.smtp_password_masked || '未配置');
        el('mailImapPasswordHint').textContent = 'IMAP 密码/授权码状态：' + (result.mail_config.imap_password_masked || '未配置');
        el('mailConfigStatus').textContent = '邮件配置已保存。';
        el('mailConfigStatus').className = 'status ok';
      } catch (err) {
        el('mailConfigStatus').textContent = explainMailLoginError(err.message);
        el('mailConfigStatus').className = 'status err';
      }
    });

    el('testMailConfig').addEventListener('click', async () => {
      el('mailConfigStatus').textContent = '正在测试邮箱配置...';
      el('mailConfigStatus').className = 'status';
      try {
        await api('/api/mail-config', {
          method: 'POST',
          headers: { 'Content-Type': 'application/json' },
          body: JSON.stringify(mailConfigPayload())
        });
        const result = await api('/api/test-mail-config', {
          method: 'POST',
          headers: { 'Content-Type': 'application/json' },
          body: JSON.stringify({})
        });
        el('mailSmtpPassword').value = '';
        el('mailImapPassword').value = '';
        el('mailConfigStatus').textContent = result.message || '邮箱配置测试成功。';
        el('mailConfigStatus').className = 'status ok';
        await loadMailConfig();
      } catch (err) {
        el('mailConfigStatus').textContent = err.message;
        el('mailConfigStatus').className = 'status err';
      }
    });

    el('configModelSelect').addEventListener('change', () => {
      el('configModel').value = el('configModelSelect').value;
    });

    el('loadModels').addEventListener('click', async () => {
      el('configTestStatus').textContent = '正在获取模型列表...';
      el('configTestStatus').className = 'status';
      try {
        const result = await api('/api/admin-models', {
          method: 'POST',
          headers: { 'Content-Type': 'application/json' },
          body: JSON.stringify(adminConfigPayload())
        });
        renderModelOptions(result.models, el('configModel').value || result.models[0]);
        el('configModel').value = el('configModelSelect').value;
        el('configTestStatus').textContent = result.warning || `已获取 ${result.models.length} 个模型。`;
        el('configTestStatus').className = result.warning ? 'status err' : 'status ok';
      } catch (err) {
        el('configTestStatus').textContent = err.message;
        el('configTestStatus').className = 'status err';
      }
    });

    el('testModel').addEventListener('click', async () => {
      el('configTestStatus').textContent = '正在测试 API Key 和模型...';
      el('configTestStatus').className = 'status';
      try {
        const result = await api('/api/admin-test-model', {
          method: 'POST',
          headers: { 'Content-Type': 'application/json' },
          body: JSON.stringify(adminConfigPayload())
        });
        el('configTestStatus').textContent = `测试成功：${result.model} 返回 ${result.message}`;
        el('configTestStatus').className = 'status ok';
      } catch (err) {
        el('configTestStatus').textContent = '测试失败：' + err.message;
        el('configTestStatus').className = 'status err';
      }
    });

    el('addUser').addEventListener('click', async () => {
      const isSuper = state.user?.is_superadmin;
      const statusEl = el('userManageStatus') || el('configTestStatus');
      statusEl.textContent = '正在新增用户...';
      statusEl.className = 'status';
      try {
        const payload = {
          username: el('newUserName').value,
          password: el('newUserPassword').value,
          name: el('newDisplayName').value
        };
        if (isSuper) payload.role = el('newUserRole').value;
        const result = await apiPost('/api/admin-users', payload);
        renderUsers(result.users || []);
        el('newUserName').value = '';
        el('newDisplayName').value = '';
        el('newUserPassword').value = '';
        statusEl.textContent = '用户已新增。';
        statusEl.className = 'status ok';
      } catch (err) {
        statusEl.textContent = err.message;
        statusEl.className = 'status err';
      }
    });

    el('addSummary').addEventListener('click', () => addWorkRow('summary'));
    el('addFollow').addEventListener('click', () => addWorkRow('follow'));
    el('addNext').addEventListener('click', () => addWorkRow('next'));
    el('loadLatestWeekly').addEventListener('click', () => loadWeeklyPrefill(true));
    el('loadLatestTrip').addEventListener('click', () => loadTripPrefill(true));
    el('weeklyStart').addEventListener('change', () => {
      syncWeeklyPeriod();
      el('diarySumStart').value = el('weeklyStart').value;
    });
    el('weeklyEnd').addEventListener('change', () => {
      syncWeeklyPeriod();
      el('diarySumEnd').value = el('weeklyEnd').value;
    });
    window.addEventListener('beforeunload', () => {
      saveCurrentModal();
      saveFormDraft();
    });
    el('historyKind').addEventListener('change', renderHistoryReports);
    if (el('profileButton')) el('profileButton').addEventListener('click', openProfileModal);
    if (el('profileClose')) el('profileClose').addEventListener('click', closeProfileModal);
    if (el('profileModal')) el('profileModal').addEventListener('click', event => {
      if (event.target === el('profileModal')) closeProfileModal();
    });
    if (el('profileAvatarFile')) el('profileAvatarFile').addEventListener('change', () => {
      const file = el('profileAvatarFile').files[0];
      if (file) el('profileAvatarPreview').src = URL.createObjectURL(file);
    });
    if (el('profileUseAssistantAvatar')) el('profileUseAssistantAvatar').addEventListener('click', () => saveProfile('assistant'));
    if (el('profileSave')) el('profileSave').addEventListener('click', () => saveProfile());

    el('uploadButton').addEventListener('click', async () => {
      const selected = [...el('uploadFiles').files];
      if (!selected.length) {
        el('status').textContent = '请先选择要上传的历史报告文件。';
        el('status').className = 'status err';
        return;
      }
      el('uploadButton').disabled = true;
      el('uploadButton').textContent = '上传中...';
      el('status').textContent = '正在读取并上传文件...';
      el('status').className = 'status';
      try {
        const files = [];
        for (const file of selected) {
          files.push({ name: file.name, data: await readFileAsBase64(file) });
        }
        const result = await api('/api/upload-history', {
          method: 'POST',
          headers: { 'Content-Type': 'application/json' },
          body: JSON.stringify({ kind: el('uploadKind').value, files })
        });
        el('uploadList').innerHTML = result.uploaded.map(item => `
          <div class="upload-item">已上传：${escapeHtml(item.name)}</div>
        `).join('');
        el('status').textContent = `上传完成，共 ${result.uploaded.length} 个文件。`;
        el('status').className = 'status ok';
        el('uploadFiles').value = '';
        await loadReports({ preserveSelection: true });
        if (state.task === 'weekly' || state.task === 'trip') {
          renderHistoryReports();
          renderReports();
        } else {
          setTask('weekly');
        }
      } catch (err) {
        el('status').textContent = err.message;
        el('status').className = 'status err';
      } finally {
        el('uploadButton').disabled = false;
        el('uploadButton').textContent = '上传到历史报告库';
      }
    });

    // ===== 工作日记 =====
    function initDiaryPanel() {
      const today = new Date().toISOString().split('T')[0];
      el('diaryDate').value = today;
      loadDiaryToForm(today);
      renderDiaryList();
    }
    function setDiaryTab(tab) {
      document.querySelectorAll('.diary-tab').forEach(t => t.classList.toggle('active', t.dataset.diarytab === tab));
      el('diaryWriteView').classList.toggle('hidden', tab !== 'write');
      el('diaryBrowseView').classList.toggle('hidden', tab !== 'browse');
      if (tab === 'browse') renderDiaryList();
    }
    async function loadDiaryToForm(dateStr) {
      try {
        const data = await api('/api/diary/get?date=' + encodeURIComponent(dateStr));
        const d = data.diary || {};
        el('diaryTodayWork').value = d.today_work || '';
        el('diaryTomorrowPlan').value = d.tomorrow_plan || '';
        el('diaryThoughts').value = d.thoughts || '';
      } catch (err) {
        el('diaryTodayWork').value = '';
        el('diaryTomorrowPlan').value = '';
        el('diaryThoughts').value = '';
      }
    }
    async function saveDiary() {
      const payload = {
        date: el('diaryDate').value,
        today_work: el('diaryTodayWork').value,
        tomorrow_plan: el('diaryTomorrowPlan').value,
        thoughts: el('diaryThoughts').value,
      };
      if (!payload.date) {
        el('diaryStatus').textContent = '请选择日期';
        el('diaryStatus').className = 'status err';
        return;
      }
      try {
        el('diaryStatus').textContent = '保存中...';
        el('diaryStatus').className = 'status';
        await apiPost('/api/diary/save', payload);
        el('diaryStatus').textContent = '保存成功';
        el('diaryStatus').className = 'status ok';
        setTimeout(() => { el('diaryStatus').textContent = ''; }, 2000);
      } catch (err) {
        el('diaryStatus').textContent = err.message;
        el('diaryStatus').className = 'status err';
      }
    }
    async function renderDiaryList() {
      const list = el('diaryList');
      list.innerHTML = '<div class="diary-empty">加载中...</div>';
      try {
        const data = await api('/api/diary/list?limit=100');
        const items = data.diaries || [];
        if (!items.length) {
          list.innerHTML = '<div class="diary-empty">暂无日记，去记录第一篇吧～</div>';
          return;
        }
        list.innerHTML = items.map(item => `
          <div class="diary-item" data-date="${escapeHtml(item.date)}">
            <div class="diary-item-date">${escapeHtml(item.date)}</div>
            <div class="diary-item-preview">${escapeHtml(item.today_work_preview || '无内容')}</div>
            <span class="icon" data-icon="chevron-right"></span>
          </div>
        `).join('');
        renderIcons(list);
        list.querySelectorAll('.diary-item').forEach(item => {
          item.addEventListener('click', () => openDiaryDetail(item.dataset.date));
        });
      } catch (err) {
        list.innerHTML = `<div class="diary-empty">加载失败：${escapeHtml(err.message)}</div>`;
      }
    }
    async function openDiaryDetail(dateStr) {
      try {
        const data = await api('/api/diary/get?date=' + encodeURIComponent(dateStr));
        const d = data.diary || {};
        el('diaryDetailDate').textContent = (d.date || dateStr) + ' 日记详情';
        el('diaryDetailToday').textContent = d.today_work || '（无内容）';
        el('diaryDetailTomorrow').textContent = d.tomorrow_plan || '（无内容）';
        el('diaryDetailThoughts').textContent = d.thoughts || '（无内容）';
        el('diaryDetailEdit').onclick = () => {
          el('diaryDetailModal').classList.add('hidden');
          setDiaryTab('write');
          el('diaryDate').value = dateStr;
          loadDiaryToForm(dateStr);
        };
        el('diaryDetailDelete').onclick = async () => {
          if (!confirm('确定删除 ' + dateStr + ' 的日记吗？')) return;
          try {
            await apiPost('/api/diary/delete', { date: dateStr });
            el('diaryDetailModal').classList.add('hidden');
            renderDiaryList();
          } catch (err) {
            alert('删除失败：' + err.message);
          }
        };
        el('diaryDetailModal').classList.remove('hidden');
      } catch (err) {
        alert('读取日记失败：' + err.message);
      }
    }

    // 日记事件绑定
    document.querySelectorAll('.diary-tab').forEach(tab => {
      tab.addEventListener('click', () => setDiaryTab(tab.dataset.diarytab));
    });
    el('diaryLoadToday').addEventListener('click', () => {
      const today = new Date().toISOString().split('T')[0];
      el('diaryDate').value = today;
      loadDiaryToForm(today);
    });
    el('diaryDate').addEventListener('change', () => loadDiaryToForm(el('diaryDate').value));
    el('diarySaveButton').addEventListener('click', saveDiary);
    el('diaryClearButton').addEventListener('click', () => {
      el('diaryTodayWork').value = '';
      el('diaryTomorrowPlan').value = '';
      el('diaryThoughts').value = '';
      el('diaryStatus').textContent = '';
    });
    function openDiaryEditor(fieldId, title) {
      const value = el(fieldId).value;
      openModal(title, [
        { key: 'content', label: title, multiline: true, value: value }
      ], (values) => {
        el(fieldId).value = values.content || '';
        saveDiary();
      });
    }
    el('diaryTodayWork').addEventListener('click', () => openDiaryEditor('diaryTodayWork', '今日工作内容'));
    el('diaryTomorrowPlan').addEventListener('click', () => openDiaryEditor('diaryTomorrowPlan', '明日工作计划'));
    el('diaryThoughts').addEventListener('click', () => openDiaryEditor('diaryThoughts', '思路与想法'));
    el('diaryAgentButton').addEventListener('click', () => startAgent('diary'));
    el('diaryRefreshList').addEventListener('click', renderDiaryList);

    // 工作日记智能总结 → 周报
    async function summarizeDiariesForWeekly() {
      const start = el('diarySumStart').value;
      const end = el('diarySumEnd').value;
      if (!start || !end) {
        el('status').textContent = '请选择日记日期范围';
        el('status').className = 'status err';
        return;
      }
      if (start > end) {
        el('status').textContent = '开始日期不能晚于结束日期';
        el('status').className = 'status err';
        return;
      }
      el('status').textContent = '正在总结工作日记，请稍候...';
      el('status').className = 'status';
      try {
        const data = await apiPost('/api/diary/summarize', { start_date: start, end_date: end });
        if (!data.ok) {
          el('status').textContent = data.error || '总结失败';
          el('status').className = 'status err';
          return;
        }
        if (data.mode === 'empty') {
          el('status').textContent = data.warning || '该范围内无日记';
          el('status').className = 'status err';
          return;
        }
        parseAndApplyDiarySummary(data.summary || '');
        el('status').textContent = '日记总结已应用到周报，请检查并调整';
        el('status').className = 'status ok';
      } catch (err) {
        el('status').textContent = err.message;
        el('status').className = 'status err';
      }
    }
    function parseAndApplyDiarySummary(summary) {
      // 解析 AI 返回的总结文本，按三个部分填充
      const sections = { '本周工作总结': [], '重点工作跟进': [], '下周工作计划': [] };
      let current = null;
      summary.split('\\n').forEach(line => {
        const trimmed = line.trim();
        if (!trimmed) return;
        if (trimmed.includes('本周工作总结') || trimmed.includes('本周工作')) {
          current = '本周工作总结';
        } else if (trimmed.includes('重点工作跟进') || trimmed.includes('重点工作')) {
          current = '重点工作跟进';
        } else if (trimmed.includes('下周工作计划') || trimmed.includes('下周工作')) {
          current = '下周工作计划';
        } else if (current && /^[\d一二三四五六七八九十]+[.、.．\s]/.test(trimmed)) {
          const content = trimmed.replace(/^[\d一二三四五六七八九十]+[.、.．\s]+/, '').trim();
          if (content) sections[current].push(content);
        } else if (current && trimmed.startsWith('- ')) {
          const content = trimmed.replace(/^- /, '').trim();
          if (content) sections[current].push(content);
        } else if (current && trimmed.startsWith('* ')) {
          const content = trimmed.replace(/^\* /, '').trim();
          if (content) sections[current].push(content);
        }
      });
      // 填充到周报表单
      const map = {
        '本周工作总结': 'summary',
        '重点工作跟进': 'follow',
        '下周工作计划': 'next'
      };
      Object.entries(map).forEach(([title, key]) => {
        const items = sections[title] || [];
        // 清空现有
        state[`weekly_${key}`] = [];
        items.forEach(content => {
          state[`weekly_${key}`].push({
            category: '',
            content: content,
            ...(key === 'summary' ? { status: '已完成' } : key === 'follow' ? { progress: '推进中' } : { difficulty: '正常' })
          });
        });
        renderWeeklyRows(key);
      });
      saveFormDraft();
    }
    el('diarySummarizeBtn').addEventListener('click', summarizeDiariesForWeekly);

    // ===== 金点子论坛 =====
    function forumSourceName(source) {
      if (source === 'ai') return '智能体';
      return '成员发起';
    }
    async function loadForumTopics(selectFirst = false) {
      const list = el('forumTopicList');
      list.innerHTML = '<div class="forum-empty">加载话题中...</div>';
      try {
        const data = await api('/api/forum/topics');
        const topics = data.topics || [];
        el('forumTopicCount').textContent = `${topics.length} 个话题`;
        if (!topics.length) {
          list.innerHTML = '<div class="forum-empty">暂无话题，先发起一个金点子吧。</div>';
          el('forumTopicDetail').innerHTML = '<div class="forum-empty">选择一个话题查看详情和讨论。</div>';
          return;
        }
        list.innerHTML = topics.map(topic => `
          <div class="forum-topic-item ${state.forumSelected === topic.id ? 'active' : ''}" data-id="${escapeHtml(topic.id)}">
            <div class="forum-topic-title">${escapeHtml(topic.title)}</div>
            <div class="forum-topic-meta">${escapeHtml(forumSourceName(topic.source))} · ${escapeHtml(topic.author || '')} · ${escapeHtml(topic.created_at || '')}</div>
            <div class="forum-topic-stats">
              <span class="forum-stat">热度 ${topic.heat || 0}</span>
              <span class="forum-stat">赞 ${topic.like_count || 0}</span>
              <span class="forum-stat">评 ${topic.comment_count || 0}</span>
              <span class="forum-stat">看 ${topic.view_count || 0}</span>
            </div>
          </div>
        `).join('');
        list.querySelectorAll('.forum-topic-item').forEach(item => {
          item.addEventListener('click', () => openForumTopic(item.dataset.id));
        });
        if (selectFirst || !state.forumSelected) {
          openForumTopic(topics[0].id);
        }
      } catch (err) {
        list.innerHTML = `<div class="forum-empty">加载失败：${escapeHtml(err.message)}</div>`;
      }
    }
    async function openForumTopic(topicId) {
      state.forumSelected = topicId;
      state.forumCommentPage = 1;
      document.querySelectorAll('.forum-topic-item').forEach(item => {
        item.classList.toggle('active', item.dataset.id === topicId);
      });
      const detail = el('forumTopicDetail');
      detail.innerHTML = '<div class="forum-empty">读取话题中...</div>';
      try {
        const data = await api('/api/forum/topic?id=' + encodeURIComponent(topicId));
        renderForumTopic(data.topic);
      } catch (err) {
        detail.innerHTML = `<div class="forum-empty">读取失败：${escapeHtml(err.message)}</div>`;
      }
    }
    function forumCommentTree(comments) {
      const byParent = {};
      comments.forEach(comment => {
        const key = comment.parent_id || '';
        (byParent[key] ||= []).push(comment);
      });
      return byParent;
    }
    function forumCommentHtml(comment, children = []) {
      return `
        <div class="forum-comment ${comment.parent_id ? 'reply' : ''}" data-comment-id="${escapeHtml(comment.id || '')}" data-author="${escapeHtml(comment.author || '')}">
          <div class="forum-comment-meta">${escapeHtml(comment.author || '')} · ${escapeHtml(comment.created_at || '')}</div>
          <div class="forum-comment-body">${escapeHtml(comment.content || '')}</div>
          <div class="forum-comment-actions">
            <button class="mini secondary forum-reply-button" type="button">回复</button>
          </div>
        </div>
        ${children.map(child => forumCommentHtml(child, [])).join('')}`;
    }
    function renderForumTopic(topic) {
      const comments = topic.comments || [];
      const pageSize = 10;
      const tree = forumCommentTree(comments);
      const topComments = tree[''] || [];
      const totalPages = Math.max(1, Math.ceil(topComments.length / pageSize));
      state.forumCommentPage = Math.min(Math.max(1, state.forumCommentPage || 1), totalPages);
      const pageTopComments = topComments.slice((state.forumCommentPage - 1) * pageSize, state.forumCommentPage * pageSize);
      el('forumTopicDetail').innerHTML = `
        <div class="forum-topic-title" style="font-size:18px;">${escapeHtml(topic.title)}</div>
        <div class="forum-topic-meta">${escapeHtml(forumSourceName(topic.source))} · ${escapeHtml(topic.author || '')} · ${escapeHtml(topic.created_at || '')}</div>
        <div class="forum-topic-stats">
          <span class="forum-stat">热度 ${topic.heat || 0}</span>
          <span class="forum-stat">点赞 ${topic.like_count || 0}</span>
          <span class="forum-stat">评论 ${topic.comment_count || 0}</span>
          <span class="forum-stat">浏览 ${topic.view_count || 0}</span>
        </div>
        <div class="forum-topic-body" style="margin-top:12px;">${escapeHtml(topic.body)}</div>
        <div class="toolbar">
          <button type="button" id="forumLikeButton" class="secondary"><span class="icon" data-icon="thumbs-up"></span> 点赞</button>
        </div>
        <div class="forum-comments">
          ${pageTopComments.length ? pageTopComments.map(comment => forumCommentHtml(comment, tree[comment.id] || [])).join('') : '<div class="forum-empty">还没有讨论，来写第一条观点。</div>'}
        </div>
        <div class="forum-pagination">
          <button type="button" class="mini secondary" id="forumPrevPage" ${state.forumCommentPage <= 1 ? 'disabled' : ''}>上一页</button>
          <span>第 ${state.forumCommentPage} / ${totalPages} 页 · 共 ${comments.length} 条评论</span>
          <button type="button" class="mini secondary" id="forumNextPage" ${state.forumCommentPage >= totalPages ? 'disabled' : ''}>下一页</button>
        </div>
        <label>参与讨论</label>
        <textarea id="forumCommentInput" placeholder="写下你的观点、建议、风险提醒或下一步行动..."></textarea>
        <input id="forumCommentParent" type="hidden" value="" />
        <div class="toolbar">
          <button type="button" id="forumCommentButton">发布讨论</button>
          <button type="button" class="secondary" id="forumAiCommentButton"><span class="icon" data-icon="bot"></span> AI 潜水评论</button>
          <span id="forumCommentStatus" class="status"></span>
        </div>
      `;
      renderIcons(el('forumTopicDetail'));
      el('forumLikeButton').addEventListener('click', async () => {
        try {
          const result = await apiPost('/api/forum/like', { topic_id: topic.id });
          renderForumTopic(result.topic);
          loadForumTopics();
        } catch (err) {
          el('forumCommentStatus').textContent = err.message;
          el('forumCommentStatus').className = 'status err';
        }
      });
      el('forumPrevPage').addEventListener('click', () => {
        state.forumCommentPage -= 1;
        renderForumTopic(topic);
      });
      el('forumNextPage').addEventListener('click', () => {
        state.forumCommentPage += 1;
        renderForumTopic(topic);
      });
      document.querySelectorAll('.forum-reply-button').forEach(button => {
        button.addEventListener('click', () => {
          const item = button.closest('.forum-comment');
          el('forumCommentParent').value = item.dataset.commentId || '';
          el('forumCommentInput').value = `回复 ${item.dataset.author || '成员'}：`;
          el('forumCommentInput').focus();
        });
      });
      el('forumCommentButton').addEventListener('click', async () => {
        const content = el('forumCommentInput').value.trim();
        if (!content) {
          el('forumCommentStatus').textContent = '请先填写讨论内容';
          el('forumCommentStatus').className = 'status err';
          return;
        }
        try {
          el('forumCommentStatus').textContent = '发布中...';
          el('forumCommentStatus').className = 'status';
          const result = await apiPost('/api/forum/comment', { topic_id: topic.id, content, parent_id: el('forumCommentParent').value });
          renderForumTopic(result.topic);
          loadForumTopics();
        } catch (err) {
          el('forumCommentStatus').textContent = err.message;
          el('forumCommentStatus').className = 'status err';
        }
      });
      el('forumAiCommentButton').addEventListener('click', async () => {
        try {
          el('forumCommentStatus').textContent = 'AI 潜水员正在读帖...';
          el('forumCommentStatus').className = 'status';
          const result = await apiPost('/api/forum/ai-comment', { topic_id: topic.id });
          renderForumTopic(result.topic);
          loadForumTopics();
        } catch (err) {
          el('forumCommentStatus').textContent = err.message;
          el('forumCommentStatus').className = 'status err';
        }
      });
    }
    el('forumCreateButton').addEventListener('click', async () => {
      const payload = { title: el('forumTitle').value, body: el('forumBody').value };
      try {
        el('forumCreateStatus').textContent = '发布中...';
        el('forumCreateStatus').className = 'status';
        const result = await apiPost('/api/forum/create', payload);
        el('forumTitle').value = '';
        el('forumBody').value = '';
        el('forumCreateStatus').textContent = '话题已发布';
        el('forumCreateStatus').className = 'status ok';
        state.forumSelected = result.topic.id;
        await loadForumTopics();
        openForumTopic(result.topic.id);
      } catch (err) {
        el('forumCreateStatus').textContent = err.message;
        el('forumCreateStatus').className = 'status err';
      }
    });
    el('forumRefreshButton').addEventListener('click', () => loadForumTopics());
    el('forumToggleCreate').addEventListener('click', () => {
      const panel = el('forumCreatePanel');
      const hidden = panel.classList.toggle('hidden');
      el('forumToggleCreate').textContent = hidden ? '发起话题' : '收起发起';
    });
    el('forumAiButton').addEventListener('click', async () => {
      try {
        el('forumAiStatus').textContent = '智能体正在起题...';
        el('forumAiStatus').className = 'status';
        const files = [];
        for (const file of [...el('forumAiFiles').files]) {
          files.push({ name: file.name, data: await readFileAsBase64(file) });
        }
        const result = await apiPost('/api/forum/ai-topic', {
          seed: el('forumAiSeed').value,
          chat: el('forumAiChat').value,
          files
        });
        el('forumAiSeed').value = '';
        el('forumAiChat').value = '';
        el('forumAiFiles').value = '';
        el('forumAiStatus').textContent = '智能话题已发布';
        el('forumAiStatus').className = 'status ok';
        state.forumSelected = result.topic.id;
        await loadForumTopics();
        openForumTopic(result.topic.id);
      } catch (err) {
        el('forumAiStatus').textContent = err.message;
        el('forumAiStatus').className = 'status err';
      }
    });

    // ===== 每日资讯 =====
    function renderNewsSources(sources = []) {
      const box = el('newsSources');
      const rows = sources.length ? sources : [{ name: '', url: '' }];
      box.innerHTML = rows.map((source, index) => `
        <div class="news-source-row" data-index="${index}">
          <input class="news-source-name" placeholder="名称" value="${escapeHtml(source.name || '')}" />
          <input class="news-source-url" placeholder="https://example.com/news" value="${escapeHtml(source.url || '')}" />
          <button class="mini danger news-remove-source" type="button">删除</button>
        </div>
      `).join('');
      box.querySelectorAll('.news-remove-source').forEach(button => {
        button.addEventListener('click', () => {
          button.closest('.news-source-row').remove();
          if (!box.querySelector('.news-source-row')) renderNewsSources();
        });
      });
    }
    function collectNewsSources() {
      return [...document.querySelectorAll('.news-source-row')].map(row => ({
        name: row.querySelector('.news-source-name').value.trim(),
        url: row.querySelector('.news-source-url').value.trim()
      })).filter(item => item.url);
    }
    function renderNewsIssue(issue) {
      if (!issue) {
        el('newsTitle').textContent = '暂无每日资讯';
        el('newsMeta').textContent = state.user?.is_superadmin ? '配置资讯源后生成今日简报。' : '暂无今日简报，请稍后刷新。';
        el('newsSummary').textContent = '暂无摘要。';
        el('newsItems').innerHTML = '<div class="forum-empty">暂无资讯内容。</div>';
        el('newsKeywords').innerHTML = '';
        return;
      }
      el('newsTitle').textContent = issue.title || `${issue.date || ''} 轨道交通每日资讯`;
      el('newsMeta').textContent = `${issue.date || ''} · ${issue.generated_at || ''} · ${issue.generated_by || ''}`;
      el('newsSummary').textContent = issue.summary || '暂无摘要。';
      const items = issue.items || [];
      el('newsItems').innerHTML = items.length ? items.map(item => `
        <div class="news-item">
          <div class="news-item-title">${escapeHtml(item.title || '未命名资讯')}</div>
          <div class="news-item-meta">${escapeHtml(item.source || '')}${item.url ? ` · <a href="${escapeHtml(item.url)}" target="_blank">来源链接</a>` : ''}</div>
          <p><strong>影响/价值：</strong>${escapeHtml(item.impact || '')}</p>
          <p><strong>建议动作：</strong>${escapeHtml(item.action || '')}</p>
        </div>
      `).join('') : '<div class="forum-empty">暂无资讯条目。</div>';
      el('newsKeywords').innerHTML = (issue.keywords || []).map(k => `<span class="news-keyword">${escapeHtml(k)}</span>`).join('');
    }
    async function loadNews() {
      try {
        const data = await api('/api/news/latest');
        if (state.user?.is_superadmin) {
          const cfg = data.config || {};
          renderNewsSources(cfg.sources || []);
          el('newsSearchQuery').value = cfg.search_query || '';
          el('newsAutoSearch').checked = !!cfg.auto_search;
          el('newsAutoPush').checked = !!cfg.auto_push;
          el('newsPushTime').value = cfg.push_time || '08:30';
        }
        renderNewsIssue(data.issue);
      } catch (err) {
        el('newsConfigStatus').textContent = err.message;
        el('newsConfigStatus').className = 'status err';
      }
    }
    el('newsAddSource').addEventListener('click', () => {
      const current = collectNewsSources();
      current.push({ name: '', url: '' });
      renderNewsSources(current);
    });
    el('newsRefresh').addEventListener('click', loadNews);
    el('newsSaveConfig').addEventListener('click', async () => {
      try {
        el('newsConfigStatus').textContent = '保存中...';
        el('newsConfigStatus').className = 'status';
        await apiPost('/api/news/config', {
          sources: collectNewsSources(),
          search_query: el('newsSearchQuery').value,
          auto_search: el('newsAutoSearch').checked,
          auto_push: el('newsAutoPush').checked,
          push_time: el('newsPushTime').value
        });
        el('newsConfigStatus').textContent = '资讯配置已保存';
        el('newsConfigStatus').className = 'status ok';
      } catch (err) {
        el('newsConfigStatus').textContent = err.message;
        el('newsConfigStatus').className = 'status err';
      }
    });
    el('newsGenerateNow').addEventListener('click', async () => {
      try {
        el('newsConfigStatus').textContent = '正在抓取网页并调用大模型生成资讯...';
        el('newsConfigStatus').className = 'status';
        const result = await apiPost('/api/news/generate', {
          search_query: el('newsSearchQuery').value,
          auto_search: el('newsAutoSearch').checked
        });
        renderNewsIssue(result.issue);
        el('newsConfigStatus').textContent = '今日资讯已生成';
        el('newsConfigStatus').className = 'status ok';
      } catch (err) {
        el('newsConfigStatus').textContent = err.message;
        el('newsConfigStatus').className = 'status err';
      }
    });

    el('modalSave').addEventListener('click', saveAndCloseModal);
    el('modalCancel').addEventListener('click', saveAndCloseModal);
    el('modalClose').addEventListener('click', saveAndCloseModal);
    el('modalFields').addEventListener('input', event => {
      if (event.target.closest('[data-modal-key]')) {
        saveCurrentModal();
      }
    });
    el('editModal').addEventListener('click', event => {
      if (event.target === el('editModal')) saveAndCloseModal();
    });
    document.addEventListener('keydown', event => {
      if (event.key === 'Escape' && !el('editModal').classList.contains('hidden')) {
        saveAndCloseModal();
      }
    });
    el('dashboardAskButton').addEventListener('click', () => {
      const text = el('dashboardAsk').value.trim();
      agentKind = agentKindForTask();
      updateAgentChrome(agentKind);
      toggleAgent(true);
      if (!agentMessages.length) {
        agentMessages.push({ role: 'assistant', content: `你好，我是${agentTitle(agentKind)}。你可以直接把当前页面相关的问题交给我。` });
      }
      renderAgentMessages();
      if (text) sendAgentMessage(text);
    });
    el('dashboardAsk').addEventListener('keydown', event => {
      if (event.key === 'Enter') {
        event.preventDefault();
        el('dashboardAskButton').click();
      }
    });
    async function optimizeField(button) {
      const fieldWrap = button.closest('.wide');
      const input = fieldWrap?.querySelector('textarea[data-modal-key]');
      const promptInput = fieldWrap?.querySelector('.field-assist-prompt');
      const fieldStatus = fieldWrap?.querySelector('.assist-status');
      if (!input || !promptInput || !fieldStatus) return;
      const original = input.value.trim();
      if (!original) {
        fieldStatus.textContent = '请先填写需要优化的内容。';
        fieldStatus.className = 'assist-status err';
        return;
      }
      button.disabled = true;
      button.textContent = '优化中...';
      fieldStatus.textContent = '正在调用 MiniMax-M2.7 优化，请稍等...';
      fieldStatus.className = 'assist-status';
      try {
        const result = await api('/api/optimize', {
          method: 'POST',
          headers: { 'Content-Type': 'application/json' },
          body: JSON.stringify({
            text: original,
            prompt: promptInput.value
          })
        });
        input.value = result.text || original;
        fieldStatus.textContent = result.warning || '优化完成，已写入当前输入框。';
        fieldStatus.className = result.warning ? 'assist-status err' : 'assist-status ok';
      } catch (err) {
        fieldStatus.textContent = '优化失败：' + err.message;
        fieldStatus.className = 'assist-status err';
      } finally {
        button.disabled = false;
        button.textContent = '辅助优化';
      }
    }

    el('modalFields').addEventListener('click', event => {
      const button = event.target.closest('.assist-optimize');
      if (button) optimizeField(button);
    });

    el('refresh').addEventListener('click', () => {
      const kind = state.task === 'mail' ? el('mailKind').value : el('kind').value;
      loadDraft(kind, state.selected);
    });
    el('copy').addEventListener('click', async () => {
      await navigator.clipboard.writeText(el('body').value);
      el('status').textContent = '正文已复制';
      el('status').className = 'status ok';
    });
    el('body').addEventListener('input', () => {
      state.bodyHtml = '';
      renderBodyPreview();
      clearSendReview();
    });
    ['to', 'cc', 'subject', 'attachment'].forEach(id => {
      el(id).addEventListener('input', clearSendReview);
      el(id).addEventListener('change', clearSendReview);
    });

    function sendPayload() {
      return {
        to: el('to').value.trim(),
        cc: el('cc').value.trim(),
        subject: el('subject').value.trim(),
        body: el('body').value,
        body_html: state.bodyHtml || '',
        attachment: el('attachment').value.trim()
      };
    }

    function sendBlockers(payload) {
      const blockers = [];
      if (!payload.to) blockers.push('收件人为空');
      if (!payload.subject) blockers.push('主题为空');
      if (!payload.attachment) blockers.push('未选择附件');
      const badWords = ['跟进内容', '计划内容', '很长内容', '总结5', '总结6'];
      const found = badWords.filter(word => payload.body.includes(word));
      if (found.length) blockers.push('正文疑似含测试残留：' + found.join('、'));
      return blockers;
    }

    function clearSendReview() {
      const box = el('sendReview');
      if (!box) return;
      box.classList.add('hidden');
      box.innerHTML = '';
    }

    function renderSendReview(payload, blockers = []) {
      const box = el('sendReview');
      if (!box) return;
      box.classList.remove('hidden');
      box.innerHTML = `
        <div class="send-review-title">发送前确认</div>
        <div class="send-review-grid">
          <div class="label">收件人</div><div>${escapeHtml(payload.to || '未填写')}</div>
          <div class="label">抄送</div><div>${escapeHtml(payload.cc || '无')}</div>
          <div class="label">主题</div><div>${escapeHtml(payload.subject || '未填写')}</div>
          <div class="label">附件</div><div>${escapeHtml(payload.attachment || '未选择')}</div>
        </div>
        ${blockers.length ? `<div class="send-review-warning">${escapeHtml(blockers.join('；'))}</div>` : ''}
        <div class="send-review-actions">
          <button type="button" class="secondary" id="cancelSendReview">取消</button>
          <button type="button" id="openMailConfigFromSend" class="secondary">邮件配置</button>
          <button type="button" class="warn" id="confirmSendReview" ${blockers.length ? 'disabled' : ''}>确认发送</button>
        </div>
      `;
      el('cancelSendReview').addEventListener('click', clearSendReview);
      el('openMailConfigFromSend').addEventListener('click', () => setTask('mailconfig'));
      el('confirmSendReview').addEventListener('click', () => doSendMail(payload));
    }

    function explainSendError(message) {
      const text = String(message || '发送失败');
      if (text.includes('SMTP 用户名')) return text + ' 建议在“邮件配置”中把 SMTP 用户名填写为发件邮箱。';
      if (text.includes('SMTP 密码')) return text + ' 请确认已填写邮箱授权码，不要使用网页登录密码。';
      if (text.includes('收件人')) return text + ' 请检查收件人邮箱，多个邮箱用分号分隔。';
      return text;
    }

    async function doSendMail(payload) {
      el('status').textContent = '处理中...';
      el('status').className = 'status';
      try {
        const result = await api('/api/send', {
          method: 'POST',
          headers: { 'Content-Type': 'application/json' },
          body: JSON.stringify(payload)
        });
        el('status').textContent = result.message;
        el('status').className = 'status ok';
        clearSendReview();
      } catch (err) {
        el('status').textContent = explainSendError(err.message);
        el('status').className = 'status err';
        renderSendReview(payload, []);
      }
    }

    el('send').addEventListener('click', async () => {
      const payload = sendPayload();
      const blockers = sendBlockers(payload);
      renderSendReview(payload, blockers);
    });

    function agentKindForTask(task = state.task) {
      if (['weekly', 'trip', 'diary', 'forum', 'news', 'mailassistant'].includes(task)) return task;
      if (task === 'calendar') return 'diary';
      return 'dashboard';
    }

    function agentTitle(kind) {
      return {
        weekly: '周报助手',
        trip: '出差报告助手',
        diary: '日记助手',
        forum: '论坛助手',
        news: '资讯助手',
        mailassistant: '智能邮件助手',
        dashboard: 'AI 办公总助手'
      }[kind] || '犇犇';
    }

    function agentIcon(kind) {
      return {
        weekly: 'file-spreadsheet',
        trip: 'briefcase-business',
        diary: 'book-open',
        forum: 'messages-square',
        news: 'newspaper',
        mailassistant: 'inbox',
        dashboard: 'layout-dashboard'
      }[kind] || 'bot';
    }

    function agentIntro(kind, hasContent = false) {
      const intro = {
        weekly: {
          title: '周报助手怎么用',
          copy: hasContent ? '我看到当前周报页已有内容。你可以先让我检查问题，也可以直接给我本周工作素材重写。' : '直接把本周做过的事、重点项目、下周计划发给我，我会自动整理成周报草稿 → 生成预览 → 发送邮件。',
          tips: ['我会自动获取当前日期和本周周期', '先对话收集信息，再一键生成草稿', '确认预览后再发送，全程不用手动填表'],
          prompts: [
            ['帮我写本周周报', '帮我写本周周报。'],
            ['根据素材写周报', '请根据我接下来提供的工作素材，生成正式周报并填入表单。'],
            ['优化语言表达', '请优化当前周报表单里的表达，让内容更正式、具体、清晰。']
          ]
        },
        mailassistant: {
          title: '邮件助手怎么用',
          copy: '我可以帮你总结邮件、起草回复、润色正文。真正发送前，页面会让你确认收件人、主题和附件。',
          tips: ['SMTP 用于发送，IMAP 用于收件箱', 'SMTP 用户名通常就是发件邮箱', '授权码不是网页登录密码'],
          prompts: [
            ['检查邮件配置', '请告诉我当前邮件配置还缺什么，以及下一步怎么修。'],
            ['优化当前正文', '请帮我把当前邮件正文优化得更清晰、礼貌、适合发送。'],
            ['起草一封邮件', '请根据我的要求起草一封普通邮件。']
          ]
        },
        dashboard: {
          title: '犇犇可以做什么',
          copy: '选择一个工作场景，或直接说你要处理的事。我会尽量把结果写回页面，而不是只给建议。',
          tips: ['周报、出差报告、日记、邮件都可以处理', '复杂操作会先生成预览或确认卡', '发送、发布等外部动作会二次确认'],
          prompts: [
            ['写本周周报', '帮我写本周周报。', 'weekly'],
            ['处理邮件', '帮我处理一封邮件。', 'mailassistant'],
            ['记录工作日记', '帮我记录今天的工作日记。', 'diary']
          ]
        }
      };
      return intro[kind] || {
        title: `${agentTitle(kind)}怎么用`,
        copy: '告诉我你要完成的目标，我会结合当前页面内容帮你整理、填写或生成。',
        tips: ['先说明目标', '确认内容后再执行外部动作', '可以随时要求重写或优化'],
        prompts: [['分析当前页面', '请分析当前页面内容，并告诉我下一步怎么做。']]
      };
    }

    function updateAgentChrome(kind = agentKindForTask()) {
      const title = agentTitle(kind);
      const header = document.querySelector('#agentWindow .agent-header span');
      if (header) {
        header.innerHTML = `<img class="agent-avatar" src="/assets/ai-assistant-avatar.png" alt="" /> ${title}`;
      }
      el('agentToggle').title = title;
      el('agentActions').innerHTML = `
        <button class="agent-action agent-kind-action" type="button" data-agent-kind="${kind}"><span class="icon" data-icon="${agentIcon(kind)}"></span> 当前</button>
        <button class="agent-action agent-kind-action" type="button" data-agent-kind="weekly"><span class="icon" data-icon="file-spreadsheet"></span> 周报</button>
        <button class="agent-action agent-kind-action" type="button" data-agent-kind="mailassistant"><span class="icon" data-icon="inbox"></span> 邮件</button>
        <button class="agent-action" type="button" id="agentClear"><span class="icon" data-icon="rotate-ccw"></span> 清空</button>
      `;
      el('agentActions').querySelectorAll('.agent-kind-action').forEach(btn => {
        btn.addEventListener('click', () => startAgent(btn.dataset.agentKind, true));
      });
      el('agentClear')?.addEventListener('click', () => startAgent(agentKind, true));
      renderIcons(el('agentWindow'));
    }

    function syncAgentToTask(task = state.task) {
      const nextKind = agentKindForTask(task);
      const changed = agentKind !== nextKind;
      agentKind = nextKind;
      updateAgentChrome(nextKind);
      if (changed) {
        agentMessages = [];
        renderAgentMessages();
      }
      const win = el('agentWindow');
      if (changed && win && !win.classList.contains('hidden')) {
        startAgent(nextKind, true);
      }
    }

    function toggleAgent(show) {
      const win = el('agentWindow');
      const shouldOpen = show === true || (show === undefined && win.classList.contains('hidden'));
      if (show === true) win.classList.remove('hidden');
      else if (show === false) win.classList.add('hidden');
      else if (shouldOpen) win.classList.remove('hidden');
      bringAgentToFront();
    }

    function bringAgentToFront() {
      const agent = el('agentFloat');
      if (!agent) return;
      agent.style.zIndex = '3000';
    }

    function agentStageSteps(kind = agentKind) {
      if (kind === 'weekly') return ['分析', '生成', '预览', '发送'];
      if (kind === 'trip') return ['分析', '整理', '生成', '发送'];
      if (kind === 'mailassistant') return ['读取', '起草', '确认', '发送'];
      return ['理解', '处理', '确认', '完成'];
    }

    function setAgentStage(index = 0) {
      state.agentStage = index;
      const box = el('agentProgress');
      if (!box) return;
      const steps = agentStageSteps(agentKind);
      box.innerHTML = steps.map((step, idx) => `
        <div class="agent-step ${idx < index ? 'done' : idx === index ? 'active' : ''}">${escapeHtml(step)}</div>
      `).join('');
    }

    function compactAssistantText(content) {
      const text = String(content || '').trim();
      if (!text) return '';
      const raw = text.replace(/^```json\s*/i, '').replace(/```$/i, '').trim();
      if (raw.startsWith('{') && raw.endsWith('}')) {
        try {
          const data = JSON.parse(raw);
          return data.reply || '已完成当前步骤，页面内容已同步更新。';
        } catch (err) {
          return text;
        }
      }
      return text;
    }

    function renderMarkdown(text) {
      if (!text) return '';
      // Protect code blocks
      const codeBlocks = [];
      text = text.replace(/```([\s\S]*?)```/g, (match, code) => {
        codeBlocks.push(escapeHtml(code.replace(/^.*?\\n/, '')));
        return '__CODEBLOCK_' + (codeBlocks.length - 1) + '__';
      });
      // Protect inline code
      const inlineCodes = [];
      text = text.replace(/`([^`]+)`/g, (match, code) => {
        inlineCodes.push(escapeHtml(code));
        return '__INLINECODE_' + (inlineCodes.length - 1) + '__';
      });
      // Escape remaining HTML
      text = escapeHtml(text);
      // Process blocks
      const rawLines = text.split('\\n');
      const blocks = [];
      let currentBlock = [];
      const flushBlock = () => {
        if (currentBlock.length === 0) return;
        if (currentBlock[0].trim().startsWith('|')) {
          blocks.push(renderMarkdownTable(currentBlock));
        } else {
          blocks.push(renderMarkdownBlock(currentBlock.join('\\n')));
        }
        currentBlock = [];
      };
      for (let line of rawLines) {
        if (line.trim() === '') {
          flushBlock();
        } else {
          currentBlock.push(line);
        }
      }
      flushBlock();
      text = blocks.join('\\n');
      // Restore inline code
      text = text.replace(/__INLINECODE_(\d+)__/g, (match, idx) => {
        return '<code style="background:#0f172a;padding:2px 5px;border-radius:4px;font-size:13px;color:#7dd3fc;">' + inlineCodes[idx] + '</code>';
      });
      // Restore code blocks
      text = text.replace(/__CODEBLOCK_(\d+)__/g, (match, idx) => {
        return '<pre style="background:#0f172a;padding:12px;border-radius:8px;overflow-x:auto;font-size:13px;line-height:1.5;border:1px solid #1e293b;margin:8px 0;"><code>' + codeBlocks[idx] + '</code></pre>';
      });
      return text;
    }
    function renderMarkdownBlock(block) {
      if (block.startsWith('# ')) return '<h1 style="margin:14px 0 10px;font-size:20px;color:#e2e8f0;font-weight:700;">' + block.slice(2) + '</h1>';
      if (block.startsWith('## ')) return '<h2 style="margin:12px 0 8px;font-size:18px;color:#e2e8f0;font-weight:700;border-bottom:1px solid #334155;padding-bottom:4px;">' + block.slice(3) + '</h2>';
      if (block.startsWith('### ')) return '<h3 style="margin:10px 0 6px;font-size:16px;color:#e2e8f0;font-weight:600;">' + block.slice(4) + '</h3>';
      if (block.trim() === '---') return '<hr style="border:none;border-top:1px solid #334155;margin:10px 0;">';
      let html = block.replace(/\*\*(.*?)\*\*/g, '<strong style="color:#e2e8f0;">$1</strong>');
      const lines = html.split('\\n');
      const isBullet = lines.every(l => l.trim().startsWith('- ') || l.trim().startsWith('* '));
      const isNumbered = lines.every(l => /^\s*\d+\./.test(l.trim()));
      if (isBullet) {
        const items = lines.map(l => '<li style="margin:4px 0;">' + l.trim().replace(/^[-*] /, '') + '</li>').join('');
        return '<ul style="margin:6px 0;padding-left:20px;">' + items + '</ul>';
      }
      if (isNumbered) {
        const items = lines.map(l => '<li style="margin:4px 0;">' + l.trim().replace(/^\d+\.\s*/, '') + '</li>').join('');
        return '<ol style="margin:6px 0;padding-left:20px;">' + items + '</ol>';
      }
      return '<p style="margin:6px 0;">' + html.replace(/\\n/g, '<br>') + '</p>';
    }
    function renderMarkdownTable(lines) {
      const rows = lines.map(line => {
        const trimmed = line.trim();
        if (!trimmed.startsWith('|')) return null;
        return trimmed.slice(1).split('|').map(c => c.trim());
      }).filter(r => r !== null);
      if (rows.length < 2) return lines.join('<br>');
      const isSep = rows[1].every(c => /^:?-+:?$/.test(c));
      if (!isSep) return lines.join('<br>');
      let html = '<table style="width:100%;border-collapse:collapse;margin:8px 0;font-size:13px;">';
      html += '<thead><tr>';
      rows[0].forEach(cell => {
        html += '<th style="border:1px solid #334155;padding:8px 10px;background:#1e293b;color:#94a3b8;text-align:left;font-weight:600;">' + cell + '</th>';
      });
      html += '</tr></thead><tbody>';
      for (let i = 2; i < rows.length; i++) {
        html += '<tr>';
        rows[i].forEach((cell, idx) => {
          const sep = rows[1][idx] || '';
          let align = 'left';
          if (sep.startsWith(':') && sep.endsWith(':')) align = 'center';
          else if (sep.endsWith(':')) align = 'right';
          html += '<td style="border:1px solid #334155;padding:8px 10px;color:#cbd5e1;text-align:' + align + ';">' + cell + '</td>';
        });
        html += '</tr>';
      }
      html += '</tbody></table>';
      return html;
    }

    function agentMessageHtml(m) {
      if (m.type === 'intro') {
        return `
          <div class="agent-msg assistant">
            <div class="agent-intro">
              <div class="agent-intro-title">${escapeHtml(m.title)}</div>
              <div class="agent-intro-copy">${escapeHtml(m.copy)}</div>
              <div class="agent-quick-grid">
                ${(m.prompts || []).map(([label, prompt, targetKind]) => `<button type="button" class="agent-quick" data-agent-prompt="${escapeHtml(prompt)}" ${targetKind ? `data-agent-target-kind="${escapeHtml(targetKind)}"` : ''}>${escapeHtml(label)}</button>`).join('')}
              </div>
              ${(m.tips || []).length ? `<ul class="agent-help-list">${m.tips.map(tip => `<li>${escapeHtml(tip)}</li>`).join('')}</ul>` : ''}
            </div>
          </div>`;
      }
      if (m.type === 'skill') {
        return `
          <div class="agent-msg assistant">
            <div class="agent-card">
              <div class="agent-card-title">${escapeHtml(m.title || 'Skill 调用完成')}</div>
              ${m.metrics ? `<div class="agent-card-grid">${m.metrics.map(item => `
                <div class="agent-card-metric"><strong>${escapeHtml(item.value)}</strong><span>${escapeHtml(item.label)}</span></div>
              `).join('')}</div>` : ''}
              ${m.items && m.items.length ? `<ul class="agent-card-list">${m.items.map(item => `<li>${escapeHtml(item)}</li>`).join('')}</ul>` : ''}
              ${m.note ? `<div class="agent-card-note">${escapeHtml(m.note)}</div>` : ''}
              ${m.image_url ? `<div><img class="agent-img" src="${escapeHtml(m.image_url)}" style="max-width:100%;border:1px solid #dbe5f1;border-radius:8px;background:#fff;cursor:zoom-in;" onclick="openAgentLightbox('${escapeHtml(m.image_url)}')" /></div>` : ''}
            </div>
          </div>`;
      }
      const content = m.role === 'assistant' ? compactAssistantText(m.content) : m.content;
      const htmlContent = m.role === 'assistant' ? renderMarkdown(content) : escapeHtml(content);
      return `<div class="agent-msg ${m.role}">
        ${htmlContent}
        ${m.image_url ? `<div style="margin-top:10px"><img class="agent-img" src="${escapeHtml(m.image_url)}" style="max-width:100%;border:1px solid #dbe5f1;border-radius:8px;background:#fff;cursor:zoom-in;" onclick="openAgentLightbox('${escapeHtml(m.image_url)}')" /></div>` : ''}
      </div>`;
    }

    function renderAgentMessages() {
      const box = el('agentMessages');
      box.innerHTML = agentMessages.map(agentMessageHtml).join('');
      box.querySelectorAll('[data-agent-prompt]').forEach(btn => {
        btn.addEventListener('click', () => {
          const targetKind = btn.dataset.agentTargetKind;
          if (targetKind && targetKind !== agentKind) startAgent(targetKind, true);
          sendAgentMessage(btn.dataset.agentPrompt || '');
        });
      });
      box.scrollTop = box.scrollHeight;
      setAgentStage(state.agentStage || 0);
    }

    function agentPayloadMessages() {
      return agentMessages
        .filter(m => ['user', 'assistant'].includes(m.role) && m.content && !m.type)
        .map(m => ({ role: m.role, content: String(m.content || '') }));
    }

    function appendSkillResult(result) {
      if (!result) return;
      // 支持多次 Skill 调用（skill_calls 数组）和单次调用（skill_call 对象）
      const calls = result.skill_calls || (result.skill_call ? [{ name: result.skill_call.name, result: result.skill_result }] : []);
      if (!calls.length) return;
      calls.forEach(call => {
        const name = call.name || '';
        const data = call.result || {};
        if (name === 'weekly.compose' && data.draft) {
          if (hasRowContent(data.draft.weekly_summary)) renderWorkRows('summary', data.draft.weekly_summary);
          if (hasRowContent(data.draft.weekly_follow)) renderWorkRows('follow', data.draft.weekly_follow);
          if (hasRowContent(data.draft.weekly_next)) renderWorkRows('next', data.draft.weekly_next);
          agentMessages.push({
            role: 'assistant',
            type: 'skill',
            title: '周报草稿已生成并填入表单',
            metrics: [
              { label: '工作总结', value: (data.draft.weekly_summary || []).length },
              { label: '重点跟进', value: (data.draft.weekly_follow || []).length },
              { label: '下周计划', value: (data.draft.weekly_next || []).length }
            ],
            items: (data.draft.weekly_summary || []).slice(0, 3).map(item => `${item.category || '事项'}：${item.content || ''}`),
            note: '草稿已填入“填写报告”页面，你可以继续补充或告诉我直接生成预览。'
          });
          setAgentStage(2);
          renderAgentMessages();
          return;
        }
        if (name === 'weekly.preview') {
          agentMessages.push({
            role: 'assistant',
            type: 'skill',
            title: '周报预览已生成',
            metrics: [
              { label: '附件', value: data.file || '1' },
              { label: '周期', value: (data.mail_draft && data.mail_draft.period) || '' }
            ].filter(m => m.value),
            items: [
              data.file && `文件：${data.file}`,
              data.mail_draft && data.mail_draft.subject && `主题：${data.mail_draft.subject}`
            ].filter(Boolean),
            note: '预览图已生成，确认无误后告诉我“发送”即可。',
            image_url: data.preview_image_url || ''
          });
          setAgentStage(3);
          renderAgentMessages();
          return;
        }
        if (name === 'weekly.send_confirmed') {
          agentMessages.push({
            role: 'assistant',
            type: 'skill',
            title: '周报邮件已发送',
            metrics: data.mode ? [{ label: '状态', value: data.mode }] : null,
            items: [data.message].filter(Boolean),
            note: '本周周报已发送完成，如需继续处理其他工作随时告诉我。'
          });
          setAgentStage(4);
          renderAgentMessages();
          return;
        }
        if (name === 'utils.get_date') {
          // 日期获取不展示卡片，由大模型在回复中自然描述
          return;
        }
        agentMessages.push({
          role: 'assistant',
          type: 'skill',
          title: `${name} 已完成`,
          metrics: data.file ? [{ label: '附件', value: '1' }] : null,
          items: [data.file, data.message, data.mode ? `结果：${data.mode}` : ''].filter(Boolean),
          note: data.mail_draft ? '邮件草稿已生成，请确认收件人、主题、附件和正文后再发送。' : '',
          image_url: data.preview_image_url || ''
        });
        renderAgentMessages();
      });
    }

    async function sendAgentContext(contextData) {
      const btn = el('agentSend');
      btn.disabled = true;
      btn.textContent = '分析中...';
      try {
        const payloadMessages = agentPayloadMessages();
        payloadMessages.push({
          role: 'user',
          content: `[系统提示：当前智能体类型为 ${agentTitle(agentKind)}，以下是当前页面上下文。请严格按这个智能体的职责处理，不要沿用其他模块逻辑。]\\n` + JSON.stringify(contextData, null, 2)
        });
        const result = await api('/api/agent', {
          method: 'POST',
          headers: { 'Content-Type': 'application/json' },
          body: JSON.stringify({ kind: agentKind, messages: payloadMessages })
        });
        agentMessages.push({ role: 'assistant', content: result.reply });
        appendSkillResult(result);
        renderAgentMessages();
        try {
          const json = JSON.parse(result.reply);
          applyAgentResult(json);
          if (json.done) {
            const doneMsg = agentKind === 'diary' ? '日记已根据对话内容自动填充并保存，你可以在页面上查看和修改。' : '已根据对话内容自动生成并填充到表单中，你可以切换到"填写报告"查看和修改。';
            agentMessages.push({ role: 'assistant', content: doneMsg });
            renderAgentMessages();
          }
        } catch (e) { /* 继续对话 */ }
      } catch (err) {
        agentMessages.push({ role: 'assistant', content: '出错：' + err.message });
        renderAgentMessages();
      } finally {
        btn.disabled = false;
        btn.textContent = '发送';
      }
    }

    function startCurrentAgent() {
      startAgent(agentKindForTask(), true);
    }

    function openAgentFromAvatar() {
      el('agentWindow')?.classList.remove('hidden');
      bringAgentToFront();
      startCurrentAgent();
    }

    function startAgent(kind, reset = true) {
      const previousKind = agentKind;
      agentKind = kind;
      state.agentStage = 0;
      updateAgentChrome(kind);
      if (reset || !agentMessages.length || previousKind !== kind) agentMessages = [];
      toggleAgent(true);
      const currentData = getCurrentFormData(kind);
      const hasContent = kind === 'weekly'
        ? [currentData.weekly_summary, currentData.weekly_follow, currentData.weekly_next].some(hasRowContent)
        : kind === 'forum'
        ? Object.values(currentData).some(v => typeof v === 'string' && v.trim())
        : kind === 'news'
        ? Object.values(currentData).some(v => typeof v === 'string' && v.trim())
        : kind === 'mailassistant'
        ? Object.values(currentData).some(v => typeof v === 'string' && v.trim())
        : Object.values(currentData).some(v => v && String(v).trim());
      const intro = agentIntro(kind, hasContent);
      agentMessages.push({ role: 'assistant', type: 'intro', ...intro });
      renderAgentMessages();
    }

    async function sendAgentMessage(text) {
      if (!text.trim()) return;
      agentMessages.push({ role: 'user', content: text.trim() });
      renderAgentMessages();
      el('agentInput').value = '';
      const btn = el('agentSend');
      btn.disabled = true;
      btn.textContent = '思考中...';
      try {
        const currentData = getCurrentFormData(agentKind);
        const payloadMessages = agentPayloadMessages();
        if (Object.values(currentData).some(v => Array.isArray(v) ? v.length : v)) {
          payloadMessages.push({
            role: 'user',
            content: `[系统提示：当前智能体类型为 ${agentTitle(agentKind)}，以下是当前页面上下文。请严格按这个智能体的职责处理，不要沿用其他模块逻辑。]\\n` + JSON.stringify(currentData, null, 2)
          });
        }
        const result = await api('/api/agent', {
          method: 'POST',
          headers: { 'Content-Type': 'application/json' },
          body: JSON.stringify({ kind: agentKind, messages: payloadMessages })
        });
        agentMessages.push({ role: 'assistant', content: result.reply });
        appendSkillResult(result);
        renderAgentMessages();
        try {
          const json = JSON.parse(result.reply);
          applyAgentResult(json);
          if (json.done) {
            const doneMsg = agentKind === 'diary' ? '日记已根据对话内容自动填充并保存，你可以在页面上查看和修改。' : '已根据对话内容自动生成并填充到表单中，你可以切换到"填写报告"查看和修改。';
            agentMessages.push({ role: 'assistant', content: doneMsg });
            renderAgentMessages();
          }
        } catch (e) { /* 继续对话 */ }
      } catch (err) {
        agentMessages.push({ role: 'assistant', content: '出错：' + err.message });
        renderAgentMessages();
      } finally {
        btn.disabled = false;
        btn.textContent = '发送';
      }
    }

    function getCurrentFormData(kind) {
      if (kind === 'weekly') {
        return {
          weekly_summary: collectWorkRows('summary'),
          weekly_follow: collectWorkRows('follow'),
          weekly_next: collectWorkRows('next')
        };
      }
      if (kind === 'diary') {
        return {
          today_work: el('diaryTodayWork').value,
          tomorrow_plan: el('diaryTomorrowPlan').value,
          thoughts: el('diaryThoughts').value
        };
      }
      if (kind === 'forum') {
        return {
          selected_topic_id: state.forumSelected || '',
          topic_title: document.querySelector('#forumTopicDetail .forum-topic-title')?.textContent || '',
          topic_detail: el('forumTopicDetail')?.innerText || '',
          comment_draft: el('forumCommentInput')?.value || '',
          create_title: el('forumTitle')?.value || '',
          create_body: el('forumBody')?.value || ''
        };
      }
      if (kind === 'news') {
        return {
          title: el('newsTitle')?.textContent || '',
          meta: el('newsMeta')?.textContent || '',
          summary: el('newsSummary')?.textContent || '',
          items: el('newsItems')?.innerText || '',
          keywords: el('newsKeywords')?.innerText || ''
        };
      }
      if (kind === 'mailassistant') {
        return {
          selected_mail: el('mailDetail')?.innerText || '',
          compose_to: el('assistantMailTo')?.value || '',
          compose_cc: el('assistantMailCc')?.value || '',
          compose_subject: el('assistantMailSubject')?.value || '',
          compose_body: el('assistantMailBody')?.value || ''
        };
      }
      if (kind === 'dashboard') {
        return {
          current_task: state.task || 'dashboard',
          current_page: el('taskTitle')?.textContent || '工作台',
          user: state.user?.name || state.user?.username || ''
        };
      }
      return {
        reporter: el('tripReporter').value,
        department: el('tripDepartment').value,
        location: el('tripLocation').value,
        trip_start: el('tripStart').value,
        trip_end: el('tripEnd').value,
        purpose: el('tripPurpose').value,
        itinerary: el('tripItinerary').value,
        details: el('tripDetails').value,
        issues: el('tripIssues').value,
        suggestions: el('tripSuggestions').value
      };
    }

    function hasRowContent(arr) {
      return Array.isArray(arr) && arr.some(r => Object.values(r).some(v => v && String(v).trim()));
    }

    function applyAgentResult(data) {
      if (agentKind === 'weekly') {
        if (hasRowContent(data.weekly_summary)) renderWorkRows('summary', data.weekly_summary);
        if (hasRowContent(data.weekly_follow)) renderWorkRows('follow', data.weekly_follow);
        if (hasRowContent(data.weekly_next)) renderWorkRows('next', data.weekly_next);
      } else if (agentKind === 'trip') {
        if (data.reporter !== undefined) el('tripReporter').value = data.reporter;
        if (data.department !== undefined) el('tripDepartment').value = data.department;
        if (data.location !== undefined) el('tripLocation').value = data.location;
        if (data.trip_start !== undefined) el('tripStart').value = data.trip_start;
        if (data.trip_end !== undefined) el('tripEnd').value = data.trip_end;
        if (data.purpose !== undefined) el('tripPurpose').value = data.purpose;
        if (data.itinerary !== undefined) el('tripItinerary').value = data.itinerary;
        if (data.details !== undefined) el('tripDetails').value = data.details;
        if (data.issues !== undefined) el('tripIssues').value = data.issues;
        if (data.suggestions !== undefined) el('tripSuggestions').value = data.suggestions;
        renderTripCards();
      } else if (agentKind === 'diary') {
        if (data.today_work !== undefined) el('diaryTodayWork').value = data.today_work;
        if (data.tomorrow_plan !== undefined) el('diaryTomorrowPlan').value = data.tomorrow_plan;
        if (data.thoughts !== undefined) el('diaryThoughts').value = data.thoughts;
        if (data.done) saveDiary();
      }
      highlightUpdatedFields();
    }

    function highlightUpdatedFields() {
      document.querySelectorAll('input, textarea, .work-block').forEach(el => {
        el.style.transition = 'box-shadow .4s ease';
        el.style.boxShadow = '0 0 0 2px #17736a40';
        setTimeout(() => { el.style.boxShadow = ''; }, 800);
      });
    }

    boot();
  </script>
  <div class="modal hidden" id="profileModal">
    <div class="modal-box" style="max-width:720px;">
      <div class="modal-head">
        <div class="modal-title">个人资料</div>
        <button class="mini secondary" id="profileClose" type="button">关闭</button>
      </div>
      <div class="profile-grid">
        <div>
          <img class="profile-avatar-large" id="profileAvatarPreview" src="/assets/ai-assistant-avatar.png" alt="" />
          <div class="profile-actions">
            <button class="mini secondary" id="profileUseAssistantAvatar" type="button">使用助手头像</button>
          </div>
        </div>
        <div>
          <label>对外显示名称</label>
          <input id="profileName" placeholder="你的姓名或昵称" />
          <label>上传头像</label>
          <input id="profileAvatarFile" type="file" accept="image/png,image/jpeg,image/webp" />
          <label>对外介绍</label>
          <textarea id="profileBio" placeholder="例如：负责项目、擅长方向、工作职责等" style="min-height:96px;"></textarea>
          <label>个人爱好</label>
          <textarea id="profileHobbies" placeholder="例如：轨道交通、AI 工具、阅读、跑步..." style="min-height:80px;"></textarea>
          <div class="toolbar">
            <button id="profileSave" type="button">保存资料</button>
            <span id="profileStatus" class="status"></span>
          </div>
        </div>
      </div>
    </div>
  </div>
  <div class="agent-float hidden" id="agentFloat">
    <button class="agent-toggle" id="agentToggle" type="button" title="犇犇" onclick="openAgentFromAvatar()"><img class="agent-avatar" src="/assets/ai-assistant-avatar.png" alt="犇犇" /></button>
    <div class="agent-window hidden" id="agentWindow">
      <div class="agent-header">
        <span><img class="agent-avatar" src="/assets/ai-assistant-avatar.png" alt="" /> 犇犇</span>
        <button class="agent-close" id="agentClose" type="button">✕</button>
      </div>
      <div class="agent-body">
        <div class="agent-progress" id="agentProgress"></div>
        <div class="agent-messages" id="agentMessages">
          <div class="agent-msg assistant">你好！我是犇犇，你的智能办公助手。点击下方的快捷按钮，我可以帮你自动生成周报或出差报告。</div>
        </div>
        <div class="agent-actions" id="agentActions">
          <button class="agent-action" type="button" data-agent-kind="weekly"><span class="icon" data-icon="file-spreadsheet"></span> 生成周报</button>
          <button class="agent-action" type="button" data-agent-kind="trip"><span class="icon" data-icon="briefcase-business"></span> 生成出差报告</button>
          <button class="agent-action" type="button" data-agent-kind="diary"><span class="icon" data-icon="book-open"></span> 记录日记</button>
        </div>
        <div class="agent-input-wrap">
          <textarea id="agentInput" placeholder="输入消息...（Ctrl+Enter 发送）"></textarea>
          <button id="agentSend" type="button">发送</button>
        </div>
      </div>
    </div>
    <div class="agent-lightbox hidden" id="agentLightbox" onclick="closeAgentLightbox()">
      <img id="agentLightboxImg" src="" alt="预览" />
    </div>
  </div>
  <script>
    function openAgentLightbox(src) {
      const box = document.getElementById('agentLightbox');
      const img = document.getElementById('agentLightboxImg');
      if (box && img) { img.src = src; box.classList.remove('hidden'); }
    }
    function closeAgentLightbox() {
      const box = document.getElementById('agentLightbox');
      if (box) box.classList.add('hidden');
    }
    (function() {
      const el = id => document.getElementById(id);
      const agentFloat = el('agentFloat');
      const agentToggle = el('agentToggle');
      const agentWindow = el('agentWindow');
      const agentHeader = agentWindow.querySelector('.agent-header');
      let suppressClickUntil = 0;

      agentToggle.addEventListener('click', event => {
        event.stopPropagation();
        if (Date.now() < suppressClickUntil) return;
        openAgentFromAvatar();
      });
      el('agentClose').addEventListener('click', () => toggleAgent(false));
      el('agentSend').addEventListener('click', () => sendAgentMessage(el('agentInput').value));
      el('agentInput').addEventListener('keydown', e => {
        if (e.key === 'Enter' && (e.ctrlKey || e.metaKey)) {
          e.preventDefault();
          sendAgentMessage(el('agentInput').value);
        }
      });
      if (typeof updateAgentChrome === 'function') updateAgentChrome(typeof agentKindForTask === 'function' ? agentKindForTask() : 'dashboard');
      if (typeof renderIcons === 'function') renderIcons(document);

      // 拖拽逻辑：拖按钮或标题栏都会移动整个 AI 助手，点击按钮只负责唤醒。
      (function() {
        let dragging = false;
        let moved = false;
        let startX = 0;
        let startY = 0;
        let initLeft = 0;
        let initTop = 0;
        let activePointerId = null;
        let pointerStartedOnToggle = false;

        function pointerDown(e) {
          if (e.target.closest('.agent-close') || e.target.closest('textarea') || e.target.closest('.agent-action') || e.target.closest('#agentSend')) return;
          bringAgentToFront();
          dragging = true;
          moved = false;
          activePointerId = e.pointerId;
          pointerStartedOnToggle = !!e.target.closest('#agentToggle');
          startX = e.clientX; startY = e.clientY;
          const rect = agentFloat.getBoundingClientRect();
          initLeft = rect.left; initTop = rect.top;
          agentFloat.style.right = 'auto';
          agentFloat.style.bottom = 'auto';
          agentFloat.style.left = initLeft + 'px';
          agentFloat.style.top = initTop + 'px';
          agentToggle.style.cursor = 'grabbing';
          agentHeader.style.cursor = 'grabbing';
          if (e.currentTarget.setPointerCapture) e.currentTarget.setPointerCapture(e.pointerId);
        }

        agentToggle.addEventListener('pointerdown', pointerDown);
        agentHeader.addEventListener('pointerdown', pointerDown);

        document.addEventListener('pointermove', e => {
          if (!dragging) return;
          if (activePointerId !== null && e.pointerId !== activePointerId) return;
          const dx = e.clientX - startX;
          const dy = e.clientY - startY;
          if (Math.abs(dx) > 4 || Math.abs(dy) > 4) {
            moved = true;
          }
          const rect = agentFloat.getBoundingClientRect();
          const maxLeft = window.innerWidth - Math.max(80, rect.width);
          const maxTop = window.innerHeight - Math.max(80, rect.height);
          const nextLeft = Math.max(8, Math.min(initLeft + dx, maxLeft));
          const nextTop = Math.max(8, Math.min(initTop + dy, maxTop));
          agentFloat.style.left = nextLeft + 'px';
          agentFloat.style.top = nextTop + 'px';
        });

        document.addEventListener('pointerup', e => {
          if (!dragging) return;
          if (activePointerId !== null && e.pointerId !== activePointerId) return;
          dragging = false;
          activePointerId = null;
          agentToggle.style.cursor = 'grab';
          agentHeader.style.cursor = 'move';
          if (moved) {
            suppressClickUntil = Date.now() + 250;
          } else if (pointerStartedOnToggle) {
            openAgentFromAvatar();
          }
          pointerStartedOnToggle = false;
        });
      })();
    })();
  </script>
</body>
</html>"""


class Handler(BaseHTTPRequestHandler):
    def log_message(self, fmt, *args):
        sys.stderr.write("[%s] %s\n" % (self.log_date_time_string(), fmt % args))

    def send_json(self, data, status=200, headers=None):
        raw = json.dumps(data, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(raw)))
        for key, value in (headers or {}).items():
            self.send_header(key, value)
        self.end_headers()
        self.wfile.write(raw)

    def cookie_value(self, name):
        cookie = self.headers.get("Cookie", "")
        for part in cookie.split(";"):
            if "=" not in part:
                continue
            key, value = part.strip().split("=", 1)
            if key == name:
                return value
        return ""

    def current_user(self):
        token = self.cookie_value("pws_session")
        username = SESSIONS.get(token)
        return find_user(username) if username else None

    def require_user(self):
        user = self.current_user()
        if not user:
            self.send_json({"ok": False, "error": "请先登录"}, status=401)
            return None
        return user

    def require_admin(self):
        user = self.require_user()
        if not user:
            return None
        if user.get("role") not in ("admin", "superadmin"):
            self.send_json({"ok": False, "error": "只有管理员可以修改系统配置"}, status=403)
            return None
        return user

    def require_superadmin(self):
        user = self.require_user()
        if not user:
            return None
        if user.get("role") != "superadmin":
            self.send_json({"ok": False, "error": "只有超级管理员可以进行该操作"}, status=403)
            return None
        return user

    def do_GET(self):
        parsed = urllib.parse.urlparse(self.path)
        if parsed.path in ("/", APP_RELATIVE_PATH, APP_RELATIVE_PATH + "/"):
            raw = app_html().encode("utf-8")
            self.send_response(200)
            self.send_header("Content-Type", "text/html; charset=utf-8")
            self.send_header("Content-Length", str(len(raw)))
            self.end_headers()
            self.wfile.write(raw)
            return
        if parsed.path.startswith("/assets/"):
            asset_name = Path(urllib.parse.unquote(parsed.path)).name
            path = BASE_DIR / "assets" / asset_name
            if not path.exists() or not path.is_file():
                self.send_json({"error": "Not found"}, status=404)
                return
            raw = path.read_bytes()
            ctype, _ = mimetypes.guess_type(str(path))
            self.send_response(200)
            self.send_header("Content-Type", ctype or "application/octet-stream")
            self.send_header("Content-Length", str(len(raw)))
            self.send_header("Cache-Control", "public, max-age=86400")
            self.end_headers()
            self.wfile.write(raw)
            return
        if parsed.path.startswith("/user-avatar/"):
            user = self.require_user()
            if not user:
                return
            name = Path(urllib.parse.unquote(parsed.path)).stem
            target = find_user(name)
            if not target:
                self.send_json({"error": "Not found"}, status=404)
                return
            path = user_profile_dir(name) / "avatar.png"
            if not path.exists():
                self.send_json({"error": "Not found"}, status=404)
                return
            raw = path.read_bytes()
            self.send_response(200)
            self.send_header("Content-Type", "image/png")
            self.send_header("Content-Length", str(len(raw)))
            self.send_header("Cache-Control", "private, max-age=3600")
            self.end_headers()
            self.wfile.write(raw)
            return
        if parsed.path == "/api/session":
            user = self.current_user()
            self.send_json(
                {
                    "authenticated": bool(user),
                    "user": public_user(user) if user else None,
                    "assistant_prompt": read_config().get("assistant_prompt", DEFAULT_ASSISTANT_PROMPT),
                }
            )
            return
        if parsed.path == "/api/admin-config":
            if not self.require_admin():
                return
            self.send_json(admin_config_payload())
            return
        if parsed.path == "/api/admin-users-list":
            if not self.require_admin():
                return
            self.send_json({"ok": True, "users": user_list(self.current_user().get("username", ""))})
            return
        if parsed.path == "/skill-docs":
            raw = (
                "<!doctype html><html lang='zh-CN'><head><meta charset='utf-8'>"
                "<title>智能办公助手 Skill 文档</title>"
                "<style>body{font-family:-apple-system,BlinkMacSystemFont,'PingFang SC','Microsoft YaHei',sans-serif;max-width:980px;margin:0 auto;padding:32px;line-height:1.7;color:#172033;background:#f4f7fb}"
                "pre{white-space:pre-wrap;background:#fff;border:1px solid #dbe5f1;border-radius:10px;padding:18px;box-shadow:0 8px 24px rgba(15,23,42,.06)}"
                "a{display:inline-block;margin-bottom:18px;color:#2563eb;font-weight:700}</style></head><body>"
                "<a href='/download-skill-doc'>下载 Markdown 文档</a>"
                "<pre>"
                + escape(skill_doc_markdown())
                + "</pre></body></html>"
            ).encode("utf-8")
            self.send_response(200)
            self.send_header("Content-Type", "text/html; charset=utf-8")
            self.send_header("Content-Length", str(len(raw)))
            self.end_headers()
            self.wfile.write(raw)
            return
        if parsed.path == "/download-skill-doc":
            raw = skill_doc_markdown().encode("utf-8")
            self.send_response(200)
            self.send_header("Content-Type", "text/markdown; charset=utf-8")
            self.send_header("Content-Length", str(len(raw)))
            self.send_header("Content-Disposition", "attachment; filename*=UTF-8''ai-office-skills.md")
            self.end_headers()
            self.wfile.write(raw)
            return
        if parsed.path.startswith("/api/") and not self.require_user():
            return
        if parsed.path == "/api/skills":
            if not self.require_superadmin():
                return
            self.send_json(public_skill_docs())
            return
        if parsed.path == "/api/agent-orchestration":
            if not self.require_superadmin():
                return
            self.send_json(agent_orchestration())
            return
        if parsed.path == "/api/agent-config":
            if not self.require_superadmin():
                return
            self.send_json({"ok": True, "config": read_agent_config()})
            return
        if parsed.path == "/api/mail-config":
            username = self.current_user().get("username", "")
            data = user_mail_config(username)
            data["smtp_password_masked"] = "已配置" if data.get("smtp_password") else "未配置"
            data["imap_password_masked"] = "已配置" if data.get("imap_password") else "未配置"
            data.pop("smtp_password", None)
            data.pop("imap_password", None)
            self.send_json(data)
            return
        if parsed.path == "/api/mailbox":
            qs = urllib.parse.parse_qs(parsed.query)
            limit = qs.get("limit", ["20"])[0]
            refresh = qs.get("refresh", ["0"])[0] == "1"
            try:
                self.send_json(list_inbox_messages(self.current_user().get("username", ""), limit, refresh))
            except Exception as exc:
                self.send_json({"ok": False, "error": str(exc)}, status=400)
            return
        if parsed.path == "/api/mailbox-detail":
            qs = urllib.parse.parse_qs(parsed.query)
            uid = qs.get("uid", [""])[0]
            refresh = qs.get("refresh", ["0"])[0] == "1"
            try:
                self.send_json(get_inbox_message(self.current_user().get("username", ""), uid, refresh))
            except Exception as exc:
                self.send_json({"ok": False, "error": str(exc)}, status=400)
            return
        if parsed.path == "/api/reports":
            username = self.current_user().get("username", "")
            files = sorted(all_files(username), key=lambda item: (item["generated"] if "generated" in item else False, item["kind"], item["sort_key"], item["mtime"]), reverse=True)
            self.send_json(
                {
                    "reports": files,
                    "latest_weekly": (newest("weekly", username, fallback_shared=True) or {}).get("name", ""),
                    "latest_trip": (newest("trip", username, fallback_shared=True) or {}).get("name", ""),
                }
            )
            return
        if parsed.path == "/api/draft":
            qs = urllib.parse.parse_qs(parsed.query)
            kind = qs.get("kind", ["weekly"])[0]
            file_name = qs.get("file", [""])[0]
            draft = compose_draft(kind, file_name, self.current_user().get("username", ""))
            status = 400 if "error" in draft else 200
            self.send_json(draft, status=status)
            return
        if parsed.path == "/api/weekly-prefill":
            self.send_json(weekly_prefill(self.current_user().get("username", "")))
            return
        if parsed.path == "/api/trip-prefill":
            self.send_json(trip_prefill(self.current_user().get("username", "")))
            return
        if parsed.path == "/api/diary/list":
            qs = urllib.parse.parse_qs(parsed.query)
            payload = {k: v[0] for k, v in qs.items()}
            self.send_json(list_diaries(payload, self.current_user().get("username", "")))
            return
        if parsed.path == "/api/diary/get":
            qs = urllib.parse.parse_qs(parsed.query)
            date_str = qs.get("date", [""])[0]
            self.send_json(get_diary(date_str, self.current_user().get("username", "")))
            return
        if parsed.path == "/api/forum/topics":
            self.send_json(forum_list_topics(self.current_user().get("username", "")))
            return
        if parsed.path == "/api/forum/topic":
            qs = urllib.parse.parse_qs(parsed.query)
            self.send_json(forum_get_topic({"id": qs.get("id", [""])[0]}, self.current_user().get("username", "")))
            return
        if parsed.path == "/api/news/latest":
            user = self.current_user() or {}
            self.send_json(news_latest(user.get("role") == "superadmin"))
            return
        if parsed.path == "/api/news/config":
            if not self.require_superadmin():
                return
            self.send_json({"ok": True, "config": news_config_payload()})
            return
        if parsed.path == "/download":
            qs = urllib.parse.parse_qs(parsed.query)
            file_name = qs.get("file", [""])[0]
            path = attachment_path_by_name(file_name, self.current_user().get("username", ""))
            if path is None:
                self.send_json({"error": "文件不存在"}, status=404)
                return
            raw = path.read_bytes()
            ctype, encoding = mimetypes.guess_type(str(path))
            self.send_response(200)
            self.send_header("Content-Type", ctype or "application/octet-stream")
            self.send_header("Content-Length", str(len(raw)))
            encoded_name = urllib.parse.quote(path.name)
            self.send_header("Content-Disposition", f"attachment; filename*=UTF-8''{encoded_name}")
            self.end_headers()
            self.wfile.write(raw)
            return
        if parsed.path == "/preview-image":
            if not self.require_user():
                return
            qs = urllib.parse.parse_qs(parsed.query)
            file_name = Path(qs.get("file", [""])[0]).name
            path = user_generated_dir(self.current_user().get("username", "")) / file_name
            if path.suffix.lower() != ".png" or not path.exists() or not path.is_file():
                self.send_json({"error": "预览图片不存在"}, status=404)
                return
            raw = path.read_bytes()
            self.send_response(200)
            self.send_header("Content-Type", "image/png")
            self.send_header("Content-Length", str(len(raw)))
            self.send_header("Cache-Control", "private, max-age=300")
            self.end_headers()
            self.wfile.write(raw)
            return
        self.send_json({"error": "Not found"}, status=404)

    def do_POST(self):
        parsed = urllib.parse.urlparse(self.path)
        try:
            length = int(self.headers.get("Content-Length", "0"))
            payload = json.loads(self.rfile.read(length).decode("utf-8"))
            if parsed.path == "/api/login":
                user = find_user(payload.get("username", ""), payload.get("password", ""))
                if not user:
                    self.send_json({"ok": False, "error": "用户名或密码错误"}, status=401)
                    return
                ensure_user_space(user.get("username", ""))
                token = secrets.token_urlsafe(24)
                SESSIONS[token] = user.get("username", "")
                self.send_json(
                    {
                        "ok": True,
                        "user": public_user(user),
                        "assistant_prompt": read_config().get("assistant_prompt", DEFAULT_ASSISTANT_PROMPT),
                    },
                    headers={"Set-Cookie": f"pws_session={token}; Path=/; HttpOnly; SameSite=Lax"},
                )
                return
            if parsed.path == "/api/logout":
                token = self.cookie_value("pws_session")
                if token in SESSIONS:
                    del SESSIONS[token]
                self.send_json({"ok": True}, headers={"Set-Cookie": "pws_session=; Path=/; Max-Age=0; HttpOnly; SameSite=Lax"})
                return
            if parsed.path.startswith("/api/") and not self.require_user():
                return
            username = self.current_user().get("username", "")
            if parsed.path == "/api/send":
                result = send_mail(payload, username)
            elif parsed.path == "/api/mail-send":
                payload["body_html"] = ""
                result = send_mail(payload, username)
            elif parsed.path == "/api/generate":
                result = generate_document(payload, username)
            elif parsed.path == "/api/optimize":
                result = optimize_text(payload)
            elif parsed.path == "/api/agent":
                result = agent_chat(payload, username)
            elif parsed.path == "/api/skill-test":
                if not self.require_superadmin():
                    return
                result = skill_test(payload, username)
            elif parsed.path == "/api/agent-config":
                if not self.require_superadmin():
                    return
                result = save_agent_config(payload)
            elif parsed.path == "/api/upload-history":
                result = upload_history_reports(payload, username)
            elif parsed.path == "/api/delete-report":
                result = delete_report_file(payload, username)
            elif parsed.path == "/api/delete-history":
                result = delete_history_report(payload, username)
            elif parsed.path == "/api/change-password":
                result = change_password(payload, username)
            elif parsed.path == "/api/profile":
                result = save_user_profile(payload, username)
            elif parsed.path == "/api/mail-config":
                result = save_user_mail_config(username, payload)
            elif parsed.path == "/api/test-mail-config":
                result = test_user_mail_config(username)
            elif parsed.path == "/api/admin-config":
                if not self.require_admin():
                    return
                result = save_admin_config(payload)
            elif parsed.path == "/api/server-config":
                if not self.require_superadmin():
                    return
                result = save_server_config(payload)
            elif parsed.path == "/api/admin-models":
                if not self.require_admin():
                    return
                result = list_admin_models(payload)
            elif parsed.path == "/api/admin-test-model":
                if not self.require_admin():
                    return
                result = test_admin_model(payload)
            elif parsed.path == "/api/admin-users":
                if not self.require_admin():
                    return
                result = add_user(payload, username)
            elif parsed.path == "/api/admin-users-delete":
                if not self.require_admin():
                    return
                result = delete_user(payload, username)
            elif parsed.path == "/api/admin-users-update":
                if not self.require_admin():
                    return
                result = update_user(payload, username)
            elif parsed.path == "/api/diary/save":
                result = save_diary(payload, username)
            elif parsed.path == "/api/diary/delete":
                result = delete_diary(payload.get("date", ""), username)
            elif parsed.path == "/api/diary/summarize":
                result = summarize_diaries_for_weekly(payload, username)
            elif parsed.path == "/api/forum/create":
                result = forum_create_topic(payload, username)
            elif parsed.path == "/api/forum/comment":
                result = forum_add_comment(payload, username)
            elif parsed.path == "/api/forum/like":
                result = forum_toggle_like(payload, username)
            elif parsed.path == "/api/forum/ai-topic":
                result = forum_ai_topic(payload, username)
            elif parsed.path == "/api/forum/ai-comment":
                result = forum_ai_comment(payload, username)
            elif parsed.path == "/api/news/config":
                if not self.require_superadmin():
                    return
                result = save_news_config(payload)
            elif parsed.path == "/api/news/generate":
                if not self.require_superadmin():
                    return
                result = generate_news_issue(payload, username)
            else:
                self.send_json({"error": "Not found"}, status=404)
                return
            self.send_json(result)
        except Exception as exc:
            self.send_json({"ok": False, "error": str(exc)}, status=400)


def main():
    load_agent_config()
    port = int(os.getenv("PORT", "8765"))
    server = ThreadingHTTPServer(("127.0.0.1", port), Handler)
    threading.Thread(target=news_auto_worker, daemon=True).start()
    print(f"个人工作报告邮件助手已启动：http://127.0.0.1:{port}")
    print(f"英文相对访问地址：{APP_RELATIVE_PATH}")
    print(f"完整访问地址：http://127.0.0.1:{port}{APP_RELATIVE_PATH}")
    print(f"用户数据目录：{USER_DATA_DIR}")
    print(f"共享模板目录：{REPORT_DIR}")
    print("按 Ctrl+C 停止服务")
    server.serve_forever()


if __name__ == "__main__":
    main()
