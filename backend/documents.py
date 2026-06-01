"""Uploads, prefill logic, weekly/trip document generation, drafts, and sending.

This module is loaded by backend.runtime into one shared application namespace.
Keep feature code here grouped by responsibility; cross-feature functions remain available
through the runtime during this incremental modularization.
"""

def safe_uploaded_name(name, kind):
    name = Path(name or "").name.strip()
    if not name:
        raise ValueError("上传文件缺少文件名")
    suffix = Path(name).suffix.lower()
    if kind == "weekly":
        if suffix not in {".xlsx", ".xls"}:
            raise ValueError("周报只支持 .xlsx 或 .xls 文件")
        if "工作周报" not in name:
            name = "工作周报-" + name
    elif kind == "trip":
        if suffix not in {".docx", ".md"}:
            raise ValueError("出差报告只支持 .docx 或 .md 文件")
        if not name.startswith("出差报告"):
            name = "出差报告-" + name
    else:
        raise ValueError("请选择上传类型")
    return re.sub(r"[/:\\\\]+", "_", name)


def unique_report_path(name, username=None):
    kind = report_kind_from_name(name)
    base = user_report_kind_dir(username, kind) if username else REPORT_DIR / safe_report_kind(kind)
    base.mkdir(parents=True, exist_ok=True)
    return base / name


def generated_report_path(name, username=None, kind=None):
    report_kind = kind or report_kind_from_name(name)
    base = user_generated_kind_dir(username, report_kind) if username else GENERATED_DIR / safe_report_kind(report_kind)
    base.mkdir(parents=True, exist_ok=True)
    return base / name


def configured_template_path(kind, username=None):
    path = report_template_path(kind)
    return path if path.exists() and path.is_file() else None


def template_for_report(kind, username=None, include_generated=False):
    configured = configured_template_path(kind, username)
    if configured:
        return configured
    finder = newest_any if include_generated else newest
    item = finder(kind, username, fallback_shared=True)
    path = Path((item or {}).get("path", ""))
    return path if path.exists() and path.is_file() else None


def report_kind_from_name(name):
    name = Path(name or "").name
    suffix = Path(name).suffix.lower()
    if "工作周报" in name and suffix in {".xlsx", ".xls"}:
        return "weekly"
    if "出差报告" in name and suffix in {".docx", ".md"}:
        return "trip"
    return "other"


def normalize_report_period_name(period):
    text = str(period or "").strip()

    def repl(match):
        return f"{int(match.group(1))}.{int(match.group(2))}.{int(match.group(3))}"

    text = re.sub(r"(20[0-9]{2})[.\-/年]0?([0-9]{1,2})[.\-/月]0?([0-9]{1,2})(?:日)?", repl, text)
    return re.sub(r"[^0-9A-Za-z.\-\u4e00-\u9fff]+", "", text)


def report_template_info(username=None):
    templates = {}
    for kind in ("weekly", "trip"):
        path = configured_template_path(kind, username)
        templates[kind] = {
            "kind": kind,
            "configured": bool(path),
            "name": path.name if path else "",
            "mtime": path.stat().st_mtime if path else None,
            "download_url": f"/download-template?kind={kind}",
        }
    return {"ok": True, "templates": templates}


def save_report_template(payload, username=None):
    kind = safe_report_kind(payload.get("kind"))
    if kind not in {"weekly", "trip"}:
        raise ValueError("请选择模板类型")
    file_item = payload.get("file") or {}
    name = Path(file_item.get("name", "")).name
    suffix = Path(name).suffix.lower()
    if kind == "weekly" and suffix not in {".xlsx", ".xls"}:
        raise ValueError("周报模板只支持 .xlsx 或 .xls 文件")
    if kind == "trip" and suffix != ".docx":
        raise ValueError("出差报告模板只支持 .docx 文件")
    data = file_item.get("data", "")
    if "," in data:
        data = data.split(",", 1)[1]
    raw = base64.b64decode(data)
    if not raw:
        raise ValueError("模板文件内容为空")
    path = report_template_path(kind)
    if kind == "weekly" and suffix == ".xls":
        path = path.with_suffix(".xls")
    path.parent.mkdir(parents=True, exist_ok=True)
    for old in path.parent.glob("template.*"):
        if old != path:
            old.unlink()
    path.write_bytes(raw)
    return {"ok": True, "template": {"kind": kind, "name": path.name, "path": str(path), "mtime": path.stat().st_mtime}}


def delete_report_template(payload, username=None):
    kind = safe_report_kind(payload.get("kind"))
    removed = []
    for path in report_template_path(kind).parent.glob("template.*"):
        if path.is_file():
            removed.append(path.name)
            path.unlink()
    return {"ok": True, "deleted": removed, "templates": report_template_info(username)["templates"]}


def upload_history_reports(payload, username=None):
    kind = payload.get("kind")
    files = payload.get("files") or []
    if not files:
        raise ValueError("请选择需要上传的文件")
    uploaded = []
    for item in files:
        name = safe_uploaded_name(item.get("name", ""), kind)
        data = item.get("data", "")
        if "," in data:
            data = data.split(",", 1)[1]
        raw = base64.b64decode(data)
        if not raw:
            raise ValueError(f"{name} 文件内容为空")
        path = unique_report_path(name, username)
        path.write_bytes(raw)
        uploaded.append({"name": path.name, "path": str(path), "size": len(raw)})
    return {"ok": True, "uploaded": uploaded}


def delete_history_report(payload, username=None):
    path = history_path_by_name(payload.get("name", ""), username)
    if path is None:
        raise ValueError("历史报告文件不存在")
    name = path.name
    path.unlink()
    return {"ok": True, "deleted": name}


def delete_report_file(payload, username=None):
    path = attachment_path_by_name(payload.get("name", ""), username)
    if path is None:
        raise ValueError("报告文件不存在，或不是当前账号的可删除文件")
    name = path.name
    path.unlink()
    return {"ok": True, "deleted": name}


def promote_sent_report(file_name, username=None):
    safe_name = Path(file_name or "").name
    if not safe_name:
        return None
    kind = report_kind_from_name(safe_name)
    source = user_generated_kind_dir(username, kind) / safe_name if username else GENERATED_DIR / safe_report_kind(kind) / safe_name
    if not source.exists() or not source.is_file():
        legacy_source = user_generated_dir(username) / safe_name if username else GENERATED_DIR / safe_name
        if legacy_source.exists() and legacy_source.is_file():
            source = legacy_source
        else:
            return None
    target = user_report_kind_dir(username, kind) / safe_name if username else REPORT_DIR / safe_report_kind(kind) / safe_name
    target.parent.mkdir(parents=True, exist_ok=True)
    if target.exists():
        target.unlink()
    shutil.move(str(source), str(target))
    return {"kind": kind, "name": target.name, "path": str(target)}


def estimated_text_lines(text, width_chars):
    lines = 0
    for part in str(text or "").splitlines() or [""]:
        visual_len = sum(2 if ord(ch) > 127 else 1 for ch in part)
        lines += max(1, (visual_len + width_chars - 1) // width_chars)
    return lines


def adjust_row_height(ws, row_idx, columns):
    width_chars = {
        2: 40,
        3: 30,
        4: 72,
        5: 24,
        6: 72,
        (3, 4): 68,
        (4, 5): 92,
        (3, 4, 5): 112,
    }
    max_lines = 1
    for col in columns:
        width = width_chars.get(col, 32)
        for merged in ws.merged_cells.ranges:
            if merged.min_row == row_idx and merged.max_row == row_idx and merged.min_col == col:
                key = tuple(range(merged.min_col, merged.max_col + 1))
                width = width_chars.get(key, width)
                break
        max_lines = max(max_lines, estimated_text_lines(ws.cell(row_idx, col).value, width))
    ws.row_dimensions[row_idx].height = min(320, max(34, max_lines * 16 + 8))


def dependency_error_message(package_name, feature_name):
    return f"{feature_name}功能缺少依赖库 {package_name}，请在服务端安装后重试"


def xlsx_column_index(cell_ref):
    letters = re.sub(r"[^A-Z]", "", str(cell_ref or "").upper())
    value = 0
    for letter in letters:
        value = value * 26 + (ord(letter) - ord("A") + 1)
    return value


def xlsx_xml_cell_text(cell, shared_strings):
    ns = {"a": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
    cell_type = cell.attrib.get("t", "")
    if cell_type == "inlineStr":
        return xml_text(cell.find("a:is", ns)).strip()
    value = cell.find("a:v", ns)
    if value is None or value.text is None:
        return ""
    if cell_type == "s":
        try:
            idx = int(value.text)
        except ValueError:
            return ""
        return shared_strings[idx].strip() if 0 <= idx < len(shared_strings) else ""
    return value.text.strip()


def load_xlsx_sheet_rows(path):
    ns = {"a": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
    with zipfile.ZipFile(path) as book:
        shared_strings = []
        if "xl/sharedStrings.xml" in book.namelist():
            root = ET.fromstring(book.read("xl/sharedStrings.xml"))
            shared_strings = [xml_text(item).strip() for item in root.findall(".//a:si", ns)]
        sheets = sorted(
            [name for name in book.namelist() if name.startswith("xl/worksheets/sheet") and name.endswith(".xml")]
        )
        if not sheets:
            return {}, 0
        root = ET.fromstring(book.read(sheets[0]))
        rows = {}
        for row in root.findall(".//a:sheetData/a:row", ns):
            row_idx = int(row.attrib.get("r", "0") or 0)
            rows[row_idx] = {}
            for cell in row.findall("a:c", ns):
                col_idx = xlsx_column_index(cell.attrib.get("r", ""))
                if col_idx:
                    rows[row_idx][col_idx] = xlsx_xml_cell_text(cell, shared_strings)
        return rows, (max(rows) if rows else 0)


def build_weekly_prefill(max_row, read_cell, source_name):
    def find_section_row(keyword):
        for row in range(1, max_row + 1):
            text = "".join(str(read_cell(row, col) or "") for col in range(1, 9))
            if keyword in text:
                return row
        return 0

    def next_section_row(after_row):
        candidates = [
            row
            for row in (find_section_row("重点工作跟进"), find_section_row("下周工作计划"))
            if row > after_row
        ]
        return min(candidates) if candidates else max_row + 1

    follow_section = find_section_row("重点工作跟进")
    next_section = find_section_row("下周工作计划")

    summary = []
    summary_rows = []
    next_start = next_section + 2 if next_section else 18
    next_end = next_section_row(next_section) - 1 if next_section else min(max_row, 26)
    for row in range(next_start, min(max_row, next_end) + 1):
        category = str(read_cell(row, 2) or "").strip()
        content = normalize_numbered_text(read_cell(row, 3))
        if category in {"工作分类", "三、下周工作计划"} or content == "工作内容":
            continue
        if category or content:
            summary.append(xlsx_row_text([category, content, "", ""]))
            summary_rows.append({"category": category, "content": content, "status": "", "plan": ""})

    follow = []
    follow_rows = []
    follow_start = follow_section + 2 if follow_section else 11
    follow_end = next_section - 1 if next_section else min(max_row, 15)
    for row in range(follow_start, min(max_row, follow_end) + 1):
        category = str(read_cell(row, 2) or "").strip()
        content = normalize_numbered_text(read_cell(row, 3))
        progress = normalize_numbered_text(read_cell(row, 4))
        difficulty = normalize_numbered_text(read_cell(row, 6))
        if category in {"工作分类", "二、重点工作跟进", "三、下周工作计划"} or content == "工作内容":
            continue
        if category or content or progress or difficulty:
            follow.append(xlsx_row_text([category, content, progress, difficulty]))
            follow_rows.append(
                {
                    "category": category,
                    "content": content,
                    "progress": progress,
                    "difficulty": difficulty,
                }
            )

    # 下周工作计划：工作分类默认与本周工作总结一致，工作内容留空待填。
    next_rows = []
    next_lines = []
    for item in summary_rows:
        category = item.get("category", "")
        if not category:
            continue
        next_rows.append({"category": category, "content": "", "difficulty": ""})
        next_lines.append(xlsx_row_text([category, "", ""]))

    return {
        "weekly_summary": "\n".join(summary),
        "weekly_follow": "\n".join(follow),
        "weekly_next": "\n".join(next_lines),
        "summary_rows": summary_rows,
        "follow_rows": follow_rows,
        "next_rows": next_rows,
        "source": source_name,
    }


def empty_weekly_prefill():
    return {
        "weekly_summary": "",
        "weekly_follow": "",
        "weekly_next": "",
        "summary_rows": [],
        "follow_rows": [],
        "next_rows": [],
        "source": "",
    }


def user_weekly_prefill_source(username=None):
    if not username:
        return None

    def newest_weekly(items):
        weekly = [item for item in items if item.get("kind") == "weekly"]
        if not weekly:
            return None
        return sorted(weekly, key=lambda item: (item["sort_key"], item["mtime"]), reverse=True)[0]

    # 优先取已归档的真实历史周报（reports/weekly，发送后由 promote_sent_report 移入），
    # 避免被本周尚未发送的 generated 草稿盖过；没有历史时才回退到 generated 输出。
    return newest_weekly(report_files(username)) or newest_weekly(generated_files(username))


def weekly_prefill(username=None):
    item = user_weekly_prefill_source(username)
    if not item:
        return empty_weekly_prefill()
    template = Path(item.get("path", ""))
    if not template.exists():
        return empty_weekly_prefill()

    try:
        from openpyxl import load_workbook
    except ModuleNotFoundError as exc:
        if exc.name != "openpyxl":
            raise
        if template.suffix.lower() != ".xlsx":
            return {
                "weekly_summary": "",
                "weekly_follow": "",
                "weekly_next": "",
                "summary_rows": [],
                "follow_rows": [],
                "next_rows": [],
                "source": template.name,
                "error": dependency_error_message("openpyxl", "历史周报预填"),
            }
        rows, max_row = load_xlsx_sheet_rows(template)
        return build_weekly_prefill(max_row, lambda row, col: rows.get(row, {}).get(col, ""), template.name)

    wb = load_workbook(template, data_only=True)
    ws = wb[wb.sheetnames[0]]
    return build_weekly_prefill(ws.max_row, lambda row, col: cell_value(ws, row, col), template.name)


def table_cell_text(table, row, col):
    try:
        return normalize_numbered_text(table.cell(row, col).text.strip())
    except Exception:
        return ""


def split_trip_date(date_text):
    text = date_text or ""
    match = re.search(r"(\d{4})年(\d{1,2})月(\d{1,2})日\s*至\s*(\d{4})年(\d{1,2})月(\d{1,2})日", text)
    if match:
        start = f"{int(match.group(1)):04d}-{int(match.group(2)):02d}-{int(match.group(3)):02d}"
        end = f"{int(match.group(4)):04d}-{int(match.group(5)):02d}-{int(match.group(6)):02d}"
        return start, end
    match = re.search(r"(\d{4})[-/.](\d{1,2})[-/.](\d{1,2}).*?(\d{4})[-/.](\d{1,2})[-/.](\d{1,2})", text)
    if match:
        start = f"{int(match.group(1)):04d}-{int(match.group(2)):02d}-{int(match.group(3)):02d}"
        end = f"{int(match.group(4)):04d}-{int(match.group(5)):02d}-{int(match.group(6)):02d}"
        return start, end
    return "", ""


def format_trip_date_text(start, end):
    def fmt(value):
        value = str(value or "").strip()
        match = re.match(r"(\d{4})[-/.]?(\d{1,2})[-/.]?(\d{1,2})$", value)
        if not match:
            return value
        return f"{int(match.group(1)):04d}年{int(match.group(2)):02d}月{int(match.group(3)):02d}日"
    start_text = fmt(start)
    end_text = fmt(end)
    if start_text and end_text:
        return f"{start_text}至{end_text}"
    return start_text or end_text


def trip_prefill(username=None):
    template = Path((newest("trip", username, fallback_shared=True) or {}).get("path", ""))
    if not template.exists() or template.suffix.lower() != ".docx":
        return {"source": ""}

    try:
        from docx import Document
    except ModuleNotFoundError as exc:
        if exc.name != "docx":
            raise
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        with zipfile.ZipFile(template) as docx_file:
            root = ET.fromstring(docx_file.read("word/document.xml"))
        table = root.find(".//w:tbl", ns)
        if table is None:
            return {"source": template.name}
        rows = []
        for row in table.findall("w:tr", ns):
            rows.append([normalize_numbered_text(xml_text(cell).strip()) for cell in row.findall("w:tc", ns)])

        def get_cell(row_idx, col_idx):
            try:
                return rows[row_idx][col_idx]
            except Exception:
                return ""

    else:
        doc = Document(template)
        if not doc.tables:
            return {"source": template.name}
        table = doc.tables[0]

        def get_cell(row_idx, col_idx):
            return table_cell_text(table, row_idx, col_idx)

    date_text = get_cell(1, 1)
    start, end = split_trip_date(date_text)
    return {
        "source": template.name,
        "reporter": get_cell(0, 1) or "周颖超",
        "department": get_cell(0, 3) or "场景研究院",
        "location": get_cell(0, 5),
        "trip_start": start,
        "trip_end": end,
        "trip_date_text": date_text,
        "purpose": get_cell(2, 1),
        "itinerary": get_cell(3, 1),
        "details": get_cell(4, 1),
        "issues": get_cell(5, 1),
        "suggestions": get_cell(6, 1),
    }


def generate_weekly(payload, username=None):
    try:
        from openpyxl import Workbook, load_workbook
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    except ModuleNotFoundError as exc:
        if exc.name == "openpyxl":
            raise ValueError(dependency_error_message("openpyxl", "周报生成")) from exc
        raise

    template = Path((newest_any("weekly", username, fallback_shared=True) or {}).get("path", ""))

    period = (payload.get("period") or datetime.now().strftime("%Y.%m.%d-%Y.%m.%d")).strip()
    safe_period = normalize_report_period_name(period)
    output = generated_report_path(f"{safe_display_name(username)}工作周报{safe_period}.xlsx", username, "weekly")
    if template.exists() and template.is_file():
        if template.resolve() == output.resolve():
            fallback = newest("weekly", username, fallback_shared=True)
            if fallback and Path(fallback.get("path", "")).resolve() != output.resolve():
                template = Path(fallback["path"])
        shutil.copy2(template, output)
    else:
        output.parent.mkdir(parents=True, exist_ok=True)
        Workbook().save(output)

    wb = load_workbook(output)
    ws = wb[wb.sheetnames[0]]

    for mc in list(ws.merged_cells.ranges):
        ws.unmerge_cells(str(mc))
    ws.delete_rows(1, max(ws.max_row, 1))

    ws.sheet_view.showGridLines = False
    ws.freeze_panes = None
    widths = {"A": 3, "B": 30, "C": 34, "D": 34, "E": 25, "F": 32}
    for col, width in widths.items():
        ws.column_dimensions[col].width = width

    summary_rows = parse_table_lines(payload.get("weekly_summary", ""), 4) or [
        ["", "", "", ""]
    ]
    follow_rows = parse_table_lines(payload.get("weekly_follow", ""), 4) or [
        ["", "", "", ""]
    ]
    next_rows = parse_table_lines(payload.get("weekly_next", ""), 3) or [
        ["", "", ""]
    ]

    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    title_fill = PatternFill("solid", fgColor="00A9D6")
    white_fill = PatternFill("solid", fgColor="FFFFFF")
    header_font = Font(name="宋体", size=12, bold=True)
    body_font = Font(name="宋体", size=11, bold=True)
    title_font = Font(name="宋体", size=20, bold=True, color="000000")

    def merge_safe(min_r, min_c, max_r, max_c):
        try:
            ws.merge_cells(start_row=min_r, start_column=min_c, end_row=max_r, end_column=max_c)
        except Exception:
            pass

    def write_and_style(row_idx, col_idx, value, halign="left", valign="center", font=None, fill=None, normalize=True):
        cell = ws.cell(row_idx, col_idx)
        cell.value = normalize_numbered_text(value) if normalize else value
        cell.border = thin_border
        cell.font = font or body_font
        if fill:
            cell.fill = fill
        cell.alignment = Alignment(wrap_text=True, horizontal=halign, vertical=valign)

    def write_section_title(row_idx, title):
        for col_idx in range(2, 7):
            write_and_style(row_idx, col_idx, "", halign="left", font=header_font, fill=white_fill)
        write_and_style(row_idx, 2, title, halign="left", font=header_font, fill=white_fill)
        merge_safe(row_idx, 2, row_idx, 6)
        ws.row_dimensions[row_idx].height = 32

    def write_header(row_idx, labels, merges=None):
        for col_idx in range(2, 7):
            write_and_style(row_idx, col_idx, "", halign="center", font=header_font, fill=white_fill, normalize=False)
        for col_idx, label in labels:
            write_and_style(row_idx, col_idx, label, halign="center", font=header_font, fill=white_fill, normalize=False)
        for merge in merges or []:
            merge_safe(row_idx, merge[0], row_idx, merge[1])
        ws.row_dimensions[row_idx].height = 28

    def write_title(row_idx):
        for col_idx in range(2, 7):
            write_and_style(row_idx, col_idx, "", halign="center", font=title_font, fill=title_fill, normalize=False)
        write_and_style(row_idx, 2, f"工作周报（{period}）", halign="center", font=title_font, fill=title_fill, normalize=False)
        merge_safe(row_idx, 2, row_idx, 6)
        ws.row_dimensions[row_idx].height = 44

    write_title(1)

    summary_title = 2
    summary_header = 3
    summary_start = 4
    write_section_title(summary_title, "一、本周工作总结")
    write_header(summary_header, [(2, "工作分类"), (3, "工作内容"), (5, "上周内容完成情况"), (6, "后续计划")], [(3, 4)])

    # 本周工作总结: B=工作分类 C=工作内容 E=完成情况 F=后续计划
    for idx, row in enumerate(summary_rows, start=summary_start):
        write_and_style(idx, 2, row[0], halign="center")
        write_and_style(idx, 3, row[1])
        write_and_style(idx, 5, row[2])
        write_and_style(idx, 6, row[3])
        adjust_row_height(ws, idx, (2, 3, 5, 6))
        merge_safe(idx, 3, idx, 4)

    follow_title = summary_start + len(summary_rows)
    follow_header = follow_title + 1
    follow_start = follow_header + 1
    write_section_title(follow_title, "二、重点工作跟进")
    write_header(follow_header, [(2, "工作分类"), (3, "工作内容"), (4, "当前进展"), (6, "困难与求助")], [(4, 5)])

    # 重点工作跟进: B=工作分类 C=工作内容 D=当前进展 F=困难与求助
    for idx, row in enumerate(follow_rows, start=follow_start):
        write_and_style(idx, 2, row[0], halign="center")
        write_and_style(idx, 3, row[1])
        write_and_style(idx, 4, row[2])
        write_and_style(idx, 6, row[3])
        adjust_row_height(ws, idx, (2, 3, 4, 6))
        merge_safe(idx, 4, idx, 5)

    next_title = follow_start + len(follow_rows)
    next_header = next_title + 1
    next_start = next_header + 1
    write_section_title(next_title, "三、下周工作计划")
    write_header(next_header, [(2, "工作分类"), (3, "工作内容"), (6, "困难与求助")], [(3, 5)])

    # 下周工作计划: B=工作分类 C=工作内容 F=困难与求助
    for idx, row in enumerate(next_rows, start=next_start):
        write_and_style(idx, 2, row[0], halign="center")
        write_and_style(idx, 3, row[1])
        write_and_style(idx, 6, row[2])
        adjust_row_height(ws, idx, (2, 3, 6))
        merge_safe(idx, 3, idx, 5)

    merge_safe(2, 2, 2, 6)

    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.print_area = f"B1:F{next_start + len(next_rows) - 1}"

    wb.save(output)
    return output


def clear_docx_row_height(row):
    from docx.oxml.ns import qn

    tr_pr = row._tr.get_or_add_trPr()
    for height in list(tr_pr.findall(qn("w:trHeight"))):
        tr_pr.remove(height)


def set_cell_text(cell, text, bold=False, compact=True):
    from docx.shared import Pt
    from docx.oxml.ns import qn

    text = normalize_numbered_text(text)
    parts = text.splitlines() or [""]
    while len(cell.paragraphs) < len(parts):
        cell.add_paragraph()
    for index, part in enumerate(parts):
        paragraph = cell.paragraphs[index]
        for run in list(paragraph.runs):
            paragraph._p.remove(run._r)
        run = paragraph.add_run(part)
        run.bold = bold
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(11)
        paragraph.paragraph_format.line_spacing = 1.08 if compact else 1.2
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2 if compact else 4)
    for paragraph in cell.paragraphs[len(parts):]:
        paragraph._element.getparent().remove(paragraph._element)


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_NS = "http://www.w3.org/XML/1998/namespace"
ET.register_namespace("w", W_NS)


def w_tag(name):
    return f"{{{W_NS}}}{name}"


def set_docx_xml_cell_text(cell, text):
    text = normalize_numbered_text(text)
    lines = text.splitlines() or [""]
    paragraphs = cell.findall(w_tag("p"))
    paragraph_template = deepcopy(paragraphs[0]) if paragraphs else ET.Element(w_tag("p"))
    paragraph_props = paragraph_template.find(w_tag("pPr"))
    run_template = paragraph_template.find(".//" + w_tag("r"))
    run_props = deepcopy(run_template.find(w_tag("rPr"))) if run_template is not None and run_template.find(w_tag("rPr")) is not None else None

    for paragraph in paragraphs:
        cell.remove(paragraph)

    for line in lines:
        paragraph = ET.Element(w_tag("p"))
        if paragraph_props is not None:
            paragraph.append(deepcopy(paragraph_props))
        run = ET.SubElement(paragraph, w_tag("r"))
        if run_props is not None:
            run.append(deepcopy(run_props))
        text_node = ET.SubElement(run, w_tag("t"))
        text_node.set(f"{{{XML_NS}}}space", "preserve")
        text_node.text = line
        cell.append(paragraph)


def generate_trip_from_docx_template(template, output, values):
    with zipfile.ZipFile(template, "r") as source:
        document_xml = source.read("word/document.xml")
        root = ET.fromstring(document_xml)
        table = root.find(".//" + w_tag("tbl"))
        if table is None:
            raise ValueError("出差报告模板中没有找到表格")
        rows = table.findall(w_tag("tr"))

        def cell(row_idx, actual_col_idx):
            return rows[row_idx].findall(w_tag("tc"))[actual_col_idx]

        cell_map = {
            (0, 1): values.get("reporter", ""),
            (0, 3): values.get("department", "场景研究院"),
            (0, 5): values.get("location", ""),
            (1, 1): values.get("date_text", ""),
            (2, 1): values.get("purpose", ""),
            (3, 1): values.get("itinerary", ""),
            (4, 1): values.get("details", ""),
            (5, 1): values.get("issues", ""),
            (6, 1): values.get("suggestions", ""),
        }
        for (row_idx, col_idx), value in cell_map.items():
            set_docx_xml_cell_text(cell(row_idx, col_idx), value)

        output.parent.mkdir(parents=True, exist_ok=True)
        temp_output = output.with_suffix(output.suffix + ".tmp")
        with zipfile.ZipFile(temp_output, "w", zipfile.ZIP_DEFLATED) as target:
            for item in source.infolist():
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True) if item.filename == "word/document.xml" else source.read(item.filename)
                target.writestr(item, data)
        temp_output.replace(output)


def generate_trip_default_docx(output, values):
    def text_node(parent, text):
        node = ET.SubElement(parent, w_tag("t"))
        node.set(f"{{{XML_NS}}}space", "preserve")
        node.text = str(text or "")

    def paragraph(text, bold=False):
        p = ET.Element(w_tag("p"))
        r = ET.SubElement(p, w_tag("r"))
        if bold:
            r_pr = ET.SubElement(r, w_tag("rPr"))
            ET.SubElement(r_pr, w_tag("b"))
        text_node(r, text)
        return p

    def cell(text, grid_span=1, bold=False):
        tc = ET.Element(w_tag("tc"))
        tc_pr = ET.SubElement(tc, w_tag("tcPr"))
        if grid_span > 1:
            span = ET.SubElement(tc_pr, w_tag("gridSpan"))
            span.set(w_tag("val"), str(grid_span))
        tc.append(paragraph(text, bold))
        return tc

    def row(*cells):
        tr = ET.Element(w_tag("tr"))
        for item in cells:
            tr.append(item)
        return tr

    document = ET.Element(w_tag("document"))
    body = ET.SubElement(document, w_tag("body"))
    body.append(paragraph("出差报告", True))
    table = ET.SubElement(body, w_tag("tbl"))
    tbl_pr = ET.SubElement(table, w_tag("tblPr"))
    borders = ET.SubElement(tbl_pr, w_tag("tblBorders"))
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = ET.SubElement(borders, w_tag(side))
        border.set(w_tag("val"), "single")
        border.set(w_tag("sz"), "4")
        border.set(w_tag("space"), "0")
        border.set(w_tag("color"), "BFC7D1")
    grid = ET.SubElement(table, w_tag("tblGrid"))
    for width in ("1300", "1300", "1000", "1700", "1300", "2100"):
        col = ET.SubElement(grid, w_tag("gridCol"))
        col.set(w_tag("w"), width)

    table.extend(
        [
            row(cell("报告人", bold=True), cell(values.get("reporter", "")), cell("部门", bold=True), cell(values.get("department", "")), cell("出差地点", bold=True), cell(values.get("location", ""))),
            row(cell("出差时间", bold=True), cell(values.get("date_text", ""), 5)),
            row(cell("出差目的", bold=True), cell(values.get("purpose", ""), 5)),
            row(cell("行程概览", bold=True), cell(values.get("itinerary", ""), 5)),
            row(cell("工作详情", bold=True), cell(values.get("details", ""), 5)),
            row(cell("问题与反馈", bold=True), cell(values.get("issues", ""), 5)),
            row(cell("总结与建议", bold=True), cell(values.get("suggestions", ""), 5)),
        ]
    )
    sect_pr = ET.SubElement(body, w_tag("sectPr"))
    pg_sz = ET.SubElement(sect_pr, w_tag("pgSz"))
    pg_sz.set(w_tag("w"), "11906")
    pg_sz.set(w_tag("h"), "16838")
    pg_mar = ET.SubElement(sect_pr, w_tag("pgMar"))
    pg_mar.set(w_tag("top"), "1440")
    pg_mar.set(w_tag("right"), "1440")
    pg_mar.set(w_tag("bottom"), "1440")
    pg_mar.set(w_tag("left"), "1440")

    output.parent.mkdir(parents=True, exist_ok=True)
    temp_output = output.with_suffix(output.suffix + ".tmp")
    with zipfile.ZipFile(temp_output, "w", zipfile.ZIP_DEFLATED) as target:
        target.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
            "</Types>",
        )
        target.writestr(
            "_rels/.rels",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
            "</Relationships>",
        )
        target.writestr("word/_rels/document.xml.rels", '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>')
        target.writestr("word/document.xml", ET.tostring(document, encoding="utf-8", xml_declaration=True))
    temp_output.replace(output)


def generate_trip(payload, username=None):
    template = Path(template_for_report("trip", username) or "")

    start = (payload.get("trip_start") or datetime.now().strftime("%Y%m%d")).replace("-", "")
    end = (payload.get("trip_end") or start).replace("-", "")[-4:]
    reporter = display_name_for_user(username)
    output = generated_report_path(f"出差报告-{start}-{end}-{safe_display_name(username)}.docx", username, "trip")
    date_text = payload.get("trip_date_text") or format_trip_date_text(payload.get("trip_start", ""), payload.get("trip_end", ""))
    values = {
        "reporter": reporter,
        "department": payload.get("department", "场景研究院"),
        "location": payload.get("location", ""),
        "date_text": date_text,
        "purpose": payload.get("purpose", ""),
        "itinerary": payload.get("itinerary", ""),
        "details": payload.get("details", ""),
        "issues": payload.get("issues", ""),
        "suggestions": payload.get("suggestions", ""),
    }
    if template.exists() and template.suffix.lower() == ".docx":
        generate_trip_from_docx_template(template, output, values)
    else:
        generate_trip_default_docx(output, values)
    return output


def generate_document(payload, username=None):
    kind = payload.get("kind")
    if kind == "weekly":
        path = generate_weekly(payload, username)
    elif kind == "trip":
        path = generate_trip(payload, username)
    else:
        raise ValueError("未知生成类型")
    draft = compose_draft(kind, path.name, username)
    return {
        "ok": True,
        "file": path.name,
        "path": str(path),
        "draft": draft,
    }


def save_draft(payload, username=None):
    draft_dir = user_draft_dir(username) if username else DRAFT_DIR
    draft_dir.mkdir(parents=True, exist_ok=True)
    msg = build_message(payload, username)
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    safe_subject = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff_-]+", "_", payload.get("subject", "draft"))[:40]
    path = draft_dir / f"{stamp}-{safe_subject}.eml"
    path.write_bytes(bytes(msg))
    return path


def send_mail(payload, username=None):
    account = str(payload.get("mail_account") or "company").strip()
    settings = smtp_settings_for_account(username, account)
    if not settings["host"]:
        draft = save_draft(payload, username)
        return {"ok": True, "mode": "draft", "message": f"邮件未发出：当前账号未配置 SMTP 服务器，已生成邮件草稿：{draft}"}
    validate_smtp_ready(settings)

    msg = build_message(payload, username, account)
    recipients = split_addresses(payload.get("to", "")) + split_addresses(payload.get("cc", ""))
    if not recipients:
        raise ValueError("请填写收件人邮箱")

    if settings["use_ssl"]:
        with smtplib.SMTP_SSL(settings["host"], settings["port"], context=ssl.create_default_context(), timeout=20) as smtp:
            if settings["user"]:
                smtp.login(settings["user"], settings["password"])
            smtp.send_message(msg, to_addrs=recipients)
    else:
        with smtplib.SMTP(settings["host"], settings["port"], timeout=20) as smtp:
            if settings["use_tls"]:
                smtp.starttls(context=ssl.create_default_context())
            if settings["user"]:
                smtp.login(settings["user"], settings["password"])
            smtp.send_message(msg, to_addrs=recipients)
    promoted = promote_sent_report(payload.get("attachment", ""), username)
    result = {"ok": True, "mode": "sent", "message": "邮件已发送"}
    if promoted:
        result["promoted"] = promoted
    return result
