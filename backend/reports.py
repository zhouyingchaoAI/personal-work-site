"""Report discovery, document previews, and draft composition.

This module is loaded by backend.runtime into one shared application namespace.
Keep feature code here grouped by responsibility; cross-feature functions remain available
through the runtime during this incremental modularization.
"""

def parse_weekly_date(name):
    match = re.search(r"(?:([0-9]{2})年)?([0-9]{1,2})月([0-9]{1,2})日-([0-9]{1,2})月([0-9]{1,2})日", name)
    if match:
        year = 2000 + int(match.group(1) or "25")
        month1 = int(match.group(2))
        day1 = int(match.group(3))
        month2 = int(match.group(4))
        day2 = int(match.group(5))
        return (year, month2, day2, month1 * 100 + day1)
    match = re.search(r"(20[0-9]{2})[.\-/年]([0-9]{1,2})[.\-/月]([0-9]{1,2})(?:日)?[-~至]+(?:20[0-9]{2}[.\-/年])?([0-9]{1,2})[.\-/月]([0-9]{1,2})(?:日)?", name)
    if not match:
        return (0, 0, 0, 0)
    year = int(match.group(1))
    month1 = int(match.group(2))
    day1 = int(match.group(3))
    month2 = int(match.group(4))
    day2 = int(match.group(5))
    return (year, month2, day2, month1 * 100 + day1)


def parse_trip_date(name):
    match = re.search(r"出差报告-([0-9]{8})-([0-9]{4})", name)
    if not match:
        return (0, 0)
    return (int(match.group(1)), int(match.group(2)))


def report_storage_dirs(root, include_legacy=True):
    dirs = []
    for kind in ("weekly", "trip"):
        path = root / kind
        if path.exists():
            dirs.append((kind, path))
    if include_legacy:
        dirs.append((None, root))
    return dirs


def is_legacy_auto_numbered_report(name):
    return bool(re.search(r"-(?:生成|上传)[0-9]+(?=\.[^.]+$)", str(name or "")))


def report_files(username=None):
    files = []
    base = user_report_dir(username) if username else REPORT_DIR
    if not base.exists():
        return files
    seen = set()
    for expected_kind, directory in report_storage_dirs(base):
        for path in directory.iterdir():
            if not path.is_file() or path.name.startswith(".") or path in seen:
                continue
            seen.add(path)
            if is_legacy_auto_numbered_report(path.name):
                continue
            lower = path.suffix.lower()
            if "工作周报" in path.name and lower in {".xlsx", ".xls"}:
                kind = "weekly"
                sort_key = parse_weekly_date(path.name)
            elif path.name.startswith("出差报告") and lower in {".docx", ".md"}:
                kind = "trip"
                sort_key = parse_trip_date(path.name)
            else:
                continue
            if expected_kind and expected_kind != kind:
                continue
            files.append(
                {
                    "kind": kind,
                    "name": path.name,
                    "path": str(path),
                    "mtime": path.stat().st_mtime,
                    "sort_key": sort_key,
                    "deletable": bool(username),
                }
            )
    return files


def generated_files(username=None):
    files = []
    base = user_generated_dir(username) if username else GENERATED_DIR
    if not base.exists():
        return files
    seen = set()
    for expected_kind, directory in report_storage_dirs(base):
        for path in directory.iterdir():
            if not path.is_file() or path.name.startswith(".") or path in seen:
                continue
            seen.add(path)
            lower = path.suffix.lower()
            if "工作周报" in path.name and lower in {".xlsx", ".xls"}:
                kind = "weekly"
            elif "出差报告" in path.name and lower in {".docx", ".md"}:
                kind = "trip"
            else:
                continue
            if expected_kind and expected_kind != kind:
                continue
            files.append(
                {
                    "kind": kind,
                    "name": path.name,
                    "path": str(path),
                    "mtime": path.stat().st_mtime,
                    "sort_key": (9999, int(path.stat().st_mtime)),
                    "generated": True,
                    "deletable": bool(username),
                }
            )
    return files


def all_files(username=None):
    user_files = generated_files(username) + report_files(username)
    if username and not any(not f.get("generated") for f in user_files):
        user_files = user_files + report_files(None)
    return user_files


def newest(kind, username=None, fallback_shared=False):
    items = [item for item in report_files(username) if item["kind"] == kind]
    if not items and fallback_shared and username:
        items = [item for item in report_files(None) if item["kind"] == kind]
    if not items and fallback_shared and username:
        for fallback_user in ("zhouyingchao", "admin"):
            if fallback_user == username:
                continue
            items = [item for item in report_files(fallback_user) if item["kind"] == kind]
            if items:
                break
    if not items:
        return None
    return sorted(items, key=lambda item: (item["sort_key"], item["mtime"]), reverse=True)[0]


def newest_any(kind, username=None, fallback_shared=False):
    items = [item for item in all_files(username) if item["kind"] == kind]
    if not items and fallback_shared and username:
        items = [item for item in all_files(None) if item["kind"] == kind]
    if not items and kind == "weekly":
        for fallback_user in ("zhouyingchao", "admin"):
            items = [item for item in all_files(fallback_user) if item["kind"] == kind]
            if items:
                break
    if not items:
        return None
    return sorted(items, key=lambda item: (item["sort_key"], item["mtime"]), reverse=True)[0]


def xml_text(node):
    if node is None:
        return ""
    return "".join(node.itertext())


def preview_docx(path, max_chars=900):
    try:
        with zipfile.ZipFile(path) as docx:
            xml = docx.read("word/document.xml")
        root = ET.fromstring(xml)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        paragraphs = [xml_text(p).strip() for p in root.findall(".//w:p", ns)]
        text = "\n".join(p for p in paragraphs if p)
        return text[:max_chars]
    except Exception:
        return ""


def preview_xlsx(path, max_chars=900):
    try:
        with zipfile.ZipFile(path) as book:
            shared = []
            if "xl/sharedStrings.xml" in book.namelist():
                root = ET.fromstring(book.read("xl/sharedStrings.xml"))
                ns = {"a": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
                shared = [xml_text(si).strip() for si in root.findall(".//a:si", ns)]
            sheet_names = [n for n in book.namelist() if n.startswith("xl/worksheets/sheet") and n.endswith(".xml")]
            texts = []
            for sheet in sheet_names[:2]:
                root = ET.fromstring(book.read(sheet))
                for cell in root.iter("{http://schemas.openxmlformats.org/spreadsheetml/2006/main}c"):
                    value = cell.find("{http://schemas.openxmlformats.org/spreadsheetml/2006/main}v")
                    if value is None or value.text is None:
                        continue
                    if cell.attrib.get("t") == "s":
                        idx = int(value.text)
                        if idx < len(shared):
                            texts.append(shared[idx])
                    else:
                        texts.append(value.text)
            return "\n".join(t for t in texts if t)[:max_chars]
    except Exception:
        return ""


def table_style_html():
    return (
        "width:100%;border-collapse:collapse;margin:8px 0 14px 0;"
        "table-layout:fixed;"
        "font-family:-apple-system,BlinkMacSystemFont,Segoe UI,PingFang SC,Microsoft YaHei,sans-serif;"
        "font-size:13px;background:#fff;"
    )


def xlsx_original_table_html(path):
    try:
        from openpyxl import load_workbook
        from openpyxl.utils import get_column_letter

        wb = load_workbook(path, data_only=True)
        ws = wb[wb.sheetnames[0]]
        min_row, max_row = 1, ws.max_row
        min_col, max_col = 1, ws.max_column

        while max_row > 1 and not any(ws.cell(max_row, col).value not in (None, "") for col in range(min_col, max_col + 1)):
            max_row -= 1
        while max_col > 1 and not any(ws.cell(row, max_col).value not in (None, "") for row in range(min_row, max_row + 1)):
            max_col -= 1

        merged_starts = {}
        merged_covered = set()
        for area in ws.merged_cells.ranges:
            if area.max_row < min_row or area.min_row > max_row or area.max_col < min_col or area.min_col > max_col:
                continue
            row_span = area.max_row - area.min_row + 1
            col_span = area.max_col - area.min_col + 1
            merged_starts[(area.min_row, area.min_col)] = (row_span, col_span)
            for r in range(area.min_row, area.max_row + 1):
                for c in range(area.min_col, area.max_col + 1):
                    if (r, c) != (area.min_row, area.min_col):
                        merged_covered.add((r, c))

        parts = [f'<table class="raw-table" style="{table_style_html()}">']
        colgroup = ["<colgroup>"]
        for col in range(min_col, max_col + 1):
            letter = get_column_letter(col)
            width = ws.column_dimensions[letter].width or 12
            colgroup.append(f'<col style="width:{min(max(width * 8, 72), 220)}px">')
        colgroup.append("</colgroup>")
        parts.extend(colgroup)

        for row in range(min_row, max_row + 1):
            parts.append("<tr>")
            for col in range(min_col, max_col + 1):
                if (row, col) in merged_covered:
                    continue
                cell = ws.cell(row, col)
                value = "" if cell.value is None else str(cell.value)
                row_span, col_span = merged_starts.get((row, col), (1, 1))
                attrs = []
                if row_span > 1:
                    attrs.append(f'rowspan="{row_span}"')
                if col_span > 1:
                    attrs.append(f'colspan="{col_span}"')
                align = cell.alignment.horizontal or "left"
                valign = cell.alignment.vertical or "middle"
                weight = "font-weight:700;" if cell.font and cell.font.bold else ""
                fill = ""
                fg = getattr(cell.fill, "fgColor", None)
                if fg and fg.type == "rgb" and fg.rgb and fg.rgb not in {"00000000", "FFFFFFFF"}:
                    fill = f"background:#{fg.rgb[-6:]};"
                style = (
                    "border:1px solid #cfd6df;padding:7px 9px;white-space:pre-wrap;word-break:break-word;"
                    f"text-align:{align};vertical-align:{valign};{weight}{fill}"
                )
                parts.append(f'<td {" ".join(attrs)} style="{style}">{html_escape(value)}</td>')
            parts.append("</tr>")
        parts.append("</table>")
        return "".join(parts)
    except Exception:
        return ""


def docx_original_table_html(path):
    try:
        ns = {"w": W_NS}
        parts = []
        with zipfile.ZipFile(path) as docx:
            root = ET.fromstring(docx.read("word/document.xml"))
        for table in root.findall(".//w:tbl", ns):
            parts.append(f'<table class="raw-table" style="{table_style_html()}">')
            grid_cols = []
            for grid_col in table.findall("w:tblGrid/w:gridCol", ns):
                try:
                    grid_cols.append(int(grid_col.attrib.get(w_tag("w"), "0") or 0))
                except ValueError:
                    grid_cols.append(0)
            if grid_cols and sum(grid_cols) > 0:
                total = sum(grid_cols)
                parts.append("<colgroup>")
                for width in grid_cols:
                    parts.append(f'<col style="width:{width / total * 100:.3f}%">')
                parts.append("</colgroup>")
            for row in table.findall("w:tr", ns):
                parts.append("<tr>")
                for cell in row.findall("w:tc", ns):
                    grid_span = cell.find("w:tcPr/w:gridSpan", ns)
                    colspan = grid_span.attrib.get(w_tag("val"), "1") if grid_span is not None else "1"
                    paragraphs = []
                    for paragraph in cell.findall("w:p", ns):
                        text = "".join(node.text or "" for node in paragraph.findall(".//w:t", ns)).strip()
                        if text:
                            paragraphs.append(text)
                    text = html_escape("\n".join(paragraphs)).replace("\n", "<br>")
                    span_attr = f' colspan="{html_escape(colspan)}"' if colspan and colspan != "1" else ""
                    cell_style = (
                        "border:1px solid #cfd6df;padding:7px 9px;vertical-align:top;"
                        "white-space:pre-wrap;word-break:break-word;"
                    )
                    parts.append(
                        f'<td{span_attr} style="{cell_style}">'
                        f"{text}</td>"
                    )
                parts.append("</tr>")
            parts.append("</table>")
        if parts:
            return "".join(parts)
        paragraphs = []
        for paragraph in root.findall(".//w:p", ns):
            text = "".join(node.text or "" for node in paragraph.findall(".//w:t", ns)).strip()
            if text:
                paragraphs.append(html_escape(text))
        return "".join(f"<p>{text}</p>" for text in paragraphs)
    except Exception:
        return ""


def preview_file_html(path):
    suffix = path.suffix.lower()
    if suffix in {".xlsx", ".xls"}:
        return xlsx_original_table_html(path)
    if suffix == ".docx":
        return docx_original_table_html(path)
    if suffix == ".md":
        text = path.read_text(encoding="utf-8", errors="ignore")
        return f'<div style="white-space:pre-wrap;">{html_escape(text)}</div>'
    return ""


def preview_file(path):
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return preview_docx(path)
    if suffix == ".xlsx":
        return preview_xlsx(path)
    if suffix == ".md":
        return path.read_text(encoding="utf-8", errors="ignore")[:900]
    return ""


def html_escape(text):
    if not text:
        return ""
    return (
        str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def format_weekly_body(path):
    """从 xlsx 周报提取结构化摘要，返回 (纯文本, HTML)"""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(path, data_only=True)
        ws = wb[wb.sheetnames[0]]

        def get(r, c):
            v = ws.cell(r, c).value
            return str(v).strip() if v else ""

        def find_row(label, fallback):
            for r in range(1, ws.max_row + 1):
                if label in get(r, 2):
                    return r
            return fallback

        def valid_data_row(cat, content):
            if not cat and not content:
                return False
            invalid = {"工作分类", "工作内容", "一、本周工作总结", "二、重点工作跟进", "三、下周工作计划"}
            return cat not in invalid and content not in invalid

        summary_title_row = find_row("一、本周工作总结", 3)
        follow_title_row = find_row("二、重点工作跟进", 9)
        next_title_row = find_row("三、下周工作计划", 16)
        summary_range = range(summary_title_row + 2, follow_title_row)
        follow_range = range(follow_title_row + 2, next_title_row)
        next_range = range(next_title_row + 2, ws.max_row + 1)

        # ===== 纯文本版本 =====
        lines = []
        title = get(2, 2)
        if title:
            lines.append(title)
            lines.append("")

        # 一、本周工作总结: B=工作分类 C=工作内容 E=完成情况 F=后续计划
        lines.append("一、本周工作总结")
        for row in summary_range:
            cat = get(row, 2)
            content = get(row, 3)
            status = get(row, 5)
            plan = get(row, 6)
            if not valid_data_row(cat, content):
                continue
            if cat:
                lines.append(f"【{cat}】")
            if content:
                for ln in content.split("\n"):
                    ln = ln.strip()
                    if ln:
                        lines.append(f"  {ln}")
            if status:
                lines.append(f"  完成情况：{status}")
            if plan:
                lines.append(f"  后续计划：{plan}")
            lines.append("")

        # 二、重点工作跟进: B=工作分类 C=工作内容 D=当前进展 F=困难与求助
        lines.append("二、重点工作跟进")
        for row in follow_range:
            cat = get(row, 2)
            content = get(row, 3)
            progress = get(row, 4)
            difficulty = get(row, 6)
            if not valid_data_row(cat, content):
                continue
            if cat:
                lines.append(f"【{cat}】")
            if content:
                for ln in content.split("\n"):
                    ln = ln.strip()
                    if ln:
                        lines.append(f"  {ln}")
            if progress:
                lines.append(f"  当前进展：{progress}")
            if difficulty:
                lines.append(f"  困难与求助：{difficulty}")
            lines.append("")

        # 三、下周工作计划: B=工作分类 C=工作内容 F=困难与求助
        lines.append("三、下周工作计划")
        for row in next_range:
            cat = get(row, 2)
            content = get(row, 3)
            difficulty = get(row, 6)
            if not valid_data_row(cat, content):
                continue
            if cat:
                lines.append(f"【{cat}】")
            if content:
                for ln in content.split("\n"):
                    ln = ln.strip()
                    if ln:
                        lines.append(f"  {ln}")
            if difficulty:
                lines.append(f"  困难与求助：{difficulty}")
            lines.append("")

        text_body = "\n".join(lines).strip()

        # ===== HTML 表格版本 =====
        h = []
        if title:
            h.append(f'<h3 style="margin:0 0 12px 0;font-size:16px;">{html_escape(title)}</h3>')

        th_style = "border:1px solid #ccc;padding:6px 8px;text-align:left;background:#f5f5f5;font-size:13px;"
        td_style = "border:1px solid #ccc;padding:6px 8px;text-align:left;font-size:13px;vertical-align:top;"
        table_style = "width:100%;border-collapse:collapse;margin-bottom:16px;font-family:-apple-system,BlinkMacSystemFont,Segoe UI,PingFang SC,Microsoft YaHei,sans-serif;"

        def make_table(headers, rows):
            parts = ['<table style="' + table_style + '">']
            parts.append('<tr>')
            for hd in headers:
                parts.append(f'<th style="{th_style}">{html_escape(hd)}</th>')
            parts.append('</tr>')
            for row in rows:
                if not any(row):
                    continue
                parts.append('<tr>')
                for cell in row:
                    cell_html = html_escape(cell).replace("\n", "<br>")
                    parts.append(f'<td style="{td_style}">{cell_html}</td>')
                parts.append('</tr>')
            parts.append('</table>')
            return "".join(parts)

        # 本周工作总结
        h.append('<p style="font-size:14px;font-weight:bold;margin:12px 0 6px 0;">一、本周工作总结</p>')
        summary_rows = []
        for row in summary_range:
            cat = get(row, 2)
            content = get(row, 3)
            status = get(row, 5)
            plan = get(row, 6)
            if valid_data_row(cat, content):
                summary_rows.append([cat, content, status, plan])
        h.append(make_table(["工作分类", "工作内容", "完成情况", "后续计划"], summary_rows))

        # 重点工作跟进
        h.append('<p style="font-size:14px;font-weight:bold;margin:12px 0 6px 0;">二、重点工作跟进</p>')
        follow_rows = []
        for row in follow_range:
            cat = get(row, 2)
            content = get(row, 3)
            progress = get(row, 4)
            difficulty = get(row, 6)
            if valid_data_row(cat, content):
                follow_rows.append([cat, content, progress, difficulty])
        h.append(make_table(["工作分类", "工作内容", "当前进展", "困难与求助"], follow_rows))

        # 下周工作计划
        h.append('<p style="font-size:14px;font-weight:bold;margin:12px 0 6px 0;">三、下周工作计划</p>')
        next_rows = []
        for row in next_range:
            cat = get(row, 2)
            content = get(row, 3)
            difficulty = get(row, 6)
            if valid_data_row(cat, content):
                next_rows.append([cat, content, difficulty])
        h.append(make_table(["工作分类", "工作内容", "困难与求助"], next_rows))

        html_body = "".join(h)
        return text_body, html_body
    except Exception:
        return "", ""


def format_trip_body(path):
    """从 docx 出差报告提取结构化摘要，返回 (纯文本, HTML)"""
    try:
        from docx import Document
        doc = Document(path)
        if not doc.tables:
            return "", ""
        table = doc.tables[0]

        def cell_text(r, c):
            try:
                return table.cell(r, c).text.strip()
            except Exception:
                return ""

        reporter = cell_text(0, 1)
        department = cell_text(0, 3)
        location = cell_text(0, 5)
        date_text = cell_text(1, 1)
        purpose = cell_text(2, 1)
        itinerary = cell_text(3, 1)
        details = cell_text(4, 1)
        issues = cell_text(5, 1)
        suggestions = cell_text(6, 1)

        # ===== 纯文本版本 =====
        lines = []
        if reporter:
            lines.append(f"报告人：{reporter}")
        if department:
            lines.append(f"部门：{department}")
        if location:
            lines.append(f"出差地点：{location}")
        if date_text:
            lines.append(f"出差时间：{date_text}")
        lines.append("")

        if purpose:
            lines.append("【出差目的】")
            lines.append(purpose)
            lines.append("")
        if itinerary:
            lines.append("【行程概览】")
            lines.append(itinerary)
            lines.append("")
        if details:
            lines.append("【工作详情】")
            lines.append(details)
            lines.append("")
        if issues:
            lines.append("【问题与困难】")
            lines.append(issues)
            lines.append("")
        if suggestions:
            lines.append("【改进建议】")
            lines.append(suggestions)
            lines.append("")

        text_body = "\n".join(lines).strip()

        # ===== HTML 表格版本 =====
        h = []
        h.append('<table style="width:100%;border-collapse:collapse;margin-bottom:16px;font-family:-apple-system,BlinkMacSystemFont,Segoe UI,PingFang SC,Microsoft YaHei,sans-serif;font-size:13px;">')

        info_rows = [
            ["报告人", reporter, "部门", department, "出差地点", location],
            ["出差时间", date_text, "", "", "", ""],
        ]
        td_label = "border:1px solid #ccc;padding:6px 8px;background:#f5f5f5;font-weight:bold;width:12%;"
        td_value = "border:1px solid #ccc;padding:6px 8px;"

        for row in info_rows:
            h.append('<tr>')
            h.append(f'<td style="{td_label}">{html_escape(row[0])}</td>')
            h.append(f'<td style="{td_value}" colspan="5">{html_escape(row[1])}</td>')
            h.append('</tr>')
        h.append('</table>')

        section_style = "font-size:14px;font-weight:bold;margin:12px 0 6px 0;"
        content_style = "font-size:13px;line-height:1.6;margin:0 0 12px 0;"

        if purpose:
            h.append(f'<p style="{section_style}">【出差目的】</p>')
            h.append(f'<p style="{content_style}">{html_escape(purpose).replace(chr(10), "<br>")}</p>')
        if itinerary:
            h.append(f'<p style="{section_style}">【行程概览】</p>')
            h.append(f'<p style="{content_style}">{html_escape(itinerary).replace(chr(10), "<br>")}</p>')
        if details:
            h.append(f'<p style="{section_style}">【工作详情】</p>')
            h.append(f'<p style="{content_style}">{html_escape(details).replace(chr(10), "<br>")}</p>')
        if issues:
            h.append(f'<p style="{section_style}">【问题与困难】</p>')
            h.append(f'<p style="{content_style}">{html_escape(issues).replace(chr(10), "<br>")}</p>')
        if suggestions:
            h.append(f'<p style="{section_style}">【改进建议】</p>')
            h.append(f'<p style="{content_style}">{html_escape(suggestions).replace(chr(10), "<br>")}</p>')

        html_body = "".join(h)
        return text_body, html_body
    except Exception:
        return "", ""


def compose_draft(kind, file_name=None, username=None):
    config = read_config()
    mail = user_mail_config(username)
    report = None
    if file_name:
        for item in all_files(username):
            if item["name"] == file_name:
                report = item
                break
    if report is None:
        report = newest(kind, username, fallback_shared=True)
    if report is None:
        return {"error": "没有找到对应报告"}

    path = Path(report["path"])
    preview = preview_file(path)
    preview_html = preview_file_html(path)
    today_text = datetime.now().strftime("%Y年%m月%d日")
    if kind == "weekly":
        subject = f"【周报】{path.stem}"
        summary_text, summary_html = format_weekly_body(path)
        weekly_table_html = preview_html or summary_html
        body = (
            f"{config.get('weekly_greeting', '领导您好：')}\n\n"
            f"附件为我的本周工作周报《{path.name}》，请查收。\n\n"
        )
        if summary_text:
            body += summary_text + "\n\n"
        body += f"{today_text}"
        body = append_email_signature(body, username)
        body_html = f'<p>{html_escape(config.get("weekly_greeting", "领导您好："))}</p><p>附件为我的本周工作周报《{html_escape(path.name)}》，请查收。</p>'
        if weekly_table_html:
            body_html += weekly_table_html
        body_html += f'<p>{html_escape(today_text)}</p>'
        body_html = append_email_signature_html(body_html, username)
        to_addr = mail.get("weekly_to", "")
        cc_addr = mail.get("weekly_cc", "")
    else:
        subject = f"【出差报告】{path.stem}"
        summary_text, summary_html = format_trip_body(path)
        trip_table_html = preview_html or summary_html
        body = (
            f"{config.get('trip_greeting', '领导您好：')}\n\n"
            f"附件为我的出差报告《{path.name}》，请查收。\n\n"
        )
        if summary_text:
            body += summary_text + "\n\n"
        body += f"{today_text}"
        body = append_email_signature(body, username)
        body_html = f'<p>{html_escape(config.get("trip_greeting", "领导您好："))}</p><p>附件为我的出差报告《{html_escape(path.name)}》，请查收。</p>'
        if trip_table_html:
            body_html += trip_table_html
        body_html += f'<p>{html_escape(today_text)}</p>'
        body_html = append_email_signature_html(body_html, username)
        to_addr = mail.get("trip_to", "")
        cc_addr = mail.get("trip_cc", "")

    return {
        "kind": kind,
        "subject": subject,
        "body": body,
        "body_html": body_html,
        "to": to_addr,
        "cc": cc_addr,
        "attachment": report["name"],
        "attachment_path": report["path"],
        "download_url": "/download?file=" + urllib.parse.quote(report["name"]),
        "preview": preview,
        "preview_html": preview_html,
        "email_signature": email_signature_for_user(username),
    }
